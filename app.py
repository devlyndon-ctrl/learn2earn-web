from flask import (
    Flask, render_template, redirect, url_for, request,
    flash, jsonify, session, send_from_directory
)
from werkzeug.utils import secure_filename
from supabase import create_client
from flask_mail import Mail, Message

import os
import re
import string
import random
import secrets
import uuid
import requests
import nltk
import threading
import spacy
from textblob import TextBlob
from collections import Counter
import numpy as np
from nltk.sentiment import SentimentIntensityAnalyzer
from nltk.tokenize import word_tokenize, sent_tokenize
from nltk.corpus import stopwords, words
from langdetect import detect, LangDetectException
from datetime import datetime, timedelta, timezone
import pytz
import json
import language_tool_python
import hashlib
from functools import lru_cache
from concurrent.futures import ThreadPoolExecutor
import logging
import traceback
from typing import Tuple, Dict, Any, Optional, List
from difflib import SequenceMatcher
import PyPDF2
from docx import Document
import tempfile

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = Flask(__name__)
import os
from dotenv import load_dotenv

load_dotenv()

# Load environment variables
SUPABASE_URL = os.getenv('SUPABASE_URL')
SUPABASE_KEY = os.getenv('SUPABASE_KEY')
HTTPSMS_API_KEY = os.getenv('HTTPSMS_API_KEY')
FLASK_SECRET_KEY = os.getenv('FLASK_SECRET_KEY')

app = Flask(__name__)
app.secret_key = FLASK_SECRET_KEY

app.config['MAIL_SERVER'] = os.getenv('MAIL_SERVER')
app.config['MAIL_PORT'] = int(os.getenv('MAIL_PORT', 587))
app.config['MAIL_USE_TLS'] = os.getenv('MAIL_USE_TLS', 'True').lower() == 'true'
app.config['MAIL_USERNAME'] = os.getenv('MAIL_USERNAME')
app.config['MAIL_PASSWORD'] = os.getenv('MAIL_PASSWORD')
mail = Mail(app)

# Initialize Supabase client
supabase = create_client(SUPABASE_URL, SUPABASE_KEY)

# Import science datasets and quiz generator
try:
    from datasets.science.science_grade7_dataset import SCIENCE_GRADE7_QUESTIONS
    from datasets.science.science_grade8_dataset import SCIENCE_GRADE8_QUESTIONS
    from datasets.science.science_grade9_dataset import SCIENCE_GRADE9_QUESTIONS
    from datasets.science.science_grade10_dataset import SCIENCE_GRADE10_QUESTIONS
    from datasets.science.quiz_generator import ScienceQuizGenerator
    quiz_generator = ScienceQuizGenerator(SCIENCE_GRADE7_QUESTIONS, SCIENCE_GRADE8_QUESTIONS, SCIENCE_GRADE9_QUESTIONS, SCIENCE_GRADE10_QUESTIONS)
except ImportError as e:
    SCIENCE_GRADE7_QUESTIONS = []
    SCIENCE_GRADE8_QUESTIONS = []
    SCIENCE_GRADE9_QUESTIONS = []
    SCIENCE_GRADE10_QUESTIONS = []
    quiz_generator = None

# Import English datasets and quiz generator
try:
    from datasets.english.english_grade7_dataset import ENGLISH_GRADE7_QUESTIONS
    from datasets.english.english_grade8_dataset import ENGLISH_GRADE8_QUESTIONS
    from datasets.english.english_grade9_dataset import ENGLISH_GRADE9_QUESTIONS
    from datasets.english.english_grade10_dataset import ENGLISH_GRADE10_QUESTIONS
    from datasets.science.quiz_generator import EnglishQuizGenerator
    from datasets.science.subject_detector import detect_subject_from_lesson
    english_quiz_generator = EnglishQuizGenerator(ENGLISH_GRADE7_QUESTIONS, ENGLISH_GRADE8_QUESTIONS, ENGLISH_GRADE9_QUESTIONS, ENGLISH_GRADE10_QUESTIONS)
except ImportError as e:
    ENGLISH_GRADE7_QUESTIONS = []
    ENGLISH_GRADE8_QUESTIONS = []
    ENGLISH_GRADE9_QUESTIONS = []
    ENGLISH_GRADE10_QUESTIONS = []
    english_quiz_generator = None
    detect_subject_from_lesson = None

philippines_tz = pytz.timezone('Asia/Manila')
now_ph = datetime.now(philippines_tz).isoformat()

UPLOAD_FOLDER = os.path.join(os.path.dirname(__file__), 'static', 'uploads')
ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif', 'pdf', 'docx', 'txt'}
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size

# Thread pool for parallel processing
executor = ThreadPoolExecutor(max_workers=4)

# Download required NLTK data
try:
    nltk.data.find('sentiment/vader_lexicon.zip')
except LookupError:
    nltk.download('vader_lexicon')

try:
    nltk.data.find('corpora/words')
except LookupError:
    nltk.download('words')

try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt')

try:
    nltk.data.find('corpora/stopwords')
except LookupError:
    nltk.download('stopwords')

# Load spaCy model for advanced NLP
try:
    nlp = spacy.load("en_core_web_sm")
except:
    import subprocess
    subprocess.run(["python", "-m", "spacy", "download", "en_core_web_sm"])
    nlp = spacy.load("en_core_web_sm")

from datetime import datetime

otp_storage = {}

def generate_otp(length=6):
    """Generate a random OTP"""
    return ''.join(random.choices(string.digits, k=length))

def get_teacher_name_with_prefix(teacher_data):
    """Get teacher name with appropriate gender prefix"""
    prefix = 'Mr.'
    if teacher_data.get('gender', '').lower() == 'female':
        prefix = 'Mrs.'
    elif teacher_data.get('gender', '').lower() == 'other':
        prefix = 'Mx.'
    
    return f"{prefix} {teacher_data['first_name']} {teacher_data['last_name']}".strip()

def safe_execute(query):
    try:
        return query.execute()
    except Exception as e:
        print(f"Safe execute error: {e}")
        return type('obj', (object,), {'data': []})()

def format_date(dt_str):
    if not dt_str:
        return ''
    try:
        # Parse the datetime string (assume UTC if no timezone info)
        dt = datetime.fromisoformat(dt_str.replace('Z', '+00:00'))
        # Convert to Asia/Manila timezone
        manila_tz = pytz.timezone('Asia/Manila')
        if dt.tzinfo is None:
            dt = pytz.utc.localize(dt)
        dt = dt.astimezone(manila_tz)
        return dt.strftime('%b %d, %Y • %I:%M %p')
    except Exception:
        return dt_str

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF file"""
    text = ""
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                except Exception as e:
                    logger.warning(f"Error extracting page text: {e}")
        return text
    except Exception as e:
        logger.error(f"PDF extraction error: {e}")
        raise

def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX file"""
    text = ""
    try:
        doc = Document(file_path)
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + "\n"
        
        return text
    except Exception as e:
        logger.error(f"DOCX extraction error: {e}")
        raise

def extract_text_from_file(file) -> str:
    """Extract text from uploaded file"""
    filename = file.filename.lower()
    
    # Save temporarily
    temp_dir = tempfile.gettempdir()
    temp_path = os.path.join(temp_dir, secure_filename(file.filename))
    file.save(temp_path)
    
    try:
        if filename.endswith('.pdf'):
            text = extract_text_from_pdf(temp_path)
        elif filename.endswith('.docx'):
            text = extract_text_from_docx(temp_path)
        elif filename.endswith('.txt'):
            with open(temp_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
        else:
            raise ValueError("Unsupported file format")
        
        # Cleanup
        if os.path.exists(temp_path):
            os.remove(temp_path)
        
        return ' '.join(text.split())  # Normalize whitespace
        
    except Exception as e:
        if os.path.exists(temp_path):
            os.remove(temp_path)
        logger.error(f"File extraction error: {e}")
        raise

def generate_password(length=12):
    """Generate a secure random password"""
    # Define character sets
    lowercase = string.ascii_lowercase
    uppercase = string.ascii_uppercase
    numbers = string.digits
    special_chars = "!@#$%^&*()_+-=[]{}|;:,.<>?"
    
    # Ensure at least one of each character type
    password = [
        secrets.choice(lowercase),
        secrets.choice(uppercase),
        secrets.choice(numbers),
        secrets.choice(special_chars)
    ]
    
    # Fill the rest of the password
    all_chars = lowercase + uppercase + numbers + special_chars
    password.extend(secrets.choice(all_chars) for _ in range(length - 4))
    
    # Shuffle the password
    secrets.SystemRandom().shuffle(password)
    
    return ''.join(password)

def send_password_reset_email(user_email, user_name, new_password):
    """Send password reset notification email to user"""
    try:
        subject = "🔐 Password Reset - Masico National High School Learn2Earn"
        
        email_body = f"""Dear {user_name},

Your password has been reset by the administration.

Here are your new login credentials:
• Email: {user_email}
• New Password: {new_password}

For security reasons, we recommend that you:
1. Log in to your account immediately
2. Change your password after logging in
3. Keep your login credentials secure

You can access your account through:
• Web Portal: [Your Website URL]
• Mobile App: Learn2Earn Mobile Application

If you did not request this password reset or have any concerns, please contact the administration immediately.

Stay secure and happy learning!

Best regards,
Masico National High School
Learn2Earn Administration Team"""

        msg = Message(
            subject=subject,
            sender=app.config['MAIL_USERNAME'],
            recipients=[user_email]
        )
        msg.body = email_body
        
        mail.send(msg)
        app.logger.info(f"Password reset email sent to {user_email}")
        return True
        
    except Exception as e:
        app.logger.error(f"Failed to send password reset email to {user_email}: {str(e)}")
        return False

def send_welcome_email(email, password, first_name, role):
    try:
        subject = "Welcome to Learn2Earn!"
        if role.lower() == "student":
            body = f"""
            Hi {first_name},

            Your account has been created for Learn2Earn.

            Email: {email}
            Password: {password}

            Please change your password after your first login.
            """
        else:
            body = f"""
            Hi {first_name},

            Your account has been created for Learn2Earn.

            Email: {email}
            Password: {password}

            Please change your password after your first login.
            """

        msg = Message(subject, recipients=[email], body=body)
        mail.send(msg)
    except Exception as e:
        app.logger.error(f"Failed to send welcome email: {str(e)}")

def send_approval_email(student):
    """Send approval email to student"""
    try:
        student_name = f"{student['first_name']} {student['last_name']}"
        grade_level = student.get('year_level', 'Unknown')
        section_name = student.get('section', 'Unknown')
        
        subject = "🎉 Welcome to Masico National High School - Learn2earn!"
        
        email_body = f"""Dear {student_name},

We are pleased to inform you that your registration for Masico National High School - Learn2Earn has been approved!

Welcome to our learning community! Here are your details:
• Grade: {grade_level}
• Section: {section_name}
• Status: Active

You can now access your student account and begin your educational journey with us.

For mobile app access:
1. Download the Learn2Earn mobile app
2. Login with your registered email
3. Start Earning points!

If you have any questions or need assistance, please don't hesitate to contact your class advisor or visit the administration office.

We're excited to have you as part of our school family!

Best regards,
Masico National High School
Learn2Earn Administration Team"""

        msg = Message(
            subject=subject,
            sender=app.config['MAIL_USERNAME'],
            recipients=[student['email']]
        )
        msg.body = email_body
        
        mail.send(msg)
        app.logger.info(f"Approval email sent to {student['email']}")
        
    except Exception as e:
        app.logger.error(f"Failed to send approval email to {student['email']}: {str(e)}")
        raise

def send_rejection_email(student):
    """Send rejection email to student"""
    try:
        student_name = f"{student['first_name']} {student['last_name']}"
        
        subject = "Masico National High School - Learn2Earn Registration Update"
        
        email_body = f"""Dear {student_name},

Thank you for your interest in joining Masico Masico National High School - Learn2Earn .

After careful review, we regret to inform you that your registration application has not been approved at this time.

This decision may be due to various factors including:
• Documentation requirements
• Grade level capacity
• Other administrative considerations

If you believe this decision was made in error or would like to inquire about the specific reasons, please contact our administration office during school hours.

We appreciate your understanding and encourage you to explore other educational opportunities that may better suit your needs.

Thank you for considering Masico National High School.

Sincerely,
Masico National High School
Admissions Office"""

        msg = Message(
            subject=subject,
            sender=app.config['MAIL_USERNAME'],
            recipients=[student['email']]
        )
        msg.body = email_body
        
        mail.send(msg)
        app.logger.info(f"Rejection email sent to {student['email']}")
        
    except Exception as e:
        app.logger.error(f"Failed to send rejection email to {student['email']}: {str(e)}")
        raise

############################# QUARTERLY REWARD REDEMPTION HELPERS #############################

def get_current_quarter():
    """Get the current quarter based on the current date"""
    try:
        now = datetime.now(philippines_tz)
        now_str = now.isoformat()
        
        # Get all quarters and find the active one.
        # Priority:
        # 1) Explicit DB status == active
        # 2) Date-range match fallback
        quarters_result = supabase.table('quarters').select('*').order('start_date').execute()
        
        if not quarters_result.data:
            app.logger.warning("No quarters found in database")
            return None

        # 1) Prefer explicit status flag when present
        explicit_active = [q for q in quarters_result.data if str(q.get('status', '')).lower() == 'active']
        if explicit_active:
            active_quarter = explicit_active[0]
            active_quarter['status'] = 'active'  # Ensure status is set
            app.logger.info(
                f"Found active quarter by status: {active_quarter.get('quarter_name')} "
                f"({active_quarter.get('school_year')})"
            )
            return active_quarter
        
        # 2) Fallback: find active quarter by checking date ranges
        for quarter in quarters_result.data:
            start_date = quarter.get('start_date', '')
            end_date = quarter.get('end_date', '')
            
            # Normalize date strings for comparison (remove time component)
            now_date = now_str.split('T')[0]  # Get just the date part
            start_date_part = start_date.split('T')[0] if isinstance(start_date, str) else str(start_date)
            end_date_part = end_date.split('T')[0] if isinstance(end_date, str) else str(end_date)
            
            # Compare dates as strings (YYYY-MM-DD format works for string comparison)
            if start_date_part <= now_date <= end_date_part:
                app.logger.info(f"Found active quarter: {quarter.get('quarter_name')} ({start_date_part} to {end_date_part})")
                quarter['status'] = 'active'  # Ensure status is set
                return quarter
        
        app.logger.warning(f"No active quarter found for date {now_date}")
        return None
    except Exception as e:
        app.logger.error(f"Error getting current quarter: {str(e)}")
        return None

def get_current_school_year():
    """Get the current school year intelligently handling gap periods between school years"""
    try:
        active_quarter = get_current_quarter()
        if active_quarter and active_quarter.get('school_year'):
            return active_quarter.get('school_year')

        now = datetime.now(philippines_tz)
        now_date = now.strftime('%Y-%m-%d')
        
        # Get all quarters ordered by school_year descending (latest first)
        quarters_result = supabase.table('quarters').select('*').order('school_year', desc=True).execute()
        
        if not quarters_result.data:
            app.logger.warning("No quarters found in database")
            return None
        
        # Try to find active quarter first
        for quarter in quarters_result.data:
            start_date = quarter.get('start_date', '').split('T')[0]
            end_date = quarter.get('end_date', '').split('T')[0]
            
            if start_date <= now_date <= end_date:
                app.logger.info(f"Found active quarter in school year: {quarter['school_year']}")
                return quarter['school_year']
        
        # If no active quarter found (gap period), get the NEXT upcoming school year
        # or the latest one if we're before all quarters
        quarters_sorted_by_start = sorted(quarters_result.data, key=lambda q: q.get('start_date', ''))
        
        for quarter in quarters_sorted_by_start:
            start_date = quarter.get('start_date', '').split('T')[0]
            if start_date > now_date:
                # Found upcoming quarter, return its school year
                app.logger.info(f"In gap period, using upcoming school year: {quarter['school_year']}")
                return quarter['school_year']
        
        # If current date is after all quarters, use the latest school year
        latest_quarter = quarters_result.data[0]
        app.logger.info(f"Past all quarters, using latest school year: {latest_quarter['school_year']}")
        return latest_quarter['school_year']
        
    except Exception as e:
        app.logger.error(f"Error getting current school year: {str(e)}")
        return None

def get_redemption_count_for_quarter(student_id, reward_id, quarter_id):
    """
    Get the number of times a student has redeemed a specific reward in a quarter
    Returns: integer count
    """
    try:
        # Get the quarter details
        quarter_result = supabase.table('quarters').select('*').eq('quarter_id', quarter_id).execute()
        if not quarter_result.data:
            return 0
        
        quarter = quarter_result.data[0]
        start_date = quarter['start_date']
        end_date = quarter['end_date']
        
        # Ensure dates are in ISO format for proper comparison
        # Parse if they're strings
        if isinstance(start_date, str):
            start_date = start_date.split('T')[0] + 'T00:00:00'
        if isinstance(end_date, str):
            # Set end_date to end of day
            end_date = end_date.split('T')[0] + 'T23:59:59'
        
        app.logger.info(f"Checking redemptions for student {student_id}, reward {reward_id}, from {start_date} to {end_date}")
        
        # Count redemptions for this student and reward within the quarter period
        redemptions = supabase.table('reward_redemptions') \
            .select('redemption_id', count='exact') \
            .eq('student_id', student_id) \
            .eq('reward_id', reward_id) \
            .gte('processed_at', start_date) \
            .lte('processed_at', end_date) \
            .execute()
        
        count = redemptions.count if redemptions.count is not None else 0
        app.logger.info(f"Redemption count for student {student_id} on reward {reward_id}: {count}")
        
        return count
    except Exception as e:
        app.logger.error(f"Error getting redemption count: {str(e)}")
        return 0

def check_can_redeem_reward(student_id, reward_id):
    """
    Check if a student can still redeem a reward in the current quarter
    Returns: dict with 'can_redeem' (bool), 'count' (current count), 'remaining' (3 - count), 'quarter_id'
    """
    try:
        current_quarter = get_current_quarter()
        if not current_quarter:
            # If no active quarter, allow redemption
            return {
                'can_redeem': True,
                'count': 0,
                'remaining': 3,
                'quarter_id': None,
                'quarter_name': 'N/A'
            }
        
        quarter_id = current_quarter['quarter_id']
        count = get_redemption_count_for_quarter(student_id, reward_id, quarter_id)
        remaining = max(0, 3 - count)
        
        return {
            'can_redeem': remaining > 0,
            'count': count,
            'remaining': remaining,
            'quarter_id': quarter_id,
            'quarter_name': current_quarter.get('quarter_name', 'N/A')
        }
    except Exception as e:
        app.logger.error(f"Error checking can_redeem_reward: {str(e)}")
        return {
            'can_redeem': True,
            'count': 0,
            'remaining': 3,
            'quarter_id': None,
            'quarter_name': 'N/A'
        }

############################# ENHANCED NLP NOTIFICATION SYSTEM #############################

# Load English word dictionary
ENGLISH_WORDS = set(words.words())
stop_words = set(stopwords.words('english'))

# Expanded content filtering lists with variations
SEXUAL_TERMS = {
    'sexual', 'sex', 'porn', 'porno', 'xxx', 'nude', 'naked', 'erotic', 'erotica',
    'fuck', 'fucking', 'fucker', 'shit', 'asshole', 'bitch', 'bastard', 'dick',
    'cock', 'pussy', 'whore', 'slut', 'cunt', 'nigga', 'nigger', 'fag', 'faggot',
    'dyke', 'tranny', 'retard', 'chink', 'spic', 'kike', 'penis', 'vagina', 'boobs',
    'boobies', 'tits', 'titties', 'breasts', 'anal', 'blowjob', 'handjob', 'masterbate',
    'masturbate', 'orgasm', 'cum', 'semen', 'ejaculate', 'horny', 'aroused',
    'milf', 'incest', 'pedo', 'child porn', 'lolita', 'barely legal'
}

RACIST_TERMS = {
    'nigger', 'nigga', 'chink', 'gook', 'spic', 'wetback', 'kike', 'heeb',
    'raghead', 'towelhead', 'cameljockey', 'sandnigger', 'beaner', 'coon',
    'darkie', 'porchmonkey', 'redskin', 'squaw', 'yellow', 'cracker', 'honky',
    'whitey', 'gyp', 'gypsy', 'jap', 'nip', 'mongoloid', 'oriental', 'ape',
    'monkey', 'slave', 'master race', 'white power', 'blackie'
}

HOMOPHOBIC_TERMS = {
    'fag', 'faggot', 'dyke', 'queer', 'homo', 'lesbo', 'tranny', 'shemale',
    'he-she', 'fruit', 'fairy', 'butch', 'femme', 'genderbender', 'it', 'thing',
    'transvestite', 'crossdresser', 'sodomite', 'bugger'
}

PROFANITY_TERMS = {
    'fuck', 'shit', 'ass', 'asshole', 'bitch', 'bastard', 'dick', 'cock',
    'pussy', 'whore', 'slut', 'cunt', 'damn', 'hell', 'crap', 'piss',
    'dickhead', 'motherfucker', 'bullshit', 'bollocks', 'wanker', 'twat',
    'arse', 'arsehole', 'bugger', 'sod', 'bloody', 'git', 'screw', 'screwing',
    'goddamn', 'christ sake', 'jesus christ'
}

# Combined set for quick lookups
ALL_RESTRICTED_TERMS = SEXUAL_TERMS | RACIST_TERMS | HOMOPHOBIC_TERMS | PROFANITY_TERMS

# Common evasive spellings and variations
EVASIVE_SPELLINGS = {
    'fuck': ['phuck', 'fuk', 'f u c k', 'f*ck', 'f**k', 'f--k', 'f_u_c_k'],
    'shit': ['shyt', 'shet', 'sh!t', 'sh*t', 's**t', 's--t', 's_h_i_t'],
    'ass': ['azz', 'as$', 'a$$', 'a**', 'a--s', 'a_s_s'],
    'bitch': ['biatch', 'b!tch', 'b*tch', 'b**ch', 'b--ch', 'b_i_t_c_h'],
    'nigger': ['n!gger', 'n*gger', 'n**ger', 'n--ger', 'n_i_g_g_e_r', 'nigga'],
}

class EnhancedFeedbackAnalyzer:
    def __init__(self):
        self.sia = SentimentIntensityAnalyzer()
        self.stop_words = set(stopwords.words('english'))
        self.english_words = set(words.words())
        
        # Initialize language tool for grammar checking
        try:
            self.language_tool = language_tool_python.LanguageTool('en-US')
        except:
            self.language_tool = None
        
        # Enhanced motivational message templates
        self.motivation_templates = {
            'exceptional': [
                "🌟 Exceptional work! Your deep understanding and thoughtful analysis set a brilliant example. Keep pushing boundaries!",
                "🎓 Outstanding performance! Your mastery of this topic is truly impressive. You're on a path to greatness!",
                "💫 Phenomenal effort! The clarity and depth of your thinking shows remarkable growth. Keep shining!",
                "🏆 You've exceeded expectations! Your dedication to excellence is inspiring to everyone around you."
            ],
            'very_positive': [
                "✨ Excellent progress! Your consistent effort and positive attitude are yielding great results. Keep it up!",
                "🚀 You're doing fantastic! Your engagement and enthusiasm are making a real difference in your learning.",
                "💪 Strong work! You're building important skills and showing real improvement. This momentum is powerful!",
                "🎯 Great job! Your focus and determination are paying off. Stay on this path to success!"
            ],
            'positive': [
                "📚 Good progress! Your effort is showing in your work. Keep building on this foundation!",
                "🌱 You're making nice improvements! Every step forward counts toward your goals. Keep going!",
                "⭐ Solid work! With continued practice and focus, you'll achieve even more. Stay motivated!",
                "🎈 You're on the right track! Your positive attitude is helping you learn and grow. Keep it up!"
            ],
            'slightly_positive': [
                "📝 You're making progress! A little more attention to detail will take you even further.",
                "🔍 Good start! Try to dig a bit deeper and you'll uncover even more understanding.",
                "📈 You're moving in the right direction! Keep building on what you've learned.",
                "🎯 You've got the basics down! Now focus on refining your understanding."
            ],
            'neutral': [
                "⚖️ Steady progress! You're building knowledge at a good pace. Stay consistent!",
                "📊 You're meeting expectations. Challenge yourself to go beyond the basics.",
                "🎯 You're on track! Consider how you can add more depth to your work.",
                "📚 Keep going! Every bit of effort adds to your learning journey."
            ],
            'slightly_negative': [
                "🔄 You're close! With a bit more effort on these areas, you'll improve significantly.",
                "🎯 Let's focus on strengthening these concepts. You have the ability to do better!",
                "💡 I see potential here. Let's work together on understanding these points more deeply.",
                "🌱 This is a learning opportunity. Identify what's challenging and tackle it step by step."
            ],
            'negative': [
                "🤔 Let's take a step back and review these concepts. Everyone faces challenges - it's how we grow!",
                "💪 You can do this! Let's break down what's difficult and work through it together.",
                "🎓 Don't be discouraged. Use this feedback to guide your improvement. You're capable of more!",
                "🌟 Remember: challenges help us grow. Let's identify the gaps and fill them together."
            ],
            'very_negative': [
                "🆘 This work needs significant improvement. Let's schedule a one-on-one session to discuss how we can get you back on track.",
                "⚠️ Your performance is below expectations. This is a critical moment to refocus and seek help.",
                "❗ Important: You're struggling with fundamental concepts. Let's create an improvement plan immediately.",
                "🔴 Your work shows serious gaps. Please reach out to me so we can address these issues right away."
            ]
        }
        
        # Constructive feedback templates based on common issues
        self.constructive_templates = {
            'too_brief': [
                "Your response is quite brief. Try to elaborate more on your thoughts and provide examples.",
                "Consider expanding your answer with more details and explanations.",
                "This is a good start! Add more depth by explaining your reasoning."
            ],
            'lacks_depth': [
                "You've touched on important points. Now dig deeper into the 'why' and 'how'.",
                "Good foundation! Build on it by connecting ideas and showing relationships.",
                "You're on the right track. Add more analysis and critical thinking."
            ],
            'unclear': [
                "Some of your ideas are unclear. Try organizing your thoughts more logically.",
                "Your message could be clearer. Break down complex ideas into simpler parts.",
                "I'm having trouble following your argument. Can you restate it more simply?"
            ],
            'off_topic': [
                "Your response seems to drift from the main topic. Let's refocus on the key question.",
                "Make sure you're addressing the specific question or prompt.",
                "Stay focused on the main topic. Your ideas are good but need to be more relevant."
            ],
            'grammar_issues': [
                "Some grammatical errors make your writing hard to follow. Proofread carefully.",
                "Work on sentence structure and grammar to communicate more effectively.",
                "Consider using simpler sentences to express your ideas clearly."
            ]
        }
    
    def analyze_sentiment_nuanced(self, feedback):
        """
        Enhanced sentiment analysis with more nuanced categories
        """
        # Get basic VADER scores
        vader_scores = self.sia.polarity_scores(feedback)
        compound = vader_scores['compound']
        
        # Use TextBlob for additional analysis
        blob = TextBlob(feedback)
        textblob_polarity = blob.sentiment.polarity
        textblob_subjectivity = blob.sentiment.subjectivity
        
        # Combined sentiment score (weighted average)
        combined_sentiment = (compound * 0.7 + textblob_polarity * 0.3)
        
        # More nuanced sentiment classification (9 levels)
        if combined_sentiment >= 0.8:
            sentiment_level = "Exceptional"
            sentiment_color = "primary"
            template_key = 'exceptional'
        elif combined_sentiment >= 0.6:
            sentiment_level = "Very Positive"
            sentiment_color = "success"
            template_key = 'very_positive'
        elif combined_sentiment >= 0.3:
            sentiment_level = "Positive"
            sentiment_color = "success"
            template_key = 'positive'
        elif combined_sentiment >= 0.1:
            sentiment_level = "Slightly Positive"
            sentiment_color = "info"
            template_key = 'slightly_positive'
        elif combined_sentiment > -0.1:
            sentiment_level = "Neutral"
            sentiment_color = "secondary"
            template_key = 'neutral'
        elif combined_sentiment > -0.3:
            sentiment_level = "Slightly Negative"
            sentiment_color = "warning"
            template_key = 'slightly_negative'
        elif combined_sentiment > -0.6:
            sentiment_level = "Negative"
            sentiment_color = "danger"
            template_key = 'negative'
        else:
            sentiment_level = "Very Negative"
            sentiment_color = "dark"
            template_key = 'very_negative'
        
        # Get appropriate motivational message
        motivation = random.choice(self.motivation_templates[template_key])
        
        # Analyze specific aspects
        aspects = self._analyze_aspects(feedback)
        
        # Get constructive feedback if needed
        constructive = self._get_constructive_feedback(feedback, aspects)
        
        # Extract key phrases
        key_phrases = self.extract_key_phrases(feedback)
        
        # Generate summary if feedback is long
        summary = None
        word_count = len(feedback.split())
        if word_count > 50:
            summary = self.generate_summary(feedback)
        
        # Calculate readability score
        readability = self._calculate_readability(feedback)
        
        return {
            'feedback': feedback,
            'sentiment': sentiment_level,
            'sentiment_color': sentiment_color,
            'scores': {
                'vader_positive': round(vader_scores['pos'], 3),
                'vader_neutral': round(vader_scores['neu'], 3),
                'vader_negative': round(vader_scores['neg'], 3),
                'vader_compound': round(vader_scores['compound'], 3),
                'textblob_polarity': round(textblob_polarity, 3),
                'textblob_subjectivity': round(textblob_subjectivity, 3),
                'combined_score': round(combined_sentiment, 3)
            },
            'aspects': aspects,
            'constructive_feedback': constructive,
            'motivation': motivation,
            'key_phrases': key_phrases,
            'summary': summary,
            'word_count': word_count,
            'sentence_count': len(sent_tokenize(feedback)),
            'complexity_score': self._calculate_complexity(feedback),
            'readability_score': readability,
            'feedback_hidden': sentiment_level in ['Negative', 'Very Negative']
        }
    
    def _analyze_aspects(self, text):
        """
        Analyze specific aspects of the feedback
        """
        aspects = {
            'effort': False,
            'understanding': False,
            'improvement': False,
            'engagement': False,
            'creativity': False,
            'critical_thinking': False
        }
        
        # Keywords for different aspects
        effort_keywords = ['effort', 'try', 'attempt', 'work hard', 'dedication', 'commitment', 'persistence']
        understanding_keywords = ['understand', 'comprehend', 'grasp', 'know', 'learn', 'concept', 'idea']
        improvement_keywords = ['improve', 'better', 'progress', 'growth', 'develop', 'advance', 'increase']
        engagement_keywords = ['engage', 'participate', 'interest', 'excitement', 'enthusiasm', 'passion']
        creativity_keywords = ['creative', 'innovative', 'imaginative', 'original', 'unique', 'novel']
        critical_keywords = ['analyze', 'evaluate', 'synthesize', 'critique', 'reason', 'logic', 'argument']
        
        text_lower = text.lower()
        
        for keyword in effort_keywords:
            if keyword in text_lower:
                aspects['effort'] = True
                break
        
        for keyword in understanding_keywords:
            if keyword in text_lower:
                aspects['understanding'] = True
                break
        
        for keyword in improvement_keywords:
            if keyword in text_lower:
                aspects['improvement'] = True
                break
        
        for keyword in engagement_keywords:
            if keyword in text_lower:
                aspects['engagement'] = True
                break
        
        for keyword in creativity_keywords:
            if keyword in text_lower:
                aspects['creativity'] = True
                break
        
        for keyword in critical_keywords:
            if keyword in text_lower:
                aspects['critical_thinking'] = True
                break
        
        return aspects
    
    def _get_constructive_feedback(self, text, aspects):
        """
        Generate constructive feedback based on analysis
        """
        issues = []
        
        # Check length
        words_list = text.split()
        if len(words_list) < 10:
            issues.append('too_brief')
        elif len(words_list) < 20:
            # Brief but acceptable
            pass
        
        # Check for complexity
        if self._calculate_complexity(text) < 0.3 and len(words_list) > 30:
            issues.append('lacks_depth')
        
        # Check for clarity (using spaCy if available)
        try:
            doc = nlp(text)
            if len(doc) > 0:
                # Check for vague language
                vague_terms = ['thing', 'stuff', 'something', 'somewhere', 'someone']
                if any(term in text.lower() for term in vague_terms):
                    issues.append('unclear')
        except:
            pass
        
        # Grammar check
        if self.language_tool:
            matches = self.language_tool.check(text)
            if len(matches) > 3:
                issues.append('grammar_issues')
        
        constructive = {}
        for issue in set(issues):  # Remove duplicates
            if issue in self.constructive_templates:
                constructive[issue] = random.choice(self.constructive_templates[issue])
        
        return constructive
    
    def _calculate_complexity(self, text):
        """
        Calculate text complexity score (0-1)
        """
        words_list = text.split()
        if not words_list:
            return 0
        
        # Average word length as simple complexity measure
        avg_word_length = sum(len(word) for word in words_list) / len(words_list)
        
        # Normalize to 0-1 (assuming max average length ~10)
        complexity = min(avg_word_length / 10, 1.0)
        
        return round(complexity, 2)
    
    def _calculate_readability(self, text):
        """
        Calculate Flesch Reading Ease score (simplified)
        """
        sentences = sent_tokenize(text)
        words_list = word_tokenize(text)
        
        if not sentences or not words_list:
            return 0
        
        # Average words per sentence
        avg_words_per_sentence = len(words_list) / len(sentences)
        
        # Average syllables per word (simplified)
        syllables = 0
        for word in words_list:
            syllables += self._count_syllables(word)
        avg_syllables_per_word = syllables / len(words_list)
        
        # Flesch Reading Ease formula: 206.835 - 1.015*(words/sentences) - 84.6*(syllables/words)
        readability = 206.835 - 1.015 * avg_words_per_sentence - 84.6 * avg_syllables_per_word
        
        # Normalize to 0-100
        readability = max(0, min(100, readability))
        
        return round(readability, 2)
    
    def _count_syllables(self, word):
        """
        Simple syllable counter
        """
        word = word.lower()
        count = 0
        vowels = 'aeiouy'
        if word[0] in vowels:
            count += 1
        for index in range(1, len(word)):
            if word[index] in vowels and word[index-1] not in vowels:
                count += 1
        if word.endswith('e'):
            count -= 1
        if word.endswith('le') and len(word) > 2 and word[-3] not in vowels:
            count += 1
        if count == 0:
            count += 1
        return count
    
    def extract_key_phrases(self, text):
        """
        Extract key phrases for summary
        """
        try:
            doc = nlp(text)
            
            # Extract noun chunks as key phrases
            key_phrases = []
            for chunk in doc.noun_chunks:
                if len(chunk.text.split()) <= 3:  # Limit to short phrases
                    key_phrases.append(chunk.text)
            
            # Get unique phrases
            key_phrases = list(set(key_phrases))
            
            # Limit to top 5
            return key_phrases[:5]
        except:
            return []
    
    def generate_summary(self, text, max_sentences=2):
        """
        Generate a summary of the feedback
        """
        sentences = sent_tokenize(text)
        if len(sentences) <= max_sentences:
            return text
        
        # Simple extractive summarization using sentence scoring
        word_freq = Counter()
        for sentence in sentences:
            words_list = word_tokenize(sentence.lower())
            for word in words_list:
                if word not in self.stop_words and word.isalpha():
                    word_freq[word] += 1
        
        # Score sentences
        sentence_scores = {}
        for sentence in sentences:
            words_list = word_tokenize(sentence.lower())
            score = 0
            for word in words_list:
                if word in word_freq:
                    score += word_freq[word]
            sentence_scores[sentence] = score / max(1, len(words_list))
        
        # Get top sentences
        top_sentences = sorted(sentence_scores.items(), key=lambda x: x[1], reverse=True)[:max_sentences]
        summary = ' '.join([s[0] for s in sorted(top_sentences, key=lambda x: sentences.index(x[0]))])
        
        return summary

# Initialize the enhanced analyzer
enhanced_analyzer = EnhancedFeedbackAnalyzer()

def normalize_text_advanced(text):
    """
    Advanced text normalization with better evasion detection
    """
    original = text
    text = text.lower()
    
    # Handle visual similarity replacements
    visual_replacements = {
        '|': 'i', '¡': 'i', '!': 'i',
        '@': 'a', '€': 'e', '£': 'e',
        '$': 's', '§': 's', '₿': 'b',
        '¢': 'c', '©': 'c', '®': 'r',
        '™': 't', '°': 'o', '±': 't',
        '0': 'o', '1': 'i', '2': 'z', '3': 'e',
        '4': 'a', '5': 's', '6': 'g', '7': 't',
        '8': 'b', '9': 'g'
    }
    
    # Apply visual replacements
    for symbol, letter in visual_replacements.items():
        text = text.replace(symbol, letter)
    
    # Handle common obfuscations
    obfuscation_patterns = [
        (r'ph(?=[uck]+)', 'f'),  # phuck -> fuck
        (r'zh(?=[it]+)', 'sh'),  # zhit -> shit
        (r'kz(?=s)', 'x'),       # kzs -> xs (sex)
        (r'cks(?=s)', 'x'),      # cks -> x (cocks -> cox)
        (r'kk', 'k'),             # fukk -> fuk
        (r'cc', 'c'),             # fucc -> fuc
        (r'vva', 'w'),            # vvank -> wank
    ]
    
    for pattern, replacement in obfuscation_patterns:
        text = re.sub(pattern, replacement, text)
    
    # Remove separators between letters (f.u.c.k, f-u-c-k, f*u*c*k)
    text = re.sub(r'[.\-*_~]\s*', '', text)
    
    # Handle character repetition with meaning (loooove -> love)
    text = re.sub(r'(.)\1{3,}', r'\1\1', text)
    
    # Handle bidirectional text and reverse psychology
    text = handle_bidirectional_text(text)
    
    # Remove excessive spaces
    text = re.sub(r'\s+', ' ', text).strip()
    
    return text

def handle_bidirectional_text(text):
    """Handle bidirectional Unicode text and reverse psychology"""
    # Remove RTL/LTR markers
    text = re.sub(r'[\u200e\u200f\u202a\u202b\u202c\u202d\u202e]', '', text)
    
    # Check for reversed words (like "kcuf" -> "fuck")
    words = text.split()
    processed_words = []
    
    for word in words:
        # Check if reversed version is a restricted term
        if word[::-1] in ALL_RESTRICTED_TERMS:
            processed_words.append(word[::-1])
        else:
            processed_words.append(word)
    
    return ' '.join(processed_words)

def contains_inappropriate_content(text):
    """
    Robust check for inappropriate content with multiple detection methods
    Returns (has_inappropriate, reason)
    """
    original_text = text
    text_lower = text.lower()
    normalized_text = normalize_text_advanced(text)
    
    # Method 1: Direct word matching
    all_restricted_terms = SEXUAL_TERMS | RACIST_TERMS | HOMOPHOBIC_TERMS | PROFANITY_TERMS
    
    # Check normalized text for direct matches
    words_in_normalized = set(re.findall(r'\b\w+\b', normalized_text))
    direct_matches = words_in_normalized.intersection(all_restricted_terms)
    
    if direct_matches:
        # Categorize the matches
        sexual_found = direct_matches.intersection(SEXUAL_TERMS)
        racist_found = direct_matches.intersection(RACIST_TERMS)
        homophobic_found = direct_matches.intersection(HOMOPHOBIC_TERMS)
        profanity_found = direct_matches.intersection(PROFANITY_TERMS)
        
        reasons = []
        if sexual_found:
            reasons.append(f"sexual content: {', '.join(sexual_found)}")
        if racist_found:
            reasons.append(f"racist content: {', '.join(racist_found)}")
        if homophobic_found:
            reasons.append(f"homophobic content: {', '.join(homophobic_found)}")
        if profanity_found:
            reasons.append(f"profanity: {', '.join(profanity_found)}")
        
        return True, f"Inappropriate content detected ({'; '.join(reasons)})"
    
    # Method 2: Check for partial matches (like "fuck" inside "fucking")
    for term in all_restricted_terms:
        if len(term) > 3:  # Only check longer terms to avoid false positives
            pattern = r'\b\w*' + re.escape(term) + r'\w*\b'
            if re.search(pattern, normalized_text):
                return True, f"Term '{term}' detected within longer word"
    
    # Method 3: Character repetition detection (like "fffffuck")
    repetition_pattern = r'(\w)\1{2,}.*\b(' + '|'.join(re.escape(t) for t in all_restricted_terms) + r')\b'
    if re.search(repetition_pattern, normalized_text):
        return True, "Disguised inappropriate content detected"
    
    # Method 4: Aggressive/harassment patterns
    aggressive_patterns = [
        r'\b(?:kill|die|hurt|harm)\s+(?:yourself|urself|u|you|them|him|her)\b',
        r'\b(?:go\s+)?(?:fuck|kill)\s+(?:yourself|off|urself)\b',
        r'\b(?:suck|lick)\s+my\s+(?:dick|cock|balls|ass)\b',
        r'\b(?:beat|kick|punch|hit)\s+(?:your|ur)\s+(?:ass|face|head|mom|mother)\b',
        r'\b(?:i\s+hope\s+you|wish\s+you)\s+(?:die|fail|suffer|rot)\b',
        r'\b(?:you\s+should)\s+(?:die|kill yourself|jump off)\b',
        r'\b(?:your\s+so\s+)(?:stupid|dumb|retarded|ugly)\b',
    ]
    
    for pattern in aggressive_patterns:
        if re.search(pattern, normalized_text, re.IGNORECASE):
            return True, "Aggressive/harassment content detected"
    
    # Method 5: Check for word combinations that create inappropriate meanings
    inappropriate_combinations = [
        (r'\b(?:suck|blow)\b.*\b(?:dick|cock|penis)\b', "sexual content"),
        (r'\b(?:eat|lick)\b.*\b(?:ass|pussy|cunt|vagina)\b', "sexual content"),
        (r'\b(?:fuck)\b.*\b(?:you|u|off|me)\b', "profanity"),
        (r'\b(?:stupid|dumb|retarded)\b.*\b(?:nigger|fag|dyke|bitch)\b', "discriminatory content"),
        (r'\b(?:rape|molest)\b', "violent sexual content"),
    ]
    
    for pattern, content_type in inappropriate_combinations:
        if re.search(pattern, normalized_text, re.IGNORECASE):
            return True, f"Inappropriate combination detected ({content_type})"
    
    # Method 6: Check for excessive special characters trying to evade detection
    if len(text) > 0:
        special_char_ratio = len(re.findall(r'[\*\-\_\~\+\=\!\@\#\$\%\^\&\*\(\)]', text)) / len(text)
        if special_char_ratio > 0.3:  # More than 30% special characters
            # Check if it contains suspicious patterns even with special chars
            suspicious_pattern = r'[fps][\*\-\_]*[uckh][\*\-\_]*[it]'
            if re.search(suspicious_pattern, normalized_text):
                return True, "Evasion attempt detected (excessive special characters)"
    
    # Method 7: Check for space-separated inappropriate words (f u c k)
    space_separated_patterns = [
        r'\bf\s+u\s+c\s+k\b',
        r'\bs\s+h\s+i\s+t\b',
        r'\ba\s+s\s+s\b',
        r'\bb\s+i\s+t\s+c\s+h\b',
        r'\bn\s+i\s+g\s+g\s+e\s+r\b',
    ]
    
    for pattern in space_separated_patterns:
        if re.search(pattern, text_lower):
            return True, "Space-separated inappropriate content detected"
    
    return False, None

def is_valid_english_text(text):
    """
    Enhanced validation with robust content filtering
    Returns (is_valid, error_message)
    """
    # First check for inappropriate content with robust detection
    has_inappropriate, reason = contains_inappropriate_content(text)
    if has_inappropriate:
        return False, f"Content violates community guidelines: {reason}"
    
    # Remove punctuation and split into words
    cleaned_text = re.sub(r'[^\w\s]', '', text.lower())
    words_list = cleaned_text.split()
    
    # Check minimum word count
    if len(words_list) < 3:
        return False, "Feedback is too short. Please provide at least 3 words."
    
    # Check maximum length (prevent abuse)
    if len(text) > 1000:
        return False, "Feedback is too long. Please limit to 1000 characters."
    
    # Check for excessive numbers
    number_count = sum(1 for word in words_list if word.isdigit())
    if len(words_list) > 0 and number_count / len(words_list) > 0.3:
        return False, "Feedback contains too many numbers. Please use proper sentences."
    
    # Check for valid English words (at least 60% should be real words)
    valid_word_count = 0
    for word in words_list:
        # Skip very short words and numbers
        if len(word) <= 2 or word.isdigit():
            continue
        # Check if word is in dictionary or has common English suffixes
        if word in ENGLISH_WORDS or word.endswith(('ing', 'ed', 'ly', 's', 'es', 'er', 'est', 'tion')):
            valid_word_count += 1
    
    # Calculate percentage of valid words (excluding short words)
    significant_words = [w for w in words_list if len(w) > 2 and not w.isdigit()]
    if len(significant_words) == 0:
        return False, "Please provide meaningful feedback with real words."
    
    valid_percentage = valid_word_count / len(significant_words)
    
    if valid_percentage < 0.6:
        return False, "Feedback appears to contain invalid or gibberish text. Please use real English words."
    
    # Check for excessive character repetition (like "awfawfawf")
    if re.search(r'(.{3,})\1{2,}', text.lower()):
        return False, "Feedback contains repeated character patterns. Please use proper words."
    
    # Check for random keyboard mashing
    keyboard_mash_pattern = r'\b(?:asdf|jkl|qwert|zxcv|mnb|qaz|wsx|edc|rfv|tgb|yhn|ujm|ik,|ol.|p;/)\w*\b'
    if len(re.findall(keyboard_mash_pattern, text.lower())) > 2:
        return False, "Feedback appears to contain random text. Please provide meaningful feedback."
    
    # Check for URL shorteners or suspicious links
    url_pattern = r'bit\.ly|tinyurl|goo\.gl|ow\.ly|is\.gd|buff\.ly|short\.link'
    if re.search(url_pattern, text.lower()):
        return False, "URL shorteners are not allowed. Please provide full URLs if necessary."
    
    return True, None

@app.route('/analyze', methods=['POST'])
def analyze():
    data = request.get_json()
    feedback = data.get('feedback', '').strip()
    
    if not feedback:
        return jsonify({'success': False, 'error': 'Feedback cannot be empty'})

    # Language detection: Only allow English
    try:
        if len(feedback.split()) < 3:
            return jsonify({'success': False, 'error': 'Feedback is too short. Please provide at least 3 words.'})
            
        lang = detect(feedback)
        if lang != 'en':
            return jsonify({'success': False, 'error': 'Please provide feedback in English.'})
    except LangDetectException:
        # Fallback: Check if most words are English
        words_list = re.findall(r'\b[a-zA-Z]+\b', feedback)
        if words_list:
            pass  # Continue with validation

    # Validate text content with robust filtering
    is_valid, error_message = is_valid_english_text(feedback)
    if not is_valid:
        return jsonify({'success': False, 'error': error_message})

    result = enhanced_analyzer.analyze_sentiment_nuanced(feedback)
    
    if result['sentiment'] in ['Negative', 'Very Negative']:
        result['raw_feedback'] = result['feedback']  # Keep for admin/teacher view
        result['feedback'] = None  # Hide from student view
        result['feedback_hidden'] = True
    else:
        result['feedback_hidden'] = False
    
    return jsonify({'success': True, 'data': result})

@app.route('/api/analyze-enhanced', methods=['POST'])
def analyze_enhanced():
    """
    Enhanced analysis endpoint with more features
    """
    if 'user_id' not in session:
        return jsonify({'success': False, 'error': 'Please log in first'}), 401
    
    data = request.get_json()
    feedback = data.get('feedback', '').strip()
    student_id = data.get('student_id')
    
    if not feedback:
        return jsonify({'success': False, 'error': 'Feedback cannot be empty'})
    
    # Validate content
    is_valid, error_message = is_valid_english_text(feedback)
    if not is_valid:
        return jsonify({'success': False, 'error': error_message})
    
    # Perform enhanced analysis
    result = enhanced_analyzer.analyze_sentiment_nuanced(feedback)
    
    # Add improvement suggestions if needed
    if result['constructive_feedback']:
        result['improvement_suggestions'] = list(result['constructive_feedback'].values())
    
    return jsonify({'success': True, 'data': result})

@app.route('/api/analyze-batch', methods=['POST'])
def analyze_batch():
    """
    Batch analysis endpoint for multiple feedbacks
    """
    if 'user_id' not in session:
        return jsonify({'success': False, 'error': 'Please log in first'}), 401
    
    data = request.get_json()
    feedbacks = data.get('feedbacks', [])
    
    if not feedbacks:
        return jsonify({'success': False, 'error': 'No feedbacks provided'})
    
    results = []
    for feedback in feedbacks:
        # Validate content
        is_valid, error_message = is_valid_english_text(feedback)
        if not is_valid:
            results.append({
                'feedback': feedback,
                'error': error_message,
                'skipped': True
            })
            continue
        
        # Analyze
        result = enhanced_analyzer.analyze_sentiment_nuanced(feedback)
        results.append(result)
    
    # Calculate statistics
    stats = {
        'total': len(results),
        'positive': sum(1 for r in results if r.get('sentiment') in ['Exceptional', 'Very Positive', 'Positive', 'Slightly Positive']),
        'neutral': sum(1 for r in results if r.get('sentiment') == 'Neutral'),
        'negative': sum(1 for r in results if r.get('sentiment') in ['Negative', 'Very Negative', 'Slightly Negative']),
        'skipped': sum(1 for r in results if r.get('skipped', False))
    }
    
    return jsonify({
        'success': True,
        'results': results,
        'statistics': stats
    })

# Test endpoint to verify filtering (remove in production)
@app.route('/test_filter', methods=['POST'])
def test_filter():
    """Endpoint to test various filtering attempts"""
    data = request.get_json()
    test_cases = data.get('test_cases', [])
    
    results = []
    for test_case in test_cases:
        is_valid, message = is_valid_english_text(test_case)
        has_inappropriate, reason = contains_inappropriate_content(test_case)
        results.append({
            'text': test_case,
            'is_valid': is_valid,
            'validation_message': message,
            'has_inappropriate': has_inappropriate,
            'inappropriate_reason': reason
        })
    
    return jsonify({'results': results})

@app.route('/api/nlp-notification', methods=['POST'])
def api_nlp_notification():
    if 'user_id' not in session or session.get('role') != 'Teacher':
        return jsonify({'success': False, 'message': 'Unauthorized'}), 403

    data = request.get_json()
    
    # Prepare notification data
    notification_data = {
        'teacher_id': session['user_id'],
        'student_id': data.get('student_id'),
        'feedback': data.get('feedback', ''),
        'sentiment': data.get('sentiment', ''),
        'sentiment_col': data.get('sentiment_color', ''),
        'positive_score': float(data.get('positive_score', 0)),
        'negative_score': float(data.get('negative_score', 0)),
        'neutral_score': float(data.get('neutral_score', 0)),
        'compound_score': float(data.get('compound_score', 0)),
        'motivation': data.get('motivation', ''),
        'feedback_hidden': data.get('feedback_hidden', False),
        'word_count': int(data.get('word_count', 0)),
        'complexity_score': float(data.get('complexity_score', 0)),
        'readability_score': float(data.get('readability_score', 0)),
        'key_phrases': data.get('key_phrases', []),
        'improvement_suggestions': data.get('improvement_suggestions', []),
    }

    # Insert into nlp_notifications table
    try:
        result = supabase.table('nlp_notifications').insert(notification_data).execute()
        return jsonify({'success': True})
    except Exception as e:
        logger.error(f"Error inserting NLP notification: {str(e)}")
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/api/nlp-notifications', methods=['GET'])
def api_nlp_notifications():
    if 'user_id' not in session or session.get('role') != 'Student':
        return jsonify({'success': False, 'notifications': []})
    student_id = session['user_id']
    try:
        result = supabase.table('nlp_notifications') \
            .select('*') \
            .eq('student_id', student_id) \
            .order('created_at', desc=True) \
            .limit(5) \
            .execute()
        notifications = result.data if result.data else []
        
        # Format for display
        for notif in notifications:
            notif['formatted_date'] = format_date(notif.get('created_at', ''))
        
        return jsonify({'success': True, 'notifications': notifications})
    except Exception as e:
        return jsonify({'success': False, 'notifications': []})

@app.route('/api/feedback-history/<int:student_id>', methods=['GET'])
def get_feedback_history(student_id):
    """
    Get feedback history for a student
    """
    if 'user_id' not in session or session.get('role') != 'Teacher':
        return jsonify({'success': False, 'message': 'Unauthorized'}), 403
    
    try:
        # Get feedback history
        result = supabase.table('nlp_notifications') \
            .select('*') \
            .eq('student_id', student_id) \
            .order('created_at', desc=True) \
            .limit(20) \
            .execute()
        
        feedback_history = result.data if result.data else []
        
        # Format dates
        for item in feedback_history:
            item['formatted_date'] = format_date(item.get('created_at', ''))
        
        # Calculate trends
        if feedback_history:
            sentiments = [f['sentiment'] for f in feedback_history if f.get('sentiment')]
            sentiment_trend = {
                'positive': sum(1 for s in sentiments if s in ['Exceptional', 'Very Positive', 'Positive', 'Slightly Positive']),
                'neutral': sentiments.count('Neutral'),
                'negative': sum(1 for s in sentiments if s in ['Negative', 'Very Negative', 'Slightly Negative'])
            }
            
            # Calculate average compound score trend
            compound_scores = [float(f.get('compound_score', 0)) for f in feedback_history if f.get('compound_score')]
            avg_compound = sum(compound_scores) / len(compound_scores) if compound_scores else 0
            
            # Calculate trend direction
            trend_direction = 'stable'
            if len(compound_scores) >= 3:
                first_avg = sum(compound_scores[:len(compound_scores)//2]) / (len(compound_scores)//2) if compound_scores[:len(compound_scores)//2] else 0
                last_avg = sum(compound_scores[-len(compound_scores)//2:]) / (len(compound_scores)//2) if compound_scores[-len(compound_scores)//2:] else 0
                
                if last_avg - first_avg > 0.2:
                    trend_direction = 'improving'
                elif first_avg - last_avg > 0.2:
                    trend_direction = 'declining'
            
            trend = {
                'sentiment_distribution': sentiment_trend,
                'average_compound': round(avg_compound, 3),
                'total_feedback': len(feedback_history),
                'trend_direction': trend_direction
            }
        else:
            trend = {
                'sentiment_distribution': {'positive': 0, 'neutral': 0, 'negative': 0},
                'average_compound': 0,
                'total_feedback': 0,
                'trend_direction': 'insufficient_data'
            }
        
        return jsonify({
            'success': True,
            'feedback_history': feedback_history,
            'trend': trend
        })
        
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/api/student-wellness-check/<int:student_id>', methods=['GET'])
def student_wellness_check(student_id):
    """
    Perform wellness check based on feedback history
    """
    if 'user_id' not in session or session.get('role') != 'Teacher':
        return jsonify({'success': False, 'message': 'Unauthorized'}), 403
    
    try:
        # Get recent feedback (last 30 days)
        thirty_days_ago = (datetime.now() - timedelta(days=30)).isoformat()
        
        result = supabase.table('nlp_notifications') \
            .select('*') \
            .eq('student_id', student_id) \
            .gte('created_at', thirty_days_ago) \
            .order('created_at', desc=True) \
            .execute()
        
        recent_feedback = result.data if result.data else []
        
        if not recent_feedback:
            return jsonify({
                'success': True,
                'wellness_status': 'insufficient_data',
                'message': 'Not enough feedback data for wellness check',
                'feedback_count': 0,
                'very_negative_count': 0,
                'avg_compound_score': 0
            })
        
        # Analyze sentiment trend
        compound_scores = [float(f.get('compound_score', 0)) for f in recent_feedback if f.get('compound_score')]
        
        wellness_status = 'stable'
        message = 'Student sentiment is stable.'
        
        if len(compound_scores) >= 3:
            # Check for declining trend
            first_third = compound_scores[:len(compound_scores)//3]
            last_third = compound_scores[-len(compound_scores)//3:]
            
            avg_first = sum(first_third) / len(first_third) if first_third else 0
            avg_last = sum(last_third) / len(last_third) if last_third else 0
            
            decline_threshold = -0.3
            if avg_last - avg_first < decline_threshold:
                wellness_status = 'declining'
                message = '⚠️ Student shows declining sentiment trend. Consider checking in personally.'
            elif avg_last - avg_first > 0.3:
                wellness_status = 'improving'
                message = '📈 Student shows improving sentiment trend. Keep up the good support!'
        
        # Check for extremely negative feedback
        very_negative_count = sum(1 for f in recent_feedback if f.get('sentiment') == 'Very Negative')
        
        if very_negative_count >= 2:
            wellness_status = 'concerning'
            message = '🚨 Multiple very negative feedback entries detected. Immediate attention recommended.'
        
        # Calculate average compound score
        avg_compound = sum(compound_scores) / len(compound_scores) if compound_scores else 0
        
        return jsonify({
            'success': True,
            'wellness_status': wellness_status,
            'message': message,
            'feedback_count': len(recent_feedback),
            'very_negative_count': very_negative_count,
            'avg_compound_score': round(avg_compound, 3)
        })
        
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/api/nlp-notifications/<int:notif_id>/read', methods=['POST'])
def mark_nlp_notification_read(notif_id):
    if 'user_id' not in session or session.get('role') != 'Student':
        return jsonify({'success': False, 'message': 'Unauthorized'}), 403
    try:
        supabase.table('nlp_notifications').update({'status': 'Read'}).eq('id', notif_id).execute()
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/api/feedback-stats/<int:student_id>', methods=['GET'])
def get_feedback_stats(student_id):
    """
    Get statistical summary of feedback for a student
    """
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': 'Unauthorized'}), 403
    
    try:
        # Get all feedback for student
        result = supabase.table('nlp_notifications') \
            .select('*') \
            .eq('student_id', student_id) \
            .order('created_at', desc=True) \
            .execute()
        
        feedback_list = result.data if result.data else []
        
        if not feedback_list:
            return jsonify({
                'success': True,
                'stats': {
                    'total_count': 0,
                    'avg_compound': 0,
                    'sentiment_breakdown': {'positive': 0, 'neutral': 0, 'negative': 0},
                    'avg_word_count': 0,
                    'avg_complexity': 0
                }
            })
        
        # Calculate stats
        sentiments = [f.get('sentiment', '') for f in feedback_list]
        compound_scores = [float(f.get('compound_score', 0)) for f in feedback_list if f.get('compound_score')]
        word_counts = [int(f.get('word_count', 0)) for f in feedback_list if f.get('word_count')]
        complexity_scores = [float(f.get('complexity_score', 0)) for f in feedback_list if f.get('complexity_score')]
        
        stats = {
            'total_count': len(feedback_list),
            'avg_compound': round(sum(compound_scores) / len(compound_scores), 3) if compound_scores else 0,
            'sentiment_breakdown': {
                'positive': sum(1 for s in sentiments if s in ['Exceptional', 'Very Positive', 'Positive', 'Slightly Positive']),
                'neutral': sentiments.count('Neutral'),
                'negative': sum(1 for s in sentiments if s in ['Negative', 'Very Negative', 'Slightly Negative'])
            },
            'avg_word_count': round(sum(word_counts) / len(word_counts), 1) if word_counts else 0,
            'avg_complexity': round(sum(complexity_scores) / len(complexity_scores), 2) if complexity_scores else 0
        }
        
        return jsonify({'success': True, 'stats': stats})
        
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/api/teacher-feedback-summary', methods=['GET'])
def teacher_feedback_summary():
    """
    Get summary of all feedback given by teacher
    """
    if 'user_id' not in session or session.get('role') != 'Teacher':
        return jsonify({'success': False, 'message': 'Unauthorized'}), 403
    
    teacher_id = session['user_id']
    
    try:
        # Get all feedback from this teacher
        result = supabase.table('nlp_notifications') \
            .select('*') \
            .eq('teacher_id', teacher_id) \
            .order('created_at', desc=True) \
            .limit(100) \
            .execute()
        
        feedback_list = result.data if result.data else []
        
        # Calculate summary stats
        total_feedback = len(feedback_list)
        if total_feedback == 0:
            return jsonify({
                'success': True,
                'summary': {
                    'total_feedback': 0,
                    'sentiment_breakdown': {'positive': 0, 'neutral': 0, 'negative': 0},
                    'avg_word_count': 0,
                    'most_common_sentiment': 'N/A'
                }
            })
        
        sentiments = [f.get('sentiment', '') for f in feedback_list]
        
        sentiment_counts = {
            'Exceptional': sentiments.count('Exceptional'),
            'Very Positive': sentiments.count('Very Positive'),
            'Positive': sentiments.count('Positive'),
            'Slightly Positive': sentiments.count('Slightly Positive'),
            'Neutral': sentiments.count('Neutral'),
            'Slightly Negative': sentiments.count('Slightly Negative'),
            'Negative': sentiments.count('Negative'),
            'Very Negative': sentiments.count('Very Negative')
        }
        
        # Find most common sentiment
        most_common = max(sentiment_counts.items(), key=lambda x: x[1]) if sentiment_counts else ('N/A', 0)
        
        word_counts = [int(f.get('word_count', 0)) for f in feedback_list if f.get('word_count')]
        
        summary = {
            'total_feedback': total_feedback,
            'sentiment_breakdown': sentiment_counts,
            'most_common_sentiment': most_common[0],
            'avg_word_count': round(sum(word_counts) / len(word_counts), 1) if word_counts else 0
        }
        
        return jsonify({'success': True, 'summary': summary})
        
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500

####################################################################
@app.route('/partials/showError')
def show_error_partial():
    return render_template('partials/showError.html')

@app.route('/')
def index():
    return redirect(url_for('login'))

@app.route('/api/notifications', methods=['GET'])
def get_notifications():
    try:
        user_id = session.get('user_id')
        if not user_id:
            return jsonify({'success': False, 'notifications': []}), 401
        # Fetch notifications for the user
        result = supabase.table('notifications') \
            .select('*') \
            .eq('user_id', user_id) \
            .order('created_at', desc=True) \
            .limit(10) \
            .execute()
        notifications = result.data if result.data else []

        # Collect all sender_ids
        sender_ids = list({n['sender_id'] for n in notifications if n.get('sender_id')})
        sender_map = {}
        if sender_ids:
            # Fetch sender info (include role and gender)
            senders = supabase.table('user_info').select('id, first_name, last_name, role, gender').in_('id', sender_ids).execute()
            if senders.data:
                for s in senders.data:
                    sender_map[s['id']] = {
                        'first_name': s.get('first_name', ''),
                        'last_name': s.get('last_name', ''),
                        'role': s.get('role', ''),
                        'gender': s.get('gender', '')
                    }
            # Fetch sender profile pictures
            pics = supabase.table('profile_pictures').select('user_id, file_path').in_('user_id', sender_ids).execute()
            if pics.data:
                for p in pics.data:
                    if p['user_id'] in sender_map:
                        sender_map[p['user_id']]['pic'] = p['file_path']

        # Attach sender info to notifications
        for n in notifications:
            sid = n.get('sender_id')
            if sid and sid in sender_map:
                sender = sender_map[sid]
                role = (sender.get('role') or '').lower()
                last_name = sender.get('last_name', '')
                first_name = sender.get('first_name', '')
                gender = (sender.get('gender') or '').lower()
                if role == 'admin':
                    n['sender_name'] = f"Admin {last_name}"
                elif role == 'teacher':
                    prefix = 'Mr.'
                    if gender == 'female':
                        prefix = 'Mrs.'
                    elif gender == 'other':
                        prefix = 'Mx.'
                    n['sender_name'] = f"{prefix} {last_name}"
                elif role == 'student':
                    n['sender_name'] = f"{first_name} {last_name}"
                else:
                    n['sender_name'] = 'System'
                n['sender_pic'] = sender.get('pic', '/static/image/default-avatar.png')
            else:
                n['sender_name'] = 'System'
                n['sender_pic'] = '/static/image/default-avatar.png'

        return jsonify({'success': True, 'notifications': notifications})
    except Exception as e:
        return jsonify({'success': False, 'notifications': []}), 500


#############################LOGIN.HTML#############################
@app.route('/logout')
def logout():
    user_id = session.get('user_id')
    role = session.get('role')
    last_name = session.get('last_name', '')
    email = session.get('email', '')
    
    # Get client IP address
    ip_address = request.headers.get('X-Forwarded-For', request.remote_addr)
    
    # Detect device type
    user_agent = request.user_agent.string.lower()
    if 'mobile' in user_agent or 'android' in user_agent or 'iphone' in user_agent:
        device_type = 'mobile'
    elif 'tablet' in user_agent or 'ipad' in user_agent:
        device_type = 'tablet'
    elif 'windows' in user_agent:
        device_type = 'Windows computer'
    elif 'mac' in user_agent or 'macos' in user_agent:
        device_type = 'Mac computer'
    elif 'linux' in user_agent:
        device_type = 'Linux computer'
    else:
        device_type = 'computer'
    
    # Calculate session duration
    session_duration = "Unknown"
    if 'login_time' in session:
        login_time = session['login_time']
        try:
            # If login_time is a timestamp (float)
            if isinstance(login_time, (int, float)):
                login_dt = datetime.fromtimestamp(login_time)
            # If login_time is a string
            elif isinstance(login_time, str):
                login_dt = datetime.fromisoformat(login_time)
            else:
                login_dt = login_time
                
            logout_time = datetime.now()
            duration_seconds = (logout_time - login_dt).total_seconds()
            
            # Format duration
            if duration_seconds < 60:
                session_duration = f"{int(duration_seconds)} seconds"
            elif duration_seconds < 3600:
                minutes = int(duration_seconds // 60)
                seconds = int(duration_seconds % 60)
                session_duration = f"{minutes} minutes {seconds} seconds"
            else:
                hours = int(duration_seconds // 3600)
                minutes = int((duration_seconds % 3600) // 60)
                session_duration = f"{hours} hours {minutes} minutes"
                
        except Exception as e:
            print(f"Error calculating session duration: {e}")
            session_duration = "Calculation error"
    
    # Record logout in admin_history
    if user_id and role:
        record_admin_activity_log(
            user_id=user_id,
            action='Logout',
            activity='Authentication',
            description=f"{role} {last_name} logged out from {device_type}",
            user_role=role,
            details=f"Email: {email}\nIP Address: {ip_address.split(',')[0].strip()}\nSession Duration: {session_duration}"
        )
    session.clear()
    return redirect(url_for('login'))

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        is_js = request.headers.get('X-Requested-With') == 'XMLHttpRequest' or request.is_json

        email = request.form.get('email') or (request.json and request.json.get('email'))
        password = request.form.get('password') or (request.json and request.json.get('password'))

        # Get client IP address
        ip_address = request.headers.get('X-Forwarded-For', request.remote_addr)
        
        # Detect device type
        user_agent = request.user_agent.string.lower()
        if 'mobile' in user_agent or 'android' in user_agent or 'iphone' in user_agent:
            device_type = 'mobile'
        elif 'tablet' in user_agent or 'ipad' in user_agent:
            device_type = 'tablet'
        elif 'windows' in user_agent:
            device_type = 'Windows computer'
        elif 'mac' in user_agent or 'macos' in user_agent:
            device_type = 'Mac computer'
        elif 'linux' in user_agent:
            device_type = 'Linux computer'
        else:
            device_type = 'computer'

        # Debug: print received email and password
        print(f"Login attempt: email={email}, password={'*' * len(password) if password else None}")

        result = supabase.table('user_info').select('*').eq('email', email).execute()

        if not result.data:
            msg = 'Invalid email or password (no user found)'
            print(f"Login error: {msg}")
            if is_js:
                return jsonify({'success': False, 'message': msg}), 401
            flash(msg, 'danger')
            return redirect(url_for('login'))

        user = result.data[0]
        print(f"User found: {user}")

        if user['password'] != password:
            msg = 'Invalid email or password (wrong password)'
            print(f"Login error: {msg}")
            if is_js:
                return jsonify({'success': False, 'message': msg}), 401
            flash(msg, 'danger')
            return redirect(url_for('login'))

        role = user['role'].strip()
        print(f"User role: {role}")

        if role == 'Student':
            status = user.get('status', '').lower()
            print(f"Student status: {status}")
            if status != 'active':
                msg = f"Your account status is '{user.get('status', '')}'. Please wait for admin approval or contact your teacher."
                print(f"Login error: {msg}")
                if is_js:
                    return jsonify({'success': False, 'message': msg}), 403
                flash(msg, 'danger')
                return redirect(url_for('login'))

            # Try to set session and redirect
            try:
                session['user_id'] = user['id']
                session['email'] = user['email']
                session['role'] = role
                session['first_name'] = user.get('first_name', '')
                session['last_name'] = user.get('last_name', '')
                session['gender'] = user.get('gender', '')
                session['year_level'] = user.get('year_level', '')
                session['section'] = user.get('section', '')
                
                # Set login time for session duration tracking
                session['login_time'] = datetime.now().timestamp()
                
                # Set profile picture for student
                pic_result = supabase.table('profile_pictures').select('file_path').eq('user_id', user['id']).execute()
                if pic_result.data and len(pic_result.data) > 0:
                    session['profile_picture'] = pic_result.data[0]['file_path']
                else:
                    session['profile_picture'] = None
                
                # Set student points in session
                session['points'] = user.get('total_points', 0)
                
                # Set student streak in session
                streak_value = user.get('streak', 0) or 0
                session['streak'] = streak_value
                
                # Determine streak color class based on streak value
                if streak_value < 7:
                    streak_class = ""  # Default grey
                elif streak_value < 15:
                    streak_class = "streak-yellow"
                elif streak_value < 22:
                    streak_class = "streak-orange"
                elif streak_value < 29:
                    streak_class = "streak-red"
                elif streak_value < 36:
                    streak_class = "streak-green"
                elif streak_value < 50:
                    streak_class = "streak-blue"
                else:
                    streak_class = "streak-purple streak-max"
                
                session['streak_class'] = streak_class
                
                print(f"Student login successful: {user['email']}, Streak: {streak_value}, Class: {streak_class}")
                
                # Record login activity with enhanced details
                record_admin_activity_log(
                    user_id=user['id'],
                    action='Login',
                    activity='Authentication',
                    description=f"Student {user.get('last_name', '')} logged in from {device_type}",
                    user_role=role,
                    details=f"Email: {user['email']}\nIP Address: {ip_address.split(',')[0].strip()}\nStreak: {streak_value}"
                )
                
                if is_js:
                    return jsonify({'success': True, 'role': role}), 200
                return redirect(url_for('student_home'))
            except Exception as e:
                msg = f"Unexpected error during login: {str(e)}"
                print(f"Login error: {msg}")
                if is_js:
                    return jsonify({'success': False, 'message': msg}), 500
                flash(msg, 'danger')
                return redirect(url_for('login'))

        elif role in ['Admin', 'Teacher']:
            session['user_id'] = user['id']
            session['email'] = user['email']
            session['role'] = role
            session['last_name'] = user.get('last_name', '')
            session['gender'] = user.get('gender', '')
            # Set login time for session duration tracking
            session['login_time'] = datetime.now().timestamp()
            if role == 'Teacher':
                session['subject'] = user.get('subject', '')
            
            # Record login activity with enhanced details
            record_admin_activity_log(
                user_id=user['id'],
                action='Login',
                activity='Authentication',
                description=f"{role} {user.get('last_name', '')} logged in from {device_type}",
                user_role=role,
                details=f"Email: {user['email']}\nIP Address: {ip_address.split(',')[0].strip()}"
            )
            
            if is_js:
                return jsonify({'success': True, 'role': role}), 200
            if role == 'Admin':
                return redirect(url_for('admin_dashboard'))
            elif role == 'Teacher':
                return redirect(url_for('dashboard'))
        else:
            msg = f'Invalid role: "{role}". Expected "Admin", "Teacher", or "Student".'
            print(f"Login error: {msg}")
            if is_js:
                return jsonify({'success': False, 'message': msg}), 401
            flash(msg, 'danger')
            return redirect(url_for('login'))

    return render_template('login.html')


@app.route('/api/register-student', methods=['POST'])
def api_register_student():
    try:
        data = request.json
        required_fields = ['first_name', 'last_name', 'email', 'password', 'year_level', 'section', 'gender']
        missing = [f for f in required_fields if not data.get(f)]
        if missing:
            return jsonify({'success': False, 'message': f'Missing fields: {", ".join(missing)}'}), 400

        student_data = {
            "first_name": data['first_name'],
            "middle_name": data.get('middle_name', ''),
            "last_name": data['last_name'],
            "gender": data['gender'],
            "email": data['email'],
            "mobile_no": data.get('phone', ''),
            "role": "Student",
            "year_level": data['year_level'],
            "section": data['section'],
            "password": data['password'],  # In production, hash this!
            "status": "Pending"
        }
        result = supabase.table('user_info').insert(student_data).execute()
        if hasattr(result, 'error') and result.error:
            raise Exception(result.error.message)
        return jsonify({'success': True, 'message': 'Registration submitted! Await admin approval.'}), 201
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/forgot-password', methods=['POST'])
def forgot_password():
    """Send OTP to user's email for password reset"""
    try:
        data = request.get_json()
        email = data.get('email')
        
        if not email:
            return jsonify({'success': False, 'message': 'Email is required'}), 400
        
        # Check if email exists in database
        user_response = supabase.table('user_info').select('*').eq('email', email).execute()
        
        if not user_response.data:
            return jsonify({'success': False, 'message': 'Email not found'}), 404
        
        # Generate OTP
        otp = generate_otp()
        expiration_time = datetime.utcnow() + timedelta(minutes=10)
        
        # Store OTP with expiration
        otp_storage[email] = {
            'otp': otp,
            'expires_at': expiration_time.isoformat(),
            'verified': False
        }
        
        # Send OTP email
        send_otp_email(email, otp)
        
        return jsonify({
            'success': True, 
            'message': 'OTP sent to your email'
        })
        
    except Exception as e:
        print(f"Error in forgot_password: {e}")
        return jsonify({'success': False, 'message': 'An error occurred. Please try again.'}), 500

def send_otp_email(email, otp):
    """Send OTP email to user"""
    try:
        msg = Message(
            subject='Learn2Earn - Password Reset OTP',
            sender='glamourgaze33@gmail.com',
            recipients=[email]
        )
        
        msg.body = f"""
Dear User,

You have requested to reset your password for your Learn2Earn account.

Your One-Time Password (OTP) is: {otp}

This OTP will expire in 10 minutes.

If you did not request this password reset, please ignore this email.

Best regards,
Masico National High School
Learn2Earn Administration Team
"""
        
        mail.send(msg)
        print(f"OTP email sent to {email}: {otp}")
    except Exception as e:
        print(f"Error sending email: {e}")

@app.route('/verify-otp', methods=['POST'])
def verify_otp():
    """Verify OTP entered by user"""
    try:
        data = request.get_json()
        email = data.get('email')
        otp = data.get('otp')
        
        if not email or not otp:
            return jsonify({'success': False, 'message': 'Email and OTP are required'}), 400
        
        # Check if OTP exists and is valid
        stored_otp_data = otp_storage.get(email)
        
        if not stored_otp_data:
            return jsonify({'success': False, 'message': 'OTP not found or expired'}), 404
        
        # Check if OTP has expired
        expires_at = datetime.fromisoformat(stored_otp_data['expires_at'])
        if datetime.utcnow() > expires_at:
            # Clean up expired OTP
            del otp_storage[email]
            return jsonify({'success': False, 'message': 'OTP has expired'}), 400
        
        if stored_otp_data['otp'] != otp:
            return jsonify({'success': False, 'message': 'Invalid OTP'}), 400
        
        # Mark OTP as verified
        stored_otp_data['verified'] = True
        
        return jsonify({
            'success': True, 
            'message': 'OTP verified successfully'
        })
        
    except Exception as e:
        print(f"Error in verify_otp: {e}")
        return jsonify({'success': False, 'message': 'An error occurred. Please try again.'}), 500

@app.route('/reset-password', methods=['POST'])
def reset_password():
    """Reset password after OTP verification"""
    try:
        data = request.get_json()
        email = data.get('email')
        new_password = data.get('new_password')
        confirm_password = data.get('confirm_password')
        
        if not email or not new_password or not confirm_password:
            return jsonify({'success': False, 'message': 'All fields are required'}), 400
        
        if new_password != confirm_password:
            return jsonify({'success': False, 'message': 'Passwords do not match'}), 400
        
        # Check if OTP was verified
        stored_otp_data = otp_storage.get(email)
        
        if not stored_otp_data or not stored_otp_data.get('verified'):
            return jsonify({'success': False, 'message': 'OTP verification required'}), 400
        
        # Update password in database
        update_response = supabase.table('user_info')\
            .update({'password': new_password})\
            .eq('email', email)\
            .execute()
        
        if not update_response.data:
            return jsonify({'success': False, 'message': 'Failed to update password'}), 500
        
        # Clean up OTP storage
        del otp_storage[email]
        
        # Send confirmation email
        send_password_reset_confirmation(email)
        
        return jsonify({
            'success': True, 
            'message': 'Password reset successfully'
        })
        
    except Exception as e:
        print(f"Error in reset_password: {e}")
        return jsonify({'success': False, 'message': 'An error occurred. Please try again.'}), 500

def send_password_reset_confirmation(email):
    """Send password reset confirmation email"""
    try:
        msg = Message(
            subject='Learn2Earn - Password Reset Successful',
            sender='glamourgaze33@gmail.com',
            recipients=[email]
        )
        
        msg.body = f"""
Dear User,

Your password has been successfully reset.

If you did not perform this action, please contact the administration immediately.

You can now login with your new password.

Best regards,
Masico National High School
Learn2Earn Administration Team
"""
        
        mail.send(msg)
        print(f"Password reset confirmation sent to {email}")
    except Exception as e:
        print(f"Error sending confirmation email: {e}")

#############################STUDENT ROUTES#############################
@app.route('/student_home')
def student_home():
    if 'user_id' not in session or session.get('role') != 'Student':
        return redirect(url_for('login'))
    user_id = session['user_id']

    recent_activities = []

    # --- Recent completed/denied tasks ---
    tasks_resp = safe_execute(
        supabase.table('task_assignments').select(
            'task, points, status, completed_at, due_date'
        ).eq('student_id', user_id).in_('status', ['Completed', 'Denied']).order('completed_at', desc=True).limit(5)
    )
    if tasks_resp.data:
        for t in tasks_resp.data:
            recent_activities.append({
                'type': 'task',
                'title': t['task'],
                'points': t['points'],
                'status': t['status'],
                'time': format_date(t.get('completed_at', t.get('due_date', ''))),
            })

    # --- Recent reward redemptions ---
    try:
        rewards_resp = safe_execute(
            supabase.table('reward_redemptions').select(
                'reward_id, points_deducted, processed_at, rewards(reward_name)'
            ).eq('student_id', user_id).order('processed_at', desc=True).limit(5)
        )
    except Exception as e:
        print(f"Error in reward_redemptions query: {e}")
        rewards_resp = type('obj', (object,), {'data': []})()
    if rewards_resp.data:
        for r in rewards_resp.data:
            reward_name = ''
            if r.get('rewards') and r['rewards'].get('reward_name'):
                reward_name = r['rewards']['reward_name']
            recent_activities.append({
                'type': 'reward',
                'title': f"Redeemed: {reward_name}",
                'points': -abs(r['points_deducted']),
                'status': 'Redeemed',
                'time': format_date(r.get('processed_at', '')),
            })

    # --- Recent teacher awards (e.g., points) ---
    stars_resp = safe_execute(
        supabase.table('points').select(
            'points, received_at, note'
        ).eq('student_id', user_id).eq('status', 'approved').order('received_at', desc=True).limit(5)
    )
    if stars_resp.data:
        for s in stars_resp.data:
            recent_activities.append({
                'type': 'star',
                'title': s.get('note', 'Teacher Award'),
                'points': s['points'],
                'status': 'Awarded',
                'time': format_date(s.get('received_at', '')),
            })

    # --- Recent milestone claims ---
    milestone_claims_resp = safe_execute(
        supabase.table('milestone_claims')
        .select('milestone_type, milestone, points_awarded, claimed_at')
        .eq('student_id', user_id)
        .order('claimed_at', desc=True)
        .limit(5)
    )
    if milestone_claims_resp.data:
        for m in milestone_claims_resp.data:
            recent_activities.append({
                'type': 'milestone',
                'title': f"Claimed {m['milestone_type'].capitalize()} Milestone: {m['milestone']}",
                'points': m.get('points_awarded', 0),
                'status': 'Claimed',
                'time': format_date(m.get('claimed_at', '')),
                'is_me': True,
            })     

    # --- Streak milestones (e.g., every 7, 30, 50 days) ---
    streak_milestone_notifs = safe_execute(
        supabase.table('notifications')
        .select('title, message, created_at')
        .eq('user_id', user_id)
        .eq('title', 'Streak Milestone Unlocked!')
        .order('created_at', desc=True)
    )

    if streak_milestone_notifs.data:
        for notif in streak_milestone_notifs.data:
            import re
            milestone_match = re.search(r'the (\d+) streak milestone', notif.get('message', ''))
            milestone_num = milestone_match.group(1) if milestone_match else '?'
            recent_activities.append({
                'type': 'streak',
                'title': f"Streak Milestone: {milestone_num} days!",
                'points': 0,
                'status': 'Milestone',
                'time': format_date(notif.get('created_at', '')),
            })
        
    # Sort by time, most recent first (if not already)
    recent_activities.sort(key=lambda x: x.get('time') or '', reverse=True)
    recent_activities = recent_activities[:10]

    final_activities = sorted(
        [a for a in recent_activities if a.get('time')],
        key=lambda x: x.get('time') or '',
        reverse=True
    )[:5]

    return render_template('student/student_home.html', activities=final_activities)

@app.route('/student_activities')
def student_activities():
    if 'user_id' not in session or session.get('role') != 'Student':
        return redirect(url_for('login'))
    return render_template('student/student_activities.html')


@app.route('/student_rewards')
def student_rewards():
    if 'user_id' not in session or session.get('role') != 'Student':
        return redirect(url_for('login'))
    return render_template('student/student_rewards.html')


@app.route('/student_rewards_history')
def student_rewards_history():
    if 'user_id' not in session or session.get('role') != 'Student':
        return redirect(url_for('login'))
    return render_template('student/student_rewards_history.html')


@app.route('/student_leaderboard')
def student_leaderboard():
    if 'user_id' not in session or session.get('role') != 'Student':
        return redirect(url_for('login'))
    return render_template('student/student_leaderboard.html', current_user_id=session['user_id'])


@app.route('/student_profile')
def student_profile():
    if 'user_id' not in session or session.get('role') != 'Student':
        return redirect(url_for('login'))
    return render_template('student/student_profile.html')


@app.route('/student_all_activity')
def student_all_activity():
    if 'user_id' not in session or session.get('role') != 'Student':
        return redirect(url_for('login'))
    return render_template('student/student_all_activity.html')


@app.route('/student_quiz')
def student_quiz():
    if 'user_id' not in session or session.get('role') != 'Student':
        return redirect(url_for('login'))
    return render_template('student/student_quiz.html')

#############################TEACHER ROUTES#############################
@app.route('/dashboard')
def dashboard():
    if 'user_id' in session:
        pic_result = supabase.table('profile_pictures').select('file_path').eq('user_id', session['user_id']).execute()
        if pic_result.data and len(pic_result.data) > 0:
            session['profile_picture'] = pic_result.data[0]['file_path']
        else:
            session['profile_picture'] = None
    return render_template('teacher/dashboard.html')

@app.route('/award_points_view_all')
def award_points_view_all():
    return render_template('teacher/award_points_view_all.html')

@app.route('/students')
def student_list():
    # Only allow teachers to view their students
    if 'user_id' not in session or session.get('role') != 'Teacher':
        flash('Unauthorized access.', 'danger')
        return redirect(url_for('login'))

    teacher_id = session['user_id']

    try:
        # Get current school year
        current_school_year = get_current_school_year()
        if not current_school_year:
            flash('Unable to determine current school year', 'danger')
            return redirect(url_for('dashboard'))
        
        # 1. Get all classrooms assigned to this teacher for CURRENT school year
        assignments_result = supabase.table('teacher_class_assignments') \
            .select('grade_level, section') \
            .eq('teacher_id', teacher_id) \
            .eq('school_year', current_school_year) \
            .eq('status', 'active') \
            .execute()
        assignments = assignments_result.data if assignments_result.data else []

        students = []
        # 2. For each classroom, get all students in that grade_level and section
        for assignment in assignments:
            grade_level = assignment['grade_level']
            section = assignment['section']
            students_result = supabase.table('user_info') \
                .select('*') \
                .eq('role', 'Student') \
                .eq('year_level', grade_level) \
                .eq('section', section) \
                .eq('status', 'Active') \
                .execute()
            if students_result.data:
                students.extend(students_result.data)

        # 3. Optionally, remove duplicates if a student is in multiple classes
        unique_students = {student['id']: student for student in students}.values()

        # 4. Prepare data for the template
        student_list = []
        for student in unique_students:
            student_list.append({
                'id': student['id'],
                'name': f"{student.get('first_name', '')} {student.get('last_name', '')}",
                'grade': student.get('year_level', ''),
                'section': student.get('section', ''),
                'points': student.get('total_points', 0),
            })

        return render_template('students.html', students=student_list)

    except Exception as e:
        app.logger.error(f"Error fetching students: {str(e)}")
        return render_template('teacher/students.html', students=[])

@app.route('/activities')
def activities():
    return render_template('teacher/activities.html')

@app.route('/rewards')
def reward_list():
    # Make sure the user is logged in and is a teacher
    if 'user_id' not in session or session.get('role') != 'Teacher':
        flash('Unauthorized access.', 'danger')
        return redirect(url_for('login'))

    # Just render the template - JavaScript will load the data
    return render_template('teacher/rewards.html')

@app.route('/reward-history-redemptions')
def reward_history_redemptions():
    # Only allow teachers
    if 'user_id' not in session or session.get('role') != 'Teacher':
        flash('Unauthorized access.', 'danger')
        return redirect(url_for('login'))
    return render_template('teacher/reward_history_redemption.html')

@app.route('/analytics')
def analytics():
    if 'user_id' not in session or session.get('role') != 'Teacher':
        flash('Unauthorized access.', 'danger')
        return redirect(url_for('login'))

    teacher_id = session['user_id']

    try:
        # Get current school year
        current_school_year = get_current_school_year()
        if not current_school_year:
            flash('Unable to determine current school year', 'danger')
            return redirect(url_for('dashboard'))
        
        # Get teacher's classrooms and students for CURRENT school year
        assignments_result = supabase.table('teacher_class_assignments') \
            .select('grade_level, section') \
            .eq('teacher_id', teacher_id) \
            .eq('school_year', current_school_year) \
            .eq('status', 'active') \
            .execute()
        assignments = assignments_result.data if assignments_result.data else []

        student_ids = []
        for assignment in assignments:
            students_result = supabase.table('user_info') \
                .select('id') \
                .eq('role', 'Student') \
                .eq('year_level', assignment['grade_level']) \
                .eq('section', assignment['section']) \
                .execute()
            if students_result.data:
                student_ids.extend([s['id'] for s in students_result.data])

        # Remove duplicates
        student_ids = list(set(student_ids))
        active_students_count = len(student_ids)

        # --- Points calculations (by this teacher) ---
        total_points = 0
        category_distribution = {}
        
        if student_ids:
            # Get points awarded by this teacher to these students
            points_result = supabase.table('points') \
                .select('student_id, points, received_at, point_category') \
                .eq('teacher_id', teacher_id) \
                .in_('student_id', student_ids) \
                .execute()
            
            if points_result.data:
                for row in points_result.data:
                    points = row['points'] or 0
                    total_points += points

                    # Track category distribution
                    cat = row.get('point_category', 'Other')
                    if cat:
                        cat = cat.lower()
                        category_distribution[cat] = category_distribution.get(cat, 0) + points

        # --- Average Points Per Student (by this teacher) ---
        avg_points = round(total_points / len(student_ids), 2) if student_ids else 0

        # --- Task Completion Rate (tasks assigned by this teacher) ---
        task_result = supabase.table('task_assignments') \
            .select('student_id, status') \
            .eq('teacher_id', teacher_id) \
            .in_('student_id', student_ids) \
            .execute()
        
        total_tasks = 0
        completed_tasks = 0
        if task_result.data:
            for row in task_result.data:
                total_tasks += 1
                if row['status'] and row['status'].lower() == 'completed':
                    completed_tasks += 1
        
        task_completion_rate = round((completed_tasks / total_tasks) * 100, 2) if total_tasks else 0

        from datetime import datetime, timedelta, timezone
        now = datetime.now(timezone.utc)
        
        # --- Monthly Points Trend for this Teacher ---
        months = []
        month_labels = []
        for i in range(5, -1, -1):  # Last 6 months
            month = (now.replace(day=1) - timedelta(days=30*i)).replace(day=1)
            months.append(month)
            month_labels.append(month.strftime('%b %Y'))

        # Get all points awarded by this teacher (for trend charts)
        points_result_teacher = supabase.table('points') \
            .select('points, received_at') \
            .eq('teacher_id', teacher_id) \
            .execute()
        
        from collections import OrderedDict
        monthly_points = OrderedDict((label, 0) for label in month_labels)
        
        if points_result_teacher.data:
            for row in points_result_teacher.data:
                received_at = row.get('received_at')
                points = row.get('points', 0)
                if received_at:
                    try:
                        dt = datetime.fromisoformat(str(received_at).replace('Z', '+00:00'))
                        if dt.tzinfo is None:
                            dt = dt.replace(tzinfo=timezone.utc)
                        else:
                            dt = dt.astimezone(timezone.utc)
                        for i, month in enumerate(months):
                            if dt.year == month.year and dt.month == month.month:
                                monthly_points[month_labels[i]] += points
                                break
                    except Exception:
                        continue

        # --- Weekly Points Trend for this Teacher (last 4 weeks) ---
        week_labels = []
        weekly_points = []
        for i in range(3, -1, -1):
            week_start = (now - timedelta(days=now.weekday())) - timedelta(weeks=i)
            week_end = week_start + timedelta(days=6)
            week_label = f"{week_start.strftime('%b %d')}-{week_end.strftime('%b %d')}"
            week_labels.append(week_label)
            weekly_points.append(0)
        
        if points_result_teacher.data:
            for row in points_result_teacher.data:
                received_at = row.get('received_at')
                points = row.get('points', 0)
                if received_at:
                    try:
                        dt = datetime.fromisoformat(str(received_at).replace('Z', '+00:00'))
                        if dt.tzinfo is None:
                            dt = dt.replace(tzinfo=timezone.utc)
                        else:
                            dt = dt.astimezone(timezone.utc)
                        for i in range(4):
                            week_start = (now - timedelta(days=now.weekday())) - timedelta(weeks=3-i)
                            week_end = week_start + timedelta(days=6)
                            if week_start <= dt <= week_end:
                                weekly_points[i] += points
                                break
                    except Exception:
                        continue

        return render_template(
            'teacher/analytics.html',
            points_all_time=total_points,
            avg_points=avg_points,
            task_completion_rate=task_completion_rate,
            active_students=active_students_count,  # Using the count from teacher's classes
            category_distribution=category_distribution,
            monthly_points_labels=list(monthly_points.keys()),
            monthly_points_data=list(monthly_points.values()),
            weekly_points_labels=week_labels,
            weekly_points_data=weekly_points
        )
        
    except Exception as e:
        app.logger.error(f"Error fetching analytics: {str(e)}")
        flash('Error loading analytics. Please try again later.', 'danger')
        return render_template(
            'teacher/analytics.html',
            points_all_time=0,
            avg_points=0,
            task_completion_rate=0,
            active_students=0,
            category_distribution={},
            monthly_points_labels=[],
            monthly_points_data=[],
            weekly_points_labels=[],
            weekly_points_data=[]
        )

@app.route('/award-points', methods=['GET', 'POST'])
def award_points():
    if request.method == 'POST':
        flash('Points awarded successfully!', 'success')
        return redirect(url_for('student_list'))
    return render_template('teacher/award_points.html')

@app.route('/quiz-maker')
def quiz_maker():
    if session.get('role') != 'Teacher':
        flash('You do not have permission to access this page.', 'danger')
        return redirect(url_for('login'))
    return render_template('teacher/quiz_maker.html')


@app.route('/settings')
def settings():
    return render_template('settings.html')


#############################ADMIN ROUTES#############################
@app.route('/admin_activity_log')
def admin_activity_log():
    return render_template('admin/admin_activity_log.html')

@app.route('/user_management')
def user_management():
    return render_template('admin/user_management.html')

@app.route('/admin/classroom')
def admin_classroom():
    return render_template('admin/admin_subject.html')

@app.route('/admin_dashboard')
def admin_dashboard():
    try:
        # Set admin profile picture in session
        if 'user_id' in session and session.get('role') == 'Admin':
            pic_result = supabase.table('profile_pictures').select('file_path').eq('user_id', session['user_id']).execute()
            if pic_result.data and len(pic_result.data) > 0:
                session['profile_picture'] = pic_result.data[0]['file_path']
            else:
                session['profile_picture'] = None

        # Query the database for the total number of teachers
        teachers_result = supabase.table('user_info').select('id').eq('role', 'Teacher').execute()
        total_teachers = len(teachers_result.data) if teachers_result.data else 0

        # Query the database for the total number of students
        students_result = supabase.table('user_info').select('id').eq('role', 'Student').execute()
        total_students = len(students_result.data) if students_result.data else 0

        # Query the database for the total points awarded
        points_result = supabase.table('points').select('points').execute()
        total_points_awarded = sum(record['points'] for record in points_result.data) if points_result.data else 0

        # Pass the data to the template
        return render_template(
            'admin/admin_dashboard.html',
            total_teachers=total_teachers,
            total_students=total_students,
            total_points_awarded=total_points_awarded
        )
    except Exception as e:
        app.logger.error(f"Error fetching dashboard data: {str(e)}")
        flash('Error loading dashboard data. Please try again later.', 'danger')
        return render_template(
            'admin/admin_dashboard.html',
            total_teachers=0,
            total_students=0,
            total_points_awarded=0
        )
    

#############################ADMIN_ANALYTICS.HTML#############################
@app.route('/admin_analytics')
def admin_analytics():
    try:
        # --- USERS ---
        users_result = supabase.table('user_info').select('id, role, status, created_at').execute()
        users = users_result.data if users_result.data else []
        students = [u for u in users if u['role'] == 'Student']
        total_students = len(students)
        total_teachers = len([u for u in users if u['role'] == 'Teacher'])
        total_admins = len([u for u in users if u['role'] == 'Admin'])
        active_users = len([u for u in users if u.get('status', '').lower() == 'active'])

        # --- TOTAL POINTS AWARDED (across all teachers, including tasks) ---
        total_points_awarded = 0
        points_result = supabase.table('points').select('points').execute()
        if points_result.data:
            total_points_awarded = sum(row.get('points', 0) for row in points_result.data)

        # --- STUDENT ENGAGEMENT: students who received points or were assigned a task in the last 7 days ---
        now = datetime.now(timezone.utc)
        week_ago = now - timedelta(days=7)
        student_ids = [s['id'] for s in students]
        engaged_students_set = set()

        # 1. Students who received points in the last 7 days
        if student_ids:
            points_result = supabase.table('points').select('student_id, received_at').in_('student_id', student_ids).execute()
            if points_result.data:
                for row in points_result.data:
                    received_at = row.get('received_at')
                    if received_at:
                        try:
                            dt = datetime.fromisoformat(str(received_at).replace('Z', '+00:00'))
                            if dt >= week_ago:
                                engaged_students_set.add(row['student_id'])
                        except Exception:
                            continue

        # 2. Students who completed a task in the last 7 days
        if student_ids:
            tasks_result = supabase.table('task_assignments').select('student_id, completed_at, status').in_('student_id', student_ids).execute()
            if tasks_result.data:
                for row in tasks_result.data:
                    status = row.get('status', '').lower()
                    completed_at = row.get('completed_at')
                    if status == 'completed' and completed_at:
                        try:
                            dt = datetime.fromisoformat(str(completed_at).replace('Z', '+00:00'))
                            if dt >= week_ago:
                                engaged_students_set.add(row['student_id'])
                        except Exception:
                            continue

        student_engagement = round((len(engaged_students_set) / total_students) * 100, 1) if total_students else 0

        # --- MONTHLY GROWTH: % increase in users this month vs last month ---
        this_month = now.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
        last_month = (this_month - timedelta(days=1)).replace(day=1)
        users_this_month = 0
        users_last_month = 0
        for u in users:
            created_at = u.get('created_at')
            if created_at:
                try:
                    dt = datetime.fromisoformat(str(created_at).replace('Z', '+00:00'))
                    if dt >= this_month:
                        users_this_month += 1
                    elif last_month <= dt < this_month:
                        users_last_month += 1
                except Exception:
                    continue
        if users_last_month:
            monthly_growth = round(((users_this_month - users_last_month) / users_last_month) * 100, 1)
        else:
            monthly_growth = 100.0 if users_this_month else 0.0

        # --- SYSTEM ALERTS: users with status Pending or Rejected ---
        system_alerts = len([u for u in users if u.get('status', '').lower() in ['pending', 'rejected']])

        # --- MONTHLY REDEEMED REWARDS (COUNT) ---
        monthly_redeemed_rewards = 0
        start_of_month = now.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
        redemptions_result = supabase.table('reward_redemptions').select('processed_at').execute()
        if redemptions_result.data:
            for r in redemptions_result.data:
                processed_at = r.get('processed_at')
                if processed_at:
                    try:
                        dt = datetime.fromisoformat(str(processed_at).replace('Z', '+00:00'))
                        if dt >= start_of_month:
                            monthly_redeemed_rewards += 1  # Count each redemption
                    except Exception:
                        continue

        # --- DAILY LOGINS (last 7 days, using admin_activity_log for Authentication/Login) ---
        day_labels = []
        daily_logins = [0] * 7
        # Fetch all login activities in the last 7 days
        activity_result = supabase.table('admin_activity_log') \
            .select('created_at, action, activity') \
            .gte('created_at', (now - timedelta(days=6)).strftime('%Y-%m-%dT00:00:00Z')) \
            .eq('activity', 'Authentication') \
            .eq('action', 'Login') \
            .execute()
        activities = activity_result.data if activity_result.data else []
        for i in range(6, -1, -1):
            day = now - timedelta(days=i)
            day_labels.append(day.strftime('%A'))
            count = 0
            for act in activities:
                created_at = act.get('created_at')
                if created_at:
                    try:
                        dt = datetime.fromisoformat(str(created_at).replace('Z', '+00:00'))
                        if dt.date() == day.date():
                            count += 1
                    except Exception:
                        continue
            daily_logins[6 - i] = count

        # --- ENGAGEMENT BREAKDOWN ---
        engagement_breakdown = [total_teachers, total_students, total_admins]

        # --- NEW REGISTRATIONS (last 7 days) ---
        new_registrations = 0
        for u in users:
            created_at = u.get('created_at')
            if created_at:
                try:
                    dt = datetime.fromisoformat(str(created_at).replace('Z', '+00:00'))
                    if dt >= week_ago:
                        new_registrations += 1
                except Exception:
                    continue

        # --- POINTS REDEEMED (last 7 days) ---
        points_redeemed = 0
        if redemptions_result.data:
            for r in redemptions_result.data:
                processed_at = r.get('processed_at')
                if processed_at:
                    try:
                        dt = datetime.fromisoformat(str(processed_at).replace('Z', '+00:00'))
                        if dt >= week_ago:
                            points_redeemed += int(r.get('points_deducted', 0))
                    except Exception:
                        continue

        # --- MONTHLY POINTS TREND (all teachers, last 6 months) ---
        from collections import OrderedDict
        months = []
        month_labels = []
        now_month = now.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
        for i in range(5, -1, -1):  # Last 6 months
            month = (now_month - timedelta(days=30*i)).replace(day=1)
            months.append(month)
            month_labels.append(month.strftime('%b %Y'))

        # Fetch all points awarded (all teachers)
        points_result_all = supabase.table('points').select('points, received_at').execute()
        monthly_points = OrderedDict((label, 0) for label in month_labels)
        if points_result_all.data:
            for row in points_result_all.data:
                received_at = row.get('received_at')
                points = row.get('points', 0)
                if received_at:
                    try:
                        dt = datetime.fromisoformat(str(received_at).replace('Z', '+00:00'))
                        for i, month in enumerate(months):
                            if dt.year == month.year and dt.month == month.month:
                                monthly_points[month_labels[i]] += points
                                break
                    except Exception:
                        continue

        # --- SYSTEM UPTIME (dummy, always 99.9%) ---
        system_uptime = 99.9

        # --- RECENT TRENDS CHANGES (compare last 7 days vs previous 7 days) ---
        prev_week_ago = now - timedelta(days=13)
        prev_week_start = now - timedelta(days=13)
        prev_week_end = now - timedelta(days=7)
        prev_week_logins = 0
        for i in range(13, 6, -1):
            day = now - timedelta(days=i)
            count = 0
            for act in activities:
                created_at = act.get('created_at')
                if created_at:
                    try:
                        dt = datetime.fromisoformat(str(created_at).replace('Z', '+00:00'))
                        if dt.date() == day.date():
                            count += 1
                    except Exception:
                        continue
            prev_week_logins += count
        daily_logins_change = f"{round(((sum(daily_logins) - prev_week_logins) / prev_week_logins * 100), 1) if prev_week_logins else 100 if sum(daily_logins) else 0}%"

        prev_week_regs = 0
        for u in users:
            created_at = u.get('created_at')
            if created_at:
                try:
                    dt = datetime.fromisoformat(str(created_at).replace('Z', '+00:00'))
                    if prev_week_ago <= dt < week_ago:
                        prev_week_regs += 1
                except Exception:
                    continue
        new_registrations_change = f"{round(((new_registrations - prev_week_regs) / prev_week_regs * 100), 1) if prev_week_regs else 100 if new_registrations else 0}%"

        prev_week_points = 0
        if redemptions_result.data:
            for r in redemptions_result.data:
                processed_at = r.get('processed_at')
                if processed_at:
                    try:
                        dt = datetime.fromisoformat(str(processed_at).replace('Z', '+00:00'))
                        if prev_week_ago <= dt < week_ago:
                            prev_week_points += int(r.get('points_deducted', 0))
                    except Exception:
                        continue
        points_redeemed_change = f"{round(((points_redeemed - prev_week_points) / prev_week_points * 100), 1) if prev_week_points else 100 if points_redeemed else 0}%"

        system_uptime_change = "+0%"

        return render_template(
            'admin/admin_analytics.html',
            active_users=active_users,
            student_engagement=student_engagement,
            monthly_growth=monthly_growth,
            monthly_redeemed_rewards=monthly_redeemed_rewards,  # Now count of redemptions
            system_alerts=system_alerts,
            daily_logins=daily_logins,
            day_labels=day_labels,
            engagement_breakdown=engagement_breakdown,
            new_registrations=new_registrations,
            points_redeemed=points_redeemed,
            system_uptime=system_uptime,
            daily_logins_change=daily_logins_change,
            new_registrations_change=new_registrations_change,
            points_redeemed_change=points_redeemed_change,
            system_uptime_change=system_uptime_change,
            total_points_awarded=total_points_awarded,
            monthly_points_labels=list(monthly_points.keys()),
            monthly_points_data=list(monthly_points.values())
        )
    except Exception as e:
        app.logger.error(f"Error fetching admin analytics: {str(e)}")
        flash('Error loading analytics. Please try again later.', 'danger')
        return render_template(
            'admin/admin_analytics.html',
            active_users=0,
            student_engagement=0,
            monthly_growth=0,
            monthly_redeemed_rewards=0,
            system_alerts=0,
            daily_logins=[0]*7,
            day_labels=["Monday","Tuesday","Wednesday","Thursday","Friday","Saturday","Sunday"],
            engagement_breakdown=[0,0,0],
            new_registrations=0,
            points_redeemed=0,
            system_uptime=99.9,
            daily_logins_change="+0%",
            new_registrations_change="+0%",
            points_redeemed_change="+0%",
            system_uptime_change="+0%",
            total_points_awarded=0,
            monthly_points_labels=["Jan 2025", "Feb 2025", "Mar 2025", "Apr 2025", "May 2025", "Jun 2025"],
            monthly_points_data=[0, 0, 0, 0, 0, 0]
        )

@app.route('/api/admin-analytics-summary')
def admin_analytics_summary():
    try:
        now = datetime.now(timezone.utc)
        # --- DAILY LOGINS (last 7 days) ---
        day_labels = []
        daily_logins = [0] * 7
        activity_result = supabase.table('admin_activity_log') \
            .select('created_at, action, activity') \
            .gte('created_at', (now - timedelta(days=6)).strftime('%Y-%m-%dT00:00:00Z')) \
            .eq('activity', 'Authentication') \
            .eq('action', 'Login') \
            .execute()
        activities = activity_result.data if activity_result.data else []
        for i in range(6, -1, -1):
            day = now - timedelta(days=i)
            day_labels.append(day.strftime('%A'))
            count = 0
            for act in activities:
                created_at = act.get('created_at')
                if created_at:
                    try:
                        dt = datetime.fromisoformat(str(created_at).replace('Z', '+00:00'))
                        if dt.date() == day.date():
                            count += 1
                    except Exception:
                        continue
            daily_logins[6 - i] = count

        # --- MONTHLY POINTS TREND (last 6 months) ---
        from collections import OrderedDict
        months = []
        month_labels = []
        now_month = now.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
        for i in range(5, -1, -1):  # Last 6 months
            month = (now_month - timedelta(days=30*i)).replace(day=1)
            months.append(month)
            month_labels.append(month.strftime('%b %Y'))

        points_result_all = supabase.table('points').select('points, received_at').execute()
        monthly_points = OrderedDict((label, 0) for label in month_labels)
        if points_result_all.data:
            for row in points_result_all.data:
                received_at = row.get('received_at')
                points = row.get('points', 0)
                if received_at:
                    try:
                        dt = datetime.fromisoformat(str(received_at).replace('Z', '+00:00'))
                        for i, month in enumerate(months):
                            if dt.year == month.year and dt.month == month.month:
                                monthly_points[month_labels[i]] += points
                                break
                    except Exception:
                        continue

        # --- ENGAGEMENT BREAKDOWN ---
        users_result = supabase.table('user_info').select('id, role').execute()
        users = users_result.data if users_result.data else []
        total_teachers = len([u for u in users if u['role'] == 'Teacher'])
        total_students = len([u for u in users if u['role'] == 'Student'])
        total_admins = len([u for u in users if u['role'] == 'Admin'])
        engagement_breakdown = [total_teachers, total_students, total_admins]

        return jsonify({
            'day_labels': day_labels,
            'daily_logins': daily_logins,
            'monthly_points_labels': list(monthly_points.keys()),
            'monthly_points_data': list(monthly_points.values()),
            'engagement_breakdown': engagement_breakdown
        })
    except Exception as e:
        return jsonify({
            'day_labels': ["Monday","Tuesday","Wednesday","Thursday","Friday","Saturday","Sunday"],
            'daily_logins': [0]*7,
            'monthly_points_labels': [],
            'monthly_points_data': [],
            'engagement_breakdown': [0,0,0]
        }), 500

#############################USER_MANAGEMENT.HTML#############################
@app.route('/api/users', methods=['POST'])
def create_user():
    try:
        # Get request data
        data = request.json
        if not data:
            return jsonify({'success': False, 'message': 'Invalid request data'}), 400

        # Validate required fields
        required_fields = ['firstName', 'lastName', 'email', 'userRole']
        missing_fields = [field for field in required_fields if not data.get(field)]
        if missing_fields:
            return jsonify({'success': False, 'message': f'Missing fields: {", ".join(missing_fields)}'}), 400

        # Extract user data
        user_data = {
            "first_name": data.get('firstName'),
            "middle_name": data.get('middleName', ''),
            "last_name": data.get('lastName'),
            "gender": data.get('gender'),
            "email": data.get('email'),
            "mobile_no": data.get('mobileNo'),
            "role": data.get('userRole').capitalize(),  # Ensure proper capitalization
            "status": "Active"
        }

        # Add role-specific fields
        if user_data["role"] == "Teacher":
            user_data["subject"] = data.get('subject')
        elif user_data["role"] == "Student":
            user_data["year_level"] = data.get('grade')
            user_data["section"] = data.get('section')

        # Generate or use provided password
        if data.get('generatePassword', True):
            password = generate_password()
        else:
            password = data.get('password')

        if not password:
            return jsonify({'success': False, 'message': 'Password is required'}), 400

        # Add the password to user_data
        user_data["password"] = password  # In a real app, hash the password before storing it

        # Store user in Supabase
        result = supabase.table('user_info').insert(user_data).execute()

        # Check for errors in the result
        if hasattr(result, 'error') and result.error:
            raise Exception(result.error.message)

        # --- ACTIVITY LOG: Add User ---
        admin_id = session.get('user_id')
        admin_role = session.get('role')
        record_admin_activity_log(
            user_id=admin_id,
            action='Add User',
            activity='User Management',
            description=f"Added user: {user_data['first_name']} {user_data['last_name']} ({user_data['role']})",
            user_role=admin_role
        )

        # Send welcome email if requested
        if data.get('sendWelcomeEmail', True):
            send_welcome_email(
                email=user_data["email"],
                password=password,
                first_name=user_data["first_name"],
                role=user_data["role"]
            )

        response_data = {
            'success': True,
            'message': 'User created successfully',
            'user': result.data[0] if hasattr(result, 'data') and result.data else None
        }

        # Include generated password in response if applicable
        if data.get('generatePassword', True):
            response_data['generatedPassword'] = password

        return jsonify(response_data), 201

    except Exception as e:
        app.logger.error(f"Error creating user: {str(e)}")
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/api/users', methods=['GET'])
def get_users():
    try:
        # Only fetch users with Active or Inactive status (exclude Pending)
        result = supabase.table('user_info').select('*').in_('status', ['Active', 'Inactive']).order('created_at', desc=True).execute()
        
        if hasattr(result, 'error') and result.error:
            return jsonify({'success': False, 'message': result.error.message}), 500

        users = result.data if result.data else []

        # Get all profile pictures
        user_ids = [u['id'] for u in users]
        pics_map = {}
        if user_ids:
            pics_result = supabase.table('profile_pictures').select('user_id, file_path').in_('user_id', user_ids).execute()
            if pics_result.data:
                # Build public URL for each file_path
                for pic in pics_result.data:
                    # If file_path is already a full URL, use it; otherwise, build it
                    file_path = pic['file_path']
                    if file_path.startswith('http'):
                        public_url = file_path
                    else:
                        # Replace with your actual Supabase project URL
                        public_url = f"https://bdcmztafoacnsfdpudv.supabase.co/storage/v1/object/public/profile-pictures/{file_path}"
                    pics_map[pic['user_id']] = public_url

        # Attach avatar to each user
        for u in users:
            u['avatar'] = pics_map.get(u['id'], '/static/image/default-avatar.png')

        return jsonify({
            'success': True, 
            'users': users
        })
    except Exception as e:
        app.logger.error(f"Error fetching users: {str(e)}")
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/api/users/<int:user_id>', methods=['PUT'])
def update_user(user_id):
    try:
        # Get request data
        data = request.json
        if not data:
            return jsonify({'success': False, 'message': 'Invalid request data'}), 400

        # Validate required fields
        required_fields = ['firstName', 'lastName', 'email', 'userRole']
        missing_fields = [field for field in required_fields if not data.get(field)]
        if missing_fields:
            return jsonify({'success': False, 'message': f'Missing fields: {", ".join(missing_fields)}'}), 400

        # Check if user exists
        check_result = supabase.table('user_info').select('id').eq('id', user_id).execute()
        if not check_result.data:
            return jsonify({'success': False, 'message': 'User not found'}), 404

        # Check if the email is already used by another user
        email = data.get('email')
        email_check_result = supabase.table('user_info').select('id').eq('email', email).execute()
        # Only consider it a duplicate if the email belongs to a DIFFERENT user
        if email_check_result.data:
            email_belongs_to_other_user = False
            for user in email_check_result.data:
                if user['id'] != user_id:
                    email_belongs_to_other_user = True
                    break
            if email_belongs_to_other_user:
                return jsonify({'success': False, 'message': 'Email is already in use by another user'}), 400

        # Extract user data for update
        user_data = {
            "first_name": data.get('firstName'),
            "middle_name": data.get('middleName', ''),
            "last_name": data.get('lastName'),
            "gender": data.get('gender'),
            "email": data.get('email'),
            "mobile_no": data.get('mobileNo'),
            "role": data.get('userRole').capitalize()
        }

        # Add role-specific fields
        if user_data["role"] == "Teacher":
            user_data["subject"] = data.get('subject')
            # Remove any potential student-specific fields
            user_data.pop("year_level", None)
            user_data.pop("section", None)
        elif user_data["role"] == "Student":
            user_data["year_level"] = data.get('grade')
            user_data["section"] = data.get('section')
            # Remove any potential teacher-specific fields
            user_data.pop("subject", None)

        # Update password if user didn't choose to generate password
        if not data.get('generatePassword', True) and data.get('password'):
            user_data["password"] = data.get('password')

        # Update user in Supabase
        result = supabase.table('user_info').update(user_data).eq('id', user_id).execute()

        # --- ACTIVITY LOG: Edit User ---
        admin_id = session.get('user_id')
        admin_role = session.get('role')
        record_admin_activity_log(
            user_id=admin_id,
            action='Edit User',
            activity='User Management',
            description=f"Edited user: {user_data['first_name']} {user_data['last_name']} ({user_data['role']})",
            user_role=admin_role
        )

        return jsonify({
            'success': True,
            'message': 'User updated successfully',
            'user': result.data[0] if result.data else None
        }), 200

    except Exception as e:
        app.logger.error(f"Error updating user: {str(e)}")
        return jsonify({'success': False, 'message': str(e)}), 500
    
@app.route('/api/users/<int:user_id>/reset-password', methods=['POST'])
def reset_user_password(user_id):
    user = None
    new_password = None
    email_sent = False
    database_success = False
    operation_error = None
    
    try:
        # Check if user exists
        check_result = supabase.table('user_info').select('id, email, first_name, last_name, role').eq('id', user_id).execute()
        if not check_result.data:
            # LOG ACTIVITY EVEN WHEN USER NOT FOUND
            admin_id = session.get('user_id')
            admin_role = session.get('role')
            if admin_id:
                record_admin_activity_log(
                    user_id=admin_id,
                    action='Reset Password - FAILED',
                    activity='User Management',
                    description=f"Attempted to reset password for user ID {user_id} - User not found",
                    user_role=admin_role
                )
            return jsonify({'success': False, 'message': 'User not found', 'emailSent': False}), 404
        
        user = check_result.data[0]
        new_password = generate_password()
        
        # Update the user's password in the database
        result = supabase.table('user_info').update({"password": new_password}).eq('id', user_id).execute()
        
        if hasattr(result, 'error') and result.error:
            operation_error = f"Database error: {result.error.message}"
            raise Exception(operation_error)
        
        database_success = True
        
    except Exception as e:
        app.logger.error(f"Error in password reset operation: {str(e)}")
        operation_error = str(e)
        database_success = False
    
    # ALWAYS attempt to send email if we have user info and password
    if user and new_password:
        try:
            # FIXED: Call the function with the correct parameter names
            email_result = send_password_reset_email(
                user_email=user["email"],           # Changed from 'email'
                user_name=user["first_name"],       # Changed from 'first_name'  
                new_password=new_password           # Changed from 'password'
            )
            email_sent = True
            app.logger.info(f"Password reset email sent to {user['email']}")
        except Exception as email_error:
            app.logger.error(f"Failed to send password reset email: {str(email_error)}")
            email_sent = False
    
    # LOG ACTIVITY based on operation outcome
    admin_id = session.get('user_id')
    admin_role = session.get('role')
    if admin_id and user:
        if database_success:
            record_admin_activity_log(
                user_id=admin_id,
                action='Reset Password',
                activity='User Management',
                description=f"Reset password for user: {user['first_name']} {user['last_name']} ({user['role']}) - Email sent: {email_sent}",
                user_role=admin_role
            )
        else:
            record_admin_activity_log(
                user_id=admin_id,
                action='Reset Password - FAILED',
                activity='User Management',
                description=f"Failed to reset password for {user['first_name']} {user['last_name']} - Error: {operation_error}",
                user_role=admin_role
            )
    
    # Prepare response
    if database_success:
        response_data = {
            'success': True,
            'message': 'Password reset successfully',
            'newPassword': new_password,
            'email': user["email"],
            'emailSent': email_sent
        }
        return jsonify(response_data), 200
    else:
        response_data = {
            'success': False, 
            'message': operation_error or 'Database operation failed',
            'emailSent': email_sent
        }
        if new_password and user:
            response_data['newPassword'] = new_password
            response_data['email'] = user["email"]
        return jsonify(response_data), 500

@app.route('/api/users/<int:user_id>', methods=['GET'])
def get_user(user_id):
    try:
        # Get specific user from database
        result = supabase.table('user_info').select('*').eq('id', user_id).execute()
        
        if not result.data:
            return jsonify({'success': False, 'message': 'User not found'}), 404
            
        user = result.data[0]

        # Fetch profile picture path if exists
        pic_result = supabase.table('profile_pictures').select('file_path').eq('user_id', user_id).execute()
        if pic_result.data and len(pic_result.data) > 0:
            user['profile_picture'] = pic_result.data[0]['file_path']
        else:
            user['profile_picture'] = None
        
        return jsonify({
            'success': True, 
            'user': user
        })
        
    except Exception as e:
        app.logger.error(f"Error fetching user: {str(e)}")
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/api/pending-students', methods=['GET'])
def api_pending_students():
    try:
        # Only fetch users with Pending status who are students
        result = supabase.table('user_info').select('*').eq('status', 'Pending').eq('role', 'Student').order('created_at', desc=True).execute()
        
        if hasattr(result, 'error') and result.error:
            return jsonify({'success': False, 'message': result.error.message}), 500

        students = result.data if result.data else []
        return jsonify({'success': True, 'students': students})
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/api/approve-student/<int:user_id>', methods=['POST'])
def api_approve_student(user_id):
    # Check if user is logged in and has admin privileges (case-insensitive)
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': 'Unauthorized - Please log in'}), 401
    
    # Case-insensitive role check
    user_role = session.get('role', '').lower()
    if user_role != 'admin':
        return jsonify({'success': False, 'message': 'Unauthorized - Admin access required'}), 401

    try:
        # Update student status to active (only update status since approved_at doesn't exist)
        update_result = supabase.table('user_info').update({
            'status': 'Active'
            # Remove 'approved_at': 'now()' since column doesn't exist
        }).eq('id', user_id).execute()

        if not update_result.data:
            return jsonify({'success': False, 'message': 'Student not found'}), 404

        # Get student information
        student_result = supabase.table('user_info').select(
            'first_name, last_name, email, year_level, section'
        ).eq('id', user_id).execute()

        if not student_result.data:
            return jsonify({'success': False, 'message': 'Student information not found'}), 404

        student = student_result.data[0]
        
        # Send approval email
        send_approval_email(student)
        
        # Log this action
        admin_id = session.get('user_id')
        admin_role = session.get('role')
        record_admin_activity_log(
            user_id=admin_id,
            action='Approve Student',
            activity='User Management',
            description=f"Approved student: {student['first_name']} {student['last_name']} (ID: {user_id})",
            user_role=admin_role
        )
        
        return jsonify({
            'success': True, 
            'message': 'Student approved successfully and notification sent!'
        })

    except Exception as e:
        app.logger.error(f"Error approving student: {str(e)}")
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/api/reject-student/<int:user_id>', methods=['POST'])
def api_reject_student(user_id):
    # Check if user is logged in and has admin privileges (case-insensitive)
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': 'Unauthorized - Please log in'}), 401
    
    # Case-insensitive role check
    user_role = session.get('role', '').lower()
    if user_role != 'admin':
        return jsonify({'success': False, 'message': 'Unauthorized - Admin access required'}), 401

    try:
        # Get student information before deleting (for email and logging)
        student_result = supabase.table('user_info').select(
            'first_name, last_name, email, year_level, section'
        ).eq('id', user_id).execute()

        if not student_result.data:
            return jsonify({'success': False, 'message': 'Student not found'}), 404

        student = student_result.data[0]
        
        # Send rejection email first
        send_rejection_email(student)
        
        # COMPLETELY DELETE THE USER from the database
        delete_result = supabase.table('user_info').delete().eq('id', user_id).execute()

        # Log this action
        admin_id = session.get('user_id')
        admin_role = session.get('role')
        record_admin_activity_log(
            user_id=admin_id,
            action='Reject and Remove Student',
            activity='User Management',
            description=f"Rejected and permanently removed student: {student['first_name']} {student['last_name']} (ID: {user_id})",
            user_role=admin_role
        )

        return jsonify({
            'success': True, 
            'message': 'Student rejected and removed successfully!'
        })

    except Exception as e:
        app.logger.error(f"Error rejecting and removing student: {str(e)}")
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/api/users/<int:user_id>/status', methods=['PUT'])
def update_user_status(user_id):
    try:
        data = request.json
        if not data or 'status' not in data:
            return jsonify({'success': False, 'message': 'Status is required'}), 400

        new_status = data['status']
        if new_status not in ['Active', 'Inactive']:
            return jsonify({'success': False, 'message': 'Status must be either "Active" or "Inactive"'}), 400

        # Update user status in Supabase
        result = supabase.table('user_info').update({'status': new_status}).eq('id', user_id).execute()

        if hasattr(result, 'error') and result.error:
            raise Exception(result.error.message)

        # Get user info for activity log
        user_result = supabase.table('user_info').select('first_name, last_name, role').eq('id', user_id).execute()
        user_info = user_result.data[0] if user_result.data else {}

        # --- ACTIVITY LOG: Update User Status ---
        admin_id = session.get('user_id')
        admin_role = session.get('role')
        record_admin_activity_log(
            user_id=admin_id,
            action='Update User Status',
            activity='User Management',
            description=f"Changed status of {user_info.get('first_name', '')} {user_info.get('last_name', '')} to {new_status}",
            user_role=admin_role
        )

        return jsonify({
            'success': True,
            'message': f'User status updated to {new_status} successfully'
        })

    except Exception as e:
        app.logger.error(f"Error updating user status: {str(e)}")
        return jsonify({'success': False, 'message': str(e)}), 500

#############################ADMIN_ACTIVITY_LOG.HTML#############################
def record_admin_activity_log(user_id, action, activity, description, user_role=None, details=None):
    try:
        supabase.table('admin_activity_log').insert({
            'user_id': user_id,
            'action': action,
            'activity': activity,
            'description': description,
            'user_role': user_role,
            'details': details
        }).execute()
    except Exception as e:
        app.logger.error(f"Failed to record admin activity log: {str(e)}")

@app.route('/api/admin-activity-log')
def api_admin_activity_log():
    try:
        all_logs = []
        batch_size = 1000
        offset = 0
        
        while True:
            result = (
                supabase.table('admin_activity_log')
                .select('history_id, user_id, user_role, action, activity, description, details, created_at, '
                       'user_info(id, first_name, last_name, role, profile_pictures(file_path))')
                .order('created_at', desc=True)
                .range(offset, offset + batch_size - 1)
                .execute()
            )
            
            if not result.data:
                break
                
            all_logs.extend(result.data)
            
            # If we got fewer results than batch_size, we've reached the end
            if len(result.data) < batch_size:
                break
                
            offset += batch_size
        
        return jsonify({'success': True, 'logs': all_logs, 'total': len(all_logs)})
        
    except Exception as e:
        app.logger.error(f"Error fetching activity log: {str(e)}")
        return jsonify({'success': False, 'logs': []}), 500


# For export - optimized endpoint that fetches everything
@app.route('/api/admin-activity-log/export')
def api_admin_activity_log_export():
    try:
        # Get filter parameters
        activity = request.args.get('activity', 'all')
        role = request.args.get('role', 'all')
        date_from = request.args.get('date_from', '')
        date_to = request.args.get('date_to', '')
        
        # Build query with filters
        query = supabase.table('admin_activity_log').select(
            'history_id, user_id, user_role, action, activity, description, details, created_at, '
            'user_info(id, first_name, last_name, role, profile_pictures(file_path))'
        )
        
        # Apply filters
        if activity != 'all':
            query = query.ilike('activity', f'%{activity}%')
        if role != 'all':
            query = query.eq('user_role', role)
        if date_from:
            query = query.gte('created_at', date_from)
        if date_to:
            from datetime import datetime, timedelta
            end_date = datetime.strptime(date_to, '%Y-%m-%d') + timedelta(days=1)
            query = query.lt('created_at', end_date.strftime('%Y-%m-%d'))
        
        query = query.order('created_at', desc=True)
        
        # Fetch all matching records in batches
        all_logs = []
        batch_size = 1000
        offset = 0
        
        while True:
            result = query.range(offset, offset + batch_size - 1).execute()
            
            if not result.data:
                break
                
            all_logs.extend(result.data)
            
            if len(result.data) < batch_size:
                break
                
            offset += batch_size
        
        return jsonify({'success': True, 'logs': all_logs})
        
    except Exception as e:
        app.logger.error(f"Error fetching logs for export: {str(e)}")
        return jsonify({'success': False, 'logs': []}), 500
    
@app.route('/api/recent-admin-activities')
def api_recent_admin_activities():
    try:
        result = (
            supabase.table('admin_activity_log')
            .select('created_at, action, activity, description, user_info(first_name, last_name, role)')
            .order('created_at', desc=True)
            .limit(6)
            .execute()
        )
        activities = result.data if result.data else []
        return jsonify({'success': True, 'activities': activities})
    except Exception as e:
        app.logger.error(f"Error fetching recent admin activities: {str(e)}")
        return jsonify({'success': False, 'activities': []}), 500

#############################ADMIN_SUBJECT.HTML#############################
@app.route('/api/teachers', methods=['GET'])
def get_teachers():
    try:
        # Get current school year
        current_school_year = None
        try:
            quarters_result = supabase.table('quarters') \
                .select('school_year') \
                .order('school_year', desc=True) \
                .limit(1) \
                .execute()
            
            if quarters_result.data and len(quarters_result.data) > 0:
                current_school_year = quarters_result.data[0].get('school_year')
        except Exception as quarter_error:
            logger.warning(f"Could not fetch school year from quarters: {str(quarter_error)}")
        
        # Get all teachers
        teachers_result = supabase.table('user_info').select('id, first_name, last_name, subject, email').eq('role', 'Teacher').execute()
        teachers = teachers_result.data if teachers_result.data else []

        # Get assignment counts for each teacher - FILTERED by current school year and active status
        if current_school_year:
            assignments_result = supabase.table('teacher_class_assignments') \
                .select('teacher_id') \
                .eq('school_year', current_school_year) \
                .eq('status', 'active') \
                .execute()
        else:
            # Fallback: get active assignments from any school year
            assignments_result = supabase.table('teacher_class_assignments') \
                .select('teacher_id') \
                .eq('status', 'active') \
                .execute()
        
        assignment_counts = {}
        if assignments_result.data:
            for row in assignments_result.data:
                tid = row['teacher_id']
                assignment_counts[tid] = assignment_counts.get(tid, 0) + 1

        # Build teacher list with assignment count
        teacher_list = []
        for teacher in teachers:
            teacher_list.append({
                'id': teacher['id'],
                'first_name': teacher['first_name'],
                'last_name': teacher['last_name'],
                'subject': teacher.get('subject', ''),
                'email': teacher.get('email', ''),
                'assignments': assignment_counts.get(teacher['id'], 0),
                'current_school_year': current_school_year
            })

        return jsonify({'success': True, 'teachers': teacher_list, 'current_school_year': current_school_year})
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/api/classroom-assignments', methods=['POST'])
def assign_teacher_to_classroom():
    try:
        data = request.get_json()
        teacher_id = data.get('teacher_id')
        grade_level = data.get('grade_level')
        section = data.get('section')
        subject = data.get('subject')

        # DEBUG: Log incoming request
        logger.info(f"\n{'='*60}")
        logger.info(f"[ASSIGNMENT DEBUG] Incoming Request:")
        logger.info(f"  Teacher ID: {teacher_id}")
        logger.info(f"  Grade: {grade_level}")
        logger.info(f"  Section: {section}")
        logger.info(f"  Subject: {subject}")

        if not teacher_id or not grade_level or not section or not subject:
            return jsonify({'success': False, 'message': 'Missing required fields.'}), 400

        # Get the current school year
        quarters_result = supabase.table('quarters') \
            .select('school_year') \
            .order('school_year', desc=True) \
            .limit(1) \
            .execute()
        
        current_school_year = quarters_result.data[0]['school_year'] if quarters_result.data else None
        logger.info(f"  Current School Year: {current_school_year}")
        if not current_school_year:
            return jsonify({'success': False, 'message': 'No active school year found.'}), 400

        # Check if this classroom already has 8 subjects assigned (in current school year only)
        classroom_assignments = supabase.table('teacher_class_assignments') \
            .select('assignment_id') \
            .eq('grade_level', grade_level) \
            .eq('section', section) \
            .eq('status', 'active') \
            .eq('school_year', current_school_year) \
            .execute()
        logger.info(f"\n[VALIDATION 1] Classroom Subject Count:")
        logger.info(f"  Current assignments in {grade_level}-{section} ({current_school_year}): {len(classroom_assignments.data) if classroom_assignments.data else 0}")
        if classroom_assignments.data and len(classroom_assignments.data) >= 8:
            return jsonify({'success': False, 'message': 'This classroom already has 8 subjects assigned.'}), 400

        # Check if this subject in this classroom is already assigned (in current school year only)
        existing = supabase.table('teacher_class_assignments') \
            .select('assignment_id') \
            .eq('grade_level', grade_level) \
            .eq('section', section) \
            .eq('subject', subject) \
            .eq('status', 'active') \
            .eq('school_year', current_school_year) \
            .execute()
        logger.info(f"\n[VALIDATION 2] Subject Duplicate Check:")
        logger.info(f"  Existing {subject} in {grade_level}-{section} ({current_school_year}): {len(existing.data) if existing.data else 0}")
        if existing.data:
            return jsonify({'success': False, 'message': f'{subject} is already assigned in Grade {grade_level} Section {section}.'}), 400

        # Insert the new assignment with assigned_at and status
        now = datetime.now(philippines_tz).isoformat()
        insert_data = {
            'teacher_id': teacher_id,
            'grade_level': grade_level,
            'section': section,
            'subject': subject,
            'assigned_at': now,
            'status': 'active',
            'school_year': current_school_year
        }
        logger.info(f"\n[INSERT DATA]:")
        logger.info(f"  {insert_data}")
        
        result = supabase.table('teacher_class_assignments').insert(insert_data).execute()

        logger.info(f"\n[INSERT RESULT]:")
        logger.info(f"  Result data: {result.data}")
        logger.info(f"  Has error: {hasattr(result, 'error') and result.error}")
        
        if hasattr(result, 'error') and result.error:
            logger.error(f"  ERROR MESSAGE: {result.error}")
            raise Exception(str(result.error))

        # --- ACTIVITY LOG: Assign Teacher to Classroom ---
        admin_id = session.get('user_id')
        admin_role = session.get('role')
        record_admin_activity_log(
            user_id=admin_id,
            action='Assign Teacher',
            activity='Classroom Management',
            description=f"Assigned teacher ID {teacher_id} to Grade {grade_level} Section {section} for {subject}.",
            user_role=admin_role
        )

        # --- NOTIFICATION: Notify Teacher ---
        notif_title = "New Classroom Assignment"
        notif_message = f"You have been assigned to Grade {grade_level} - Section {section}."
        assignment_id = None
        if result.data and 'assignment_id' in result.data[0]:
            assignment_id = result.data[0]['assignment_id']
        supabase.table('notifications').insert({
            'user_id': teacher_id,
            'sender_id': admin_id,
            'title': notif_title,
            'message': notif_message,
            'assignment_id': assignment_id,
            'notif_type': 'Class Assignment',
            'status': 'Unread'
        }).execute()

        logger.info(f"[SUCCESS] Assignment created successfully!")
        logger.info(f"{'='*60}\n")
        return jsonify({'success': True, 'message': 'Teacher assigned to classroom successfully.'}), 201

    except Exception as e:
        logger.error(f"[ERROR] {str(e)}")
        logger.error(f"{'='*60}\n")
        return jsonify({'success': False, 'message': str(e)}), 500  

@app.route('/api/classroom-assignments', methods=['GET'])
def get_classroom_assignments():
    try:
        # Get the current school year from quarters table
        current_school_year = None
        
        try:
            quarters_result = supabase.table('quarters') \
                .select('school_year') \
                .order('school_year', desc=True) \
                .limit(1) \
                .execute()
            
            if quarters_result.data and len(quarters_result.data) > 0:
                current_school_year = quarters_result.data[0].get('school_year')
                logger.info(f"Current school year from quarters: {current_school_year}")
        except Exception as quarter_error:
            logger.warning(f"Could not fetch school year from quarters: {str(quarter_error)}")
        
        # If we couldn't get it from quarters, try to get most recent from assignments
        if not current_school_year:
            try:
                all_assignments = supabase.table('teacher_class_assignments') \
                    .select('school_year') \
                    .order('school_year', desc=True) \
                    .limit(100) \
                    .execute()
                
                if all_assignments.data:
                    # Get the first non-null school_year
                    for assignment in all_assignments.data:
                        if assignment.get('school_year'):
                            current_school_year = assignment['school_year']
                            logger.info(f"Current school year from assignments table: {current_school_year}")
                            break
            except Exception as assign_error:
                logger.warning(f"Could not fetch school year from assignments: {str(assign_error)}")
        
        if not current_school_year:
            logger.warning("No school year found in either quarters or assignments table")
        
        # Get only ACTIVE assignments for the current school year
        if current_school_year:
            assignments_result = supabase.table('teacher_class_assignments') \
                .select('*') \
                .eq('status', 'active') \
                .eq('school_year', current_school_year) \
                .execute()
        else:
            # Fallback: get active assignments from any school year
            assignments_result = supabase.table('teacher_class_assignments') \
                .select('*') \
                .eq('status', 'active') \
                .execute()
        
        assignments = assignments_result.data if assignments_result.data else []

        # Get all teachers for mapping
        teachers_result = supabase.table('user_info').select('id, first_name, last_name, subject, email').eq('role', 'Teacher').execute()
        teachers = {t['id']: t for t in teachers_result.data} if teachers_result.data else {}

        # Count assignments per teacher
        assignment_counts = {}
        # Count subjects per classroom (grade_level, section)
        classroom_subject_counts = {}
        teacher_ids = set()
        for row in assignments:
            tid = row['teacher_id']
            assignment_counts[tid] = assignment_counts.get(tid, 0) + 1
            teacher_ids.add(tid)
            key = (row['grade_level'], row['section'])
            classroom_subject_counts[key] = classroom_subject_counts.get(key, 0) + 1

        # --- Fetch ALL teacher profile pictures in one query ---
        pics_map = {}
        if teacher_ids:
            pics_result = supabase.table('profile_pictures').select('user_id, file_path').in_('user_id', list(teacher_ids)).execute()
            if pics_result.data:
                for p in pics_result.data:
                    pics_map[p['user_id']] = p['file_path']

        # Build assignment list with teacher info and profile picture
        assignment_list = []
        for assignment in assignments:
            teacher = teachers.get(assignment['teacher_id'], {})
            assigned_at = assignment.get('assigned_at', '')
            assigned_date = ''
            if assigned_at:
                try:
                    assigned_date = str(assigned_at)[:10]
                except Exception:
                    assigned_date = assigned_at  # fallback

            # Use pre-fetched profile picture
            profile_pic = pics_map.get(assignment['teacher_id'], '/static/image/default-avatar.png')

            assignment_list.append({
                'assignment_id': assignment['assignment_id'],
                'teacher_id': assignment['teacher_id'],
                'teacher_name': f"{teacher.get('first_name', '')} {teacher.get('last_name', '')}",
                'teacher_email': teacher.get('email', ''),
                'subject': assignment.get('subject', ''),
                'assignment_count': assignment_counts.get(assignment['teacher_id'], 0),
                'grade_level': assignment['grade_level'],
                'section': assignment['section'],
                'assigned_at': assigned_date,
                'profile_pic': profile_pic
            })

        # Add classroom subject counts to the response
        classroom_subject_counts_serializable = {
            f"{grade}_{section}": count
            for (grade, section), count in classroom_subject_counts.items()
        }

        return jsonify({
            'success': True,
            'assignments': assignment_list,
            'classroom_subject_counts': classroom_subject_counts_serializable,
            'current_school_year': current_school_year
        })
    except Exception as e:
        logger.error(f"Error in get_classroom_assignments: {str(e)}")
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/api/classroom-stats', methods=['GET'])
def get_classroom_stats():
    """Get statistics about classroom assignments"""
    try:
        # Get the current school year
        current_school_year = None
        quarters_result = supabase.table('quarters').select('school_year').order('school_year', desc=True).limit(1).execute()
        if quarters_result.data:
            current_school_year = quarters_result.data[0].get('school_year')
        
        # Get active assignments for current school year
        if current_school_year:
            assignments_result = supabase.table('teacher_class_assignments').select('*').eq('status', 'active').eq('school_year', current_school_year).execute()
        else:
            assignments_result = supabase.table('teacher_class_assignments').select('*').eq('status', 'active').execute()
        
        assignments = assignments_result.data if assignments_result.data else []
        
        # Calculate stats
        total_assignments = len(assignments)
        assigned_teachers = len(set(a['teacher_id'] for a in assignments)) if assignments else 0
        
        # Get all classrooms (grade 7-10 with all possible sections)
        all_classrooms = set()
        assigned_classrooms = set()
        
        for assignment in assignments:
            assigned_classrooms.add((assignment['grade_level'], assignment['section']))
        
        # Generate all possible classroom combinations
        grades = ['7', '8', '9', '10']
        sections = ['Love', 'Faith', 'Hope', 'Peace', 'Matthew', 'Mark', 'Luke', 'John', 
                   'Psalms', 'Jeremiah', 'Isaiah', 'Proverbs', 'Deuteronomy', 'Leviticus', 'Exodus', 'Genesis']
        
        for grade in grades:
            # Only specific sections per grade
            if grade == '7':
                grade_sections = ['Love', 'Faith', 'Hope', 'Peace']
            elif grade == '8':
                grade_sections = ['Matthew', 'Mark', 'Luke', 'John']
            elif grade == '9':
                grade_sections = ['Psalms', 'Jeremiah', 'Isaiah', 'Proverbs']
            elif grade == '10':
                grade_sections = ['Deuteronomy', 'Leviticus', 'Exodus', 'Genesis']
            else:
                grade_sections = []
            
            for section in grade_sections:
                all_classrooms.add((grade, section))
        
        total_classrooms = len(all_classrooms)
        unassigned_classrooms = total_classrooms - len(assigned_classrooms)
        
        return jsonify({
            'success': True,
            'total_assignments': total_assignments,
            'assigned_teachers': assigned_teachers,
            'total_classrooms': total_classrooms,
            'unassigned_classrooms': unassigned_classrooms,
            'current_school_year': current_school_year
        })
    except Exception as e:
        logger.error(f"Error in get_classroom_stats: {str(e)}")
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/api/classroom-assignments/archived', methods=['GET'])
def get_archived_classroom_assignments():
    """Get archived/ended classroom assignments for historical reference"""
    try:
        # Get all ENDED/ARCHIVED assignments (status is NOT 'active')
        # Using order by ended_date to get recently ended assignments first
        try:
            all_assignments_result = supabase.table('teacher_class_assignments') \
                .select('*') \
                .order('ended_date', desc=True) \
                .execute()
            
            # Filter to only ended/archived assignments (not 'active')
            assignments = [a for a in (all_assignments_result.data or []) if a.get('status') != 'active']
        except Exception as e:
            logger.warning(f"Error querying assignments: {str(e)}")
            assignments = []

        # Get all teachers for mapping
        teachers_result = supabase.table('user_info').select('id, first_name, last_name, subject, email').eq('role', 'Teacher').execute()
        teachers = {t['id']: t for t in teachers_result.data} if teachers_result.data else {}

        # Fetch ALL quarters ONCE (not for each assignment)
        quarters_result = supabase.table('quarters').select('school_year, end_date').execute()
        quarters_list = quarters_result.data if quarters_result.data else []

        # Fetch teacher profile pictures
        pics_map = {}
        teacher_ids = set(a['teacher_id'] for a in assignments if a.get('teacher_id'))
        if teacher_ids:
            pics_result = supabase.table('profile_pictures').select('user_id, file_path').in_('user_id', list(teacher_ids)).execute()
            if pics_result.data:
                for p in pics_result.data:
                    pics_map[p['user_id']] = p['file_path']

        # Build assignment list with teacher info and profile picture
        assignment_list = []
        for assignment in assignments:
            teacher = teachers.get(assignment['teacher_id'], {})
            assigned_at = assignment.get('assigned_at', '')
            assigned_date = ''
            if assigned_at:
                try:
                    assigned_date = str(assigned_at)[:10]
                except Exception:
                    assigned_date = assigned_at

            ended_date = assignment.get('ended_date', '')
            ended_date_formatted = ''
            if ended_date:
                try:
                    ended_date_formatted = str(ended_date)[:10]
                except Exception:
                    ended_date_formatted = ended_date

            profile_pic = pics_map.get(assignment['teacher_id'], '/static/image/default-avatar.png')

            # Get school_year from assignment if available, otherwise determine it from ended_date
            school_year = assignment.get('school_year', '')
            if not school_year and ended_date_formatted:
                # Find the school year by comparing ended_date with quarters data (in memory, NOT in DB)
                for quarter in quarters_list:
                    try:
                        quarter_end = str(quarter.get('end_date', ''))[:10]
                        if ended_date_formatted <= quarter_end:
                            school_year = quarter.get('school_year', '')
                            break
                    except Exception:
                        pass

            assignment_list.append({
                'assignment_id': assignment['assignment_id'],
                'teacher_id': assignment['teacher_id'],
                'teacher_name': f"{teacher.get('first_name', '')} {teacher.get('last_name', '')}",
                'teacher_email': teacher.get('email', ''),
                'subject': assignment.get('subject', ''),
                'grade_level': assignment['grade_level'],
                'section': assignment['section'],
                'school_year': school_year,
                'status': assignment.get('status', 'unknown'),
                'assigned_at': assigned_date,
                'ended_at': ended_date_formatted,
                'profile_pic': profile_pic
            })

        return jsonify({
            'success': True,
            'assignments': assignment_list,
            'total_count': len(assignment_list)
        })
    except Exception as e:
        logger.error(f"Error in get_archived_classroom_assignments: {str(e)}")
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/api/classroom-assignments/<int:assignment_id>', methods=['PUT'])
def update_classroom_assignment(assignment_id):
    try:
        data = request.get_json()
        teacher_id = data.get('teacher_id')
        grade_level = data.get('grade_level')
        section = data.get('section')
        subject = data.get('subject')

        if not teacher_id or not grade_level or not section or not subject:
            return jsonify({'success': False, 'message': 'Missing required fields.'}), 400

        # --- ADD THIS CHECK: Prevent duplicate assignment for the same teacher, grade, section, and subject ---
        duplicate = supabase.table('teacher_class_assignments') \
            .select('assignment_id') \
            .eq('teacher_id', teacher_id) \
            .eq('grade_level', grade_level) \
            .eq('section', section) \
            .eq('subject', subject) \
            .neq('assignment_id', assignment_id) \
            .execute()
        if duplicate.data:
            return jsonify({'success': False, 'message': 'This Subject Teacher is already assigned to the selected classroom.'}), 400

        # Update the assignment
        result = supabase.table('teacher_class_assignments').update({
            'teacher_id': teacher_id,
            'grade_level': grade_level,
            'section': section,
            'subject': subject
        }).eq('assignment_id', assignment_id).execute()

        if hasattr(result, 'error') and result.error:
            raise Exception(result.error.message)

        return jsonify({'success': True, 'message': 'Assignment updated successfully.'})
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/api/classroom-assignments/<int:assignment_id>', methods=['DELETE'])
def delete_classroom_assignment(assignment_id):
    try:
        result = supabase.table('teacher_class_assignments').delete().eq('assignment_id', assignment_id).execute()
        if hasattr(result, 'error') and result.error:
            raise Exception(result.error.message)
        return jsonify({'success': True, 'message': 'Assignment deleted.'})
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500

############################# REWARDS.HTML#############################
@app.route('/api/rewards', methods=['GET', 'POST'])
def manage_rewards():
    try:
        if request.method == 'GET':
            teacher_id = session.get('user_id')
            if not teacher_id:
                return jsonify({'success': False, 'message': 'User not authenticated'}), 401

            # Get current school year
            current_school_year = get_current_school_year()
            if not current_school_year:
                return jsonify({'success': False, 'message': 'Unable to determine current school year'}), 500
            
            # Get teacher's class assignments for CURRENT school year
            class_assignments_result = (
                supabase.table('teacher_class_assignments')
                .select('grade_level, section, subject')
                .eq('teacher_id', teacher_id)
                .eq('school_year', current_school_year)
                .eq('status', 'active')
                .execute()
            )
            
            # Extract unique classes (grade_level + section combinations)
            teacher_classes = []
            assigned_classes = []  # New list for formatted class names
            
            if class_assignments_result.data:
                for assignment in class_assignments_result.data:
                    grade = assignment['grade_level']
                    section = assignment['section']
                    class_name = f"Grade {grade} - {section}"
                    
                    teacher_classes.append({
                        'grade_level': grade,
                        'section': section
                    })
                    
                    assigned_classes.append({
                        'class_name': class_name,
                        'grade_level': grade,
                        'section': section,
                        'subject': assignment.get('subject', '')
                    })

            page = int(request.args.get('page', 1))
            per_page_raw = request.args.get('per_page', '5')
            
            if per_page_raw == 'all':
                # Fetch all rewards for this teacher
                result = (
                    supabase.table('rewards')
                    .select('*')
                    .eq('created_by', teacher_id)
                    .order('created_at', desc=True)
                    .execute()
                )
                rewards_data = []
                if result.data:
                    for reward in result.data:
                        rewards_data.append({
                            'reward_id': reward['reward_id'],
                            'reward_name': reward['reward_name'],
                            'description': reward['description'],
                            'point_cost': reward['point_cost'],
                            'available_quantity': reward['available_quantity'],
                            'category': reward['category'],
                            'status': reward['status'],
                            'created_at': reward['created_at'],
                            'created_by': reward['created_by'],
                            'grade_level': reward.get('grade_level'),
                            'section': reward.get('section')
                        })
                total_count = len(rewards_data)
                return jsonify({
                    'success': True,
                    'rewards': rewards_data,
                    'total': total_count,
                    'page': page,
                    'per_page': per_page_raw,
                    'teacher_classes': teacher_classes,
                    'assigned_classes': assigned_classes
                })
            else:
                # Handle paginated results
                try:
                    per_page = int(per_page_raw)
                except ValueError:
                    per_page = 5
                offset = (page - 1) * per_page

                # Fetch rewards with pagination
                result = (
                    supabase.table('rewards')
                    .select('*')
                    .eq('created_by', teacher_id)
                    .order('created_at', desc=True)
                    .range(offset, offset + per_page - 1)
                    .execute()
                )

                # Get total count
                total_result = (
                    supabase.table('rewards')
                    .select('reward_id')
                    .eq('created_by', teacher_id)
                    .execute()
                )
                total_count = len(total_result.data) if total_result.data else 0

                rewards_data = []
                if result.data:
                    for reward in result.data:
                        rewards_data.append({
                            'reward_id': reward['reward_id'],
                            'reward_name': reward['reward_name'],
                            'description': reward['description'],
                            'point_cost': reward['point_cost'],
                            'available_quantity': reward['available_quantity'],
                            'category': reward['category'],
                            'status': reward['status'],
                            'created_at': reward['created_at'],
                            'created_by': reward['created_by'],
                            'grade_level': reward.get('grade_level'),
                            'section': reward.get('section')
                        })

                return jsonify({
                    'success': True,
                    'rewards': rewards_data,
                    'total': total_count,
                    'page': page,
                    'per_page': per_page,
                    'teacher_classes': teacher_classes,
                    'assigned_classes': assigned_classes
                })

        elif request.method == 'POST':
            if 'user_id' not in session or session.get('role') != 'Teacher':
                return jsonify({'success': False, 'message': 'Unauthorized'}), 403
            
            teacher_id = session['user_id']
            
            # Get form data
            reward_name = request.form.get('reward_name', '').strip()
            point_cost = request.form.get('point_cost')
            available_quantity = request.form.get('available_quantity')
            category = request.form.get('category', '').strip()
            description = request.form.get('description', '').strip()
            
            # Get grade_level and section as JSON arrays
            grade_level_json = request.form.get('grade_level', '[]')
            section_json = request.form.get('section', '[]')
            
            # Parse JSON arrays
            try:
                grade_levels = json.loads(grade_level_json) if isinstance(grade_level_json, str) else grade_level_json
                sections = json.loads(section_json) if isinstance(section_json, str) else section_json
            except (json.JSONDecodeError, TypeError):
                return jsonify({'success': False, 'message': 'Invalid grade_level or section format'}), 400
            
            # Validate inputs
            if not reward_name or not point_cost or not available_quantity or not category:
                return jsonify({'success': False, 'message': 'Missing required fields'}), 400
            
            if not grade_levels or not sections:
                return jsonify({'success': False, 'message': 'Please select at least one class'}), 400
            
            # Get current school year
            current_school_year = get_current_school_year()
            if not current_school_year:
                return jsonify({'success': False, 'message': 'Unable to determine current school year'}), 500
            
            # Verify teacher is assigned to ALL selected classes for CURRENT school year
            teacher_assignments = supabase.table('teacher_class_assignments') \
                .select('grade_level, section') \
                .eq('teacher_id', teacher_id) \
                .eq('school_year', current_school_year) \
                .eq('status', 'active') \
                .execute()
            
            if not teacher_assignments.data:
                return jsonify({'success': False, 'message': 'You are not assigned to any classes'}), 403
            
            # Create a set of assigned classes for quick lookup
            assigned_classes = {(a['grade_level'], a['section']) for a in teacher_assignments.data}
            
            # Check if all selected classes are in assigned classes
            selected_classes = set(zip(grade_levels, sections))
            unassigned = selected_classes - assigned_classes
            
            if unassigned:
                unassigned_str = ', '.join([f"Grade {g} {s}" for g, s in unassigned])
                return jsonify({
                    'success': False, 
                    'message': f'You are not assigned to teach: {unassigned_str}. Please select only your assigned classes.'
                }), 403
            
            # Create reward with array data
            reward_data = {
                'reward_name': reward_name,
                'point_cost': int(point_cost),
                'available_quantity': int(available_quantity),
                'category': category,
                'description': description,
                'created_by': teacher_id,
                'status': 'Available',
                'created_at': datetime.utcnow().isoformat(),
                'grade_level': grade_levels,  # Store as array
                'section': sections  # Store as array
            }
            
            # Insert into DB
            result = supabase.table('rewards').insert(reward_data).execute()
            
            if not result.data:
                return jsonify({'success': False, 'message': 'Failed to create reward'}), 500
            
            reward_id = result.data[0]['reward_id']
            
            # Activity log
            record_admin_activity_log(
                user_id=teacher_id,
                action='Create Reward',
                activity='Rewards Management',
                description=f"Created reward: {reward_name} (Cost: {point_cost} pts, Qty: {available_quantity}) for multiple classes",
                user_role='Teacher'
            )
            
            # Notify students in all selected classrooms
            for grade_level, section in selected_classes:
                # Get students in this specific class
                students_result = supabase.table('user_info') \
                    .select('id') \
                    .eq('role', 'Student') \
                    .eq('year_level', grade_level) \
                    .eq('section', section) \
                    .execute()
                
                students = students_result.data if students_result.data else []
                
                notif_title = "New Reward Available"
                notif_message = f"Your teacher has added a new reward: {reward_name}. Check the rewards page!"
                
                for student in students:
                    supabase.table('notifications').insert({
                        'user_id': student['id'],
                        'sender_id': teacher_id,
                        'title': notif_title,
                        'message': notif_message,
                        'notif_type': 'Reward',
                        'status': 'Unread',
                        'reward_id': reward_id
                    }).execute()

            return jsonify({
                'success': True,
                'message': 'Reward added successfully!',
                'data': result.data[0]
            })
                
    except Exception as e:
        app.logger.error(f'Error in manage_rewards: {str(e)}')
        return jsonify({'success': False, 'message': f'Server error: {str(e)}'}), 500
    
@app.route('/api/rewards/<int:reward_id>', methods=['GET'])
def get_reward(reward_id):
    try:
        # Make sure the user is logged in and is a teacher
        if 'user_id' not in session or session.get('role') != 'Teacher':
            return jsonify({'success': False, 'message': 'Unauthorized access'}), 401

        teacher_id = session['user_id']

        # Fetch the reward if it belongs to the logged-in teacher
        result = supabase.table('rewards').select('*').eq('reward_id', reward_id).eq('created_by', teacher_id).execute()
        
        if not result.data:
            return jsonify({'success': False, 'message': 'Reward not found or unauthorized'}), 404

        reward = result.data[0]
        # Transform the data to include both id and reward_id
        reward_info = {
            'reward_id': reward['reward_id'],  # Make sure reward_id is included
            'reward_name': reward['reward_name'],
            'description': reward['description'],
            'point_cost': reward['point_cost'],
            'available_quantity': reward['available_quantity'],
            'category': reward['category'],
            'status': reward['status'],
            'created_at': reward['created_at'],
            'created_by': reward['created_by']
        }

        return jsonify({
            'success': True,
            'reward': reward_info
        })

    except Exception as e:
        app.logger.error(f"Error fetching reward: {str(e)}")
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/api/rewards/<int:reward_id>', methods=['PUT'])
def update_reward(reward_id):
    try:
        if 'user_id' not in session or session.get('role') != 'Teacher':
            return jsonify({'success': False, 'message': 'Unauthorized access'}), 401

        teacher_id = session['user_id']

        # Check if the reward exists and belongs to the teacher
        check_result = supabase.table('rewards').select('*').eq('reward_id', reward_id).eq('created_by', teacher_id).execute()
        if not check_result.data:
            return jsonify({'success': False, 'message': 'Reward not found or unauthorized'}), 404

        reward = check_result.data[0]
        prev_quantity = int(reward.get('available_quantity', 0))
        grade_level = reward.get('grade_level')
        section = reward.get('section')

        # Get JSON data from request
        data = request.get_json()
        if not data:
            return jsonify({'success': False, 'message': 'No data provided'}), 400

        # --- DELIST/RELIST LOGIC ---
        # If only status is being updated (delist/relist)
        if 'status' in data and data['status'] in ['Available', 'Unavailable']:
            update_data = {'status': data['status']}
            result = supabase.table('rewards').update(update_data).eq('reward_id', reward_id).execute()

            # Activity log
            record_admin_activity_log(
                user_id=teacher_id,
                action='Delist/Relist Reward',
                activity='Rewards Management',
                description=f"Changed status of reward: {reward.get('reward_name', '')} to {data['status']}",
                user_role='Teacher'
            )

            # Prepare response
            if result.data:
                reward = result.data[0]
                reward_info = {
                    'id': reward['reward_id'],
                    'reward_id': reward['reward_id'],
                    'reward_name': reward['reward_name'],
                    'description': reward['description'],
                    'point_cost': reward['point_cost'],
                    'available_quantity': reward['available_quantity'],
                    'category': reward['category'],
                    'status': reward['status'],
                    'created_at': reward['created_at'],
                    'created_by': reward['created_by']
                }
            else:
                reward_info = None

            return jsonify({
                'success': True,
                'message': f"Reward status updated to {data['status']}!",
                'reward': reward_info
            }), 200

        # --- NORMAL EDIT LOGIC ---
        # Validate required fields
        required_fields = ['reward_name', 'point_cost', 'available_quantity', 'category']
        missing = [field for field in required_fields if field not in data]
        if missing:
            return jsonify({'success': False, 'message': f'Missing required fields: {", ".join(missing)}'}), 400

        new_quantity = int(data['available_quantity'])

        # Update reward in database with enhanced status logic
        # Always set status to Unavailable if quantity is 0, otherwise use Available as default
        update_data = {
            'reward_name': data['reward_name'],
            'description': data.get('description', ''),
            'point_cost': int(data['point_cost']),
            'available_quantity': new_quantity,
            'category': data['category'],
            'status': 'Unavailable' if new_quantity == 0 else 'Available'
        }

        result = supabase.table('rewards').update(update_data).eq('reward_id', reward_id).execute()

        # --- Notify teacher if quantity becomes 0 ---
        if new_quantity == 0:
            supabase.table('notifications').insert({
                'user_id': teacher_id,
                'sender_id': teacher_id,
                'title': "Reward Out of Stock",
                'message': f"Your reward '{data['reward_name']}' is now out of stock and has been automatically unlisted.",
                'notif_type': 'Reward',
                'status': 'Unread',
                'reward_id': reward_id
            }).execute()

        # --- Notify students if reward is back in stock ---
        if prev_quantity == 0 and new_quantity > 0:
            students_query = supabase.table('user_info').select('id').eq('role', 'Student')
            if grade_level:
                students_query = students_query.eq('year_level', grade_level)
            if section:
                students_query = students_query.eq('section', section)
            students_result = students_query.execute()
            students = students_result.data if students_result.data else []

            notif_title = "Reward Back in Stock"
            notif_message = f"The reward '{data['reward_name']}' is now available again! ({new_quantity} in stock, {data['point_cost']} pts)."
            for student in students:
                supabase.table('notifications').insert({
                    'user_id': student['id'],
                    'sender_id': teacher_id,
                    'title': notif_title,
                    'message': notif_message,
                    'notif_type': 'General',
                    'status': 'Unread'
                }).execute()

            # --- ACTIVITY LOG: Reward Back in Stock ---
            record_admin_activity_log(
                user_id=teacher_id,
                action='Reward Back in Stock',
                activity='Rewards Management',
                description=f"Reward '{data['reward_name']}' is back in stock (Qty: {new_quantity}).",
                user_role='Teacher'
            )

        # --- ACTIVITY LOG: Edit Reward ---
        record_admin_activity_log(
            user_id=teacher_id,
            action='Edit Reward',
            activity='Rewards Management',
            description=f"Edited reward: {data['reward_name']} (Cost: {data['point_cost']} pts, Qty: {new_quantity})",
            user_role='Teacher'
        )

        # Prepare response
        if result.data:
            reward = result.data[0]
            reward_info = {
                'id': reward['reward_id'],
                'reward_id': reward['reward_id'],
                'reward_name': reward['reward_name'],
                'description': reward['description'],
                'point_cost': reward['point_cost'],
                'available_quantity': reward['available_quantity'],
                'category': reward['category'],
                'status': reward['status'],
                'created_at': reward['created_at'],
                'created_by': reward['created_by']
            }
        else:
            reward_info = None

        return jsonify({
            'success': True,
            'message': 'Reward updated successfully!',
            'reward': reward_info
        }), 200

    except Exception as e:
        app.logger.error(f"Error updating reward: {str(e)}")
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/api/reward-redemptions', methods=['GET'])
def api_reward_redemptions():
    if 'user_id' not in session or session.get('role') != 'Teacher':
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401

    teacher_id = session['user_id']
    # --- Query params ---
    page = int(request.args.get('page', 1))
    per_page = request.args.get('per_page', '50')
    sort_by = request.args.get('sort_by', 'date_desc')
    search = request.args.get('search', '').strip()
    classroom = request.args.get('classroom', '')
    category = request.args.get('category', '')
    date_filter = request.args.get('date', '')
    points_filter = request.args.get('points', '')
    from_date = request.args.get('from', '')
    to_date = request.args.get('to', '')
    status = request.args.get('status', '')

    try:
        # --- Get reward_ids created by this teacher ---
        rewards_result = supabase.table('rewards').select('reward_id').eq('created_by', teacher_id).execute()
        my_reward_ids = [r['reward_id'] for r in rewards_result.data] if rewards_result.data else []

        if not my_reward_ids:
            return jsonify({'success': True, 'redemptions': [], 'total': 0, 'page': page, 'per_page': per_page})

        # --- Build base query for redemptions of your rewards only ---
        # UPDATED: Added 'remarks' to the SELECT query
        query = supabase.table('reward_redemptions').select(
            'redemption_id, student_id, reward_id, processed_at, points_deducted, grade_level, section, status, notes, used_at, remarks, rewards(reward_name, category), user_info(first_name, last_name)'
        ).in_('reward_id', my_reward_ids)

        # --- Filtering ---
        if classroom:
            grade, section = classroom.split('-')
            query = query.eq('grade_level', grade).eq('section', section)
        if category:
            query = query.eq('rewards.category', category)
        if points_filter:
            if points_filter == '1-10':
                query = query.gte('points_deducted', 1).lte('points_deducted', 10)
            elif points_filter == '11-25':
                query = query.gte('points_deducted', 11).lte('points_deducted', 25)
            elif points_filter == '26-50':
                query = query.gte('points_deducted', 26).lte('points_deducted', 50)
            elif points_filter == '51+':
                query = query.gte('points_deducted', 51)
        
        # Date filter
        from datetime import datetime, timedelta, timezone
        now = datetime.now(timezone.utc)
        if date_filter == 'today':
            today = now.date().isoformat()
            query = query.gte('processed_at', today)
        elif date_filter == 'week':
            week_start = (now - timedelta(days=now.weekday())).date().isoformat()
            query = query.gte('processed_at', week_start)
        elif date_filter == 'month':
            month_start = now.replace(day=1).date().isoformat()
            query = query.gte('processed_at', month_start)
        elif date_filter == 'quarter':
            month = now.month
            quarter_start_month = ((month - 1) // 3) * 3 + 1
            quarter_start = now.replace(month=quarter_start_month, day=1).date().isoformat()
            query = query.gte('processed_at', quarter_start)
        elif date_filter == 'custom' and from_date and to_date:
            query = query.gte('processed_at', from_date).lte('processed_at', to_date)
        
        # Status filter
        if status:
            query = query.eq('status', status)
        
        # Search filter (will be applied after fetch)
        if search:
            pass

        # --- Sorting ---
        if sort_by == 'date_desc':
            query = query.order('processed_at', desc=True)
        elif sort_by == 'date_asc':
            query = query.order('processed_at', desc=False)
        elif sort_by == 'points_desc':
            query = query.order('points_deducted', desc=True)
        elif sort_by == 'points_asc':
            query = query.order('points_deducted', desc=False)
        # For student name sort, sort after fetch

        # --- Pagination with error handling ---
        if per_page != 'all':
            try:
                per_page_int = int(per_page)
                offset = (page - 1) * per_page_int
                query = query.range(offset, offset + per_page_int - 1)
            except ValueError:
                # Fallback to default if per_page is not an integer
                per_page_int = 50
                offset = (page - 1) * per_page_int
                query = query.range(offset, offset + per_page_int - 1)

        # --- Execute query ---
        result = query.execute()
        redemptions = result.data if result.data else []

        # --- Search filter (client-side) ---
        if search:
            search_lower = search.lower()
            redemptions = [
                r for r in redemptions
                if search_lower in (r.get('user_info', {}).get('first_name', '') + ' ' + r.get('user_info', {}).get('last_name', '')).lower()
                or search_lower in (r.get('rewards', {}).get('reward_name', '')).lower()
            ]

        # --- Student name sort (client-side) ---
        if sort_by == 'student_asc':
            redemptions.sort(key=lambda r: (r.get('user_info', {}).get('first_name', '') + ' ' + r.get('user_info', {}).get('last_name', '')))
        elif sort_by == 'student_desc':
            redemptions.sort(key=lambda r: (r.get('user_info', {}).get('first_name', '') + ' ' + r.get('user_info', {}).get('last_name', '')), reverse=True)

        # --- Prepare response ---
        # Get all student_ids in redemptions
        student_ids = [r.get('student_id') for r in redemptions if r.get('student_id')]
        pics_map = {}
        if student_ids:
            pics_result = supabase.table('profile_pictures').select('user_id, file_path').in_('user_id', student_ids).execute()
            if pics_result.data:
                for p in pics_result.data:
                    pics_map[p['user_id']] = p['file_path']

        redemption_list = []
        for r in redemptions:
            student_name = ""
            student_id = r.get('student_id')
            if r.get('user_info'):
                student_name = f"{r['user_info']['first_name']} {r['user_info']['last_name']}"
            reward_name = r['rewards']['reward_name'] if r.get('rewards') else ''
            reward_category = r['rewards']['category'] if r.get('rewards') else ''
            pic_url = pics_map.get(student_id, '/static/image/default-avatar.png')
            
            # UPDATED: Added 'remarks' to the response object
            redemption_list.append({
                'redemption_id': r.get('redemption_id'),
                'student': student_name,
                'student_id': student_id,
                'reward': reward_name,
                'points': r.get('points_deducted', ''),
                'date': r.get('processed_at', ''),
                'category': reward_category,
                'grade_level': r.get('grade_level', ''),
                'section': r.get('section', ''),
                'status': r.get('status', ''),
                'pic': pic_url,
                'notes': r.get('notes'),
                'used_at': r.get('used_at'),
                'remarks': r.get('remarks', ''),  # <-- ADDED THIS LINE
            })

        # --- Total count for pagination ---
        total_count = len(redemption_list) if per_page == 'all' else (result.count if hasattr(result, 'count') else None)

        return jsonify({
            'success': True,
            'redemptions': redemption_list,
            'total': total_count or len(redemption_list),
            'page': page,
            'per_page': per_page
        })
        
    except Exception as e:
        app.logger.error(f"Error fetching reward redemptions: {str(e)}")
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/api/reward-redemptions/<int:redemption_id>/use', methods=['POST'])
def use_reward_redemption(redemption_id):
    if 'user_id' not in session or session.get('role') != 'Teacher':
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401
    
    teacher_id = session['user_id']
    data = request.get_json()
    notes = data.get('notes', '')
    used_at = datetime.utcnow().isoformat()
    
    try:
        # First, get the redemption record to extract student_id and reward details
        redemption_result = supabase.table('reward_redemptions').select(
            'student_id, reward_id, points_deducted, rewards(reward_name)'
        ).eq('redemption_id', redemption_id).execute()
        
        if not redemption_result.data or len(redemption_result.data) == 0:
            return jsonify({'success': False, 'message': 'Redemption not found'}), 404
        
        redemption = redemption_result.data[0]
        student_id = redemption.get('student_id')
        reward_name = redemption.get('rewards', {}).get('reward_name', 'a reward') if redemption.get('rewards') else 'a reward'
        points = redemption.get('points_deducted', 0)
        
        # Update the reward redemption record
        result = supabase.table('reward_redemptions').update({
            'status': 'Used',
            'notes': notes,
            'used_at': used_at
        }).eq('redemption_id', redemption_id).execute()
        
        if hasattr(result, 'error') and result.error:
            raise Exception(result.error.message)
        
        # Create a notification for the student
        if student_id:
            notification_message = f'Your reward "{reward_name}" has been marked as used by your teacher.'
            
            notification_result = supabase.table('notifications').insert({
                'user_id': student_id,
                'sender_id': teacher_id,
                'title': 'Reward Used',
                'message': notification_message,
                'notif_type': 'Reward',  # Changed from 'type' to 'notif_type' and value to 'Reward'
                'status': 'Unread',
                'created_at': used_at,
                'redemption_id': redemption_id  # Added redemption_id for reference
            }).execute()
            
            if hasattr(notification_result, 'error') and notification_result.error:
                # Don't fail the main operation if notification fails, just log it
                app.logger.error(f"Failed to create notification for student {student_id}: {notification_result.error}")
        
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500  # Fixed status code (was 5, should be 500)

@app.route('/api/recent-reward-redemptions', methods=['GET'])
def api_recent_reward_redemptions():
    if 'user_id' not in session or session.get('role') != 'Teacher':
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401

    teacher_id = session['user_id']
    try:
        # Get all reward_ids created by this teacher
        rewards_result = supabase.table('rewards').select('reward_id').eq('created_by', teacher_id).execute()
        my_reward_ids = [r['reward_id'] for r in rewards_result.data] if rewards_result.data else []

        if not my_reward_ids:
            return jsonify({'success': True, 'redemptions': []})

        # Get the 5 most recent redemptions for these rewards
        redemptions_result = (
            supabase.table('reward_redemptions')
            .select('redemption_id, student_id, reward_id, processed_at, points_deducted, status, rewards(reward_name), user_info(first_name, last_name)')
            .in_('reward_id', my_reward_ids)
            .order('processed_at', desc=True)
            .limit(7)
            .execute()
        )
        redemptions = redemptions_result.data if redemptions_result.data else []

        # Get student profile pictures
        student_ids = [r['student_id'] for r in redemptions if r.get('student_id')]
        pics_map = {}
        if student_ids:
            pics_result = supabase.table('profile_pictures').select('user_id, file_path').in_('user_id', student_ids).execute()
            if pics_result.data:
                for p in pics_result.data:
                    pics_map[p['user_id']] = p['file_path']

        redemption_list = []
        for r in redemptions:
            student_name = ""
            if r.get('user_info'):
                student_name = f"{r['user_info'].get('first_name', '')} {r['user_info'].get('last_name', '')}"
            reward_name = r['rewards']['reward_name'] if r.get('rewards') else ''
            pic_url = pics_map.get(r['student_id'], '/static/image/default-avatar.png')
            redemption_list.append({
                'student': student_name,
                'student_id': r.get('student_id'),
                'reward': reward_name,
                'points': r.get('points_deducted', ''),
                'date': r.get('processed_at', ''),
                'status': r.get('status', ''),
                'pic': pic_url,
            })

        return jsonify({'success': True, 'redemptions': redemption_list})
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500
    
#############################DASHBOARD.HTML#############################
@app.route('/api/dashboard-stats', methods=['GET'])
def api_dashboard_stats():
    if 'user_id' not in session or session.get('role') != 'Teacher':
        return jsonify({'success': False}), 403

    user_id = session['user_id']
    
    try:
        # Calculate this month's date range
        now = datetime.now()
        this_month_start = now.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
        
        # Points awarded this month by this teacher
        points_result = supabase.table('points') \
            .select('points') \
            .eq('teacher_id', user_id) \
            .eq('status', 'approved') \
            .gte('received_at', this_month_start.isoformat()) \
            .execute()
        
        points_total = sum(int(p['points']) for p in points_result.data) if points_result.data else 0
        
        # Rewards redeemed this month (by students from this teacher's rewards)
        # First, get all reward IDs created by this teacher
        rewards_result = supabase.table('rewards') \
            .select('reward_id') \
            .eq('created_by', user_id) \
            .execute()
        
        teacher_reward_ids = [r['reward_id'] for r in rewards_result.data] if rewards_result.data else []
        
        # Count redemptions of these rewards this month
        rewards_count = 0
        if teacher_reward_ids:
            redemptions_result = supabase.table('reward_redemptions') \
                .select('redemption_id') \
                .in_('reward_id', teacher_reward_ids) \
                .gte('processed_at', this_month_start.isoformat()) \
                .execute()
            
            rewards_count = len(redemptions_result.data) if redemptions_result.data else 0
        
        return jsonify({
            'success': True,
            'points_awarded_this_month': points_total,
            'rewards_redeemed_this_month': rewards_count
        })
    except Exception as e:
        app.logger.error(f"Error fetching dashboard stats: {str(e)}")
        return jsonify({'success': False, 'message': str(e)}), 500

#############################STUDENTS.HTML#############################
# Diagnostic endpoint - Shows what's in the database
@app.route('/api/diagnostic/teacher-assignments', methods=['GET'])
def api_diagnostic_assignments():
    """Shows all teachers and their assignments for debugging"""
    if 'user_id' not in session or session.get('role') != 'Admin':
        return jsonify({'success': False, 'message': 'Admin only'}), 403

    try:
        current_school_year = get_current_school_year()
        
        # Get all assignments for current school year
        all_assignments = supabase.table('teacher_class_assignments') \
            .select('*') \
            .eq('school_year', current_school_year) \
            .eq('status', 'active') \
            .execute()
        
        assignments = all_assignments.data if all_assignments.data else []
        
        # Group by teacher
        by_teacher = {}
        for a in assignments:
            tid = a['teacher_id']
            if tid not in by_teacher:
                by_teacher[tid] = []
            by_teacher[tid].append(f"{a['grade_level']}-{a['section']} ({a.get('subject', 'N/A')})")
        
        # Get teacher names
        teacher_ids = list(by_teacher.keys())
        teachers_map = {}
        if teacher_ids:
            teachers_result = supabase.table('user_info') \
                .select('id, first_name, last_name') \
                .in_('id', teacher_ids) \
                .execute()
            if teachers_result.data:
                for t in teachers_result.data:
                    teachers_map[t['id']] = f"{t['first_name']} {t['last_name']}"
        
        # Format output
        diagnostic_data = {
            'current_school_year': current_school_year,
            'total_assignments': len(assignments),
            'teachers_with_assignments': len(by_teacher),
            'assignments_by_teacher': {}
        }
        
        for tid, classrooms in by_teacher.items():
            teacher_name = teachers_map.get(tid, f'ID: {tid}')
            diagnostic_data['assignments_by_teacher'][teacher_name] = {
                'teacher_id': tid,
                'count': len(classrooms),
                'classrooms': classrooms
            }
        
        logger.info(f"🔍 DIAGNOSTIC: {diagnostic_data}")
        
        return jsonify({'success': True, 'data': diagnostic_data})
    except Exception as e:
        logger.error(f"Diagnostic error: {str(e)}")
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/api/my-students', methods=['GET'])
def api_my_students():
    if 'user_id' not in session or session.get('role') != 'Teacher':
        return jsonify({'success': False, 'message': 'Unauthorized'}), 403

    teacher_id = session['user_id']
    logger.info(f"🔍 DEBUG /api/my-students - Teacher ID: {teacher_id}")

    try:
        # Get the current school year
        current_school_year = get_current_school_year()
        logger.info(f"🔍 DEBUG /api/my-students - Current School Year: {current_school_year}")
        
        if not current_school_year:
            logger.warning(f"🔍 DEBUG /api/my-students - No school year found")
            return jsonify({'success': False, 'message': 'Unable to determine current school year'}), 500

        # Get classrooms assigned to this teacher for the CURRENT school year
        logger.info(f"🔍 DEBUG /api/my-students - Querying assignments: teacher_id={teacher_id}, school_year={current_school_year}, status=active")
        
        assignments_result = supabase.table('teacher_class_assignments') \
            .select('grade_level, section, assignment_id, subject') \
            .eq('teacher_id', teacher_id) \
            .eq('school_year', current_school_year) \
            .eq('status', 'active') \
            .execute()
        assignments = assignments_result.data if assignments_result.data else []
        
        logger.info(f"🔍 DEBUG /api/my-students - Found {len(assignments)} assignments for this teacher")
        if assignments:
            logger.info(f"🔍 DEBUG /api/my-students - Assignment details: {assignments}")

        students = []
        classrooms = []
        for assignment in assignments:
            grade_level = assignment['grade_level']
            section = assignment['section']
            logger.info(f"🔍 DEBUG /api/my-students - Processing assignment: {grade_level}-{section}")
            
            students_result = supabase.table('user_info') \
                .select('*') \
                .eq('role', 'Student') \
                .eq('year_level', grade_level) \
                .eq('section', section) \
                .eq('status', 'Active') \
                .execute()
            
            found_count = len(students_result.data) if students_result.data else 0
            logger.info(f"🔍 DEBUG /api/my-students - Found {found_count} students in {grade_level}-{section}")
            
            if students_result.data:
                students.extend(students_result.data)
            classrooms.append({
                'grade': grade_level,
                'section': section,
                'label': f"Grade {grade_level} - {section}"
            })

        # Remove duplicates
        unique_students = {student['id']: student for student in students}.values()
        student_ids = [student['id'] for student in unique_students]
        
        logger.info(f"🔍 DEBUG /api/my-students - Total unique students: {len(student_ids)}, Total classrooms: {len(classrooms)}")

        # --- Fetch all reward_ids created by this teacher ---
        rewards_result = supabase.table('rewards').select('reward_id, reward_name').eq('created_by', teacher_id).execute()
        teacher_rewards = {r['reward_id']: r['reward_name'] for r in rewards_result.data} if rewards_result.data else {}

        # --- For each student, get their latest redemption of a reward created by this teacher ---
        latest_redeemed_map = {}
        # --- Count total redemptions per student ---
        total_redemptions_map = {}
        
        if student_ids and teacher_rewards:
            redemptions_result = supabase.table('reward_redemptions') \
                .select('student_id, reward_id, processed_at') \
                .in_('student_id', student_ids) \
                .in_('reward_id', list(teacher_rewards.keys())) \
                .order('processed_at', desc=True) \
                .execute()
            if redemptions_result.data:
                for row in redemptions_result.data:
                    sid = row['student_id']
                    rid = row['reward_id']
                    
                    # Track latest redemption
                    if sid not in latest_redeemed_map:
                        latest_redeemed_map[sid] = teacher_rewards.get(rid, '')
                    
                    # Count total redemptions
                    if sid not in total_redemptions_map:
                        total_redemptions_map[sid] = 0
                    total_redemptions_map[sid] += 1

        # --- Participation Calculation (per teacher, per student) ---
        from datetime import datetime, timedelta, timezone
        now = datetime.now(timezone.utc)
        week_ago = now - timedelta(days=7)
        two_weeks_ago = now - timedelta(days=14)

        # --- NLP Notification Check (this week) ---
        start_of_week = now - timedelta(days=now.weekday())
        start_of_week = start_of_week.replace(hour=0, minute=0, second=0, microsecond=0)
        
        notes_this_week = {}
        if student_ids:
            notif_result = supabase.table('nlp_notifications') \
                .select('student_id, created_at') \
                .in_('student_id', student_ids) \
                .gte('created_at', start_of_week.isoformat()) \
                .execute()
            if notif_result.data:
                for row in notif_result.data:
                    notes_this_week[row['student_id']] = True

        # Map: student_id -> {'this_week': X, 'last_week': Y}
        participation_map = {sid: {'this_week': 0, 'last_week': 0} for sid in student_ids}

        if student_ids:
            # Only count points awarded by this teacher
            points_result = supabase.table('points') \
                .select('student_id, received_at') \
                .eq('teacher_id', teacher_id) \
                .in_('student_id', student_ids) \
                .gte('received_at', two_weeks_ago.isoformat()) \
                .execute()
            if points_result.data:
                for row in points_result.data:
                    sid = row['student_id']
                    received_at = row['received_at']
                    try:
                        dt = datetime.fromisoformat(str(received_at).replace('Z', '+00:00'))
                    except Exception:
                        continue
                    if dt >= week_ago:
                        participation_map[sid]['this_week'] += 1
                    elif dt >= two_weeks_ago:
                        participation_map[sid]['last_week'] += 1

        # Fetch profile pictures
        pics_map = {}
        if student_ids:
            pics_result = supabase.table('profile_pictures').select('user_id, file_path').in_('user_id', student_ids).execute()
            if pics_result.data:
                for p in pics_result.data:
                    pics_map[p['user_id']] = p['file_path']

        # Fetch last activity for each student (only tasks assigned by this teacher)
        last_activity_map = {}
        if student_ids:
            tasks_result = supabase.table('task_assignments') \
                .select('student_id, task, status, completed_at, submitted_at, due_date, assigned_at') \
                .eq('teacher_id', teacher_id) \
                .in_('student_id', student_ids) \
                .order('completed_at', desc=True) \
                .order('submitted_at', desc=True) \
                .order('assigned_at', desc=True) \
                .execute()
            if tasks_result.data:
                for row in tasks_result.data:
                    sid = row['student_id']
                    dt = None
                    if row.get('completed_at'):
                        dt = row['completed_at']
                        activity_type = 'Completed'
                    elif row.get('submitted_at'):
                        dt = row['submitted_at']
                        activity_type = 'Submitted'
                    elif row.get('assigned_at'):
                        dt = row['assigned_at']
                        activity_type = 'Assigned'
                    else:
                        dt = row.get('due_date')
                        activity_type = 'Due'
                    if sid not in last_activity_map or (dt and dt > last_activity_map[sid]['date']):
                        last_activity_map[sid] = {
                            'desc': f"{row['task']} ({activity_type})",
                            'date': dt
                        }

        # Build student list for frontend
        student_list = []
        for student in unique_students:
            sid = student['id']
            part = participation_map.get(sid, {'this_week': 0, 'last_week': 0})
            last = part['last_week']
            curr = part['this_week']
            
            # Calculate percent change and cap at 100%
            if last > 0:
                percent = min(round((curr - last) / last * 100, 1), 100)
                friendly = f"{percent:+.0f}% vs last week"
            elif curr > 0:
                percent = 100.0
                friendly = "New this week"
            else:
                percent = 0.0
                friendly = "No activity"
            
            last_activity = last_activity_map.get(sid, {})
            
            # Get total redemption count for this student
            total_redeemed = total_redemptions_map.get(sid, 0)
            latest_reward = latest_redeemed_map.get(sid, '')
            
            student_list.append({
                'id': sid,
                'name': f"{student.get('first_name', '')} {student.get('last_name', '')}",
                'grade': student.get('year_level', ''),
                'section': student.get('section', ''),
                'points': student.get('total_points', 0),
                'behavior': student.get('behavior', 'Good'),
                'rewards_redeemed': latest_reward,
                'total_rewards_redeemed': total_redeemed,  # Total count
                'last_activity': last_activity.get('desc', ''),
                'pic': pics_map.get(sid, '/static/image/default-avatar.png'),
                'participation_percent': percent,
                'participation_friendly': friendly,
                'participation_this_week': curr,
                'participation_last_week': last,
                'streak': student.get('streak', 0),
                'note_given_this_week': notes_this_week.get(sid, False)
            })

        return jsonify({'success': True, 'students': student_list, 'classrooms': classrooms})

    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/api/class-summary', methods=['GET'])
def api_class_summary():
    if 'user_id' not in session or session.get('role') != 'Teacher':
        return jsonify({'success': False, 'message': 'Unauthorized'})

    teacher_id = session['user_id']
    grade = request.args.get('grade')
    section = request.args.get('section')

    try:
        # Get current school year
        current_school_year = get_current_school_year()
        if not current_school_year:
            return jsonify({'success': False, 'message': 'Unable to determine current school year'}), 500
        
        # Get all classrooms assigned to this teacher for CURRENT school year
        assignments_result = supabase.table('teacher_class_assignments') \
            .select('grade_level, section') \
            .eq('teacher_id', teacher_id) \
            .eq('school_year', current_school_year) \
            .eq('status', 'active') \
            .execute()
        assignments = assignments_result.data if assignments_result.data else []

        students = []
        for assignment in assignments:
            # Fetch students for each classroom
            student_result = supabase.table('user_info') \
                .select('id, first_name, last_name, year_level, section') \
                .eq('role', 'Student') \
                .eq('year_level', assignment['grade_level']) \
                .eq('section', assignment['section']) \
                .execute()
            if student_result.data:
                students.extend(student_result.data)

        # Remove duplicates
        unique_students = {student['id']: student for student in students}.values()
        student_ids = [student['id'] for student in unique_students]

        # Get profile pictures
        pics_map = {}
        if student_ids:
            pics_result = supabase.table('profile_pictures').select('user_id', 'file_path').in_('user_id', student_ids).execute()
            if pics_result.data:
                for pic in pics_result.data:
                    pics_map[pic['user_id']] = pic['file_path']

        # Get total points for each student - FILTERED BY THIS TEACHER
        points_map = {}
        if student_ids:
            points_result = supabase.table('points') \
                .select('student_id', 'points') \
                .eq('teacher_id', teacher_id) \
                .in_('student_id', student_ids) \
                .execute()
            if points_result.data:
                for row in points_result.data:
                    sid = row['student_id']
                    points_map[sid] = points_map.get(sid, 0) + row['points']

        # Prepare summary data
        student_list = []
        total_points = 0
        for student in unique_students:
            pic = pics_map.get(student['id'], '')
            points = points_map.get(student['id'], 0)
            total_points += points
            student_list.append({
                'id': student['id'],
                'name': f"{student['first_name']} {student['last_name']}",
                'grade': student['year_level'],
                'section': student['section'],
                'points': points,
                'pic': pic,
            })

        # Sort by points descending
        student_list.sort(key=lambda x: x['points'], reverse=True)
        top_students = student_list[:5]
        avg_points = total_points / len(student_list) if student_list else 0

        # Participation calculation - USE SAME LOGIC AS MY_STUDENTS
        from datetime import datetime, timedelta, timezone
        now = datetime.now(timezone.utc)
        week_ago = now - timedelta(days=7)
        two_weeks_ago = now - timedelta(days=14)

        participation = {
            'this_week': 0,
            'last_week': 0,
            'this_month': 0,
            'last_month': 0
        }

        # Count points using ROLLING PERIODS like my_students
        if student_ids:
            # Get all points from last 2 weeks
            points_result = supabase.table('points') \
                .select('student_id, received_at') \
                .eq('teacher_id', teacher_id) \
                .in_('student_id', student_ids) \
                .gte('received_at', two_weeks_ago.isoformat()) \
                .execute()
            
            if points_result.data:
                for row in points_result.data:
                    received_at = row['received_at']
                    try:
                        dt = datetime.fromisoformat(str(received_at).replace('Z', '+00:00'))
                    except Exception:
                        continue
                    if dt >= week_ago:
                        participation['this_week'] += 1
                    elif dt >= two_weeks_ago:
                        participation['last_week'] += 1

        # Month calculations (keeping calendar months for now)
        start_of_month = now.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
        last_month = (start_of_month - timedelta(days=1)).replace(day=1, hour=0, minute=0, second=0, microsecond=0)
        
        if student_ids:
            points_month = supabase.table('points') \
                .select('student_id, received_at') \
                .eq('teacher_id', teacher_id) \
                .in_('student_id', student_ids) \
                .gte('received_at', start_of_month.isoformat()) \
                .execute()
            participation['this_month'] = len(points_month.data) if points_month.data else 0
            
            points_last_month = supabase.table('points') \
                .select('student_id, received_at') \
                .eq('teacher_id', teacher_id) \
                .in_('student_id', student_ids) \
                .gte('received_at', last_month.isoformat()) \
                .lt('received_at', start_of_month.isoformat()) \
                .execute()
            participation['last_month'] = len(points_last_month.data) if points_last_month.data else 0

        # Calculate percentages with same logic as my_students
        week_increase = 0
        if participation['last_week'] > 0:
            week_increase = min(round((participation['this_week'] - participation['last_week']) / participation['last_week'] * 100, 1), 100)
        elif participation['this_week'] > 0:
            week_increase = 100.0
        
        month_increase = 0
        if participation['last_month'] > 0:
            month_increase = min(round((participation['this_month'] - participation['last_month']) / participation['last_month'] * 100, 1), 100)
        elif participation['this_month'] > 0:
            month_increase = 100.0

        # Add week_increase to each top student
        for stu in top_students:
            stu['week_increase'] = round(week_increase, 1)

        return jsonify({
            'success': True,
            'summary': {
                'total_students': len(student_list),
                'average_points': round(avg_points, 2),
                'top_students': top_students,
                'participation': {
                    'this_week': participation['this_week'],
                    'last_week': participation['last_week'],
                    'week_increase': round(week_increase, 1),
                    'this_month': participation['this_month'],
                    'last_month': participation['last_month'],
                    'month_increase': round(month_increase, 1)
                }
            }
        })
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)})

@app.route('/api/send-dismissal', methods=['POST'])
def api_send_dismissal():
    if 'user_id' not in session or session.get('role') != 'Teacher':
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401

    data = request.get_json()
    dismissal_time = data.get('dismissal_time')
    message = data.get('message', '')
    notify_email = data.get('notify_email')
    notify_sms = data.get('notify_sms')  # New SMS notification flag
    grade = data.get('grade')
    section = data.get('section')
    student_id = data.get('student_id')

    # Format time to AM/PM if possible
    from datetime import datetime
    dismissal_time_formatted = dismissal_time
    try:
        dt = None
        try:
            dt = datetime.strptime(dismissal_time, "%H:%M")
        except Exception:
            try:
                dt = datetime.strptime(dismissal_time, "%I:%M %p")
            except Exception:
                dt = None
        if dt:
            dismissal_time_formatted = dt.strftime("%I:%M %p").lstrip("0")
    except Exception:
        dismissal_time_formatted = dismissal_time

    try:
        # Get students in the selected classroom or single student
        students_query = supabase.table('user_info').select('id, first_name, last_name, year_level, section').eq('role', 'Student')
        if student_id:
            students_query = students_query.eq('id', student_id)
        else:
            if grade:
                students_query = students_query.eq('year_level', grade)
            if section:
                students_query = students_query.eq('section', section)
        students_result = students_query.execute()
        students = students_result.data if students_result.data else []

        if not students:
            return jsonify({'success': False, 'message': 'No students found.'}), 404

        # Fetch teacher name and gender - FIXED: teacher_id is defined here
        teacher_id = session.get('user_id')  # This defines teacher_id
        teacher_info = supabase.table('user_info').select('first_name, last_name, gender').eq('id', teacher_id).execute()
        teacher_name = ''
        if teacher_info.data and len(teacher_info.data) > 0:
            teacher_data = teacher_info.data[0]
            # Determine prefix based on gender
            gender = teacher_data.get('gender', '').lower()
            prefix = 'Mr.'
            if gender == 'female':
                prefix = 'Mrs.'
            elif gender == 'other':
                prefix = 'Mx.'
            
            teacher_name = f"{prefix} {teacher_data.get('first_name', '')} {teacher_data.get('last_name', '')}"

        # Try both string and integer IDs to handle data type mismatches
        student_ids_str = [str(s['id']).strip() for s in students]
        student_ids_int = []
        for sid in student_ids_str:
            try:
                student_ids_int.append(int(sid))
            except ValueError:
                pass

        # Get parents information
        parents = []
        try:
            parents_result = supabase.table('parents').select('student_id, email, first_name, last_name, mobile_no, relationship').in_('student_id', student_ids_str).execute()
            parents = parents_result.data if parents_result.data else []
        except Exception as e:
            print(f"DEBUG: String query failed: {e}")

        if not parents and student_ids_int:
            try:
                parents_result = supabase.table('parents').select('student_id, email, first_name, last_name, mobile_no, relationship').in_('student_id', student_ids_int).execute()
                parents = parents_result.data if parents_result.data else []
            except Exception as e:
                print(f"DEBUG: Integer query failed: {e}")

        if not parents:
            for student_id in student_ids_str[:2]:
                try:
                    individual_result = supabase.table('parents').select('student_id, email, first_name, last_name, mobile_no, relationship').eq('student_id', student_id).execute()
                    individual_parents = individual_result.data if individual_result.data else []
                    if individual_parents:
                        parents.extend(individual_parents)
                except Exception as e:
                    print(f"DEBUG: Individual query for {student_id} failed: {e}")

        # Create mappings for parents
        parent_email_map = {}
        parent_phone_map = {}
        for p in parents:
            sid = str(p['student_id']).strip()
            if p.get('email') and p['email'].strip():
                if sid not in parent_email_map:
                    parent_email_map[sid] = []
                parent_email_map[sid].append({
                    'email': p['email'].strip(),
                    'name': f"{p['first_name']} {p['last_name']}",
                    'relationship': p.get('relationship', 'Parent')
                })
            
            if p.get('mobile_no') and p['mobile_no'].strip():
                if sid not in parent_phone_map:
                    parent_phone_map[sid] = []
                parent_phone_map[sid].append({
                    'phone': p['mobile_no'].strip(),
                    'name': f"{p['first_name']} {p['last_name']}",
                    'relationship': p.get('relationship', 'Parent')
                })

        email_sent_count = 0
        sms_sent_count = 0
        students_without_parents = []
        email_failures = []
        sms_failures = []

        # Build the base message in the requested format
        base_message = """Dear Parent,

This is to inform you that your child, {student_name}, has been dismissed at {dismissal_time}.

Masico National High School
Grade: {grade} - Section: {section}"""

        # Add note if provided
        if message:
            base_message += f"\n\nNote: {message}"
        
        # Add closing
        base_message += f"\n\nThank you,\nLearn2Earn\nTeacher: {teacher_name}"

        # Send email notifications if enabled
        if notify_email:
            for student in students:
                sid = str(student['id']).strip()
                parent_emails = parent_email_map.get(sid, [])
                
                if parent_emails:
                    student_name = f"{student['first_name']} {student['last_name']}"
                    grade_level = student.get('year_level', 'Unknown')
                    section_name = student.get('section', 'Unknown')
                    
                    email_body = base_message.format(
                        student_name=student_name,
                        dismissal_time=dismissal_time_formatted,
                        grade=grade_level,
                        section=section_name
                    )
                    
                    for parent in parent_emails:
                        subject = "Dismissal Notification"
                        msg = Message(subject, sender=app.config['MAIL_USERNAME'], recipients=[parent['email']])
                        msg.body = email_body

                        try:
                            mail.send(msg)
                            email_sent_count += 1
                        except Exception as e:
                            app.logger.error(f"Failed to send dismissal email to {parent['email']}: {str(e)}")
                            email_failures.append(f"{student_name} - {parent['name']} ({parent['email']})")
                else:
                    students_without_parents.append(f"{student['first_name']} {student['last_name']}")

        # Send SMS notifications if enabled
        if notify_sms:
            for student in students:
                sid = str(student['id']).strip()
                parent_phones = parent_phone_map.get(sid, [])
                
                if parent_phones:
                    student_name = f"{student['first_name']} {student['last_name']}"
                    grade_level = student.get('year_level', 'Unknown')
                    section_name = student.get('section', 'Unknown')
                    
                    sms_message = base_message.format(
                        student_name=student_name,
                        dismissal_time=dismissal_time_formatted,
                        grade=grade_level,
                        section=section_name
                    )
                    
                    for parent in parent_phones:
                        parent_number = parent['phone']
                        
                        # Convert to +63 format if needed
                        if parent_number and parent_number.startswith('0'):
                            parent_number = '+63' + parent_number[1:]
                        elif parent_number and not parent_number.startswith('+'):
                            parent_number = '+63' + parent_number
                        
                        # Only send if we have a valid phone number
                        if parent_number and len(parent_number) >= 12:
                            # Send SMS using the SMS function
                            sms_result = send_sms_via_api(parent_number, sms_message, student['id'])
                            
                            if sms_result['success']:
                                sms_sent_count += 1
                            else:
                                sms_failures.append(f"{student_name} - {parent['name']} ({parent_number}): {sms_result['error']}")
                        else:
                            sms_failures.append(f"{student_name} - {parent['name']}: Invalid phone number format")
                else:
                    # Only add to students_without_parents if we haven't already for email
                    student_name = f"{student['first_name']} {student['last_name']}"
                    if student_name not in students_without_parents:
                        students_without_parents.append(student_name)

        # Build response message
        response_parts = []
        if notify_email:
            response_parts.append(f'Emails sent: {email_sent_count}')
        if notify_sms:
            response_parts.append(f'SMS sent: {sms_sent_count}')
        
        response_message = 'Dismissal notification sent! ' + ' | '.join(response_parts)
        
        if students_without_parents:
            response_message += f' (Note: {len(students_without_parents)} students have no parent contact on file)'

        # Include failures in response for debugging
        result_data = {
            'success': True,
            'email_sent_count': email_sent_count,
            'sms_sent_count': sms_sent_count,
            'message': response_message,
            'students_without_contacts': students_without_parents
        }
        
        if email_failures:
            result_data['email_failures'] = email_failures
        if sms_failures:
            result_data['sms_failures'] = sms_failures

        return jsonify(result_data)

    except Exception as e:
        app.logger.error(f"Error sending dismissal: {str(e)}")
        return jsonify({'success': False, 'message': str(e)}), 500

def send_notifications_background(emails_to_send, sms_to_send):
    """
    Background task to send emails and SMS asynchronously.
    This prevents the endpoint from blocking while waiting for external API calls.
    """
    # Send emails
    for email_data in emails_to_send:
        try:
            subject = email_data['subject']
            body = email_data['body']
            recipient = email_data['recipient']
            
            msg = Message(subject, sender=app.config['MAIL_USERNAME'], recipients=[recipient])
            msg.body = body
            mail.send(msg)
        except Exception as e:
            app.logger.error(f"Background: Failed to send email to {recipient}: {str(e)}")
    
    # Send SMS messages
    for sms_data in sms_to_send:
        try:
            phone_number = sms_data['phone']
            message = sms_data['message']
            student_id = sms_data['student_id']
            send_sms_via_api(phone_number, message, student_id)
        except Exception as e:
            app.logger.error(f"Background: Failed to send SMS to {phone_number}: {str(e)}")

def send_sms_via_api(phone_number, message, student_id):
    """Send SMS via HTTPSMS API"""
    # Your HTTPSMS API configuration
    
    headers = {
        "Content-Type": "application/json",
        "x-api-key": HTTPSMS_API_KEY,
        "Accept": "application/json"
    }

    payload = {
        "content": message,
        "from": "+639761271972",  # Your registered number
        "to": phone_number
    }

    try:
        response = requests.post(
            "https://api.httpsms.com/v1/messages/send",
            headers=headers,
            json=payload,
            timeout=30
        )

        # Save to Supabase regardless of success
        supabase.table('sms_messages').insert({
            "phone_number": phone_number,
            "message": message,
            "status": "sent" if response.status_code == 200 else "failed",
            "sent_at": datetime.now().isoformat(),
            "student_id": student_id,
            "message_type": "dismissal_notification",
            "response_data": json.dumps({
                "status_code": response.status_code,
                "response": response.text
            })
        }).execute()

        if response.status_code == 200:
            return {"success": True, "message": "SMS sent successfully"}
        else:
            return {"success": False, "error": f"HTTP {response.status_code}: {response.text}"}
            
    except Exception as e:
        # Save error to Supabase
        supabase.table('sms_messages').insert({
            "phone_number": phone_number,
            "message": message,
            "status": "failed",
            "sent_at": datetime.now().isoformat(),
            "student_id": student_id,
            "message_type": "dismissal_notification",
            "response_data": json.dumps({"error": str(e)})
        }).execute()
        
        return {"success": False, "error": str(e)}
    
@app.route('/api/teacher-classrooms')
def api_teacher_classrooms():
    if 'user_id' not in session or session.get('role') != 'Teacher':
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401
    
    teacher_id = session.get('user_id')
    
    try:
        # Get current school year
        current_school_year = get_current_school_year()
        if not current_school_year:
            return jsonify({'success': False, 'message': 'Unable to determine current school year'}), 500
        
        # Query teacher_class_assignments for this teacher for CURRENT school year
        assignments_result = supabase.table('teacher_class_assignments') \
            .select('grade_level, section') \
            .eq('teacher_id', teacher_id) \
            .eq('school_year', current_school_year) \
            .eq('status', 'active') \
            .execute()
        
        classrooms = []
        if assignments_result.data:
            # Get unique classrooms
            seen_classrooms = set()
            for assignment in assignments_result.data:
                classroom_key = f"{assignment['grade_level']}|{assignment['section']}"
                if classroom_key not in seen_classrooms:
                    seen_classrooms.add(classroom_key)
                    classrooms.append({
                        'grade_level': assignment['grade_level'],
                        'section': assignment['section']
                    })
        
        return jsonify({
            'success': True,
            'classrooms': classrooms
        })
        
    except Exception as e:
        app.logger.error(f"Error fetching teacher classrooms: {str(e)}")
        return jsonify({'success': False, 'message': str(e)}), 500   
#############################AWARD_POINTS.HTML#############################
from datetime import datetime
import pytz

@app.route('/api/award-points', methods=['POST'])
def api_award_points():
    if 'user_id' not in session or session.get('role') != 'Teacher':
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401

    data = request.get_json()
    teacher_id = session['user_id']
    student_id = data.get('student_id')
    points = data.get('points')
    point_category = data.get('category')
    note = data.get('note', '')

    # Validation
    if not student_id or not points or not point_category:
        return jsonify({'success': False, 'message': 'Missing required fields'}), 400

    try:
        # Get current time in Philippines timezone
        philippines_tz = pytz.timezone('Asia/Manila')
        now_ph = datetime.now(philippines_tz).isoformat()
        
        # Get student name for activity log
        student_info = supabase.table('user_info').select('first_name', 'last_name').eq('id', int(student_id)).execute()
        student_name = f"Student ID {student_id}"  # Default fallback
        if student_info.data and len(student_info.data) > 0:
            first_name = student_info.data[0].get('first_name', '')
            last_name = student_info.data[0].get('last_name', '')
            if first_name or last_name:
                student_name = f"{first_name} {last_name}".strip()
        
        # 1. Insert the points record
        result = supabase.table('points').insert({
            'teacher_id': teacher_id,
            'student_id': int(student_id),
            'points': int(points),
            'point_category': point_category,
            'note': note,
            'status': 'approved',
            'received_at': now_ph  # Save in PH time
        }).execute()

        if hasattr(result, 'error') and result.error:
            raise Exception(result.error.message)

        # --- Get the point_id from the insert result ---
        point_id = None
        if result.data and 'point_id' in result.data[0]:
            point_id = result.data[0]['point_id']

        # 2. Increment total_points in user_info for the student
        user_result = supabase.table('user_info').select('total_points').eq('id', int(student_id)).execute()
        current_total = 0
        if user_result.data and 'total_points' in user_result.data[0] and user_result.data[0]['total_points'] is not None:
            current_total = int(user_result.data[0]['total_points'])
        new_total = current_total + int(points)
        supabase.table('user_info').update({'total_points': new_total}).eq('id', int(student_id)).execute()

        # Update session points if teacher is awarding points to themselves
        if int(student_id) == session.get('user_id'):
            session['points'] = new_total

        # --- ACTIVITY LOG: Award Points (with student name) ---
        record_admin_activity_log(
            user_id=teacher_id,
            action='Award Points',
            activity='Points Management',
            description=f"Awarded {points} points to {student_name} for '{point_category}'.",
            user_role='Teacher'
        )

        # --- SEND NOTIFICATION TO STUDENT ---
        teacher_info = supabase.table('user_info').select('last_name', 'gender').eq('id', teacher_id).execute()
        teacher_last_name = ''
        teacher_gender = ''
        if teacher_info.data and len(teacher_info.data) > 0:
            teacher_last_name = teacher_info.data[0].get('last_name', '')
            teacher_gender = (teacher_info.data[0].get('gender', '') or '').lower()
        prefix = 'Mr.'
        if teacher_gender == 'female':
            prefix = 'Mrs.'
        elif teacher_gender == 'other':
            prefix = 'Mx.'
        teacher_name = f"{prefix} {teacher_last_name}"

        notif_title = "Points Awarded"
        notif_message = f"{teacher_name} has awarded you with {point_category}. You have received {points} points."
        supabase.table('notifications').insert({
            'user_id': int(student_id),
            'sender_id': teacher_id,
            'title': notif_title,
            'message': notif_message,
            'notif_type': 'Points',  # Use allowed value
            'status': 'Unread',
            'point_id': point_id  # Pass the point_id here
        }).execute()

        return jsonify({'success': True, 'message': 'Points awarded successfully!',  'data': result.data[0]}), 201

    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/api/bulk-award-points', methods=['POST'])
def api_bulk_award_points():
    if 'user_id' not in session or session.get('role') != 'Teacher':
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401

    data = request.get_json()
    teacher_id = session['user_id']
    student_ids = data.get('student_ids', [])
    points = data.get('points')
    point_category = data.get('category')
    note = data.get('note', '')

    if not student_ids or not points or not point_category:
        return jsonify({'success': False, 'message': 'Missing required fields'}), 400

    try:
        # Get current time in Philippines timezone
        philippines_tz = pytz.timezone('Asia/Manila')
        now_ph = datetime.now(philippines_tz).isoformat()
        
        # Get teacher info for notification message
        teacher_info = supabase.table('user_info').select('last_name', 'gender').eq('id', teacher_id).execute()
        teacher_last_name = ''
        teacher_gender = ''
        if teacher_info.data and len(teacher_info.data) > 0:
            teacher_last_name = teacher_info.data[0].get('last_name', '')
            teacher_gender = (teacher_info.data[0].get('gender', '') or '').lower()
        prefix = 'Mr.'
        if teacher_gender == 'female':
            prefix = 'Mrs.'
        elif teacher_gender == 'other':
            prefix = 'Mx.'
        teacher_name = f"{prefix} {teacher_last_name}"

        inserts = []
        for student_id in student_ids:
            inserts.append({
                'teacher_id': teacher_id,
                'student_id': int(student_id),
                'points': int(points),
                'point_category': point_category,
                'note': note,
                'status': 'approved',
                'received_at': now_ph  # Save in PH time
            })
        
        result = supabase.table('points').insert(inserts).execute()
        if hasattr(result, 'error') and result.error:
            raise Exception(result.error.message)

        # --- Update total_points and send notification for each student ---
        notif_title = "Points Awarded"
        notif_message = f"{teacher_name} has awarded you with {point_category}. You have received {points} points."
        # Use result.data to get point_id for each student
        if result.data:
            for i, student_id in enumerate(student_ids):
                point_id = result.data[i]['point_id'] if 'point_id' in result.data[i] else None
                user_result = supabase.table('user_info').select('total_points').eq('id', int(student_id)).execute()
                current_total = 0
                if user_result.data and 'total_points' in user_result.data[0] and user_result.data[0]['total_points'] is not None:
                    current_total = int(user_result.data[0]['total_points'])
                new_total = current_total + int(points)
                supabase.table('user_info').update({'total_points': new_total}).eq('id', int(student_id)).execute()

                supabase.table('notifications').insert({
                    'user_id': int(student_id),
                    'sender_id': teacher_id,
                    'title': notif_title,
                    'message': notif_message,
                    'notif_type': 'Points',
                    'status': 'Unread',
                    'point_id': point_id
                }).execute()
        # -------------------------------------------

        # --- ACTIVITY LOG: Bulk Award Points ---
        # For bulk operations, we'll just show the count instead of all names
        record_admin_activity_log(
            user_id=teacher_id,
            action='Bulk Award Points',
            activity='Points Management',
            description=f"Awarded {points} points to {len(student_ids)} students for '{point_category}'.",
            user_role='Teacher'
        )

        return jsonify({'success': True, 'message': 'Points awarded successfully!', 'count': len(inserts)}), 201
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/api/recent-awards', methods=['GET'])
def api_recent_awards():
    if 'user_id' not in session or session.get('role') != 'Teacher':
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401

    teacher_id = session['user_id']

    try:
        # Get the 5 most recent point awards given by this teacher
        result = (
            supabase.table('points')
            .select('point_id, student_id, points, point_category, note, received_at')
            .eq('teacher_id', teacher_id)
            .order('received_at', desc=True)
            .limit(5)
            .execute()
        )
        awards = result.data if result.data else []

        # Get student names, profile pictures, and streaks for display
        student_ids = [award['student_id'] for award in awards]
        students_map = {}
        if student_ids:
            students_result = supabase.table('user_info').select('id, first_name, last_name, streak, year_level, section').in_('id', student_ids).execute()
            pics_result = supabase.table('profile_pictures').select('user_id, file_path').in_('user_id', student_ids).execute()
            pics_map = {p['user_id']: p['file_path'] for p in pics_result.data} if pics_result.data else {}
            if students_result.data:
                for s in students_result.data:
                    students_map[s['id']] = {
                        'name': f"{s.get('first_name', '')} {s.get('last_name', '')}",
                        'pic': pics_map.get(s['id'], '/static/image/default-avatar.png'),
                        'streak': s.get('streak', 0),
                        'grade': s.get('year_level', ''),
                        'section': s.get('section', '')
                    }

        # Prepare awards list
        awards_list = []
        for award in awards:
            student_info = students_map.get(award['student_id'], {'name': 'Unknown', 'pic': '/static/image/default-avatar.png', 'streak': 0})
            awards_list.append({
                'student': student_info['name'],
                'student_pic': student_info['pic'],
                'streak': student_info['streak'],
                'grade': student_info.get('grade', ''),
                'section': student_info.get('section', ''),
                'points': award['points'],
                'category': award['point_category'],
                'note': award.get('note', ''),
                'date': str(award['received_at']) if award.get('received_at') else '',
            })

        return jsonify({'success': True, 'awards': awards_list})

    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500
    
@app.route('/api/all-awards', methods=['GET'])
def api_all_awards():
    if 'user_id' not in session or session.get('role') != 'Teacher':
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401

    teacher_id = session['user_id']

    try:
        # Get teacher's name
        teacher_result = supabase.table('user_info').select('first_name, last_name').eq('id', teacher_id).execute()
        teacher_name = "Teacher"
        if teacher_result.data:
            teacher_data = teacher_result.data[0]
            teacher_name = f"{teacher_data.get('first_name', '')} {teacher_data.get('last_name', '')}".strip()

        # Get all point awards given by this teacher
        result = (
            supabase.table('points')
            .select('point_id, student_id, points, point_category, note, received_at')
            .eq('teacher_id', teacher_id)
            .order('received_at', desc=True)
            .execute()
        )
        awards = result.data if result.data else []

        # DEBUG: Log raw awards data
        print(f"=== DEBUG: RAW AWARDS DATA FOR TEACHER {teacher_id} ===")
        print(f"Total awards found: {len(awards)}")
        print("Awards sample:", awards[:3] if awards else "No awards")
        
        # Get student names, profile pictures, AND STREAK for display
        student_ids = [award['student_id'] for award in awards]
        
        # DEBUG: Log student IDs from awards
        print(f"=== DEBUG: STUDENT IDS FROM AWARDS ===")
        print(f"Student IDs found: {student_ids}")
        print(f"Unique student IDs: {list(set(student_ids))}")
        print(f"Count of unique student IDs: {len(set(student_ids))}")
        
        students_map = {}
        if student_ids:
            # Include streak in the student query
            students_result = supabase.table('user_info').select('id, first_name, last_name, year_level, section, streak').in_('id', student_ids).execute()
            pics_result = supabase.table('profile_pictures').select('user_id, file_path').in_('user_id', student_ids).execute()
            pics_map = {p['user_id']: p['file_path'] for p in pics_result.data} if pics_result.data else {}
            
            # DEBUG: Log student query results
            print(f"=== DEBUG: STUDENT QUERY RESULTS ===")
            print(f"Students found in database: {len(students_result.data) if students_result.data else 0}")
            if students_result.data:
                for s in students_result.data:
                    print(f"Student: ID={s['id']}, Name={s.get('first_name', '')} {s.get('last_name', '')}")
            
            if students_result.data:
                for s in students_result.data:
                    students_map[s['id']] = {
                        'name': f"{s.get('first_name', '')} {s.get('last_name', '')}",
                        'grade': s.get('year_level', ''),
                        'section': s.get('section', ''),
                        'pic': pics_map.get(s['id'], '/static/image/default-avatar.png'),
                        'streak': s.get('streak', 0)  # Add streak data here
                    }

        # DEBUG: Log students map
        print(f"=== DEBUG: STUDENTS MAP ===")
        print(f"Students in map: {len(students_map)}")
        for student_id, info in students_map.items():
            print(f"Map entry: ID={student_id}, Name={info['name']}")

        # Prepare awards list
        awards_list = []
        for award in awards:
            student_info = students_map.get(award['student_id'], {
                'name': 'Unknown', 
                'grade': '', 
                'section': '', 
                'pic': '/static/image/default-avatar.png',
                'streak': 0  # Default streak
            })
            awards_list.append({
                'id': award['point_id'],
                'student': student_info['name'],
                'student_id': award['student_id'],  # Make sure this is included
                'grade': student_info['grade'],
                'section': student_info['section'],
                'pic': student_info['pic'],
                'points': award['points'],
                'category': award['point_category'],
                'note': award.get('note', ''),
                'date': str(award['received_at']) if award.get('received_at') else '',
                'awarded_by': teacher_name,
                'streak': student_info['streak']  # Add streak to awards data
            })

        # DEBUG: Log final awards list
        print(f"=== DEBUG: FINAL AWARDS LIST ===")
        print(f"Total awards in response: {len(awards_list)}")
        unique_students_in_response = set(a['student_id'] for a in awards_list)
        print(f"Unique students in response: {len(unique_students_in_response)}")
        print(f"Student IDs in response: {list(unique_students_in_response)}")

        return jsonify({'success': True, 'awards': awards_list, 'teacher_name': teacher_name})

    except Exception as e:
        print(f"=== DEBUG: ERROR IN API ===")
        print(f"Error: {str(e)}")
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/api/point-leaders', methods=['GET'])
def api_point_leaders():
    if 'user_id' not in session or session.get('role') != 'Teacher':
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401

    teacher_id = session['user_id']

    try:
        # Get the current school year
        current_school_year = get_current_school_year()
        if not current_school_year:
            return jsonify({'success': False, 'message': 'Unable to determine current school year'}), 500

        # Get all ACTIVE classrooms assigned to this teacher for CURRENT school year
        assignments_result = supabase.table('teacher_class_assignments') \
            .select('grade_level, section') \
            .eq('teacher_id', teacher_id) \
            .eq('school_year', current_school_year) \
            .eq('status', 'active') \
            .execute()
        assignments = assignments_result.data if assignments_result.data else []

        students = []
        student_ids = []
        for assignment in assignments:
            grade_level = assignment['grade_level']
            section = assignment['section']
            students_result = supabase.table('user_info') \
                .select('id, first_name, last_name, year_level, section, total_points, streak') \
                .eq('role', 'Student') \
                .eq('year_level', grade_level) \
                .eq('section', section) \
                .execute()
            if students_result.data:
                students.extend(students_result.data)
                student_ids.extend([student['id'] for student in students_result.data])

        # Remove duplicates
        unique_students = {student['id']: student for student in students}.values()

        # Get profile pictures for all students
        pics_map = {}
        if student_ids:
            pics_result = supabase.table('profile_pictures') \
                .select('user_id, file_path') \
                .in_('user_id', student_ids) \
                .execute()
            if pics_result.data:
                pics_map = {p['user_id']: p['file_path'] for p in pics_result.data}

        # Prepare and sort students by total_points
        student_list = []
        for student in unique_students:
            sid = student['id']
            student_list.append({
                'id': sid,
                'name': f"{student.get('first_name', '')} {student.get('last_name', '')}",
                'grade': student.get('year_level', ''),
                'section': student.get('section', ''),
                'points': int(student.get('total_points', 0)),  # Use total_points directly
                'streak': student.get('streak', 0),  # Add streak data
                'student_pic': pics_map.get(sid, '/static/image/default-avatar.png')  # Add profile picture
            })
        # Sort by points descending
        student_list.sort(key=lambda x: x['points'], reverse=True)
        top5 = student_list[:10]

        return jsonify({'success': True, 'leaders': top5})

    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500


#############################ACTIVITIES.HTML#############################
@app.route('/api/my-classrooms', methods=['GET'])
def api_my_classrooms():
    if 'user_id' not in session or session.get('role') != 'Teacher':
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401

    teacher_id = session['user_id']
    try:
        # Get current school year
        current_school_year = get_current_school_year()
        if not current_school_year:
            return jsonify({'success': False, 'message': 'Unable to determine current school year'}), 500
        
        assignments = supabase.table('teacher_class_assignments') \
            .select('grade_level, section') \
            .eq('teacher_id', teacher_id) \
            .eq('school_year', current_school_year) \
            .eq('status', 'active') \
            .execute()
        classrooms = []
        if assignments.data:
            for a in assignments.data:
                # Check if classroom has at least one student
                students_result = supabase.table('user_info') \
                    .select('id') \
                    .eq('role', 'Student') \
                    .eq('year_level', a['grade_level']) \
                    .eq('section', a['section']) \
                    .execute()
                if students_result.data and len(students_result.data) > 0:
                    classrooms.append({
                        'grade': a['grade_level'],
                        'section': a['section'],
                        'label': f"Grade {a['grade_level']} - Section {a['section']}"
                    })
        return jsonify({'success': True, 'classrooms': classrooms})
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500
    
@app.route('/api/assign-task', methods=['POST'])
def api_assign_task():
    if 'user_id' not in session or session.get('role') != 'Teacher':
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401

    data = request.get_json()
    grade_level = data.get('grade_level')
    section = data.get('section')
    task = data.get('task')
    points = data.get('points', 0)
    description = data.get('description', '')
    due_date = data.get('due_date', None)
    priority = data.get('priority', 'medium')
    template = data.get('template', 'default')

    if not grade_level or not section or not task:
        return jsonify({'success': False, 'message': 'Missing required fields'}), 400

    teacher_id = session['user_id']
    
    # Get current school year and quarter
    current_school_year = get_current_school_year()
    current_quarter = get_current_quarter()
    quarter_id = current_quarter.get('quarter_id') if current_quarter else None

    # Get all students in the selected classroom
    students_result = supabase.table('user_info') \
        .select('id') \
        .eq('role', 'Student') \
        .eq('year_level', grade_level) \
        .eq('section', section) \
        .execute()
    students = students_result.data if students_result.data else []

    assignments = []
    for student in students:
        assignments.append({
            'teacher_id': teacher_id,
            'student_id': student['id'],
            'grade_level': grade_level,
            'section': section,
            'task': task,
            'points': points,
            'description': description,
            'due_date': due_date,
            'priority': priority,
            'template': template,
            'status': 'Assigned',
            'school_year': current_school_year,
            'quarter_id': quarter_id,
            'assigned_at': datetime.now(philippines_tz).isoformat()
        })

    if assignments:
        result = supabase.table('task_assignments').insert(assignments).execute()
        if hasattr(result, 'error') and result.error:
            return jsonify({'success': False, 'message': str(result.error)}), 400

        # --- ACTIVITY LOG: Create Activities ---
        record_admin_activity_log(
            user_id=teacher_id,
            action='Create Activities',
            activity='Activities Management',
            description=f"Assigned Activity '{task}' for Grade {grade_level} Section {section} (Points: {points})",
            user_role='Teacher'
        )

        # --- SEND NOTIFICATION TO STUDENTS ---
        teacher_info = supabase.table('user_info').select('last_name', 'gender').eq('id', teacher_id).execute()
        teacher_last_name = ''
        teacher_gender = ''
        if teacher_info.data and len(teacher_info.data) > 0:
            teacher_last_name = teacher_info.data[0].get('last_name', '')
            teacher_gender = (teacher_info.data[0].get('gender', '') or '').lower()
        prefix = 'Mr.'
        if teacher_gender == 'female':
            prefix = 'Mrs.'
        elif teacher_gender == 'other':
            prefix = 'Mx.'
        teacher_name = f"{prefix} {teacher_last_name}"

        notif_title = "New Activity Assigned"
        notif_message = f"{teacher_name} has assigned you a new activity: '{task}'."

       # --- Send notifications to students ---
        if result.data:
            for i, student in enumerate(students):
                try:
                    # Get the task_id from the inserted record
                    task_record = result.data[i]
                    assigned_task_id = task_record.get('task_id')
                    
                    supabase.table('notifications').insert({
                        'user_id': student['id'],
                        'sender_id': teacher_id,
                        'title': notif_title,
                        'message': notif_message,
                        'notif_type': 'Task',
                        'status': 'Unread',
                        'task_id': assigned_task_id
                    }).execute()
                except (IndexError, KeyError) as e:
                    print(f"Error sending notification to student {student['id']}: {e}")
                    # Send notification without task_id as fallback
                    supabase.table('notifications').insert({
                        'user_id': student['id'],
                        'sender_id': teacher_id,
                        'title': notif_title,
                        'message': notif_message,
                        'notif_type': 'Task',
                        'status': 'Unread'
                    }).execute()

        return jsonify({'success': True, 'message': 'Activities assigned to classroom.'})


############################# AUTO-COMPLETE ACTIVITIES AT QUARTER END #############################

@app.route('/api/admin/auto-complete-quarter-activities', methods=['POST'])
def auto_complete_quarter_activities():
    """
    Auto-complete all unfinished activities for a specific quarter
    Award points to students who submitted work
    """
    if 'user_id' not in session or session.get('role') != 'Admin':
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401
    
    try:
        data = request.get_json()
        school_year = data.get('school_year')
        quarter_id = data.get('quarter_id')
        
        if not school_year or not quarter_id:
            return jsonify({'success': False, 'message': 'Missing school_year or quarter_id'}), 400
        
        # Get all unfinished activities for this quarter
        activities = supabase.table('task_assignments') \
            .select('*') \
            .eq('school_year', school_year) \
            .eq('quarter_id', quarter_id) \
            .neq('status', 'Completed') \
            .execute()
        
        if not activities.data:
            return jsonify({'success': True, 'message': 'No activities to complete', 'updated': 0})
        
        updated_count = 0
        
        for activity in activities.data:
            student_id = activity['student_id']
            points = activity.get('points', 0)
            task_id = activity['task_id']
            current_status = activity.get('status')
            
            # Mark activity as Completed
            supabase.table('task_assignments') \
                .update({'status': 'Completed'}) \
                .eq('task_id', task_id) \
                .eq('student_id', student_id) \
                .execute()
            
            # Award points to student if status was "Submitted" or "Completed"
            if current_status in ['Submitted', 'Completed'] and points > 0:
                # Get current student points
                student_result = supabase.table('user_info') \
                    .select('points') \
                    .eq('id', student_id) \
                    .execute()
                
                if student_result.data:
                    current_points = student_result.data[0].get('points', 0) or 0
                    new_points = current_points + points
                    
                    # Update student points
                    supabase.table('user_info') \
                        .update({'points': new_points}) \
                        .eq('id', student_id) \
                        .execute()
                    
                    # Log points award
                    supabase.table('points_log').insert({
                        'student_id': student_id,
                        'points_earned': points,
                        'activity': f'Auto-awarded for Quarter End: {school_year}',
                        'date_earned': datetime.now(philippines_tz).isoformat(),
                        'school_year': school_year,
                        'quarter_id': quarter_id
                    }).execute()
            
            updated_count += 1
        
        # Log admin activity
        record_admin_activity_log(
            user_id=session['user_id'],
            action='Auto-complete Activities',
            activity='Activities Management',
            description=f"Auto-completed {updated_count} activities for {school_year} Quarter {quarter_id}",
            user_role='Admin'
        )
        
        return jsonify({
            'success': True,
            'message': f'Completed {updated_count} activities and awarded points',
            'updated': updated_count
        })
        
    except Exception as e:
        app.logger.error(f"Error auto-completing activities: {str(e)}")
        return jsonify({'success': False, 'message': str(e)}), 500


@app.route('/api/completed-activities', methods=['GET'])
def get_completed_activities():
    """
    Get completed activities with filtering by school_year and quarter
    """
    if 'user_id' not in session or session.get('role') != 'Teacher':
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401
    
    try:
        teacher_id = session['user_id']
        current_school_year = get_current_school_year()
        auto_deny_finished_school_year_tasks(teacher_id, current_school_year)
        school_year = request.args.get('school_year')
        quarter_id = request.args.get('quarter_id')
        search_query = request.args.get('search', '').lower()
        quarter_name = None

        if quarter_id:
            quarter_result = supabase.table('quarters') \
                .select('quarter_name') \
                .eq('quarter_id', quarter_id) \
                .limit(1) \
                .execute()
            if quarter_result.data:
                quarter_name = quarter_result.data[0].get('quarter_name')
        
        # Build query
        query = supabase.table('task_assignments') \
            .select('*') \
            .eq('teacher_id', teacher_id)
        
        # Add school year filter if provided
        if school_year:
            query = query.eq('school_year', school_year)
        
        # Add quarter filter if provided
        if quarter_name:
            query = query.eq('quarter', quarter_name)
        
        activities = query.execute()

        rows = activities.data or []
        student_ids = list({r.get('student_id') for r in rows if r.get('student_id') is not None})
        student_map = {}
        if student_ids:
            students_result = supabase.table('user_info') \
                .select('id,first_name,last_name') \
                .in_('id', student_ids) \
                .execute()
            for s in students_result.data or []:
                student_map[s['id']] = f"{s.get('first_name', '')} {s.get('last_name', '')}".strip()

        # Group same activity rows together
        grouped = {}
        for t in rows:
            key = t.get('activity_group_id') or (
                t.get('task'),
                t.get('grade_level'),
                t.get('section'),
                t.get('school_year'),
                t.get('quarter'),
                str(t.get('assigned_at', ''))
            )

            if key not in grouped:
                grouped[key] = {
                    'id': f"{t['task_id']}",
                    'title': t.get('task', ''),
                    'description': t.get('description', ''),
                    'grade_level': t.get('grade_level', ''),
                    'section': t.get('section', ''),
                    'school_year': t.get('school_year'),
                    'quarter': t.get('quarter'),
                    'due_date': t.get('due_date', ''),
                    'points': t.get('points', 0),
                    'priority': t.get('priority') or 'medium',
                    'pending_students': [],
                    'approval_students': [],
                    'denied_students': [],
                    'completed_students': [],
                    'template': t.get('template') or 'default',
                    'activity_group_id': t.get('activity_group_id', ''),
                    'attachments': deserialize_attachment_array(t.get('attachments', [])),
                    'links': deserialize_link_array(t.get('links', [])),
                }

            student_id = t.get('student_id')
            student_entry = {
                'id': student_id,
                'name': student_map.get(student_id, 'Student')
            }
            status = t.get('status')
            if status == 'Assigned':
                grouped[key]['pending_students'].append(student_entry)
            elif status == 'Pending':
                student_entry['submitted_at'] = t.get('submitted_at') or t.get('received_at')
                student_entry['image_urls'] = t.get('image_urls')
                grouped[key]['approval_students'].append(student_entry)
            elif status == 'Denied':
                grouped[key]['denied_students'].append(student_entry)
            elif status == 'Completed':
                student_entry['completed_at'] = t.get('completed_at')
                grouped[key]['completed_students'].append(student_entry)

        # Keep only finished groups for "Completed Activities" page
        grouped_activities = []
        for g in grouped.values():
            has_pending = len(g['pending_students']) > 0
            has_approval = len(g['approval_students']) > 0
            has_finished = (len(g['completed_students']) + len(g['denied_students'])) > 0
            if not has_pending and not has_approval and has_finished:
                grouped_activities.append(g)

        # Filter by search query if provided
        if search_query:
            grouped_activities = [
                a for a in grouped_activities
                if search_query in (a.get('title', '') or '').lower() or
                   search_query in (a.get('description', '') or '').lower()
            ]

        return jsonify({
            'success': True,
            'activities': grouped_activities,
            'total': len(grouped_activities)
        })
        
    except Exception as e:
        app.logger.error(f"Error fetching completed activities: {str(e)}")
        return jsonify({'success': False, 'message': str(e)}), 500


@app.route('/api/quarters-history', methods=['GET'])
def get_quarters_history():
    """
    Get all historical quarters for filtering
    """
    try:
        preferred_school_year = '2025-2026'
        current_school_year = get_current_school_year()

        quarters = supabase.table('quarters') \
            .select('quarter_id, quarter_name, school_year, status, start_date, end_date') \
            .order('school_year', desc=True) \
            .order('quarter_id', desc=True) \
            .execute()
        
        if not quarters.data:
            return jsonify({
                'success': True,
                'quarters': [],
                'grouped_by_school_year': {},
                'default_school_year': current_school_year or preferred_school_year
            })
        
        # Group by school year
        grouped = {}
        for quarter in quarters.data:
            sy = quarter.get('school_year')
            if sy not in grouped:
                grouped[sy] = []
            grouped[sy].append(quarter)

        available_school_years = list(grouped.keys())
        if preferred_school_year in available_school_years:
            default_school_year = preferred_school_year
        elif current_school_year in available_school_years:
            default_school_year = current_school_year
        else:
            default_school_year = available_school_years[0] if available_school_years else None
        
        return jsonify({
            'success': True,
            'quarters': quarters.data or [],
            'grouped_by_school_year': grouped,
            'default_school_year': default_school_year
        })
        
    except Exception as e:
        app.logger.error(f"Error fetching quarters history: {str(e)}")
        return jsonify({'success': False, 'message': str(e)}), 500


# 1. Place the function above your /api/my-tasks route:
def auto_deny_past_due_tasks(teacher_id):
    # Get all tasks for this teacher that are past due and still assigned
    now = datetime.now(timezone.utc).date()
    result = supabase.table('task_assignments') \
        .select('task_id, due_date, status, student_id') \
        .eq('teacher_id', teacher_id) \
        .eq('status', 'Assigned') \
        .execute()
    if not result.data:
        return
    for row in result.data:
        due_date = row.get('due_date')
        if due_date:
            try:
                due = datetime.fromisoformat(str(due_date).replace('Z', '+00:00')).date()
                if due < now:
                    # Mark as Denied
                    supabase.table('task_assignments') \
                        .update({'status': 'Denied'}) \
                        .eq('task_id', row['task_id']) \
                        .eq('student_id', row['student_id']) \
                        .execute()
            except Exception:
                continue

def auto_deny_finished_school_year_tasks(teacher_id, current_school_year):
    """
    Deny unfinished tasks from school years that already ended.
    Unfinished = Assigned or Pending. Completed tasks are untouched.
    """
    try:
        if not current_school_year:
            return

        old_tasks = supabase.table('task_assignments') \
            .select('task_id, student_id') \
            .eq('teacher_id', teacher_id) \
            .in_('status', ['Assigned', 'Pending']) \
            .neq('school_year', current_school_year) \
            .execute()

        if not old_tasks.data:
            return

        for row in old_tasks.data:
            try:
                supabase.table('task_assignments') \
                    .update({'status': 'Denied'}) \
                    .eq('task_id', row['task_id']) \
                    .eq('student_id', row['student_id']) \
                    .execute()
            except Exception:
                continue
    except Exception as e:
        app.logger.warning(f"auto_deny_finished_school_year_tasks failed: {str(e)}")


def deserialize_attachment_array(attachments):
    """Deserialize text array of JSON strings into objects for frontend display"""
    if not attachments:
        return []
    
    result = []
    for attachment in attachments:
        if isinstance(attachment, str):
            try:
                obj = json.loads(attachment)
                # Map database field names to frontend display names
                result.append({
                    'url': obj.get('file_url'),
                    'name': obj.get('original_filename'),
                    'size': obj.get('file_size'),
                    'type': obj.get('file_type'),
                    'file_url': obj.get('file_url'),
                    'original_filename': obj.get('original_filename'),
                    'file_type': obj.get('file_type'),
                    'file_size': obj.get('file_size')
                })
            except (json.JSONDecodeError, TypeError):
                # If parsing fails, skip this attachment
                continue
        elif isinstance(attachment, dict):
            # Already an object (e.g., from older versions)
            result.append(attachment)
    
    return result


def deserialize_link_array(links):
    """Deserialize text array of JSON strings into objects for frontend display"""
    if not links:
        return []
    
    result = []
    for link in links:
        if isinstance(link, str):
            try:
                obj = json.loads(link)
                result.append(obj)
            except (json.JSONDecodeError, TypeError):
                # If parsing fails, skip this link
                continue
        elif isinstance(link, dict):
            # Already an object
            result.append(link)
    
    return result

def normalize_links(links):
    """Normalize links to ensure consistent format with url and title properties"""
    if not links:
        return []
    
    if not isinstance(links, list):
        # If it's a single link dict, wrap it in a list
        if isinstance(links, dict):
            links = [links]
        else:
            return []
    
    normalized = []
    for link in links:
        if isinstance(link, dict):
            # Ensure both url and title fields exist
            normalized_link = {
                'url': link.get('url', ''),
                'title': link.get('title', link.get('url', 'Link'))
            }
            if normalized_link['url']:  # Only add if url is not empty
                normalized.append(normalized_link)
        elif isinstance(link, str):
            # Try to parse JSON string
            try:
                link_obj = json.loads(link)
                if isinstance(link_obj, dict):
                    normalized_link = {
                        'url': link_obj.get('url', ''),
                        'title': link_obj.get('title', link_obj.get('url', 'Link'))
                    }
                    if normalized_link['url']:
                        normalized.append(normalized_link)
            except (json.JSONDecodeError, TypeError):
                # Skip invalid links
                continue
    
    return normalized

@app.route('/api/my-tasks', methods=['GET'])
def api_my_tasks():
    if 'user_id' not in session or session.get('role') != 'Teacher':
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401

    teacher_id = session['user_id']
    grade = request.args.get('grade')
    section = request.args.get('section')
    current_school_year = get_current_school_year()
    auto_deny_past_due_tasks(teacher_id)
    auto_deny_finished_school_year_tasks(teacher_id, current_school_year)
    query = supabase.table('task_assignments').select('*').eq('teacher_id', teacher_id)
    if current_school_year:
        query = query.eq('school_year', current_school_year)
    if grade:
        query = query.eq('grade_level', grade)
    if section:
        query = query.eq('section', section)
    result = query.execute()
    tasks = result.data if result.data else []

    # --- NEW: Fetch all student names in one query ---
    student_ids = list({t['student_id'] for t in tasks})
    student_map = {}
    if student_ids:
        students_result = supabase.table('user_info').select('id,first_name,last_name').in_('id', student_ids).execute()
        for s in students_result.data or []:
            student_map[s['id']] = f"{s['first_name']} {s['last_name']}"

    grouped = {}
    for t in tasks:
        key = t.get('activity_group_id') or (
            t.get('task'),
            t.get('grade_level'),
            t.get('section'),
            t.get('school_year'),
            t.get('quarter'),
            str(t.get('assigned_at', ''))
        )
        if key not in grouped:
            grouped[key] = {
                'id': f"{t['task_id']}",
                'title': t['task'],
                'description': t.get('description', ''),
                'grade_level': t['grade_level'],
                'section': t['section'],
                'school_year': t.get('school_year'),
                'quarter': t.get('quarter'),
                'due_date': t.get('due_date', ''),
                'points': t.get('points', 0),
                'priority': t.get('priority') or 'medium',
                'pending_students': [],
                'approval_students': [],
                'denied_students': [],
                'completed_students': [],
                'template': t.get('template') or 'default',
                'activity_group_id': t.get('activity_group_id', ''),  # Add activity_group_id
                'attachments': deserialize_attachment_array(t.get('attachments', [])),  # Deserialize JSON strings to objects
                'links': deserialize_link_array(t.get('links', [])),  # Deserialize JSON strings to objects
            }
        student_name = student_map.get(t['student_id'], 'Unknown')
        if t['status'] == 'Assigned':
            grouped[key]['pending_students'].append({
                'id': t['student_id'],
                'name': student_name
            })
        elif t['status'] == 'Pending':
            submitted_at = t.get('submitted_at') or t.get('received_at')
            grouped[key]['approval_students'].append({
                'id': t['student_id'],
                'name': student_name,
                'submitted_at': submitted_at,
                'image_urls': t.get('image_urls', None) 
            })
        elif t['status'] == 'Denied':
            grouped[key]['denied_students'].append({
                'id': t['student_id'],
                'name': student_name
            })
        elif t['status'] == 'Completed':
            grouped[key]['completed_students'].append({
                'id': t['student_id'],
                'name': student_name,
                'completed_at': t.get('completed_at')
            })

    # --- Add status field to each grouped task ---
    from datetime import datetime, timezone
    now = datetime.now(timezone.utc).date()
    def compute_status(group):
        due_date = group.get('due_date')
        has_pending = len(group['pending_students']) > 0
        has_approval = len(group['approval_students']) > 0
        has_denied = len(group['denied_students']) > 0
        has_completed = len(group['completed_students']) > 0
        # Due Soon: due in 3 days or less and has pending students
        if due_date:
            try:
                due = datetime.fromisoformat(str(due_date).replace('Z', '+00:00')).date()
                days_left = (due - now).days
                if has_pending and 0 <= days_left <= 3:
                    return 'Due Soon'
            except Exception:
                pass
        if has_pending:
            return 'To Do'
        if has_approval:
            return 'Pending'
        if has_denied and not has_pending and not has_approval and not has_completed:
            return 'Denied'
        if has_completed and not has_pending and not has_approval:
            return 'Awarded'
        return 'All'

    grouped_tasks = []
    for group in grouped.values():
        group['status'] = compute_status(group)
        grouped_tasks.append(group)

    return jsonify({'success': True, 'tasks': grouped_tasks})

@app.route('/api/approve-task-submissions', methods=['POST'])
def approve_task_submissions():
    if 'user_id' not in session or session.get('role') != 'Teacher':
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401

    data = request.get_json()
    task_id = data.get('taskId')
    selected_students = data.get('selectedStudents', [])
    denied_students = data.get('deniedStudents', [])

    if not task_id:
        return jsonify({'success': False, 'message': 'Missing taskId'}), 400

    # Get task info (points, etc.)
    task_result = supabase.table('task_assignments').select('task, points, teacher_id, grade_level, section').eq('task_id', task_id).limit(1).execute()
    if not task_result.data:
        return jsonify({'success': False, 'message': 'Task not found'}), 404
    task = task_result.data[0]
    points = int(task.get('points', 0))
    teacher_id = task.get('teacher_id')
    grade_level = task.get('grade_level')
    section = task.get('section')
    task_name = task.get('task')

    # Get teacher info for notification message
    teacher_info = supabase.table('user_info').select('last_name', 'gender').eq('id', teacher_id).execute()
    teacher_last_name = ''
    teacher_gender = ''
    if teacher_info.data and len(teacher_info.data) > 0:
        teacher_last_name = teacher_info.data[0].get('last_name', '')
        teacher_gender = (teacher_info.data[0].get('gender', '') or '').lower()
    prefix = 'Mr.'
    if teacher_gender == 'female':
        prefix = 'Mrs.'
    elif teacher_gender == 'other':
        prefix = 'Mx.'
    teacher_name = f"{prefix} {teacher_last_name}"

    # 1. Award points to selected students and mark as Completed
    for student_id in selected_students:
        assignment = supabase.table('task_assignments') \
            .select('task_id', 'status') \
            .eq('task', task_name) \
            .eq('student_id', int(student_id)) \
            .eq('grade_level', grade_level) \
            .eq('section', section) \
            .limit(1).execute()
        if assignment.data:
            assignment_id = assignment.data[0]['task_id']
            current_status = assignment.data[0].get('status')
            point_id = None
            if current_status == 'Pending':
                point_result = supabase.table('points').insert({
                    'teacher_id': teacher_id,
                    'student_id': int(student_id),
                    'points': points,
                    'point_category': 'Activity',
                    'note': f"Completed Activity: {task_name}",
                    'status': 'approved'
                }).execute()
                if point_result.data and 'point_id' in point_result.data[0]:
                    point_id = point_result.data[0]['point_id']
                user_result = supabase.table('user_info').select('total_points').eq('id', int(student_id)).execute()
                current_total = 0
                if user_result.data and 'total_points' in user_result.data[0] and user_result.data[0]['total_points'] is not None:
                    current_total = int(user_result.data[0]['total_points'])
                new_total = current_total + points
                supabase.table('user_info').update({'total_points': new_total}).eq('id', int(student_id)).execute()
            supabase.table('task_assignments').update({
                'status': 'Completed'
            }).eq('task_id', assignment_id).execute()

            # --- SEND COMPLETION NOTIFICATION ---
            notif_title = "Activity Approved"
            notif_message = f"{teacher_name} has approved your activity '{task_name}' and awarded you {points} points."
            supabase.table('notifications').insert({
                'user_id': int(student_id),
                'sender_id': teacher_id,
                'title': notif_title,
                'message': notif_message,
                'notif_type': 'Task',
                'status': 'Unread',
                'point_id': point_id,
                'task_id': assignment_id
            }).execute()

            # --- ACTIVITY LOG: Activity Approved ---
            record_admin_activity_log(
                user_id=teacher_id,
                action='Approve Activity',
                activity=' Management',
                description=f"Approved activity '{task_name}' for student ID {student_id} and awarded {points} points.",
                user_role='Teacher'
            )

    # 2. Mark as Denied for denied students
    for student_id in denied_students:
        assignment = supabase.table('task_assignments') \
            .select('task_id', 'status') \
            .eq('task', task_name) \
            .eq('student_id', int(student_id)) \
            .eq('grade_level', grade_level) \
            .eq('section', section) \
            .limit(1).execute()
        if assignment.data:
            assignment_id = assignment.data[0]['task_id']
            supabase.table('task_assignments').update({
                'status': 'Denied'
            }).eq('task_id', assignment_id).execute()
            notif_title = "Activity Submission Denied"
            notif_message = f"{teacher_name} has denied your submission for '{task_name}'. Please review and resubmit if needed."
            supabase.table('notifications').insert({
                'user_id': int(student_id),
                'sender_id': teacher_id,
                'title': notif_title,
                'message': notif_message,
                'notif_type': 'Task',
                'status': 'Unread',
                'task_id': assignment_id
            }).execute()

    # --- ACTIVITY LOG: Approve/Deny Task Submissions (summary) ---
    if selected_students or denied_students:
        approved_count = len(selected_students)
        denied_count = len(denied_students)
        description_parts = []
        if approved_count > 0:
            description_parts.append(f"Approved {approved_count} student(s)")
        if denied_count > 0:
            description_parts.append(f"Denied {denied_count} student(s)")
        description = f"{', '.join(description_parts)} for task '{task_name}'."
        record_admin_activity_log(
            user_id=session.get('user_id'),
            action='Process Task Submissions',
            activity='Activity Management',
            description=description,
            user_role='Teacher'
        )

    return jsonify({'success': True, 'message': 'Task submissions processed.'})

def send_task_reminders():
    now = datetime.now(timezone.utc)
    intervals = [
        timedelta(days=2),
        timedelta(days=1),
        timedelta(hours=10),
        timedelta(hours=3),
        timedelta(hours=1)
    ]
    # Get all pending/assigned tasks
    result = supabase.table('task_assignments') \
        .select('task_id, task, due_date, status, student_id, grade_level, section') \
        .in_('status', ['Assigned', 'Pending']) \
        .execute()
    if not result.data:
        return

    for row in result.data:
        due_date = row.get('due_date')
        if not due_date:
            continue
        try:
            due = datetime.fromisoformat(str(due_date).replace('Z', '+00:00'))
            if due.tzinfo is None:
                due = due.replace(tzinfo=timezone.utc)
        except Exception:
            continue

        # Calculate time until due
        time_until_due = due - now
        for interval in intervals:
            # Check if the reminder should be sent at this interval
            # Allow a 10-minute window for each interval
            if interval - timedelta(minutes=10) < time_until_due <= interval + timedelta(minutes=10):
                # Check if reminder already sent (optional: add a reminders table or check notifications)
                notif_title = "Task Due Soon"
                notif_message = f"Reminder: Your task '{row['task']}' is due in {str(interval)}. Please accomplish it before the deadline."
                supabase.table('notifications').insert({
                    'user_id': row['student_id'],
                    'sender_id': None,
                    'title': notif_title,
                    'message': notif_message,
                    'notif_type': 'Task',
                    'status': 'Unread',
                    'task_id': row['task_id']
                }).execute()
                break  # Only send one reminder per run per interval

@app.route('/api/task-submissions/<path:task_key>/<int:student_id>', methods=['GET'])
def api_task_submissions(task_key, student_id):
    if 'user_id' not in session or session.get('role') != 'Teacher':
        return jsonify({'success': False, 'message': 'Unauthorized'}), 403

    try:
        # Parse the task_key to get task name, grade, and section
        # Format: "task_name|grade_level|section"
        parts = task_key.split('|')
        if len(parts) != 3:
            return jsonify({'success': False, 'message': 'Invalid task key format'}), 400
            
        task_name, grade_level, section = parts
        
        # First, find the actual task_id for this student
        task_result = supabase.table('task_assignments') \
            .select('task_id') \
            .eq('task', task_name) \
            .eq('grade_level', grade_level) \
            .eq('section', section) \
            .eq('student_id', student_id) \
            .limit(1) \
            .execute()
        
        if not task_result.data:
            return jsonify({'success': True, 'files': []})
            
        actual_task_id = task_result.data[0]['task_id']
        
        # Fetch file submissions from task_file_submissions table
        submissions = supabase.table('task_file_submissions') \
            .select('*') \
            .eq('task_id', actual_task_id) \
            .eq('student_id', student_id) \
            .execute()
        
        files = []
        if submissions.data:
            for sub in submissions.data:
                files.append({
                    'file_name': sub.get('filename'),  # Renamed filename
                    'original_filename': sub.get('original_filename'),  # Original filename
                    'file_url': sub.get('file_url'),
                    'file_size': sub.get('file_size', 0),
                    'uploaded_at': sub.get('uploaded_at'),
                    'mime_type': sub.get('mime_type'),
                    'file_type': sub.get('file_type')
                })
        
        return jsonify({'success': True, 'files': files})
    except Exception as e:
        app.logger.error(f"Error fetching task submissions: {str(e)}")
        return jsonify({'success': False, 'message': 'Error loading files'}), 500

#############################SETTINGS.HTML#############################
@app.route('/api/users/<int:user_id>/profile', methods=['PUT'])
def update_profile(user_id):
    if 'user_id' not in session or session['user_id'] != user_id:
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401

    data = request.json
    if not data:
        return jsonify({'success': False, 'message': 'No data provided'}), 400

    # Only allow updating these fields
    update_fields = {
        'first_name': data.get('first_name', '').strip(),
        'last_name': data.get('last_name', '').strip(),
        'email': data.get('email', '').strip(),
        'mobile_no': data.get('mobile_no', '').strip()
    }

    # Basic validation
    if not update_fields['first_name'] or not update_fields['last_name'] or not update_fields['email']:
        return jsonify({'success': False, 'message': 'First name, last name, and email are required.'}), 400

    # Validate names: only letters and spaces
    if not re.match(r'^[A-Za-z\s]+$', update_fields['first_name']):
        return jsonify({'success': False, 'message': 'First name must contain only letters.'}), 400
    if not re.match(r'^[A-Za-z\s]+$', update_fields['last_name']):
        return jsonify({'success': False, 'message': 'Last name must contain only letters.'}), 400

    # Validate mobile_no: must start with 09 and be 11 digits
    if update_fields['mobile_no'] and not re.match(r'^09\d{9}$', update_fields['mobile_no']):
        return jsonify({'success': False, 'message': 'Mobile number must start with 09 and be 11 digits.'}), 400

    try:
        result = supabase.table('user_info').update(update_fields).eq('id', user_id).execute()
        if hasattr(result, 'error') and result.error:
            return jsonify({'success': False, 'message': 'Error updating profile.'}), 500

        # --- ACTIVITY LOG: Profile Update ---
        record_admin_activity_log(
            user_id=user_id,
            action='Update Profile',
            activity='Profile Management',
            description=f"Updated profile information.",
            user_role=session.get('role')
        )

        return jsonify({'success': True, 'message': 'Profile updated successfully.'})
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/api/users/<int:user_id>/change-password', methods=['POST'])
def change_password(user_id):
    if 'user_id' not in session or session['user_id'] != user_id:
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401

    data = request.json
    if not data:
        return jsonify({'success': False, 'message': 'No data provided'}), 400

    current_password = data.get('current_password', '').strip()
    new_password = data.get('new_password', '').strip()

    if not current_password or not new_password:
        return jsonify({'success': False, 'message': 'Current and new password required.'}), 400

    # Password strength validation (min 8 chars, upper/lower/digit/special)
    import re
    if not re.match(r'^.*(?=.{8,})(?=.*[a-z])(?=.*[A-Z])(?=.*\d)(?=.*[^A-Za-z0-9]).*$', new_password):
        return jsonify({'success': False, 'message': 'Password must be at least 8 characters and include upper/lower case letters, numbers, and special characters.'}), 400

    try:
        # Fetch user info
        result = supabase.table('user_info').select('password').eq('id', user_id).execute()
        if not result.data:
            return jsonify({'success': False, 'message': 'User not found.'}), 404

        user = result.data[0]
        if user['password'] != current_password:
            return jsonify({'success': False, 'message': 'Current password is incorrect.'}), 400

        # Update password
        update_result = supabase.table('user_info').update({'password': new_password}).eq('id', user_id).execute()
        if hasattr(update_result, 'error') and update_result.error:
            return jsonify({'success': False, 'message': update_result.error.message}), 500

        # Log activity
        record_admin_activity_log(
            user_id=user_id,
            action='Change Password',
            activity='Profile Management',
            description='Changed account password.',
            user_role=session.get('role')
        )

        return jsonify({'success': True, 'message': 'Password changed successfully.'})
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500
    
@app.route('/api/users/<int:user_id>/upload-profile-picture', methods=['POST'])
def upload_profile_picture(user_id):
    if 'user_id' not in session or session['user_id'] != user_id:
        return

    if 'profile_picture' not in request.files:
        return

    file = request.files['profile_picture']
    if file.filename == '':
        return

    if file and allowed_file(file.filename):
        # Add timestamp to filename to make it unique for history
        timestamp = int(datetime.utcnow().timestamp())
        filename = secure_filename(f"user_{user_id}_{timestamp}_" + file.filename)
        file_bytes = file.read()

        # Upload to Supabase Storage (use PUT, not POST)
        bucket = "profile-pictures"
        # FIX: Use SUPABASE_URL and SUPABASE_KEY directly (not config.*)
        storage_url = f"{SUPABASE_URL}/storage/v1/object/{bucket}/{filename}"
        headers = {
            "apikey": SUPABASE_KEY,
            "Authorization": f"Bearer {SUPABASE_KEY}",
            "Content-Type": file.mimetype
        }
        response = requests.put(
            storage_url,
            headers=headers,
            data=file_bytes
        )

        if response.status_code not in [200, 201]:
            return jsonify({'success': False, 'message': 'Upload failed'}), 400

        # FIX: Use SUPABASE_URL directly (not config.SUPABASE_URL)
        public_url = f"{SUPABASE_URL}/storage/v1/object/public/{bucket}/{filename}"

        # Save/Update in profile_pictures table using upsert (current avatar)
        supabase.table('profile_pictures').upsert({
            'user_id': user_id,
            'file_path': public_url,
            'uploaded_at': datetime.utcnow().isoformat()
        }, on_conflict='user_id').execute()

        # Check if this exact file_path already exists in history to avoid duplicates
        existing_check = supabase.table('profile_picture_history') \
            .select('id') \
            .eq('user_id', user_id) \
            .eq('file_path', public_url) \
            .execute()
        
        # Only insert into history if it's a new/different image
        if not existing_check.data:
            # Insert new avatar into history
            supabase.table('profile_picture_history').insert({
                'user_id': user_id,
                'file_path': public_url,
                'uploaded_at': datetime.utcnow().isoformat()
            }).execute()

            # Clean up old avatars - keep only the most recent 6
            cleanup_old_avatars(user_id)

        # Update session so navbar uses new image immediately
        session['profile_picture'] = public_url

        # --- ACTIVITY LOG: Avatar Change ---
        record_admin_activity_log(
            user_id=user_id,
            action='Change Avatar',
            activity='Profile Management',
            description="Changed profile picture.",
            user_role=session.get('role')
        )

        return jsonify({'success': True, 'file_path': public_url})
    
    return jsonify({'success': False, 'message': 'Invalid file type'}), 400

def cleanup_old_avatars(user_id):
    """Keep only the 6 most recent avatars for a user"""
    try:
        # Get all avatars for this user, ordered by most recent first
        all_avatars = supabase.table('profile_picture_history') \
            .select('id, uploaded_at') \
            .eq('user_id', user_id) \
            .order('uploaded_at', desc=True) \
            .execute()
        
        if all_avatars.data and len(all_avatars.data) > 6:
            # Get IDs of avatars to keep (first 6 most recent)
            keep_ids = [avatar['id'] for avatar in all_avatars.data[:6]]
            
            # Delete older avatars (everything beyond the first 6)
            supabase.table('profile_picture_history') \
                .delete() \
                .eq('user_id', user_id) \
                .not_.in_('id', keep_ids) \
                .execute()
            
            print(f"Cleaned up old avatars for user {user_id}, kept {len(keep_ids)} most recent ones")
            
    except Exception as e:
        print(f"Error cleaning up old avatars for user {user_id}: {e}")

@app.route('/api/users/<int:user_id>/profile-picture', methods=['GET'])
def get_profile_picture(user_id):
    if 'user_id' not in session or session['user_id'] != user_id:
        return jsonify({'success': False, 'message': 'Unauthorized'}), 403

    try:
        # Get profile picture from profile_pictures table
        result = supabase.table('profile_pictures').select('file_path').eq('user_id', user_id).execute()
        
        if result.data and len(result.data) > 0:
            return jsonify({
                'success': True, 
                'file_path': result.data[0]['file_path']
            })
        else:
            return jsonify({
                'success': False, 
                'message': 'No profile picture found'
            })
    except Exception as e:
        print(f"Error fetching profile picture: {e}")
        return jsonify({
            'success': False, 
            'message': 'Server error'
        }), 500

# NEW ENDPOINT: Get last 6 avatars from history
@app.route('/api/users/<int:user_id>/recent-avatars', methods=['GET'])
def get_recent_avatars(user_id):
    if 'user_id' not in session or session['user_id'] != user_id:
        return jsonify({'success': False, 'message': 'Unauthorized'}), 403

    try:
        # Get last 6 profile pictures from history table ordered by most recent first
        result = supabase.table('profile_picture_history') \
            .select('file_path, uploaded_at') \
            .eq('user_id', user_id) \
            .order('uploaded_at', desc=True) \
            .limit(6) \
            .execute()
        
        if result.data:
            return jsonify({
                'success': True, 
                'avatars': result.data
            })
        else:
            return jsonify({
                'success': True,  # Still success, just no data
                'avatars': []
            })
    except Exception as e:
        print(f"Error fetching recent avatars: {e}")
        return jsonify({
            'success': False, 
            'message': 'Server error'
        }), 500

# NEW ENDPOINT: Set avatar from history (prevents duplicate uploads)
@app.route('/api/users/<int:user_id>/set-avatar-from-history', methods=['POST'])
def set_avatar_from_history(user_id):
    if 'user_id' not in session or session['user_id'] != user_id:
        return jsonify({'success': False, 'message': 'Unauthorized'}), 403

    try:
        data = request.get_json()
        file_path = data.get('file_path')
        
        if not file_path:
            return jsonify({'success': False, 'message': 'No file path provided'}), 400

        # Update current profile picture
        supabase.table('profile_pictures').upsert({
            'user_id': user_id,
            'file_path': file_path,
            'uploaded_at': datetime.utcnow().isoformat()
        }, on_conflict='user_id').execute()

        # Update session
        session['profile_picture'] = file_path

        # --- ACTIVITY LOG: Avatar Change ---
        record_admin_activity_log(
            user_id=user_id,
            action='Change Avatar',
            activity='Profile Management',
            description="Changed profile picture from history.",
            user_role=session.get('role')
        )

        return jsonify({'success': True, 'file_path': file_path})
        
    except Exception as e:
        print(f"Error setting avatar from history: {e}")
        return jsonify({
            'success': False, 
            'message': 'Server error'
        }), 500

# Optional maintenance endpoint (you can call this periodically if needed)
@app.route('/api/users/<int:user_id>/cleanup-avatars', methods=['POST'])
def cleanup_avatars(user_id):
    if 'user_id' not in session or session['user_id'] != user_id:
        return jsonify({'success': False, 'message': 'Unauthorized'}), 403

    try:
        cleanup_old_avatars(user_id)
        return jsonify({
            'success': True, 
            'message': 'Avatar cleanup completed'
        })
            
    except Exception as e:
        print(f"Error cleaning up avatars: {e}")
        return jsonify({
            'success': False, 
            'message': 'Server error during cleanup'
        }), 500

# Add this endpoint after the set_avatar_from_history endpoint

@app.route('/api/users/<int:user_id>/delete-recent-avatar', methods=['DELETE'])
def delete_recent_avatar(user_id):
    if 'user_id' not in session or session['user_id'] != user_id:
        return jsonify({'success': False, 'message': 'Unauthorized'}), 403

    try:
        data = request.get_json()
        if not data or 'file_path' not in data:
            return jsonify({'success': False, 'message': 'Missing file_path'}), 400

        file_path = data.get('file_path')

        # Delete from profile_picture_history table
        result = supabase.table('profile_picture_history') \
            .delete() \
            .eq('user_id', user_id) \
            .eq('file_path', file_path) \
            .execute()

        # If no records were deleted from history, try to delete the file from storage if it exists
        if not result.data:
            # Still return success as the file might not be in history
            return jsonify({'success': True, 'message': 'Avatar deleted successfully'})

        # Optional: Delete the actual file from Supabase storage if needed
        # This depends on your storage setup
        try:
            # Extract bucket and path from file_path if it's a storage URL
            if 'storage' in file_path or 'uploads' in file_path:
                # Parse the file path to get the storage path
                # Example: /storage/v1/object/public/uploads/user_123/avatar_xyz.jpg
                storage_path = file_path.split('/object/public/')[-1] if '/object/public/' in file_path else file_path
                
                # Try to delete from storage
                supabase.storage.from_('uploads').remove([storage_path])
        except Exception as storage_error:
            # Log but don't fail if storage deletion fails
            app.logger.warning(f"Could not delete file from storage: {storage_error}")

        return jsonify({'success': True, 'message': 'Avatar deleted successfully'})

    except Exception as e:
        app.logger.error(f"Error deleting avatar: {str(e)}")
        return jsonify({'success': False, 'message': str(e)}), 500

from datetime import datetime

# Flask routes for quarters management with Supabase and activity logging

@app.route('/api/quarters', methods=['POST'])
def create_quarters():
    """Create school year quarters"""
    data = request.get_json()
    school_year = data.get('school_year')
    quarters_data = data.get('quarters', [])
    
    try:
        # Get user info for activity log
        user_id = session.get('user_id')
        user_role = session.get('user_role', 'Admin')
        
        # Get user's full name from database
        user_result = supabase.table('user_info')\
            .select('first_name, last_name')\
            .eq('id', user_id)\
            .execute()
        
        user_name = "Unknown"
        if user_result.data:
            user_data = user_result.data[0]
            user_name = f"{user_data['first_name']} {user_data['last_name']}"
        
        # Delete existing quarters for this school year using Supabase
        supabase.table('quarters').delete().eq('school_year', school_year).execute()
        
        # Insert new quarters
        quarters_to_insert = []
        quarter_names = []
        
        for quarter in quarters_data:
            quarter_name = quarter['quarter_name']
            quarters_to_insert.append({
                'quarter_name': quarter_name,
                'start_date': quarter['start_date'],
                'end_date': quarter['end_date'],
                'school_year': school_year,
                'created_by': user_id
            })
            quarter_names.append(quarter_name)
        
        # Batch insert all quarters
        result = supabase.table('quarters').insert(quarters_to_insert).execute()
        
        # Log the activity
        quarter_names_str = ", ".join(quarter_names)
        activity_description = f"Admin {user_name} has added quarters: {quarter_names_str} for school year {school_year}"
        
        supabase.table('admin_activity_log').insert({
            'user_id': user_id,
            'user_role': user_role,
            'action': 'CREATE',
            'activity': 'Quarters Management',
            'description': activity_description,
            'details': f"Added {len(quarters_data)} quarters for {school_year}",
            'created_at': datetime.now().isoformat()
        }).execute()
        
        # Get the active quarter
        today = datetime.now().date().isoformat()
        active_quarter_result = supabase.table('quarters')\
            .select('*')\
            .eq('school_year', school_year)\
            .lte('start_date', today)\
            .gte('end_date', today)\
            .execute()
        
        active_quarter = active_quarter_result.data[0] if active_quarter_result.data else None
        
        return jsonify({
            'success': True, 
            'message': 'Quarters created successfully',
            'active_quarter': active_quarter
        })
    
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/api/quarters/<int:quarter_id>', methods=['PUT'])
def update_quarter(quarter_id):
    """Update a specific quarter"""
    try:
        data = request.get_json()
        start_date = data.get('start_date')
        end_date = data.get('end_date')
        school_year = data.get('school_year')

        # Get user info for activity log
        user_id = session.get('user_id')
        user_role = session.get('user_role', 'Admin')
        
        # Get user's full name from database
        user_result = supabase.table('user_info')\
            .select('first_name, last_name')\
            .eq('id', user_id)\
            .execute()
        
        user_name = "Unknown"
        if user_result.data:
            user_data = user_result.data[0]
            user_name = f"{user_data['first_name']} {user_data['last_name']}"
        
        # Get the quarter name before updating
        quarter_result = supabase.table('quarters')\
            .select('quarter_name')\
            .eq('quarter_id', quarter_id)\
            .execute()
        
        quarter_name = "Unknown Quarter"
        if quarter_result.data:
            quarter_name = quarter_result.data[0]['quarter_name']

        # Update the quarter
        result = supabase.table('quarters')\
            .update({
                'start_date': start_date,
                'end_date': end_date
            })\
            .eq('quarter_id', quarter_id)\
            .execute()

        if result.data:
            # Log the activity
            activity_description = f"Admin {user_name} has edited {quarter_name} for school year {school_year}"
            
            supabase.table('admin_activity_log').insert({
                'user_id': user_id,
                'user_role': user_role,
                'action': 'UPDATE',
                'activity': 'Quarters Management',
                'description': activity_description,
                'details': f"Updated {quarter_name}: {start_date} to {end_date}",
                'created_at': datetime.now().isoformat()
            }).execute()
            
            return jsonify({'success': True, 'message': 'Quarter updated successfully'})
        else:
            return jsonify({'success': False, 'message': 'Quarter not found'}), 404

    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/api/quarters/current')
def get_current_quarters():
    """Get current school year quarters with improved current quarter detection"""
    try:
        today = datetime.now().date().isoformat()
        
        # Get all quarters and find the current one
        quarters_result = supabase.table('quarters')\
            .select('*')\
            .order('start_date')\
            .execute()
        
        if not quarters_result.data:
            return jsonify({'success': True, 'quarters': [], 'current_quarter': None})
        
        current_quarter = None
        quarters = []
        
        for quarter in quarters_result.data:
            quarter_data = quarter.copy()
            start_date = quarter['start_date']
            end_date = quarter['end_date']
            
            # Improved status logic
            if start_date <= today <= end_date:
                quarter_data['is_active'] = True
                quarter_data['is_current'] = True
                quarter_data['status'] = 'active'
                current_quarter = quarter_data  # This is the current quarter
            elif today > end_date:
                quarter_data['is_active'] = False
                quarter_data['is_current'] = False
                quarter_data['status'] = 'finished'
            elif today < start_date:
                quarter_data['is_active'] = False
                quarter_data['is_current'] = False
                quarter_data['status'] = 'upcoming'
            else:
                quarter_data['is_active'] = False
                quarter_data['is_current'] = False
                quarter_data['status'] = 'upcoming'
                
            quarters.append(quarter_data)
        
        # If no active quarter found, try to find the most recent finished quarter
        if not current_quarter:
            finished_quarters = [q for q in quarters if q['status'] == 'finished']
            if finished_quarters:
                current_quarter = finished_quarters[-1]  # Most recent finished quarter
        
        return jsonify({
            'success': True, 
            'quarters': quarters,
            'current_quarter': current_quarter
        })
    
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/api/quarters/<int:quarter_id>', methods=['GET'])
def get_quarter(quarter_id):
    """Get a specific quarter"""
    try:
        result = supabase.table('quarters')\
            .select('*')\
            .eq('quarter_id', quarter_id)\
            .execute()
        
        if result.data:
            return jsonify({'success': True, 'quarter': result.data[0]})
        else:
            return jsonify({'success': False, 'message': 'Quarter not found'}), 404
    
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500

# OLD DUPLICATE ROUTES REMOVED - Using newer versions at lines 10473+ instead

@app.route('/api/quarters/<string:school_year>')
def get_quarters_by_year(school_year):
    """Get quarters for a specific school year"""
    try:
        result = supabase.table('quarters')\
            .select('*')\
            .eq('school_year', school_year)\
            .order('start_date')\
            .execute()
        
        if result.data:
            # Calculate status for each quarter
            today = datetime.now().date().isoformat()
            quarters = []
            for quarter in result.data:
                quarter_data = quarter.copy()
                start_date = quarter['start_date']
                end_date = quarter['end_date']
                
                # Improved status logic with "Finished" status
                if start_date <= today <= end_date:
                    quarter_data['is_active'] = True
                    quarter_data['is_current'] = False
                    quarter_data['status'] = 'active'
                elif today > end_date:
                    quarter_data['is_active'] = False
                    quarter_data['is_current'] = False
                    quarter_data['status'] = 'finished'
                elif today < start_date:
                    quarter_data['is_active'] = False
                    quarter_data['is_current'] = (today >= start_date)
                    quarter_data['status'] = 'upcoming'
                else:
                    quarter_data['is_active'] = False
                    quarter_data['is_current'] = False
                    quarter_data['status'] = 'upcoming'
                    
                quarters.append(quarter_data)
            
            return jsonify({'success': True, 'quarters': quarters})
        else:
            return jsonify({'success': True, 'quarters': []})
    
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500

#############################SCHOOL YEAR TRANSITION#############################
def can_execute_school_year_transition(current_school_year):
    """Transition is allowed only after the 4th quarter has finished."""
    try:
        current_school_year = str(current_school_year or '').strip()
        if not current_school_year:
            return False, "Current school year is required."

        quarters_result = supabase.table('quarters') \
            .select('quarter_name, end_date') \
            .eq('school_year', current_school_year) \
            .order('start_date', desc=False) \
            .execute()

        quarters = quarters_result.data or []
        if not quarters:
            return False, f"No quarters found for {current_school_year}."

        def _parse_quarter_end_date(raw_value):
            raw = str(raw_value or '').strip()
            if not raw or raw.lower() == 'none':
                return None

            normalized = raw.replace('Z', '+00:00')
            try:
                return datetime.fromisoformat(normalized).date()
            except ValueError:
                pass

            # Fallback: extract YYYY-MM-DD from noisy datetime strings
            date_match = re.search(r'\d{4}-\d{2}-\d{2}', normalized)
            if date_match:
                try:
                    return datetime.strptime(date_match.group(0), '%Y-%m-%d').date()
                except ValueError:
                    return None
            return None

        # Prefer explicit 4th quarter name.
        fourth_quarter = next(
            (
                q for q in quarters
                if any(
                    token in str(q.get('quarter_name', '')).lower()
                    for token in ('4th', 'fourth', 'quarter 4', 'q4')
                )
            ),
            None
        )

        # If naming is inconsistent, use the quarter with latest parseable end date.
        if fourth_quarter is None:
            quarter_candidates = [
                (q, _parse_quarter_end_date(q.get('end_date')))
                for q in quarters
            ]
            quarter_candidates = [item for item in quarter_candidates if item[1] is not None]
            fourth_quarter = max(quarter_candidates, key=lambda item: item[1])[0] if quarter_candidates else quarters[-1]

        raw_status = str(fourth_quarter.get('status', '')).lower()
        if raw_status in ('finished', 'completed'):
            return True, None

        quarter_end_date = _parse_quarter_end_date(fourth_quarter.get('end_date'))
        if quarter_end_date is None:
            end_date_raw = str(fourth_quarter.get('end_date', '')).strip()
            return False, f"Cannot validate transition for {current_school_year}: 4th quarter end date is missing."

        today = datetime.now(philippines_tz).date()
        if today > quarter_end_date:
            return True, None

        return False, (
            f"School year {current_school_year} is not yet finished. "
            f"Transition is only available after 4th Quarter ends "
            f"({quarter_end_date.strftime('%Y-%m-%d')})."
        )
    except Exception as e:
        logger.error(f"Transition validation error for {current_school_year}: {str(e)}")
        logger.error(f"Full error details: {type(e).__name__}: {traceback.format_exc()}")
        return False, f"Transition validation error: {str(e)}"


@app.route('/api/admin/school-year-transition-preview', methods=['POST'])
def school_year_transition_preview():
    """Preview what will happen during school year transition (dry-run)"""
    try:
        data = request.get_json()
        current_school_year = data.get('current_school_year')
        next_school_year = data.get('next_school_year')
        
        if not current_school_year or not next_school_year:
            return jsonify({'success': False, 'message': 'School years are required'}), 400

        can_transition, validation_message = can_execute_school_year_transition(current_school_year)
        if not can_transition:
            return jsonify({'success': False, 'message': validation_message}), 400
        
        # Define section mapping for transitions
        section_mapping = {
            # Grade 7 sections → Grade 8 sections
            'Love': 'Matthew',
            'Faith': 'Mark',
            'Hope': 'Luke',
            'Peace': 'John',
            # Grade 8 sections → Grade 9 sections
            'Matthew': 'Psalms',
            'Mark': 'Jeremiah',
            'Luke': 'Isaiah',
            'John': 'Proverbs',
            # Grade 9 sections → Grade 10 sections
            'Psalms': 'Deuteronomy',
            'Jeremiah': 'Leviticus',
            'Isaiah': 'Exodus',
            'Proverbs': 'Genesis'
            # Grade 10 sections stay the same (Deuteronomy, Leviticus, Exodus, Genesis)
        }
        
        # Get all students to count transitions
        students_result = supabase.table('user_info')\
            .select('id, first_name, last_name, year_level, status, section')\
            .eq('role', 'Student')\
            .execute()
        
        if not students_result.data:
            return jsonify({'success': False, 'message': 'No students found'}), 404
        
        # Count transitions by grade
        transitions = {
            'grade_7_to_8': 0,
            'grade_8_to_9': 0,
            'grade_9_to_10': 0,
            'grade_10_to_archived': 0,
            'total_students': len(students_result.data),
            'section_transitions': {}
        }
        
        for student in students_result.data:
            if student['status'] != 'Active':
                continue
            
            year_level = student.get('year_level', '')
            section = student.get('section', '')
            
            # Handle both formats: '7' and 'Grade 7'
            if year_level in ['7', 'Grade 7']:
                transitions['grade_7_to_8'] += 1
            elif year_level in ['8', 'Grade 8']:
                transitions['grade_8_to_9'] += 1
            elif year_level in ['9', 'Grade 9']:
                transitions['grade_9_to_10'] += 1
            elif year_level in ['10', 'Grade 10']:
                transitions['grade_10_to_archived'] += 1
            
            # Track section transitions
            if section in section_mapping:
                new_section = section_mapping[section]
                if new_section not in transitions['section_transitions']:
                    transitions['section_transitions'][new_section] = 0
                transitions['section_transitions'][new_section] += 1
        
        return jsonify({
            'success': True,
            'preview': {
                'current_school_year': current_school_year,
                'next_school_year': next_school_year,
                'transitions': transitions,
                'timestamp': datetime.now().isoformat()
            }
        })
    
    except Exception as e:
        logger.error(f"Transition preview error: {str(e)}")
        return jsonify({'success': False, 'message': str(e)}), 500


@app.route('/api/admin/school-year-transition', methods=['POST'])
def execute_school_year_transition():
    """Execute school year transition - grade advancement and points reset"""
    try:
        data = request.get_json()
        current_school_year = data.get('current_school_year')
        next_school_year = data.get('next_school_year')
        admin_id = session.get('user_id')
        admin_role = session.get('user_role', 'Admin')
        
        if not current_school_year or not next_school_year:
            return jsonify({'success': False, 'message': 'School years are required'}), 400

        can_transition, validation_message = can_execute_school_year_transition(current_school_year)
        if not can_transition:
            return jsonify({'success': False, 'message': validation_message}), 400
        
        # Get admin name
        admin_result = supabase.table('user_info')\
            .select('first_name, last_name')\
            .eq('id', admin_id)\
            .execute()
        
        admin_name = "Unknown"
        if admin_result.data:
            admin_data = admin_result.data[0]
            admin_name = f"{admin_data['first_name']} {admin_data['last_name']}"
        
        # Get all active students
        students_result = supabase.table('user_info')\
            .select('id, first_name, last_name, year_level, status, section, total_points')\
            .eq('role', 'Student')\
            .eq('status', 'Active')\
            .execute()
        
        if not students_result.data:
            return jsonify({'success': False, 'message': 'No active students found'}), 404
        
        # Prepare data structures
        students_to_archive = []
        students_to_update = []
        
        # Grade mapping for transitions
        grade_mapping = {
            '7': '8',
            'Grade 7': '8',
            '8': '9',
            'Grade 8': '9',
            '9': '10',
            'Grade 9': '10',
            '10': 'Archived',
            'Grade 10': 'Archived'
        }
        
        # Section mapping for transitions
        section_mapping = {
            # Grade 7 sections → Grade 8 sections
            'Love': 'Matthew',
            'Faith': 'Mark',
            'Hope': 'Luke',
            'Peace': 'John',
            # Grade 8 sections → Grade 9 sections
            'Matthew': 'Psalms',
            'Mark': 'Jeremiah',
            'Luke': 'Isaiah',
            'John': 'Proverbs',
            # Grade 9 sections → Grade 10 sections
            'Psalms': 'Deuteronomy',
            'Jeremiah': 'Leviticus',
            'Isaiah': 'Exodus',
            'Proverbs': 'Genesis'
            # Grade 10 sections stay the same (Deuteronomy, Leviticus, Exodus, Genesis)
        }
        
        # Process each student
        for student in students_result.data:
            student_id = student['id']
            current_grade = student.get('year_level', '')
            current_section = student.get('section', '')
            new_grade = grade_mapping.get(current_grade)
            
            if not new_grade:
                continue
            
            # Normalize year_level to numeric format for archiving
            normalized_grade = current_grade
            if isinstance(normalized_grade, str):
                # Extract numeric part from "Grade X" format if present
                if normalized_grade.startswith('Grade'):
                    normalized_grade = normalized_grade.replace('Grade ', '').strip()
                # Remove any whitespace and ensure it's just a number
                normalized_grade = normalized_grade.strip()
            
            # Archive student data
            students_to_archive.append({
                'student_id': student_id,
                'school_year': current_school_year,
                'year_level': normalized_grade,
                'section': current_section,
                'status': student.get('status', 'Active'),
                'total_points': student.get('total_points', 0),
                'archived_at': datetime.now(philippines_tz).isoformat(),
                'archived_by': admin_id
            })
            
            # Determine new section for transition
            new_section = section_mapping.get(current_section, current_section)
            
            # For Grade 10 students: transition to Archived status
            # For other grades: advance to next grade and stay Active
            if new_grade == 'Archived':
                new_year_level = '10'  # Keep Grade 10 students as numeric format
                new_status = 'Archived'  # Set status to Archived when moving to Archived
                new_section = current_section  # Grade 10 sections don't transition further
            else:
                new_year_level = new_grade  # Keep numeric format (8, 9, 10)
                new_status = 'Active'
                # new_section already mapped above
            
            students_to_update.append({
                'id': student_id,
                'year_level': new_year_level,
                'status': new_status,
                'section': new_section
            })
        
        # Get current points for archiving (ALL students' points)
        points_result = supabase.table('points')\
            .select('student_id, points, received_at, point_category, note, teacher_id, status')\
            .execute()
        
        # Group points by student for archiving
        points_by_student = {}
        if points_result.data:
            for point_record in points_result.data:
                student_id = point_record['student_id']
                if student_id not in points_by_student:
                    points_by_student[student_id] = []
                points_by_student[student_id].append(point_record)
        
        # Archive points for each student
        points_to_archive = []
        for student_id, points_records in points_by_student.items():
            total_points = sum([p.get('points', 0) for p in points_records])
            points_to_archive.append({
                'student_id': student_id,
                'school_year': current_school_year,
                'points_records': points_records,
                'total_earned': total_points,
                'archived_at': datetime.now(philippines_tz).isoformat()
            })
        
        # Execute all operations
        try:
            # Archive student snapshots
            if students_to_archive:
                supabase.table('student_archives').insert(students_to_archive).execute()
            
            # Archive points
            if points_to_archive:
                supabase.table('points_archives').insert(points_to_archive).execute()
            
            # Update student grades, status, sections, and reset total_points to 0
            for student in students_to_update:
                supabase.table('user_info')\
                    .update({
                        'year_level': student['year_level'],
                        'status': student['status'],
                        'section': student['section'],
                        'total_points': 0
                    })\
                    .eq('id', student['id'])\
                    .execute()
            
            # Mark old teacher assignments as ended (legacy view: allows students to view past assignments)
            # Get all students' old grade levels for marking assignments as ended
            old_assignments = supabase.table('teacher_class_assignments')\
                .select('assignment_id, grade_level, section')\
                .eq('status', 'active')\
                .execute()
            
            if old_assignments.data:
                now_timestamp = datetime.now(philippines_tz).isoformat()
                for assignment in old_assignments.data:
                    # Only mark as ended if the assignment is for a grade that's transitioning
                    old_grade = assignment.get('grade_level')
                    if old_grade in grade_mapping:  # Student is advancing, so mark old assignment as ended
                        supabase.table('teacher_class_assignments')\
                            .update({
                                'status': 'ended',
                                'ended_date': now_timestamp,
                                'ended_by': admin_id
                            })\
                            .eq('assignment_id', assignment['assignment_id'])\
                            .execute()
            
            # Log the activity
            activity_description = f"School year transition executed: {current_school_year} → {next_school_year}. " \
                                 f"Archived {len(students_to_archive)} students and reset points."
            
            supabase.table('admin_activity_log').insert({
                'user_id': admin_id,
                'user_role': admin_role,
                'action': 'EXECUTE',
                'activity': 'School Year Transition',
                'description': activity_description,
                'details': f"Transitioned from {current_school_year} to {next_school_year}. "
                          f"Grade 7→8: {sum(1 for s in students_to_archive if s['year_level'] == 'Grade 7')}, "
                          f"Grade 8→9: {sum(1 for s in students_to_archive if s['year_level'] == 'Grade 8')}, "
                          f"Grade 9→10: {sum(1 for s in students_to_archive if s['year_level'] == 'Grade 9')}, "
                          f"Grade 10→Archived: {sum(1 for s in students_to_archive if s['year_level'] == 'Grade 10')}",
                'created_at': datetime.now(philippines_tz).isoformat()
            }).execute()
            
            return jsonify({
                'success': True,
                'message': 'School year transition completed successfully',
                'summary': {
                    'students_transitioned': len(students_to_archive),
                    'students_archived': sum(1 for s in students_to_archive if s['year_level'] == 'Grade 10'),
                    'points_archived': len(points_to_archive),
                    'timestamp': datetime.now(philippines_tz).isoformat()
                }
            })
        
        except Exception as e:
            logger.error(f"Transition execution error: {str(e)}")
            return jsonify({'success': False, 'message': f'Execution error: {str(e)}'}), 500
    
    except Exception as e:
        logger.error(f"School year transition error: {str(e)}")
        return jsonify({'success': False, 'message': str(e)}), 500


@app.route('/api/admin/school-year-transition-rollback-preview', methods=['POST'])
def school_year_transition_rollback_preview():
    """Preview what will be rolled back (dry-run)"""
    try:
        data = request.get_json()
        school_year_to_rollback = data.get('school_year')
        
        if not school_year_to_rollback:
            return jsonify({'success': False, 'message': 'School year is required'}), 400
        
        # Get all archived students for this school year
        archived_students = supabase.table('student_archives')\
            .select('*')\
            .eq('school_year', school_year_to_rollback)\
            .execute()
        
        if not archived_students.data:
            return jsonify({'success': False, 'message': f'No archived records found for {school_year_to_rollback}'}), 404
        
        # Count students by previous grade
        grade_counts = {
            'Grade 7': 0,
            'Grade 8': 0,
            'Grade 9': 0,
            'Grade 10': 0,
            'other': 0
        }
        
        section_counts = {}
        
        for archive_record in archived_students.data:
            year_level = archive_record.get('year_level', '')
            section = archive_record.get('section', '')
            
            if year_level in grade_counts:
                grade_counts[year_level] += 1
            else:
                grade_counts['other'] += 1
            
            if section:
                section_counts[section] = section_counts.get(section, 0) + 1
        
        # Get archived points count
        archived_points = supabase.table('points_archives')\
            .select('*', count='exact')\
            .eq('school_year', school_year_to_rollback)\
            .execute()
        
        return jsonify({
            'success': True,
            'preview': {
                'school_year': school_year_to_rollback,
                'students_to_restore': len(archived_students.data),
                'grade_distribution': grade_counts,
                'section_distribution': section_counts,
                'points_archives': archived_points.count if archived_points.count else 0,
                'warning': 'All current data will be replaced with archived data. This operation cannot be undone except by re-running the transition.',
                'timestamp': datetime.now().isoformat()
            }
        })
    
    except Exception as e:
        logger.error(f"Rollback preview error: {str(e)}")
        return jsonify({'success': False, 'message': str(e)}), 500


@app.route('/api/admin/school-year-transition-rollback', methods=['POST'])
def school_year_transition_rollback():
    """Rollback school year transition - restore students to their previous state"""
    try:
        data = request.get_json()
        school_year_to_rollback = data.get('school_year')
        admin_id = session.get('user_id')
        admin_role = session.get('user_role', 'Admin')
        
        if not school_year_to_rollback:
            return jsonify({'success': False, 'message': 'School year is required'}), 400
        
        # Get admin name
        admin_result = supabase.table('user_info')\
            .select('first_name, last_name')\
            .eq('id', admin_id)\
            .execute()
        
        admin_name = "Unknown"
        if admin_result.data:
            admin_data = admin_result.data[0]
            admin_name = f"{admin_data['first_name']} {admin_data['last_name']}"
        
        # Get all archived students for this school year
        archived_students = supabase.table('student_archives')\
            .select('*')\
            .eq('school_year', school_year_to_rollback)\
            .execute()
        
        if not archived_students.data:
            return jsonify({'success': False, 'message': f'No archived records found for {school_year_to_rollback}'}), 404
        
        # Prepare rollback data
        students_to_restore = []
        restore_count = 0
        
        for archive_record in archived_students.data:
            student_id = archive_record['student_id']
            
            # Normalize year_level to numeric format (7, 8, 9, 10)
            archived_year_level = archive_record['year_level']
            if archived_year_level and archived_year_level.startswith('Grade'):
                # Extract numeric part from "Grade X" format
                archived_year_level = archived_year_level.replace('Grade ', '').strip()
            
            # Restore student to archived state
            students_to_restore.append({
                'id': student_id,
                'year_level': archived_year_level,
                'status': archive_record['status'],
                'section': archive_record['section']
            })
            restore_count += 1
        
        # Get archived points for this school year
        archived_points = supabase.table('points_archives')\
            .select('*')\
            .eq('school_year', school_year_to_rollback)\
            .execute()
        
        try:
            # Restore student data to previous state
            for student in students_to_restore:
                supabase.table('user_info')\
                    .update({
                        'year_level': student['year_level'],
                        'status': student['status'],
                        'section': student['section']
                    })\
                    .eq('id', student['id'])\
                    .execute()
            
            # Restore total_points from archive for each student
            if archived_points.data:
                for archive in archived_points.data:
                    supabase.table('user_info')\
                        .update({'total_points': archive.get('total_earned', 0)})\
                        .eq('id', archive['student_id'])\
                        .execute()
            
            # Restore teacher class assignments to active status
            # Get all assignments that were marked as ended (they were active in the rolled-back year)
            ended_assignments = supabase.table('teacher_class_assignments')\
                .select('assignment_id')\
                .eq('status', 'ended')\
                .eq('school_year', school_year_to_rollback)\
                .execute()
            
            if ended_assignments.data:
                for assignment in ended_assignments.data:
                    supabase.table('teacher_class_assignments')\
                        .update({
                            'status': 'active',
                            'ended_date': None,
                            'ended_by': None
                        })\
                        .eq('assignment_id', assignment['assignment_id'])\
                        .execute()
            
            # DELETE archived records for this school year to prevent duplicates on re-transition
            supabase.table('student_archives')\
                .delete()\
                .eq('school_year', school_year_to_rollback)\
                .execute()
            
            # DELETE archived points for this school year
            supabase.table('points_archives')\
                .delete()\
                .eq('school_year', school_year_to_rollback)\
                .execute()
            
            # Log the rollback activity
            activity_description = f"School year transition ROLLED BACK: {school_year_to_rollback}. " \
                                 f"Restored {restore_count} students to previous state and deleted archive records."
            
            supabase.table('admin_activity_log').insert({
                'user_id': admin_id,
                'user_role': admin_role,
                'action': 'ROLLBACK',
                'activity': 'School Year Transition Rollback',
                'description': activity_description,
                'details': f"Rollback for {school_year_to_rollback}. Restored {restore_count} students to archived state.",
                'created_at': datetime.now(philippines_tz).isoformat()
            }).execute()
            
            logger.info(f"School year transition rolled back for {school_year_to_rollback} by {admin_name}")
            
            return jsonify({
                'success': True,
                'message': 'Transition rolled back successfully',
                'summary': {
                    'school_year': school_year_to_rollback,
                    'students_restored': restore_count,
                    'points_restored': len(archived_points.data) if archived_points.data else 0,
                    'timestamp': datetime.now(philippines_tz).isoformat()
                }
            })
        
        except Exception as e:
            logger.error(f"Rollback execution error: {str(e)}")
            return jsonify({'success': False, 'message': f'Rollback error: {str(e)}'}), 500
    
    except Exception as e:
        logger.error(f"School year transition rollback error: {str(e)}")
        return jsonify({'success': False, 'message': str(e)}), 500


@app.route('/api/admin/student-archives', methods=['GET'])
def get_student_archives():
    """Get archived student data with filtering"""
    try:
        school_year = request.args.get('school_year')
        year_level = request.args.get('year_level')
        
        query = supabase.table('student_archives').select('*')
        
        if school_year:
            query = query.eq('school_year', school_year)
        
        if year_level:
            query = query.eq('year_level', year_level)
        
        result = query.order('archived_at', desc=True).execute()
        
        # Enrich with student names
        if result.data:
            student_ids = set([r['student_id'] for r in result.data])
            students_result = supabase.table('user_info')\
                .select('id, first_name, last_name, email')\
                .in_('id', list(student_ids))\
                .execute()
            
            student_map = {s['id']: s for s in students_result.data}
            
            for archive in result.data:
                student_info = student_map.get(archive['student_id'], {})
                archive['student_name'] = f"{student_info.get('first_name', 'Unknown')} {student_info.get('last_name', '')}"
                archive['student_email'] = student_info.get('email', '')
        
        return jsonify({'success': True, 'archives': result.data or []})
    
    except Exception as e:
        logger.error(f"Get archives error: {str(e)}")
        return jsonify({'success': False, 'message': str(e)}), 500


@app.route('/api/admin/archive-school-years', methods=['GET'])
def get_archive_school_years():
    """Get only school years that have archived student records."""
    try:
        archives_result = supabase.table('student_archives') \
            .select('school_year') \
            .order('school_year', desc=True) \
            .execute()

        if not archives_result.data:
            return jsonify({'success': True, 'school_years': []})

        school_years = sorted(
            list({row.get('school_year') for row in archives_result.data if row.get('school_year')}),
            reverse=True
        )
        return jsonify({'success': True, 'school_years': school_years})
    except Exception as e:
        logger.error(f"Get archive school years error: {str(e)}")
        return jsonify({'success': False, 'message': str(e)}), 500


@app.route('/api/admin/student-archives/<int:student_id>', methods=['GET'])
def get_student_archive_history(student_id):
    """Get complete archive history for a specific student"""
    try:
        # Get student info
        student_result = supabase.table('user_info')\
            .select('id, first_name, last_name, email, year_level, status')\
            .eq('id', student_id)\
            .execute()
        
        if not student_result.data:
            return jsonify({'success': False, 'message': 'Student not found'}), 404
        
        student_info = student_result.data[0]
        
        # Get archive history
        archives_result = supabase.table('student_archives')\
            .select('*')\
            .eq('student_id', student_id)\
            .order('archived_at', desc=True)\
            .execute()
        
        # Get points archive history
        points_result = supabase.table('points_archives')\
            .select('*')\
            .eq('student_id', student_id)\
            .order('archived_at', desc=True)\
            .execute()
        
        return jsonify({
            'success': True,
            'student': student_info,
            'archive_history': archives_result.data or [],
            'points_history': points_result.data or []
        })
    
    except Exception as e:
        logger.error(f"Get student archive history error: {str(e)}")
        return jsonify({'success': False, 'message': str(e)}), 500


@app.route('/api/admin/points-archive/<int:student_id>', methods=['GET'])
def get_student_points_archive(student_id):
    """Get detailed points archive for a student by school year"""
    try:
        school_year = request.args.get('school_year')
        
        query = supabase.table('points_archives')\
            .select('*')\
            .eq('student_id', student_id)
        
        if school_year:
            query = query.eq('school_year', school_year)
        
        result = query.order('archived_at', desc=True).execute()
        
        return jsonify({'success': True, 'points_archives': result.data or []})
    
    except Exception as e:
        logger.error(f"Get points archive error: {str(e)}")
        return jsonify({'success': False, 'message': str(e)}), 500

#############################QUIZ_MAKER.HTML#############################
# Add this endpoint to your Flask app (where your other routes are)
@app.route('/api/analyze-lesson', methods=['POST'])
def analyze_lesson():
    """Analyze uploaded lesson plan and extract topics and quarter"""
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400
        
        file = request.files['file']
        
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        if not allowed_file(file.filename):
            return jsonify({'error': 'File type not allowed'}), 400
        
        # Extract text from file
        lesson_text = extract_text_from_file(file)
        
        if not lesson_text:
            return jsonify({'error': 'Could not extract text from file'}), 400
        
        # Detect subject from lesson text
        detected_subject = detect_subject_from_lesson(lesson_text) if detect_subject_from_lesson else 'Science'
        
        # Select appropriate quiz generator based on subject
        if detected_subject == 'English' and english_quiz_generator:
            current_generator = english_quiz_generator
        else:
            current_generator = quiz_generator  # Default to Science
            detected_subject = 'Science'
        
        if current_generator is None:
            return jsonify({'error': 'Quiz generator not available'}), 500
        
        is_valid_lesson, validation_error = current_generator.validate_lesson_plan(lesson_text)
        if not is_valid_lesson:
            # Log rejected files for monitoring
            logger.warning(f"Rejected invalid file: {secure_filename(file.filename)} - Reason: {validation_error}")
            return jsonify({'error': validation_error, 'rejected': True}), 400
        
        # Analyze using appropriate quiz generator
        try:
            # Extract topics
            topics = current_generator.extract_topics_from_lesson(lesson_text)
            
            # Detect quarter
            detected_quarter = current_generator.detect_quarter_from_lesson(lesson_text)
            
            # Get primary topic (highest percentage)
            primary_topic = None
            if topics:
                primary_topic = max(topics.items(), key=lambda x: x[1].get('percentage', 0))[0]
            
        except Exception as e:
            logger.error(f"Error analyzing lesson: {e}")
            return jsonify({'error': 'Error analyzing lesson'}), 500
        
        return jsonify({
            'success': True,
            'topics': topics,
            'detected_quarter': detected_quarter,
            'detected_subject': detected_subject,
            'file_name': secure_filename(file.filename)
        })
        
    except Exception as e:
        logger.error(f"Error in analyze_lesson: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/generate-quiz', methods=['POST'])
def generate_quiz():
    """Generate quiz questions from lesson plan"""
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400
        
        file = request.files['file']
        num_questions = request.form.get('num_questions', 5, type=int)
        topic_filter = request.form.get('topic_filter', None)
        quarter_filter = request.form.get('quarter', None)  # Add quarter parameter
        exclude_questions = request.form.get('exclude_questions', '[]')  # Get list of questions to exclude
        random_mode = request.form.get('random_mode', 'false').lower() == 'true'  # True for random refresh
        
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        if not allowed_file(file.filename):
            return jsonify({'error': 'File type not allowed'}), 400
        
        # Extract text from file
        lesson_text = extract_text_from_file(file)
        
        if not lesson_text:
            return jsonify({'error': 'Could not extract text from file'}), 400
        
        # Detect subject from lesson text
        detected_subject = detect_subject_from_lesson(lesson_text) if detect_subject_from_lesson else 'Science'
        
        # Select appropriate quiz generator based on subject
        if detected_subject == 'English' and english_quiz_generator:
            current_generator = english_quiz_generator
        else:
            current_generator = quiz_generator  # Default to Science
            detected_subject = 'Science'
        
        if current_generator is None:
            return jsonify({'error': 'Quiz generator not available'}), 500
        
        # Validate that the document is actually a lesson plan
        is_valid_lesson, validation_error = current_generator.validate_lesson_plan(lesson_text)
        if not is_valid_lesson:
            # Log rejected files for monitoring
            logger.warning(f"Rejected invalid file: {secure_filename(file.filename)} - Reason: {validation_error}")
            return jsonify({'error': validation_error, 'rejected': True}), 400
        
        try:
            # Parse excluded questions
            import json
            try:
                excluded = json.loads(exclude_questions)
            except:
                excluded = []
            
            # If random_mode is True, skip topic-based filtering and use random selection
            if random_mode:
                questions = current_generator._get_random_questions(
                    num_questions=num_questions,
                    quarter_filter=quarter_filter,
                    exclude_questions=excluded
                )
            else:
                questions = current_generator.find_matching_questions(
                    lesson_text,
                    num_questions=num_questions,
                    topic_filter=topic_filter if topic_filter else None,
                    quarter_filter=quarter_filter,  # Pass quarter filter
                    exclude_questions=excluded  # Pass questions to exclude
                )
        except Exception as e:
            logger.error(f"Error generating questions: {e}")
            return jsonify({'error': 'Error generating quiz'}), 500
        
        if not questions:
            return jsonify({'error': 'No questions could be generated from the lesson'}), 400
        
        return jsonify({
            'success': True,
            'quiz': questions,
            'count': len(questions),
            'detected_quarter': quarter_filter,
            'detected_subject': detected_subject
        })
        
    except Exception as e:
        logger.error(f"Error in generate_quiz: {e}")
        return jsonify({'error': str(e)}), 500


# Optional: Add endpoint to log quarter detection feedback
@app.route('/api/log-quarter-detection', methods=['POST'])
def log_quarter_detection():
    """Log quarter detection feedback for analysis"""
    try:
        data = request.get_json()
        user_id = session.get('user_id')
        
        if not user_id:
            return jsonify({'error': 'Not authenticated'}), 401
        
        # Log to database or file for analysis
        log_data = {
            'user_id': user_id,
            'detected_quarter': data.get('detected_quarter'),
            'is_correct': data.get('is_correct'),
            'corrected_quarter': data.get('corrected_quarter'),
            'file_name': data.get('file_name'),
            'timestamp': data.get('timestamp')
        }
        
        # You can save this to a database table or log file
        logger.info(f"Quarter detection feedback: {log_data}")
        
        # Optionally save to Supabase
        try:
            supabase.table('quarter_detection_logs').insert(log_data).execute()
        except:
            pass  # Continue even if logging fails
        
        return jsonify({'success': True})
        
    except Exception as e:
        logger.error(f"Error logging quarter detection: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/save-generated-quiz', methods=['POST'])
def save_generated_quiz():
    """Save temporarily generated quiz to database"""
    try:
        data = request.get_json()
        user_id = session.get('user_id')
        
        if not user_id:
            return jsonify({'error': 'Not authenticated'}), 401
        
        quiz_data = {
            'user_id': user_id,
            'quiz_data': data.get('quiz'),
            'topic_filter': data.get('topic_filter'),
            'num_questions': data.get('num_questions'),
            'expires_at': (datetime.now(timezone.utc) + timedelta(days=7)).isoformat()
        }
        
        result = supabase.table('generated_quizzes').insert(quiz_data).execute()
        
        return jsonify({
            'success': True,
            'quiz_id': result.data[0]['id'],
            'message': 'Quiz temporarily saved. You can now assign it to your class.'
        })
        
    except Exception as e:
        logger.error(f"Error saving generated quiz: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/teacher/create-quiz', methods=['POST'])
def create_teacher_quiz():
    """Teacher creates and assigns a quiz to their class, or updates an existing quiz"""
    try:
        data = request.get_json()
        teacher_id = session.get('user_id')
        
        if not teacher_id:
            logger.warning("Quiz creation attempt without authentication")
            return jsonify({'error': 'Not authenticated'}), 401
        
        # Check if this is an update or create
        quiz_id = data.get('quiz_id')
        is_update = quiz_id is not None and quiz_id != ''
        
        if is_update:
            logger.info(f"Updating quiz {quiz_id} for teacher {teacher_id}")
        else:
            logger.info(f"Creating new quiz for teacher {teacher_id}")
        
        # Validate required fields
        required_fields = ['quiz', 'quiz_title', 'topic', 'quarter', 'grade_level', 'section']
        missing_fields = [f for f in required_fields if not data.get(f)]
        if missing_fields:
            return jsonify({'error': f'Missing required fields: {", ".join(missing_fields)}'}), 400
        
        # Get teacher info
        teacher = supabase.table('user_info')\
            .select('first_name', 'last_name')\
            .eq('id', teacher_id)\
            .execute()
        
        if not teacher.data:
            logger.error(f"Teacher {teacher_id} not found in user_info")
            return jsonify({'error': 'Teacher not found'}), 404
            
        teacher_name = f"{teacher.data[0]['first_name']} {teacher.data[0]['last_name']}"
        
        # Verify teacher is assigned to this class
        assignment = supabase.table('teacher_class_assignments')\
            .select('*')\
            .eq('teacher_id', teacher_id)\
            .eq('grade_level', data.get('grade_level'))\
            .eq('section', data.get('section'))\
            .execute()
        
        if not assignment.data:
            logger.warning(f"Teacher {teacher_id} not assigned to Grade {data.get('grade_level')} - {data.get('section')}")
            return jsonify({'error': 'You are not assigned to this class'}), 403
        
        # Validate quiz data
        quiz_items = data.get('quiz', [])
        if not quiz_items or len(quiz_items) == 0:
            return jsonify({'error': 'Quiz must contain at least one question'}), 400
        
        # Calculate total points
        total_items = len(quiz_items)
        points_per_item = int(data.get('points_per_item', 1))
        time_limit = int(data.get('time_limit', 30))
        
        # Validate point and time values
        if points_per_item < 1 or points_per_item > 10:
            return jsonify({'error': 'Points per item must be between 1 and 10'}), 400
        
        if time_limit < 1 or time_limit > 180:
            return jsonify({'error': 'Time limit must be between 1 and 180 minutes'}), 400
        
        total_points = total_items * points_per_item
        
        # Clean up quiz data for storage
        cleaned_quiz = []
        for q in quiz_items:
            cleaned_q = {
                'question': q.get('question'),
                'correct_answer': int(q.get('correct_answer', 0)),
                'choices': q.get('choices', []),
                'topic': q.get('topic', ''),
                'quarter': q.get('quarter', data.get('quarter')),
                'grade': q.get('grade', data.get('grade_level', '7'))
            }
            cleaned_quiz.append(cleaned_q)
        
        # Create quiz record for database
        quiz_record = {
            'teacher_id': teacher_id,
            'teacher_name': teacher_name,
            'quiz_title': data.get('quiz_title'),
            'topic': data.get('topic'),
            'quarter': data.get('quarter'),
            'grade_level': data.get('grade_level'),
            'section': data.get('section'),
            'total_items': total_items,
            'total_points': total_points,
            'points_per_item': points_per_item,
            'time_limit_minutes': time_limit,
            'instructions': data.get('instructions', ''),
            'quiz_data': cleaned_quiz,
            'due_date': data.get('due_date'),
            'status': 'active'
        }
        
        if is_update:
            # Verify the quiz belongs to this teacher
            existing_quiz = supabase.table('teacher_quizzes')\
                .select('*')\
                .eq('id', quiz_id)\
                .eq('teacher_id', teacher_id)\
                .execute()
            
            if not existing_quiz.data:
                logger.warning(f"Quiz {quiz_id} not found or doesn't belong to teacher {teacher_id}")
                return jsonify({'error': 'Quiz not found or you do not have permission to edit it'}), 404
            
            # Update the existing quiz
            logger.info(f"Updating quiz: {quiz_record['quiz_title']} with {total_items} questions")
            result = supabase.table('teacher_quizzes')\
                .update(quiz_record)\
                .eq('id', quiz_id)\
                .eq('teacher_id', teacher_id)\
                .execute()
            
            if not result.data:
                logger.error(f"Failed to update quiz record. Response: {result}")
                return jsonify({'error': 'Failed to update quiz in database'}), 500
            
            logger.info(f"Quiz updated successfully with ID: {quiz_id}")
            
            return jsonify({
                'success': True,
                'quiz_id': quiz_id,
                'message': f'Quiz "{quiz_record["quiz_title"]}" updated successfully for Grade {quiz_record["grade_level"]} - {quiz_record["section"]}',
                'total_points': total_points,
                'time_limit': time_limit
            }), 200
        else:
            # Create new quiz
            logger.info(f"Inserting new quiz: {quiz_record['quiz_title']} with {total_items} questions")
            
            result = supabase.table('teacher_quizzes').insert(quiz_record).execute()
            
            if not result.data:
                logger.error(f"Failed to insert quiz record. Response: {result}")
                return jsonify({'error': 'Failed to save quiz to database'}), 500
            
            quiz_id = result.data[0]['id']
            logger.info(f"Quiz created successfully with ID: {quiz_id}")
            
            return jsonify({
                'success': True,
                'quiz_id': quiz_id,
                'message': f'Quiz "{quiz_record["quiz_title"]}" created for Grade {quiz_record["grade_level"]} - {quiz_record["section"]}',
                'total_points': total_points,
                'time_limit': time_limit
            }), 201
        
    except Exception as e:
        logger.error(f"Error creating/updating quiz: {str(e)}", exc_info=True)
        return jsonify({
            'error': 'Failed to save quiz',
            'details': str(e)
        }), 500

@app.route('/api/teacher/my-quizzes', methods=['GET'])
def get_my_quizzes():
    """Get all quizzes created by the current teacher"""
    try:
        teacher_id = session.get('user_id')
        
        if not teacher_id:
            return jsonify({'error': 'Not authenticated'}), 401
        
        grade_filter = request.args.get('grade_level')
        section_filter = request.args.get('section')
        
        query = supabase.table('teacher_quizzes')\
            .select('*')\
            .eq('teacher_id', teacher_id)\
            .order('created_at', desc=True)
        
        if grade_filter:
            query = query.eq('grade_level', grade_filter)
        
        if section_filter:
            query = query.eq('section', section_filter)
        
        result = query.execute()
        
        # Format the response and add submission counts
        quizzes = []
        for quiz in result.data:
            # Get submission count for this quiz
            submissions = supabase.table('student_quiz_results')\
                .select('*', count='exact')\
                .eq('teacher_quiz_id', quiz['id'])\
                .execute()
            
            # Get total students in the class
            class_students = supabase.table('user_info')\
                .select('id', count='exact')\
                .eq('role', 'Student')\
                .eq('year_level', quiz['grade_level'])\
                .eq('section', quiz['section'])\
                .execute()
            
            quiz['question_count'] = len(quiz.get('quiz_data', []))
            quiz['submission_count'] = submissions.count if hasattr(submissions, 'count') else 0
            quiz['total_students'] = class_students.count if hasattr(class_students, 'count') else 0
            quizzes.append(quiz)
        
        return jsonify({
            'success': True,
            'quizzes': quizzes,
            'count': len(quizzes)
        }), 200
        
    except Exception as e:
        logger.error(f"Error fetching teacher quizzes: {str(e)}", exc_info=True)
        return jsonify({'error': str(e)}), 500

@app.route('/api/quiz/<int:quiz_id>/can-edit', methods=['GET'])
def check_quiz_editable(quiz_id):
    """Check if a quiz can be edited (has no submissions)"""
    try:
        teacher_id = session.get('user_id')
        
        if not teacher_id:
            return jsonify({'error': 'Not authenticated'}), 401
        
        # Check if quiz exists and belongs to teacher
        quiz = supabase.table('teacher_quizzes')\
            .select('teacher_id')\
            .eq('id', quiz_id)\
            .execute()
        
        if not quiz.data:
            return jsonify({'error': 'Quiz not found'}), 404
        
        if quiz.data[0]['teacher_id'] != teacher_id:
            return jsonify({'error': 'Access denied'}), 403
        
        # Check for submissions
        submissions = supabase.table('student_quiz_results')\
            .select('*', count='exact')\
            .eq('teacher_quiz_id', quiz_id)\
            .execute()
        
        submission_count = submissions.count if hasattr(submissions, 'count') else len(submissions.data) if submissions.data else 0
        
        return jsonify({
            'success': True,
            'can_edit': submission_count == 0,
            'submission_count': submission_count
        })
        
    except Exception as e:
        logger.error(f"Error checking quiz edit status: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/teacher-info', methods=['GET'])
def get_teacher_info():
    """Get teacher information including name and class assignments"""
    try:
        teacher_id = session.get('user_id')
        
        if not teacher_id:
            return jsonify({'error': 'Not authenticated'}), 401
        
        # Get teacher basic info
        teacher = supabase.table('user_info')\
            .select('first_name, last_name')\
            .eq('id', teacher_id)\
            .execute()
        
        if not teacher.data:
            return jsonify({'error': 'Teacher not found'}), 404
        
        teacher_data = teacher.data[0]
        teacher_name = f"{teacher_data['first_name']} {teacher_data['last_name']}"
        
        # Get teacher's class assignments
        assignments = supabase.table('teacher_class_assignments')\
            .select('grade_level, section')\
            .eq('teacher_id', teacher_id)\
            .execute()
        
        assignment_list = []
        if assignments.data:
            for assignment in assignments.data:
                assignment_list.append({
                    'grade_level': assignment['grade_level'],
                    'section': assignment['section']
                })
        
        return jsonify({
            'success': True,
            'name': teacher_name,
            'assignments': assignment_list
        })
        
    except Exception as e:
        logger.error(f"Error getting teacher info: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/current-user', methods=['GET'])
def get_current_user():
    """Get current user information"""
    try:
        user_id = session.get('user_id')
        
        if not user_id:
            return jsonify({'error': 'Not authenticated'}), 401
        
        user = supabase.table('user_info')\
            .select('id, first_name, last_name, role')\
            .eq('id', user_id)\
            .execute()
        
        if not user.data:
            return jsonify({'error': 'User not found'}), 404
        
        user_data = user.data[0]
        
        return jsonify({
            'success': True,
            'user': {
                'id': user_data['id'],
                'name': f"{user_data['first_name']} {user_data['last_name']}",
                'role': user_data['role']
            }
        })
        
    except Exception as e:
        logger.error(f"Error getting current user: {e}")
        return jsonify({'error': str(e)}), 500

# Student Get Quizzes Route
@app.route('/api/student/get-quizzes', methods=['GET'])
def get_student_quizzes():
    """Get all quizzes assigned to a student's class"""
    try:
        student_id = session.get('user_id')
        
        if not student_id:
            return jsonify({'error': 'Not authenticated'}), 401
        
        # Get student info
        student = supabase.table('user_info')\
            .select('first_name', 'last_name', 'year_level', 'section')\
            .eq('id', student_id)\
            .execute()
        
        if not student.data:
            return jsonify({'error': 'Student not found'}), 404
            
        student_data = student.data[0]
        grade_level = student_data['year_level']
        section = student_data['section']
        
        # Get all active quizzes for this student's grade and section
        quizzes = supabase.table('teacher_quizzes')\
            .select('*')\
            .eq('grade_level', grade_level)\
            .eq('section', section)\
            .eq('status', 'active')\
            .order('assigned_date', desc=True)\
            .execute()
        
        # Check which quizzes are already submitted
        submitted = supabase.table('student_quiz_results')\
            .select('teacher_quiz_id')\
            .eq('student_id', student_id)\
            .execute()
        
        submitted_ids = [r['teacher_quiz_id'] for r in submitted.data]
        
        # Process quizzes and add teacher avatar info
        for quiz in quizzes.data:
            quiz['submitted'] = quiz['id'] in submitted_ids
            # Add teacher avatar (will use default or fetch from user_info later)
            quiz['teacher_avatar'] = '/static/image/default-avatar.png'
        
        return jsonify({
            'success': True,
            'quizzes': quizzes.data,
            'student': {
                'name': f"{student_data['first_name']} {student_data['last_name']}",
                'grade_level': grade_level,
                'section': section
            }
        })
        
    except Exception as e:
        logger.error(f"Error fetching quizzes: {e}")
        return jsonify({'error': str(e)}), 500


# Teacher Get Class Results Route
@app.route('/api/teacher/class-results', methods=['GET'])
def get_class_results():
    """Teacher views results for their class"""
    try:
        teacher_id = session.get('user_id')
        grade_level = request.args.get('grade_level')
        section = request.args.get('section')
        quiz_id = request.args.get('quiz_id')  # Optional specific quiz
        
        if not teacher_id:
            return jsonify({'error': 'Not authenticated'}), 401
        
        # Build query
        query = supabase.table('student_quiz_results')\
            .select('*')\
            .eq('teacher_id', teacher_id)
        
        if grade_level:
            query = query.eq('grade_level', grade_level)
        if section:
            query = query.eq('section', section)
        if quiz_id:
            query = query.eq('teacher_quiz_id', quiz_id)
        
        results = query.order('submitted_date', desc=True).execute()
        
        # Get quiz details for context
        quiz_query = supabase.table('teacher_quizzes')\
            .select('*')\
            .eq('teacher_id', teacher_id)
        
        if quiz_id:
            quiz_query = quiz_query.eq('id', quiz_id)
        
        quizzes = quiz_query.execute()
        
        return jsonify({
            'success': True,
            'results': results.data,
            'quizzes': quizzes.data
        })
        
    except Exception as e:
        logger.error(f"Error fetching class results: {e}")
        return jsonify({'error': str(e)}), 500

# Teacher Get Student Performance Route
@app.route('/api/teacher/student-performance/<int:student_id>', methods=['GET'])
def get_student_performance(student_id):
    """Get detailed performance of a specific student"""
    try:
        teacher_id = session.get('user_id')
        
        if not teacher_id:
            return jsonify({'error': 'Not authenticated'}), 401
        
        # Get student info
        student = supabase.table('user_info')\
            .select('first_name', 'last_name', 'year_level', 'section')\
            .eq('id', student_id)\
            .execute()
        
        if not student.data:
            return jsonify({'error': 'Student not found'}), 404
        
        # Get all quiz results for this student under this teacher
        results = supabase.table('student_quiz_results')\
            .select('*')\
            .eq('student_id', student_id)\
            .eq('teacher_id', teacher_id)\
            .order('submitted_date', desc=True)\
            .execute()
        
        # Calculate topic-wise performance
        topic_performance = {}
        for r in results.data:
            if r['topic'] not in topic_performance:
                topic_performance[r['topic']] = {
                    'topic': r['topic'],
                    'quizzes_taken': 0,
                    'total_score': 0,
                    'total_items': 0,
                    'average_percentage': 0
                }
            
            topic_performance[r['topic']]['quizzes_taken'] += 1
            topic_performance[r['topic']]['total_score'] += r['score']
            topic_performance[r['topic']]['total_items'] += r['total_items']
            
        # Calculate averages
        for topic in topic_performance.values():
            topic['average_percentage'] = (topic['total_score'] / topic['total_items'] * 100) if topic['total_items'] > 0 else 0
        
        # Calculate overall statistics
        total_quizzes = len(results.data)
        if total_quizzes > 0:
            total_score = sum(r['score'] for r in results.data)
            total_items = sum(r['total_items'] for r in results.data)
            overall_average = (total_score / total_items * 100) if total_items > 0 else 0
        else:
            overall_average = 0
        
        return jsonify({
            'success': True,
            'student': {
                'name': f"{student.data[0]['first_name']} {student.data[0]['last_name']}",
                'grade_level': student.data[0]['year_level'],
                'section': student.data[0]['section']
            },
            'quiz_history': results.data,
            'topic_performance': list(topic_performance.values()),
            'overall_average': overall_average,
            'total_quizzes': total_quizzes
        })
        
    except Exception as e:
        logger.error(f"Error fetching student performance: {e}")
        return jsonify({'error': str(e)}), 500


# Get Single Quiz Details Route
@app.route('/api/quiz/<int:quiz_id>', methods=['GET'])
def get_quiz_details(quiz_id):
    """Get details of a specific quiz"""
    try:
        user_id = session.get('user_id')
        
        if not user_id:
            return jsonify({'error': 'Not authenticated'}), 401
        
        quiz = supabase.table('teacher_quizzes')\
            .select('*')\
            .eq('id', quiz_id)\
            .execute()
        
        if not quiz.data:
            return jsonify({'error': 'Quiz not found'}), 404
        
        return jsonify({
            'success': True,
            'quiz': quiz.data[0]
        })
        
    except Exception as e:
        logger.error(f"Error fetching quiz details: {e}")
        return jsonify({'error': str(e)}), 500


# Get Quiz Results Route
@app.route('/api/quiz/<int:quiz_id>/results', methods=['GET'])
def get_quiz_results(quiz_id):
    """Get all results for a specific quiz"""
    try:
        teacher_id = session.get('user_id')
        
        if not teacher_id:
            return jsonify({'error': 'Not authenticated'}), 401
        
        # Get quiz details
        quiz = supabase.table('teacher_quizzes')\
            .select('*')\
            .eq('id', quiz_id)\
            .eq('teacher_id', teacher_id)\
            .execute()
        
        if not quiz.data:
            return jsonify({'error': 'Quiz not found or access denied'}), 404
        
        quiz_data = quiz.data[0]
        quiz_title = quiz_data.get('quiz_title', '')
        
        # Get all results for this quiz
        results = supabase.table('student_quiz_results')\
            .select('*')\
            .eq('teacher_quiz_id', quiz_id)\
            .order('submitted_date', desc=True)\
            .execute()
        
        # Get all retake allowed notifications for all quizzes
        retake_allowed_notifications = supabase.table('notifications')\
            .select('user_id, message')\
            .eq('title', 'Quiz Retake Allowed')\
            .execute()
        
        # Build a set of student IDs who have been allowed to retake THIS specific quiz
        retake_allowed_students = set()
        if retake_allowed_notifications.data and quiz_title:
            for notif in retake_allowed_notifications.data:
                # Check if this notification mentions this specific quiz
                if quiz_title in notif.get('message', ''):
                    retake_allowed_students.add(notif['user_id'])
        
        # Enhance results with total_points from quiz data and retake_allowed flag
        enhanced_results = []
        if results.data:
            for result in results.data:
                # Calculate total_points from quiz if not in result
                total_points = result.get('total_points') or quiz_data.get('total_points') or (result.get('total_items', 1) * quiz_data.get('points_per_item', 1))
                result['total_points'] = total_points
                
                # Check if teacher has sent retake allowed notification for this student/quiz
                result['retake_allowed_by_teacher'] = result['student_id'] in retake_allowed_students
                
                enhanced_results.append(result)
        
        return jsonify({
            'success': True,
            'quiz': quiz_data,
            'results': enhanced_results
        })
        
    except Exception as e:
        logger.error(f"Error fetching quiz results: {e}")
        return jsonify({'error': str(e)}), 500


# Get Non-Submitted Students for a Quiz
@app.route('/api/quiz/<int:quiz_id>/non-submitted', methods=['GET'])
def get_non_submitted_students(quiz_id):
    """Get list of students who haven't submitted the quiz"""
    try:
        teacher_id = session.get('user_id')
        
        if not teacher_id:
            return jsonify({'error': 'Not authenticated'}), 401
        
        # Get quiz details
        quiz = supabase.table('teacher_quizzes')\
            .select('grade_level, section, quiz_title')\
            .eq('id', quiz_id)\
            .eq('teacher_id', teacher_id)\
            .execute()
        
        if not quiz.data:
            return jsonify({'error': 'Quiz not found or access denied'}), 404
        
        quiz_data = quiz.data[0]
        
        # Get all students in the class
        class_students = supabase.table('user_info')\
            .select('id, first_name, last_name')\
            .eq('role', 'Student')\
            .eq('year_level', quiz_data['grade_level'])\
            .eq('section', quiz_data['section'])\
            .order('first_name', desc=False)\
            .execute()
        
        # Get students who have submitted
        submitted = supabase.table('student_quiz_results')\
            .select('student_id')\
            .eq('teacher_quiz_id', quiz_id)\
            .execute()
        
        submitted_ids = [s['student_id'] for s in submitted.data] if submitted.data else []
        
        # Filter for non-submitted students
        non_submitted = [s for s in class_students.data if s['id'] not in submitted_ids]
        
        return jsonify({
            'success': True,
            'quiz_title': quiz_data['quiz_title'],
            'submitted_count': len(submitted_ids),
            'total_students': len(class_students.data) if class_students.data else 0,
            'non_submitted_students': non_submitted
        })
        
    except Exception as e:
        logger.error(f"Error fetching non-submitted students: {e}")
        return jsonify({'error': str(e)}), 500

# Helper function to check if student was recently reminded
def was_student_reminded_recently(student_id, quiz_id, hours=24):
    """Check if a student was reminded about this quiz in the last N hours"""
    try:
        cutoff_time = datetime.now(philippines_tz) - timedelta(hours=hours)
        
        result = supabase.table('quiz_reminder_history')\
            .select('id')\
            .eq('student_id', student_id)\
            .eq('quiz_id', quiz_id)\
            .gte('reminded_at', cutoff_time.isoformat())\
            .execute()
        
        return len(result.data) > 0 if result.data else False
    except Exception as e:
        logger.warning(f"Error checking reminder history: {e}")
        return False

# Helper function to record a reminder
def record_reminder(student_id, quiz_id, teacher_id, reminder_type, message):
    """Record a reminder in the history table"""
    try:
        reminder_record = {
            'student_id': student_id,
            'quiz_id': quiz_id,
            'teacher_id': teacher_id,
            'reminder_type': reminder_type,
            'message': message,
            'reminded_at': datetime.now(philippines_tz).isoformat()
        }
        supabase.table('quiz_reminder_history').insert(reminder_record).execute()
        return True
    except Exception as e:
        logger.warning(f"Error recording reminder: {e}")
        return False

@app.route('/api/send-quiz-reminder', methods=['POST'])
def send_quiz_reminder():
    """Send a reminder notification to student about pending quiz"""
    try:
        data = request.get_json()
        teacher_id = session.get('user_id')
        
        if not teacher_id:
            return jsonify({'error': 'Not authenticated'}), 401
        
        student_id = data.get('student_id')
        quiz_id = data.get('quiz_id')
        message = data.get('message')
        remind_all = data.get('remind_all', False)
        
        # Get teacher info
        teacher = supabase.table('user_info')\
            .select('first_name', 'last_name', 'role')\
            .eq('id', teacher_id)\
            .execute()
        
        if not teacher.data:
            return jsonify({'error': 'Teacher not found'}), 404
            
        teacher_name = f"{teacher.data[0]['first_name']} {teacher.data[0]['last_name']}"
        teacher_role = teacher.data[0]['role']
        
        # Get quiz details
        quiz = supabase.table('teacher_quizzes')\
            .select('quiz_title, grade_level, section')\
            .eq('id', quiz_id)\
            .execute()
        
        if not quiz.data:
            return jsonify({'error': 'Quiz not found'}), 404
            
        quiz_data = quiz.data[0]
        
        # Get the assignment_id from teacher_class_assignments
        assignment = supabase.table('teacher_class_assignments')\
            .select('assignment_id')\
            .eq('teacher_id', teacher_id)\
            .eq('grade_level', quiz_data['grade_level'])\
            .eq('section', quiz_data['section'])\
            .execute()
        
        assignment_id = assignment.data[0]['assignment_id'] if assignment.data else None
        
        if remind_all:
            # Get all students in the class who haven't submitted
            class_students = supabase.table('user_info')\
                .select('id, first_name, last_name')\
                .eq('role', 'Student')\
                .eq('year_level', quiz_data['grade_level'])\
                .eq('section', quiz_data['section'])\
                .execute()
            
            # Get students who have already submitted
            submitted = supabase.table('student_quiz_results')\
                .select('student_id')\
                .eq('teacher_quiz_id', quiz_id)\
                .execute()
            
            submitted_ids = [s['student_id'] for s in submitted.data] if submitted.data else []
            
            # Send reminders to pending students (excluding those reminded in last 24 hours)
            notifications = []
            student_names = []  # For activity log
            skipped_count = 0
            
            for student in class_students.data:
                if student['id'] not in submitted_ids:
                    # Check if student was already reminded in the last 24 hours
                    if was_student_reminded_recently(student['id'], quiz_id, hours=24):
                        skipped_count += 1
                        continue  # Skip this student - they already got a reminder today
                    
                    student_name = f"{student['first_name']} {student['last_name']}"
                    student_names.append(student_name)
                    
                    notification = {
                        'user_id': student['id'],
                        'sender_id': teacher_id,
                        'title': 'Quiz Reminder',
                        'message': message or f"Please complete your quiz '{quiz_data['quiz_title']}'. It's still available for submission.",
                        'assignment_id': assignment_id,
                        'notif_type': 'Class Assignment',
                        'status': 'Unread'
                    }
                    notifications.append(notification)
                    
                    # Record the reminder
                    record_reminder(student['id'], quiz_id, teacher_id, 'bulk', 
                                  message or f"Please complete your quiz '{quiz_data['quiz_title']}'. It's still available for submission.")
            
            if notifications:
                supabase.table('notifications').insert(notifications).execute()
                
                # Log the activity with ALL required fields
                student_list = ', '.join(student_names[:5])  # First 5 students
                if len(student_names) > 5:
                    student_list += f" and {len(student_names) - 5} more"
                
                skip_info = f" ({skipped_count} already reminded today)" if skipped_count > 0 else ""
                
                activity_log = {
                    'user_id': teacher_id,
                    'user_role': teacher_role,
                    'action': 'Sent bulk quiz reminder',
                    'activity': 'Sent quiz reminder',
                    'description': f'Sent reminders to {len(notifications)} students for quiz "{quiz_data["quiz_title"]}" (Grade {quiz_data["grade_level"]} - {quiz_data["section"]}){skip_info}',
                    'details': f'Students: {student_list}'
                }
                supabase.table('admin_activity_log').insert(activity_log).execute()
            
            return jsonify({
                'success': True,
                'message': f'Reminders sent to {len(notifications)} students' + (f' ({skipped_count} already reminded today)' if skipped_count > 0 else ''),
                'count': len(notifications),
                'skipped_count': skipped_count
            })
            
        else:
            # Check if student was already reminded in the last 24 hours
            if was_student_reminded_recently(student_id, quiz_id, hours=24):
                return jsonify({
                    'success': False,
                    'error': 'This student was already reminded about this quiz in the last 24 hours. Please try again later.',
                    'cooldown_active': True
                }), 429  # 429 = Too Many Requests
            
            # Get student info for the log
            student = supabase.table('user_info')\
                .select('first_name, last_name')\
                .eq('id', student_id)\
                .execute()
            
            student_name = f"{student.data[0]['first_name']} {student.data[0]['last_name']}" if student.data else f"Student {student_id}"
            
            # Send reminder to single student
            notification = {
                'user_id': student_id,
                'sender_id': teacher_id,
                'title': 'Quiz Reminder',
                'message': message or f"Please complete your quiz '{quiz_data['quiz_title']}'. It's still available for submission.",
                'assignment_id': assignment_id,
                'notif_type': 'Class Assignment',
                'status': 'Unread'
            }
            
            result = supabase.table('notifications').insert(notification).execute()
            
            # Record the reminder
            record_reminder(student_id, quiz_id, teacher_id, 'single',
                          message or f"Please complete your quiz '{quiz_data['quiz_title']}'. It's still available for submission.")
            
            # Log the activity with ALL required fields
            activity_log = {
                'user_id': teacher_id,
                'user_role': teacher_role,
                'action': 'Sent quiz reminder',
                'activity': 'Sent quiz reminder',
                'description': f'Sent reminder to student {student_name} for quiz "{quiz_data["quiz_title"]}" (Grade {quiz_data["grade_level"]} - {quiz_data["section"]})',
                'details': f'Student ID: {student_id}, Quiz ID: {quiz_id}'
            }
            supabase.table('admin_activity_log').insert(activity_log).execute()
            
            return jsonify({
                'success': True,
                'message': f'Reminder sent successfully to {student_name}',
                'notification_id': result.data[0]['notif_id']
            })
        
    except Exception as e:
        logger.error(f"Error sending reminder: {e}")
        return jsonify({'error': str(e)}), 500

# Get Class Roster with Submission Status
@app.route('/api/quiz/<int:quiz_id>/class-roster', methods=['GET'])
def get_class_roster_with_status(quiz_id):
    """Get all students in a class with their submission status for a quiz"""
    try:
        teacher_id = session.get('user_id')
        
        if not teacher_id:
            return jsonify({'error': 'Not authenticated'}), 401
        
        # Get quiz details to know which class
        quiz = supabase.table('teacher_quizzes')\
            .select('grade_level, section')\
            .eq('id', quiz_id)\
            .execute()
        
        if not quiz.data:
            return jsonify({'error': 'Quiz not found'}), 404
            
        quiz_data = quiz.data[0]
        
        # Get all students in the class
        students = supabase.table('user_info')\
            .select('id, first_name, last_name')\
            .eq('role', 'Student')\
            .eq('year_level', quiz_data['grade_level'])\
            .eq('section', quiz_data['section'])\
            .order('first_name')\
            .execute()
        
        # Get submission status
        submissions = supabase.table('student_quiz_results')\
            .select('student_id, submitted_date')\
            .eq('teacher_quiz_id', quiz_id)\
            .execute()
        
        submitted_ids = {s['student_id']: s['submitted_date'] for s in submissions.data}
        
        # Combine data
        roster = []
        for student in students.data:
            roster.append({
                'student_id': student['id'],
                'student_name': f"{student['first_name']} {student['last_name']}",
                'submitted': student['id'] in submitted_ids,
                'submitted_date': submitted_ids.get(student['id'])
            })
        
        return jsonify({
            'success': True,
            'roster': roster,
            'total_students': len(roster),
            'submitted_count': len(submitted_ids),
            'pending_count': len(roster) - len(submitted_ids)
        })
        
    except Exception as e:
        logger.error(f"Error fetching class roster: {e}")
        return jsonify({'error': str(e)}), 500

# Get Student Notifications
@app.route('/api/student/notifications', methods=['GET'])
def get_student_notifications():
    """Get all notifications for the authenticated student"""
    try:
        student_id = session.get('user_id')
        
        if not student_id:
            return jsonify({'error': 'Not authenticated'}), 401
        
        # Get user role to ensure it's a student
        user = supabase.table('user_info')\
            .select('role')\
            .eq('id', student_id)\
            .execute()
        
        if not user.data or user.data[0]['role'] != 'Student':
            return jsonify({'error': 'Access denied'}), 403
        
        # Get notifications
        notifications = supabase.table('notifications')\
            .select('*')\
            .eq('user_id', student_id)\
            .order('created_at', desc=True)\
            .execute()
        
        # Get unread count
        unread_count = supabase.table('notifications')\
            .select('*', count='exact')\
            .eq('user_id', student_id)\
            .eq('status', 'Unread')\
            .execute()
        
        return jsonify({
            'success': True,
            'notifications': notifications.data,
            'unread_count': unread_count.count
        })
        
    except Exception as e:
        logger.error(f"Error fetching notifications: {e}")
        return jsonify({'error': str(e)}), 500

# Mark Notification as Read
@app.route('/api/notifications/<int:notif_id>/read', methods=['POST'])
def mark_notification_read(notif_id):
    """Mark a notification as read"""
    try:
        user_id = session.get('user_id')
        
        if not user_id:
            return jsonify({'error': 'Not authenticated'}), 401
        
        # Verify notification belongs to user
        notification = supabase.table('notifications')\
            .select('user_id')\
            .eq('notif_id', notif_id)\
            .execute()
        
        if not notification.data:
            return jsonify({'error': 'Notification not found'}), 404
            
        if notification.data[0]['user_id'] != user_id:
            return jsonify({'error': 'Access denied'}), 403
        
        # Update status
        result = supabase.table('notifications')\
            .update({'status': 'Read'})\
            .eq('notif_id', notif_id)\
            .execute()
        
        return jsonify({
            'success': True,
            'message': 'Notification marked as read'
        })
        
    except Exception as e:
        logger.error(f"Error marking notification read: {e}")
        return jsonify({'error': str(e)}), 500

# Mark All Notifications as Read
@app.route('/api/notifications/read-all', methods=['POST'])
def mark_all_notifications_read():
    """Mark all notifications for a user as read"""
    try:
        user_id = session.get('user_id')
        
        if not user_id:
            return jsonify({'error': 'Not authenticated'}), 401
        
        # Update all unread notifications
        result = supabase.table('notifications')\
            .update({'status': 'Read'})\
            .eq('user_id', user_id)\
            .eq('status', 'Unread')\
            .execute()
        
        return jsonify({
            'success': True,
            'message': 'All notifications marked as read'
        })
        
    except Exception as e:
        logger.error(f"Error marking all notifications read: {e}")
        return jsonify({'error': str(e)}), 500

# Delete Notification
@app.route('/api/notifications/<int:notif_id>', methods=['DELETE'])
def delete_notification(notif_id):
    """Delete a notification"""
    try:
        user_id = session.get('user_id')
        
        if not user_id:
            return jsonify({'error': 'Not authenticated'}), 401
        
        # Verify notification belongs to user
        notification = supabase.table('notifications')\
            .select('user_id')\
            .eq('notif_id', notif_id)\
            .execute()
        
        if not notification.data:
            return jsonify({'error': 'Notification not found'}), 404
            
        if notification.data[0]['user_id'] != user_id:
            return jsonify({'error': 'Access denied'}), 403
        
        # Delete notification
        result = supabase.table('notifications')\
            .delete()\
            .eq('notif_id', notif_id)\
            .execute()
        
        return jsonify({
            'success': True,
            'message': 'Notification deleted'
        })
        
    except Exception as e:
        logger.error(f"Error deleting notification: {e}")
        return jsonify({'error': str(e)}), 500

# Get Quiz Statistics for Teacher Dashboard
@app.route('/api/teacher/quiz-statistics', methods=['GET'])
def get_quiz_statistics():
    """Get overall quiz statistics for teacher dashboard"""
    try:
        teacher_id = session.get('user_id')
        
        if not teacher_id:
            return jsonify({'error': 'Not authenticated'}), 401
        
        # Get all quizzes created by teacher
        quizzes = supabase.table('teacher_quizzes')\
            .select('id, quiz_title, total_items, created_at')\
            .eq('teacher_id', teacher_id)\
            .order('created_at', desc=True)\
            .execute()
        
        # Get submission statistics for each quiz
        statistics = []
        for quiz in quizzes.data:
            # Get submissions
            submissions = supabase.table('student_quiz_results')\
                .select('score, student_id')\
                .eq('teacher_quiz_id', quiz['id'])\
                .execute()
            
            # Get class size (students in that grade/section)
            # This would require getting the quiz's grade/section first
            quiz_details = supabase.table('teacher_quizzes')\
                .select('grade_level, section')\
                .eq('id', quiz['id'])\
                .execute()
            
            if quiz_details.data:
                class_students = supabase.table('user_info')\
                    .select('id', count='exact')\
                    .eq('role', 'Student')\
                    .eq('year_level', quiz_details.data[0]['grade_level'])\
                    .eq('section', quiz_details.data[0]['section'])\
                    .execute()
                
                total_students = class_students.count if hasattr(class_students, 'count') else 0
            else:
                total_students = 0
            
            # Calculate average score
            if submissions.data:
                avg_score = sum(s['score'] for s in submissions.data) / len(submissions.data)
                avg_percentage = (avg_score / quiz['total_items']) * 100 if quiz['total_items'] > 0 else 0
            else:
                avg_percentage = 0
            
            statistics.append({
                'quiz_id': quiz['id'],
                'quiz_title': quiz['quiz_title'],
                'total_items': quiz['total_items'],
                'submissions': len(submissions.data),
                'total_students': total_students,
                'completion_rate': (len(submissions.data) / total_students * 100) if total_students > 0 else 0,
                'average_percentage': round(avg_percentage, 2),
                'created_at': quiz['created_at']
            })
        
        # Calculate overall statistics
        total_quizzes = len(quizzes.data)
        total_submissions = sum(s['submissions'] for s in statistics)
        overall_completion = (total_submissions / (total_quizzes * 40)) * 100 if total_quizzes > 0 else 0  # Assuming avg 40 students per class
        
        return jsonify({
            'success': True,
            'statistics': statistics,
            'overall': {
                'total_quizzes': total_quizzes,
                'total_submissions': total_submissions,
                'overall_completion_rate': round(overall_completion, 2)
            }
        })
        
    except Exception as e:
        logger.error(f"Error fetching quiz statistics: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/generate-question-by-filters', methods=['POST'])
def generate_question_by_filters():
    """Generate questions based on filters without uploading a file"""
    try:
        data = request.get_json()
        topic = data.get('topic')
        grade_level = data.get('grade_level')
        quarter = data.get('quarter')
        count = data.get('count', 1)
        
        if quiz_generator is None:
            return jsonify({'error': 'Quiz generator not available'}), 500
        
        # Generate questions based on filters
        questions = quiz_generator.find_questions_by_filters(
            topic=topic,
            grade_level=grade_level,
            quarter=quarter,
            limit=count
        )
        
        return jsonify({
            'success': True,
            'questions': questions,
            'count': len(questions)
        })
        
    except Exception as e:
        logger.error(f"Error generating questions by filters: {e}")
        return jsonify({'error': str(e)}), 500

#############################STUDENT_HOME.HTML#############################
@app.route('/api/student-recent-activities', methods=['GET'])
def api_student_recent_activities():
    if 'user_id' not in session or session.get('role') != 'Student':
        return jsonify({'success': False, 'activities': []}), 401
    
    user_id = session['user_id']
    recent_activities = []

    def get_teacher_name_with_prefix(teacher_data):
        """Get teacher name with appropriate gender prefix"""
        prefix = 'Mr.'
        if teacher_data.get('gender', '').lower() == 'female':
            prefix = 'Mrs.'
        elif teacher_data.get('gender', '').lower() == 'other':
            prefix = 'Mx.'
        
        return f"{prefix} {teacher_data['first_name']} {teacher_data['last_name']}".strip()

    try:
        # --- Recent completed/denied tasks ---
        tasks_resp = safe_execute(
            supabase.table('task_assignments').select(
                'task, points, status, completed_at, due_date, awarded_by'
            ).eq('student_id', user_id).in_('status', ['Completed', 'Denied']).order('completed_at', desc=True).limit(5)
        )
        if tasks_resp.data:
            # Get teacher names for tasks
            task_teacher_ids = list(set([t['awarded_by'] for t in tasks_resp.data if t.get('awarded_by')]))
            teachers_map = {}
            
            if task_teacher_ids:
                teachers_resp = safe_execute(
                    supabase.table('user_info').select('id, first_name, last_name, gender').in_('id', task_teacher_ids)
                )
                if teachers_resp.data:
                    for teacher in teachers_resp.data:
                        teacher_name = get_teacher_name_with_prefix(teacher)
                        teachers_map[teacher['id']] = teacher_name
            
            for t in tasks_resp.data:
                teacher_name = "Teacher"
                if t.get('awarded_by') and t['awarded_by'] in teachers_map:
                    teacher_name = teachers_map[t['awarded_by']]
                
                recent_activities.append({
                    'type': 'task',
                    'title': t['task'],
                    'points': t['points'],
                    'status': f'Awarded by {teacher_name}',
                    'time': format_date(t.get('completed_at', t.get('due_date', ''))),
                })

        # --- Recent reward redemptions ---
        rewards_resp = safe_execute(
            supabase.table('reward_redemptions').select(
                'reward_id, points_deducted, processed_at, rewards(reward_name)'
            ).eq('student_id', user_id).order('processed_at', desc=True).limit(5)
        )
        if rewards_resp.data:
            for r in rewards_resp.data:
                reward_name = ''
                if r.get('rewards') and r['rewards'].get('reward_name'):
                    reward_name = r['rewards']['reward_name']
                recent_activities.append({
                    'type': 'reward',
                    'title': f"Redeemed: {reward_name}",
                    'points': -abs(r['points_deducted']),
                    'status': 'Redeemed',
                    'time': format_date(r.get('processed_at', '')),
                })

        # --- Star awards with teacher names ---
        stars_resp = safe_execute(
            supabase.table('points').select(
                'points, received_at, note, point_category, teacher_id'
            ).eq('student_id', user_id).eq('status', 'approved').order('received_at', desc=True).limit(5)
        )
        if stars_resp.data:
            # Get teacher names for star awards
            star_teacher_ids = list(set([s['teacher_id'] for s in stars_resp.data if s.get('teacher_id')]))
            star_teachers_map = {}
            
            if star_teacher_ids:
                star_teachers_resp = safe_execute(
                    supabase.table('user_info').select('id, first_name, last_name, gender').in_('id', star_teacher_ids)
                )
                if star_teachers_resp.data:
                    for teacher in star_teachers_resp.data:
                        teacher_name = get_teacher_name_with_prefix(teacher)
                        star_teachers_map[teacher['id']] = teacher_name
            
            for s in stars_resp.data:
                # Prefer note, then point_category, then fallback
                if s.get('note'):
                    title = s['note']
                elif s.get('point_category'):
                    title = f"Awarded for {s['point_category'].capitalize()}"
                else:
                    title = "Teacher Award"
                
                # Get teacher name who gave the award
                teacher_name = "Teacher"
                if s.get('teacher_id') and s['teacher_id'] in star_teachers_map:
                    teacher_name = star_teachers_map[s['teacher_id']]
                
                recent_activities.append({
                    'type': 'star',
                    'title': title,
                    'points': s['points'],
                    'status': f'Awarded by {teacher_name}',
                    'time': format_date(s.get('received_at', '')),
                })

        # --- Recent milestone claims ---
        milestone_claims_resp = safe_execute(
            supabase.table('milestone_claims')
            .select('milestone_type, milestone, points_awarded, claimed_at')
            .eq('student_id', user_id)
            .order('claimed_at', desc=True)
            .limit(5)
        )
        if milestone_claims_resp.data:
            for m in milestone_claims_resp.data:
                recent_activities.append({
                    'type': 'milestone',
                    'title': f"Claimed {m['milestone_type'].capitalize()} Milestone: {m['milestone']}",
                    'points': m.get('points_awarded', 0),
                    'status': 'Claimed',
                    'time': format_date(m.get('claimed_at', '')),
                })

        # --- Streak milestones ---
        streak_milestone_notifs = safe_execute(
            supabase.table('notifications')
            .select('title, message, created_at')
            .eq('user_id', user_id)
            .eq('title', 'Streak Milestone Unlocked!')
            .order('created_at', desc=True)
        )

        if streak_milestone_notifs.data:
            for notif in streak_milestone_notifs.data:
                import re
                milestone_match = re.search(r'the (\d+) streak milestone', notif.get('message', ''))
                milestone_num = milestone_match.group(1) if milestone_match else '?'
                recent_activities.append({
                    'type': 'streak',
                    'title': f"Streak Milestone: {milestone_num} days!",
                    'points': 0,
                    'status': 'Milestone',
                    'time': format_date(notif.get('created_at', '')),
                })

        # Helper to parse the real datetime for sorting
        from datetime import datetime
        def parse_sort_time(x):
            # Try to parse the formatted date back to datetime for sorting
            try:
                # Example format: 'Oct 05, 2025 • 04:52 PM'
                return datetime.strptime(x.get('time', ''), '%b %d, %Y • %I:%M %p')
            except Exception:
                return datetime.min

        recent_activities = [a for a in recent_activities if a.get('time')]
        recent_activities.sort(key=parse_sort_time, reverse=True)
        final_activities = recent_activities[:5]

        return jsonify({'success': True, 'activities': final_activities})
        
    except Exception as e:
        return jsonify({'success': False, 'activities': [], 'message': str(e)}), 500

@app.route('/api/student-stats/<int:student_id>')
def api_student_stats(student_id):
    try:
        # Fetch total_points and streak from user_info
        user_result = supabase.table('user_info').select('total_points, streak').eq('id', student_id).execute()
        if not user_result.data:
            return jsonify({'success': False, 'message': 'Student not found'}), 404
        user = user_result.data[0]
        total_points = int(user.get('total_points') or 0)
        streak = int(user.get('streak') or 0)

        # Calculate today's points
        from datetime import datetime, timezone
        now = datetime.now(timezone.utc)
        today_str = now.strftime('%Y-%m-%d')
        points_today = 0
        points_result = supabase.table('points').select('points, received_at').eq('student_id', student_id).eq('status', 'approved').execute()
        if points_result.data:
            for row in points_result.data:
                received_at = row.get('received_at', '')
                if received_at and received_at.startswith(today_str):
                    points_today += int(row.get('points', 0))

        # Streak message
        if streak >= 7:
            streak_msg = "🔥 Incredible streak!"
        elif streak >= 3:
            streak_msg = "🎉 Amazing!"
        elif streak > 0:
            streak_msg = "Keep it up!"
        else:
            streak_msg = "Start your streak!"

        return jsonify({
            'success': True,
            'total_points': total_points,
            'points_today': points_today,
            'streak': streak,
            'streak_msg': streak_msg
        })
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500

# --- ACHIEVEMENTS PAGE ROUTE ---
@app.route('/achievements', methods=['GET'])
def get_achievements():
    from datetime import datetime, timedelta

    user_id = request.args.get('user_id')
    if not user_id:
        return jsonify({'error': 'Missing user_id'}), 400

    # Get streak and points (all-time)
    user_resp = safe_execute(supabase.table('user_info').select('streak, total_points').eq('id', user_id))
    user = user_resp.data[0] if user_resp.data else {}

    # --- ALL-TIME COMPLETED TASKS ---
    all_tasks_resp = safe_execute(
        supabase.table('task_assignments')
        .select('task_id')
        .eq('student_id', user_id)
        .eq('status', 'Completed')
    )
    all_completed_tasks = len(all_tasks_resp.data) if all_tasks_resp.data else 0

    # --- ALL-TIME REWARDS REDEEMED ---
    all_rewards_resp = safe_execute(
        supabase.table('reward_redemptions')
        .select('redemption_id')
        .eq('student_id', user_id)
    )
    all_redeemed_rewards = len(all_rewards_resp.data) if all_rewards_resp.data else 0

    # --- WEEKLY RANGE ---
    now = datetime.now()
    start_of_week = now - timedelta(days=now.weekday())
    end_of_week = start_of_week + timedelta(days=6)
    start_str = start_of_week.date().isoformat()
    end_str = end_of_week.date().isoformat()

    # --- WEEKLY COMPLETED TASKS ---
    weekly_tasks_resp = safe_execute(
        supabase.table('task_assignments')
        .select('task_id, completed_at')
        .eq('student_id', user_id)
        .eq('status', 'Completed')
        .gte('completed_at', start_str)
        .lte('completed_at', end_str)
    )
    weekly_completed_tasks = len(weekly_tasks_resp.data) if weekly_tasks_resp.data else 0

    # --- WEEKLY REWARDS REDEEMED ---
    weekly_rewards_resp = safe_execute(
        supabase.table('reward_redemptions')
        .select('redemption_id, processed_at')
        .eq('student_id', user_id)
        .gte('processed_at', start_str)
        .lte('processed_at', end_str)
    )
    weekly_redeemed_rewards = len(weekly_rewards_resp.data) if weekly_rewards_resp.data else 0

    # --- CLAIMED MILESTONES ---
    claimed_resp = safe_execute(
        supabase.table('milestone_claims')
        .select('milestone_type, milestone')
        .eq('student_id', user_id)
    )
    claimed_set = set()
    if claimed_resp.data:
        for c in claimed_resp.data:
            claimed_set.add((c['milestone_type'], c['milestone']))

    all_achievements = []

    # Helper for notification insert
    def insert_milestone_unlocked_notif(milestone_type, milestone):
        notif_exists = safe_execute(
            supabase.table('notifications')
            .select('notif_id')
            .eq('user_id', user_id)
            .eq('title', f'{milestone_type.capitalize()} Milestone Unlocked!')
            .eq('message', f'You unlocked the {milestone} {milestone_type} milestone! Claim your bonus points.')
        )
        if not notif_exists.data or len(notif_exists.data) == 0:
            safe_execute(supabase.table('notifications').insert({
                'user_id': user_id,
                'sender_id': user_id,
                'title': f'{milestone_type.capitalize()} Milestone Unlocked!',
                'message': f'You unlocked the {milestone} {milestone_type} milestone! Claim your bonus points.',
                'notif_type': milestone_type.capitalize(),
                'status': 'Unread',
            }))

    # Streak milestones (all-time)
    for milestone in [7, 15, 22, 29, 36, 43, 50]:
        is_unlocked = user.get('streak', 0) >= milestone
        is_claimed = ('streak', milestone) in claimed_set
        if is_unlocked and not is_claimed:
            insert_milestone_unlocked_notif('streak', milestone)
        all_achievements.append({
            'type': 'streak',
            'title': f'{milestone}-Day Streak',
            'emoji': '🔥',
            'description': f'Reach a {milestone}-day streak!',
            'progress': user.get('streak', 0),
            'total': milestone,
            'isUnlocked': is_unlocked,
            'isClaimed': is_claimed,
            'unlocked_at': None
        })

    # Points milestones (all-time)
    for milestone in [100, 200, 300, 500]:
        is_unlocked = user.get('total_points', 0) >= milestone
        is_claimed = ('points', milestone) in claimed_set
        if is_unlocked and not is_claimed:
            insert_milestone_unlocked_notif('points', milestone)
        all_achievements.append({
            'type': 'points',
            'title': f'{milestone} Points',
            'emoji': '🏆',
            'description': f'Earn {milestone} points!',
            'progress': user.get('total_points', 0),
            'total': milestone,
            'isUnlocked': is_unlocked,
            'isClaimed': is_claimed,
            'unlocked_at': None
        })

    # Task milestones (all-time)
    for milestone in [10, 25, 50, 100]:
        is_unlocked = all_completed_tasks >= milestone
        is_claimed = ('task', milestone) in claimed_set
        if is_unlocked and not is_claimed:
            insert_milestone_unlocked_notif('task', milestone)
        all_achievements.append({
            'type': 'task',
            'title': f'Completed {milestone} Activities',
            'emoji': '✅',
            'description': f'Complete {milestone} activities (all-time)!',
            'progress': all_completed_tasks,
            'total': milestone,
            'isUnlocked': is_unlocked,
            'isClaimed': is_claimed,
            'unlocked_at': None
        })

    # Reward milestones (all-time)
    for milestone in [1, 5, 10]:
        is_unlocked = all_redeemed_rewards >= milestone
        is_claimed = ('reward', milestone) in claimed_set
        if is_unlocked and not is_claimed:
            insert_milestone_unlocked_notif('reward', milestone)
        all_achievements.append({
            'type': 'reward',
            'title': f'Redeemed {milestone} Reward{"s" if milestone > 1 else ""}',
            'emoji': '🎁',
            'description': f'Redeem {milestone} reward{"s" if milestone > 1 else ""} (all-time)!',
            'progress': all_redeemed_rewards,
            'total': milestone,
            'isUnlocked': is_unlocked,
            'isClaimed': is_claimed,
            'unlocked_at': None
        })

    # Weekly milestone: 1 task
    is_unlocked = weekly_completed_tasks >= 1
    is_claimed = False  # or check if you want to track weekly claims
    if is_unlocked and not is_claimed:
        insert_milestone_unlocked_notif('task', 1)
    all_achievements.append({
        'type': 'task',
        'title': 'Completed 1 Activity (This Week)',
        'emoji': '📅',
        'description': 'Complete at least 1 activity this week!',
        'progress': weekly_completed_tasks,
        'total': 1,
        'isUnlocked': is_unlocked,
        'isClaimed': False,
        'unlocked_at': None
    })

    # Weekly milestone: 1 reward
    is_unlocked = weekly_redeemed_rewards >= 1
    is_claimed = False  # or check if you want to track weekly claims
    if is_unlocked and not is_claimed:
        insert_milestone_unlocked_notif('reward', 1)
    all_achievements.append({
        'type': 'reward',
        'title': 'Redeemed 1 Reward (This Week)',
        'emoji': '🗓️',
        'description': 'Redeem at least 1 reward this week!',
        'progress': weekly_redeemed_rewards,
        'total': 1,
        'isUnlocked': is_unlocked,
        'isClaimed': False,
        'unlocked_at': None
    })

    return jsonify({'success': True, 'achievements': all_achievements}), 200

@app.route('/claim_milestone', methods=['POST'])
def claim_milestone():
    try:
        data = request.get_json()
        user_id = data.get('user_id')
        milestone_type = data.get('milestone_type')
        milestone = data.get('milestone')

        if not all([user_id, milestone_type, milestone]):
            return jsonify({'success': False, 'error': 'Missing required fields'}), 400

        try:
            milestone_int = int(milestone)
        except (ValueError, TypeError):
            return jsonify({'success': False, 'error': 'Invalid milestone value'}), 400

        # Calculate points based on milestone type
        points_to_add = 0
        if milestone_type == 'streak':
            points_to_add = milestone_int * 5
        elif milestone_type == 'points':
            points_to_add = max(10, int(milestone_int * 0.1))
        elif milestone_type == 'reward':
            points_to_add = milestone_int * 10
        elif milestone_type == 'task':
            points_to_add = milestone_int * 2

        points_to_add = max(10, points_to_add)

        # Check if already claimed
        existing_claim = safe_execute(
            supabase.table('milestone_claims')
            .select('claim_id')
            .eq('student_id', user_id)
            .eq('milestone_type', milestone_type)
            .eq('milestone', milestone_int)
        )

        if existing_claim.data and len(existing_claim.data) > 0:
            return jsonify({'success': False, 'error': 'Already claimed'}), 400

        # Record the claim
        claim_resp = safe_execute(
            supabase.table('milestone_claims')
            .insert({
                'student_id': user_id,
                'milestone_type': milestone_type,
                'milestone': milestone_int,
                'claimed_at': datetime.now(timezone.utc).isoformat(),
                'points_awarded': points_to_add
            })
        )

        if not claim_resp.data:
            return jsonify({'success': False, 'error': 'Failed to record claim'}), 500

        # Update user's total points
        user_resp = safe_execute(
            supabase.table('user_info')
            .select('total_points', 'first_name', 'last_name')
            .eq('id', user_id)
        )

        # --- ACTIVITY LOG: Add entry for milestone claim ---
        if user_resp.data and len(user_resp.data) > 0:
            current_points = int(user_resp.data[0].get('total_points', 0))
            new_points = current_points + points_to_add
            first_name = user_resp.data[0].get('first_name', '')
            last_name = user_resp.data[0].get('last_name', '')
            student_name = f"{first_name} {last_name}".strip()

            # Update points
            safe_execute(
                supabase.table('user_info')
                .update({'total_points': new_points})
                .eq('id', user_id)
            )

            # Add to activity log
            safe_execute(
                supabase.table('admin_activity_log').insert({
                    'user_id': user_id,
                    'user_role': 'Student',
                    'action': 'Claim Milestone',
                    'activity': 'Milestone Achievement',
                    'description': f"{student_name} has redeemed their milestone achievement",
                    'details': f"Milestone: {milestone_type} - {milestone_int}, Points Awarded: {points_to_add}"
                })
            )
        else:
            # Fallback if user info not found
            new_points = None

        # Add to activity feed
        safe_execute(
            supabase.table('student_activity')
            .insert({
                'student_id': user_id,
                'activity_type': 'milestone',
                'description': f'Claimed {milestone_type} milestone: {milestone_int}',
                'points_earned': points_to_add,
                'timestamp': datetime.now().isoformat()
            })
        )

        # --- Always return a success JSON response ---
        return jsonify({'success': True, 'points_awarded': points_to_add}), 200

    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

#############################STUDENT_ACTIVITIES.HTML#############################
@app.route('/api/student-activities', methods=['GET'])
def api_student_activities():
    if 'user_id' not in session or session.get('role') != 'Student':
        return jsonify({'success': False, 'activities': []}), 401

    student_id = session['user_id']
    try:
        # First, get the student's grade level and section
        student_result = supabase.table('user_info') \
            .select('year_level, section') \
            .eq('id', student_id) \
            .execute()
        
        if not student_result.data:
            return jsonify({'success': False, 'activities': [], 'message': 'Student not found'}), 404
        
        student_data = student_result.data[0]
        year_level = student_data.get('year_level')
        section = student_data.get('section')
        
        # Debug: Log student info
        print(f"Student year_level: {year_level}, section: {section}")
        
        # Get teachers assigned to this student's grade level and section
        # NOTE: Only fetch ACTIVE teachers for the dropdown (legacy view pattern)
        assigned_teachers_result = supabase.table('teacher_class_assignments') \
            .select('teacher_id, subject') \
            .eq('grade_level', year_level) \
            .eq('section', section) \
            .eq('status', 'active') \
            .execute()
        
        # Debug: Log assigned teachers query result
        print(f"Assigned teachers result (active only): {assigned_teachers_result.data}")
        
        assigned_teacher_ids = set()
        assigned_teachers_map = {}
        
        if assigned_teachers_result.data:
            for assignment in assigned_teachers_result.data:
                teacher_id = assignment.get('teacher_id')
                if teacher_id:  # Only add if teacher_id exists
                    assigned_teacher_ids.add(teacher_id)
                    assigned_teachers_map[teacher_id] = {
                        'subject': assignment.get('subject', '')
                    }
        
        print(f"Assigned teacher IDs: {assigned_teacher_ids}")

        # Fetch all tasks assigned to this student
        task_result = supabase.table('task_assignments') \
            .select('task, description, status, due_date, points, priority, image_urls, task_id, teacher_id, attachments, links') \
            .eq('student_id', student_id) \
            .execute()
        activities = []
        
        teacher_ids = set()  # Use set to avoid duplicates
        now_date = datetime.now(philippines_tz).date()
        
        if task_result.data:
            for row in task_result.data:
                task_status = row.get('status')
                due_date_str = row.get('due_date')
                
                # Check if task is past due and still in "Assigned" or "Submitted" status
                if due_date_str and task_status in ['Assigned', 'Submitted']:
                    try:
                        # Parse due date - handle both date and datetime formats
                        due_date_clean = str(due_date_str).replace('Z', '+00:00').split('T')[0]
                        due_date = datetime.strptime(due_date_clean, '%Y-%m-%d').date()
                        
                        print(f"Task {row['task_id']}: due_date={due_date}, now={now_date}, is_past={due_date < now_date}")
                        
                        if due_date < now_date:
                            # Task is past due - auto deny it
                            print(f"Auto-denying task {row['task_id']} - due date {due_date} is before {now_date}")
                            task_status = 'Denied'
                            # Update in database
                            supabase.table('task_assignments') \
                                .update({'status': 'Denied'}) \
                                .eq('task_id', row['task_id']) \
                                .eq('student_id', student_id) \
                                .execute()
                    except Exception as e:
                        print(f"Error checking due date for task {row.get('task_id')}: {str(e)}")
                
                # Fetch submitted files for this task
                submitted_files = []
                try:
                    file_submissions = supabase.table('task_file_submissions') \
                        .select('original_filename, filename, file_size, mime_type, file_url, file_type, uploaded_at') \
                        .eq('task_id', row['task_id']) \
                        .eq('student_id', student_id) \
                        .execute()
                    
                    if file_submissions.data:
                        submitted_files = file_submissions.data
                except Exception as e:
                    print(f"Error fetching submitted files for task {row['task_id']}: {str(e)}")
                
                activities.append({
                    'type': 'task',
                    'task': row.get('task'),
                    'description': row.get('description'),
                    'status': task_status,
                    'due_date': row.get('due_date'),
                    'points': row.get('points'),
                    'priority': row.get('priority'),
                    'image_urls': row.get('image_urls'),
                    'task_id': row.get('task_id'),
                    'teacher_id': row.get('teacher_id'),
                    'attachments': row.get('attachments'),
                    'links': row.get('links'),
                    'submitted_files': submitted_files,
                })
                # Collect teacher_id from tasks
                if row.get('teacher_id'):
                    teacher_ids.add(row['teacher_id'])

        # Fetch all awarded points for this student
        points_result = supabase.table('points') \
            .select('point_id, points, point_category, note, status, received_at, teacher_id') \
            .eq('student_id', student_id) \
            .order('received_at', desc=True) \
            .execute()
            
        if points_result.data:
            for row in points_result.data:
                if row.get('teacher_id'):
                    teacher_ids.add(row['teacher_id'])
                activities.append({
                    'type': 'award',
                    'task': f"Award: {row.get('point_category', 'Points')}",
                    'description': row.get('note', ''),
                    'status': 'Awarded',
                    'due_date': row.get('received_at'),
                    'points': row.get('points'),
                    'priority': None,
                    'image_urls': None,
                    'point_id': row.get('point_id'),
                    'teacher_id': row.get('teacher_id'),
                })

        # Combine teacher IDs from assignments and activities
        all_teacher_ids = assigned_teacher_ids.union(teacher_ids)
        
        print(f"All teacher IDs to fetch: {all_teacher_ids}")
        
        # Get teacher info for all relevant teachers
        teacher_map = {}
        if all_teacher_ids:
            # Convert set to list for query
            teacher_ids_list = list(all_teacher_ids)
            
            # Fetch teacher basic info
            teacher_result = supabase.table('user_info') \
                .select('id, first_name, last_name, gender, subject') \
                .in_('id', teacher_ids_list) \
                .execute()
            
            print(f"Teacher info fetched: {len(teacher_result.data)} teachers")
            
            # Fetch profile pictures
            pics_map = {}
            pics_result = supabase.table('profile_pictures') \
                .select('user_id, file_path') \
                .in_('user_id', teacher_ids_list) \
                .execute()
            
            if pics_result.data:
                for p in pics_result.data:
                    pics_map[p['user_id']] = p['file_path']
            
            # Create teacher map
            if teacher_result.data:
                for t in teacher_result.data:
                    prefix = 'Mr.' if (t.get('gender', '').lower() == 'male') else 'Ms.'
                    teacher_name = f"{prefix} {t.get('last_name', 'Teacher')}"
                    
                    # Use assigned subject if available, otherwise fall back to teacher's subject
                    assigned_subject = assigned_teachers_map.get(t['id'], {}).get('subject')
                    subject = assigned_subject or t.get('subject', 'General')
                    
                    teacher_map[t['id']] = {
                        'name': teacher_name,
                        'subject': subject,
                        'avatar': pics_map.get(t['id'], '/static/image/default-avatar.png'),
                        'is_assigned': t['id'] in assigned_teacher_ids
                    }

        # Add teacher info to all activities
        for activity in activities:
            teacher_id = activity.get('teacher_id')
            if teacher_id and teacher_id in teacher_map:
                teacher_info = teacher_map[teacher_id]
                activity['teacher_name'] = teacher_info['name']
                activity['teacher_subject'] = teacher_info['subject']
                activity['teacher_avatar'] = teacher_info['avatar']
            else:
                # Set defaults if no teacher info found
                activity['teacher_name'] = 'Teacher'
                activity['teacher_subject'] = ''
                activity['teacher_avatar'] = '/static/image/default-avatar.png'

        # Return both activities and assigned teachers for the dropdown
        assigned_teachers_list = []
        for teacher_id in assigned_teacher_ids:
            if teacher_id in teacher_map:
                assigned_teachers_list.append({
                    'id': teacher_id,
                    'name': teacher_map[teacher_id]['name'],
                    'subject': teacher_map[teacher_id]['subject'],
                    'avatar': teacher_map[teacher_id]['avatar'],
                    'is_assigned': True
                })
            else:
                # Debug: Log missing teacher
                print(f"Warning: Assigned teacher {teacher_id} not found in teacher_map")

        print(f"Final assigned_teachers_list: {assigned_teachers_list}")
        
        # Optionally, sort all activities by date (due_date/received_at) descending
        def get_date(act):
            return act.get('due_date') or ''
        activities.sort(key=get_date, reverse=True)

        return jsonify({
            'success': True, 
            'activities': activities,
            'assigned_teachers': assigned_teachers_list,
            'student_info': {
                'year_level': year_level,
                'section': section
            }
        })
    except Exception as e:
        print(f"Error in api_student_activities: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'activities': [], 'message': str(e)}), 500

@app.route('/api/task-files/<int:task_id>', methods=['GET'])
def get_task_files(task_id):
    """Get all files submitted for a specific task"""
    try:
        student_id = session.get('user_id')
        if not student_id:
            return jsonify({'success': False, 'message': 'Not logged in'}), 401

        # Verify the task belongs to this student
        task_check = supabase.table('task_assignments').select('task_id').eq('task_id', task_id).eq('student_id', student_id).execute()
        if not task_check.data:
            return jsonify({'success': False, 'message': 'Task not found or access denied'}), 404

        # Fetch files from task_file_submissions
        files_result = supabase.table('task_file_submissions')\
            .select('*')\
            .eq('task_id', task_id)\
            .eq('student_id', student_id)\
            .order('uploaded_at', desc=True)\
            .execute()

        return jsonify({
            'success': True,
            'files': files_result.data if files_result.data else []
        })

    except Exception as e:
        print(f"Error fetching task files: {e}")
        return jsonify({'success': False, 'message': 'Server error'}), 500

@app.route('/api/points/<int:point_id>', methods=['GET'])
def get_point_details(point_id):
    try:
        # Fetch point data - REMOVE 'subject' from select since it's not in points table
        result = supabase.table('points') \
            .select('*, teacher_id, points, point_category, note, received_at') \
            .eq('point_id', point_id) \
            .execute()
        
        if not result.data:
            return jsonify({'success': False, 'message': 'Point record not found'}), 404
        
        point = result.data[0]
        
        # If teacher_id exists, fetch teacher info and profile picture
        if point.get('teacher_id'):
            # Fetch teacher basic info - subject comes from user_info table
            teacher_result = supabase.table('user_info') \
                .select('first_name, last_name, gender, subject') \
                .eq('id', point['teacher_id']) \
                .execute()
            
            if teacher_result.data:
                teacher = teacher_result.data[0]
                prefix = 'Mr.' if (teacher.get('gender', '').lower() == 'male') else 'Ms.'
                point['teacher_name'] = f"{prefix} {teacher['last_name']}"
                point['teacher_subject'] = teacher.get('subject', 'No Subject')  # Provide default
            
            # Fetch teacher profile picture
            pic_result = supabase.table('profile_pictures') \
                .select('file_path') \
                .eq('user_id', point['teacher_id']) \
                .execute()
            
            if pic_result.data and pic_result.data[0].get('file_path'):
                point['teacher_avatar'] = pic_result.data[0]['file_path']
            else:
                point['teacher_avatar'] = '/static/image/default-avatar.png'
        else:
            # Set defaults if no teacher_id
            point['teacher_name'] = 'Teacher'
            point['teacher_subject'] = 'No Subject'
            point['teacher_avatar'] = '/static/image/default-avatar.png'
        
        return jsonify({'success': True, 'point': point})
        
    except Exception as e:
        print(f"Error in get_point_details: {str(e)}")  # Debug logging
        return jsonify({'success': False, 'message': str(e)}), 500

############################# UNIFIED QUARTERS API ENDPOINTS #############################

@app.route('/api/current-school-year', methods=['GET'])
def api_current_school_year():
    """Get the current system school year and active quarter"""
    try:
        # Single source of truth:
        # 1) Use active quarter first (if any)
        # 2) Fallback to helper school-year logic for gap periods
        current_quarter = get_current_quarter()
        if current_quarter:
            current_school_year = current_quarter.get('school_year')
        else:
            current_school_year = get_current_school_year()

        if not current_school_year:
            return jsonify({'success': False, 'message': 'No school years found'}), 404

        return jsonify({
            'success': True,
            'school_year': current_school_year,
            'current_quarter': current_quarter if current_quarter else None
        })
    except Exception as e:
        app.logger.error(f"Error getting current school year: {str(e)}")
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/api/school-years', methods=['GET'])
def api_school_years():
    """Get all available school years"""
    try:
        quarters_result = supabase.table('quarters').select('school_year').order('school_year', desc=True).execute()
        
        if not quarters_result.data:
            return jsonify({'success': True, 'school_years': []})
        
        # Get unique school years
        school_years = list(set(q['school_year'] for q in quarters_result.data))
        school_years.sort(reverse=True)
        
        return jsonify({'success': True, 'school_years': school_years})
    except Exception as e:
        app.logger.error(f"Error fetching school years: {str(e)}")
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/api/quarters', methods=['GET'])
def api_quarters():
    """Get quarters for a specific school year"""
    try:
        school_year = request.args.get('school_year')
        if not school_year:
            return jsonify({'success': False, 'message': 'Missing school_year parameter'}), 400
        
        quarters_result = supabase.table('quarters') \
            .select('*') \
            .eq('school_year', school_year) \
            .order('start_date', desc=False) \
            .execute()
        
        if not quarters_result.data:
            return jsonify({'success': True, 'quarters': []})
        
        # Add status to each quarter.
        # Respect explicit DB status when provided; fallback to date-based computation.
        now = datetime.now(philippines_tz)
        now_date = now.strftime('%Y-%m-%d')
        
        for quarter in quarters_result.data:
            start_date = quarter.get('start_date', '').split('T')[0]
            end_date = quarter.get('end_date', '').split('T')[0]
            raw_status = str(quarter.get('status', '')).lower()
            
            if raw_status == 'active':
                quarter['status'] = 'active'
            elif raw_status in ('finished', 'completed'):
                quarter['status'] = 'finished'
            elif raw_status == 'upcoming':
                quarter['status'] = 'upcoming'
            elif start_date <= now_date <= end_date:
                quarter['status'] = 'active'
            elif end_date < now_date:
                quarter['status'] = 'finished'
            else:
                quarter['status'] = 'upcoming'
        
        return jsonify({'success': True, 'quarters': quarters_result.data})
    except Exception as e:
        app.logger.error(f"Error fetching quarters: {str(e)}")
        return jsonify({'success': False, 'message': str(e)}), 500

############################# END QUARTERS API ENDPOINTS #############################

@app.route('/api/my-points', methods=['GET'])
def api_my_points():
    if 'user_id' not in session or session.get('role') != 'Student':
        return jsonify({'success': False, 'points': 0}), 401
    user_id = session['user_id']
    result = supabase.table('user_info').select('total_points').eq('id', user_id).execute()
    points = result.data[0]['total_points'] if result.data and 'total_points' in result.data[0] else 0
    # Optionally update session['points'] here if you want
    return jsonify({'success': True, 'points': points})

@app.route('/api/update-task/<int:task_id>', methods=['PUT'])
def update_task_route(task_id):
    """Update an existing task/activity"""
    if 'user_id' not in session or session.get('role') != 'Teacher':
        return jsonify({'success': False, 'message': 'Unauthorized'}), 403

    teacher_id = session['user_id']
    data = request.get_json()

    if not data:
        return jsonify({'success': False, 'message': 'No data provided'}), 400

    # Validate required fields
    required_fields = ['task', 'priority', 'points']
    missing_fields = [field for field in required_fields if field not in data or not data[field]]
    if missing_fields:
        return jsonify({
            'success': False, 
            'message': f'Missing required fields: {", ".join(missing_fields)}'
        }), 400

    try:
        # Verify task exists and belongs to this teacher
        task_check = supabase.table('task_assignments').select('*').eq('task_id', task_id).limit(1).execute()
        
        if not task_check.data:
            return jsonify({'success': False, 'message': 'Task not found'}), 404

        task = task_check.data[0]
        
        # Verify teacher ownership
        if task['teacher_id'] != teacher_id:
            return jsonify({'success': False, 'message': 'You do not have permission to edit this task'}), 403

        # Extract and validate input
        task_name = data.get('task', '').strip()
        description = data.get('description', '').strip()
        due_date = data.get('due_date')
        priority = data.get('priority', 'medium').lower()
        points = int(data.get('points', 0))
        template = data.get('template', 'default')

        # Validate task name
        if not task_name:
            return jsonify({'success': False, 'message': 'Activity name cannot be empty'}), 400

        # Validate priority
        if priority not in ['low', 'medium', 'high']:
            return jsonify({'success': False, 'message': 'Invalid priority level'}), 400

        # Validate points with priority limits
        PRIORITY_POINT_LIMITS = {
            'low': {'min': 1, 'max': 15},
            'medium': {'min': 16, 'max': 25},
            'high': {'min': 26, 'max': 50}
        }
        
        limits = PRIORITY_POINT_LIMITS.get(priority, {'min': 1, 'max': 15})
        if points < limits['min'] or points > limits['max']:
            return jsonify({
                'success': False, 
                'message': f'For {priority} priority, points must be between {limits["min"]} and {limits["max"]}'
            }), 400

        # Validate template
        if template not in ['default', 'success-green', 'deep-indigo']:
            template = 'default'

        # Validate due date format if provided
        if due_date:
            try:
                from datetime import datetime
                # Handle both date and datetime formats
                if 'T' in due_date:
                    datetime.fromisoformat(due_date.replace('Z', '+00:00'))
                else:
                    datetime.strptime(due_date, '%Y-%m-%d')
            except (ValueError, TypeError):
                return jsonify({'success': False, 'message': 'Invalid due date format. Use YYYY-MM-DD'}), 400

        # Prepare update data - use 'task' column name, not 'title'
        update_data = {
            'task': task_name,
            'description': description,
            'priority': priority,
            'points': points,
            'template': template
        }
        
        # Only add due_date if provided (handle both formats)
        if due_date:
            if 'T' not in due_date:
                update_data['due_date'] = f"{due_date}T00:00:00"
            else:
                update_data['due_date'] = due_date

        # Update the task in database
        result = supabase.table('task_assignments').update(update_data).eq('task_id', task_id).execute()

        if not result.data:
            return jsonify({'success': False, 'message': 'Failed to update task'}), 500

        # --- ACTIVITY LOG: Edit Activity ---
        record_admin_activity_log(
            user_id=teacher_id,
            action='Edit Activity',
            activity='Activities Management',
            description=f"Updated activity: {task_name} (Priority: {priority}, Points: {points})",
            user_role='Teacher'
        )

        # --- NOTIFICATION: Notify students of task changes ---
        # Get all students assigned to this task
        students_result = supabase.table('task_assignments').select('student_id').eq('task_id', task_id).execute()
        student_ids = [s['student_id'] for s in students_result.data] if students_result.data else []

        if student_ids:
            # Create notification for students
            notifications = []
            for student_id in student_ids:
                notifications.append({
                    'user_id': student_id,
                    'sender_id': teacher_id,
                    'title': 'Activity Updated',
                    'message': f'The activity "{task_name}" has been updated with new details.',
                    'notif_type': 'Task',
                    'status': 'Unread',
                    'task_id': task_id
                })
            
            if notifications:
                try:
                    supabase.table('notifications').insert(notifications).execute()
                except Exception as notif_error:
                    app.logger.warning(f"Failed to send notifications: {str(notif_error)}")

        return jsonify({
            'success': True,
            'message': 'Activity updated successfully',
            'task': result.data[0] if result.data else update_data
        }), 200

    except ValueError as e:
        app.logger.error(f"ValueError updating task {task_id}: {str(e)}")
        return jsonify({'success': False, 'message': f'Invalid data format: {str(e)}'}), 400
    except Exception as e:
        app.logger.error(f"Error updating task {task_id}: {str(e)}")
        return jsonify({'success': False, 'message': 'An error occurred while updating the activity: ' + str(e)}), 500

@app.route('/api/upload-proof-image', methods=['POST'])
def upload_proof_image():
    if 'file' not in request.files:
        return jsonify({'success': False, 'message': 'No file uploaded'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'success': False, 'message': 'No filename'}), 400
    
    try:
        # Get original filename - this is the key part for preserving the name
        original_filename = request.form.get('original_filename', file.filename)
        
        # Get file size first (before reading the content)
        file.seek(0, os.SEEK_END)
        file_size = file.tell()
        file.seek(0)  # Reset file pointer to beginning
        
        # Generate unique filename for storage
        file_extension = os.path.splitext(file.filename)[1]
        filename = f"proof_{uuid.uuid4().hex}{file_extension}"
        
        # Read file content after getting size
        file_content = file.read()
        
        # Upload to Supabase Storage
        result = supabase.storage.from_('task-images').upload(
            filename, 
            file_content,
            {"content-type": file.content_type}
        )
        
        # Get public URL
        public_url = supabase.storage.from_('task-images').get_public_url(filename)
        
        return jsonify({
            'success': True, 
            'url': public_url,
            'filename': filename,
            'original_filename': original_filename,  # Send back the original filename
            'file_size': file_size,
            'content_type': file.content_type
        })
        
    except Exception as e:
        print(f"Error uploading file: {e}")
        return jsonify({'success': False, 'message': 'Upload failed'}), 500

@app.route('/api/submit-task-proof', methods=['POST'])
def submit_task_proof():
    import json
    import mimetypes
    data = request.json
    task_id = data.get('task_id')
    image_urls = data.get('image_urls', [])
    file_metadata = data.get('file_metadata', [])  # File metadata array

    if not task_id:
        return jsonify({'success': False, 'message': 'Missing task ID'}), 400

    if not image_urls:
        return jsonify({'success': False, 'message': 'No proof images provided'}), 400

    try:
        student_id = session.get('user_id')
        if not student_id:
            return jsonify({'success': False, 'message': 'Not logged in'}), 401

        # Fetch student info for logging
        user_result = supabase.table('user_info').select('first_name', 'last_name', 'role').eq('id', student_id).execute()
        user_info = user_result.data[0] if user_result.data else {}
        user_name = f"{user_info.get('first_name', '')} {user_info.get('last_name', '')}"
        user_role = user_info.get('role', 'Student')

        # Fetch task info for logging
        task_result = supabase.table('task_assignments').select('task', 'teacher_id').eq('task_id', task_id).eq('student_id', student_id).execute()
        task_info = task_result.data[0] if task_result.data else {}
        task_name = task_info.get('task', 'Activity')
        teacher_id = task_info.get('teacher_id')

        # Always store as JSON string in TEXT column (for backward compatibility)
        result = supabase.table('task_assignments').update({
            'image_urls': json.dumps(image_urls),  # <-- Always stringify!
            'status': 'Pending',
            'submitted_at': datetime.now(timezone.utc).isoformat()
        }).eq('task_id', task_id).eq('student_id', student_id).execute()

        # First, delete any existing file submissions for this task (if any)
        supabase.table('task_file_submissions').delete().eq('task_id', task_id).eq('student_id', student_id).execute()

        # Insert each file into task_file_submissions table with all metadata
        for i, url in enumerate(image_urls):
            # Get filename from URL
            filename = url.split('/')[-1] if '/' in url else f"file_{i}"
            
            # Get metadata for this file if available
            metadata = file_metadata[i] if i < len(file_metadata) else {}
            
            # Determine file type based on extension
            file_extension = os.path.splitext(filename)[1].lower().replace('.', '')
            file_type = 'image' if file_extension in ['jpg', 'jpeg', 'png', 'gif', 'webp', 'bmp'] else 'document'
            
            # Get mime type
            mime_type = metadata.get('mime_type') or mimetypes.guess_type(filename)[0] or 'application/octet-stream'
            
            # Get file size (try to get from metadata, otherwise null)
            file_size = metadata.get('file_size')
            
            # Get original filename
            original_filename = metadata.get('original_filename', filename)
            
            print(f"Inserting file: {original_filename}, size: {file_size}, type: {file_type}, mime: {mime_type}")
            
            # Insert into task_file_submissions with all fields
            insert_data = {
                'task_id': task_id,
                'student_id': student_id,
                'filename': filename,
                'original_filename': original_filename,
                'file_url': url,
                'file_type': file_type,
                'mime_type': mime_type,
                'uploaded_at': datetime.now(timezone.utc).isoformat()
            }
            
            # Only add file_size if it's not None
            if file_size is not None:
                insert_data['file_size'] = file_size
            
            supabase.table('task_file_submissions').insert(insert_data).execute()

        # Only log and notify if status is "Pending"
        if result.data:
            # Log to admin_activity_log
            record_admin_activity_log(
                user_id=student_id,
                action='Activity Submission',
                activity='Submits Activity',
                description=f"{user_name} has submitted the activity '{task_name}'.",
                user_role=user_role,
                details=(
                    f"{task_name} - Status: Pending Teacher Approval\n"
                    f"Submitted by Student: {user_name}\n"
                    f"Files uploaded: {len(image_urls)} file(s)\n"
                    "Action Required: Please wait for teacher review"
                ),
            )

            # Insert notification for teacher/admin
            if teacher_id:
                supabase.table('notifications').insert({
                    'user_id': teacher_id,  # teacher/admin who assigned the task
                    'sender_id': student_id,
                    'title': 'Activity Submission',
                    'message': f"{user_name} has submitted the activity '{task_name}' with {len(image_urls)} file(s). Kindly review and approve.",
                    'task_id': task_id,
                    'notif_type': 'Task',
                    'status': 'Unread',
                }).execute()

            return jsonify({'success': True, 'message': 'Proof submitted successfully'})
        else:
            return jsonify({'success': False, 'message': 'Failed to update task'}), 500

    except Exception as e:
        print(f"Error submitting proof: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'message': 'Server error'}), 500

#############################STUDENT_QUIZ.HTML#############################
# Student Submit Quiz Route
@app.route('/api/student/submit-quiz', methods=['POST'])
def submit_student_quiz():
    """Student submits quiz answers"""
    try:
        data = request.get_json()
        student_id = session.get('user_id')
        
        if not student_id:
            return jsonify({'error': 'Not authenticated'}), 401
        
        teacher_quiz_id = data.get('teacher_quiz_id')
        answers = data.get('answers')
        score = data.get('score')
        time_spent = data.get('time_spent', 0)
        
        # Validate required fields
        if not teacher_quiz_id or not answers or score is None:
            return jsonify({'error': 'Missing required fields'}), 400
        
        # Get student info
        student = supabase.table('user_info')\
            .select('first_name', 'last_name', 'year_level', 'section')\
            .eq('id', student_id)\
            .execute()
        
        if not student.data:
            return jsonify({'error': 'Student not found'}), 404
            
        student_data = student.data[0]
        student_name = f"{student_data['first_name']} {student_data['last_name']}"
        
        # Get quiz details
        quiz = supabase.table('teacher_quizzes')\
            .select('*')\
            .eq('id', teacher_quiz_id)\
            .execute()
        
        if not quiz.data:
            return jsonify({'error': 'Quiz not found'}), 404
            
        quiz_data = quiz.data[0]
        
        # Check for previous attempts (for retakes)
        existing_attempts = supabase.table('student_quiz_results')\
            .select('*')\
            .eq('teacher_quiz_id', teacher_quiz_id)\
            .eq('student_id', student_id)\
            .order('attempt_number', desc=True)\
            .execute()
        
        # Determine attempt number and get previous best score
        attempt_number = 1
        previous_best_score = 0
        previous_best_attempt_id = None
        if existing_attempts.data:
            attempt_number = max([a.get('attempt_number', 1) for a in existing_attempts.data]) + 1
            # Find the best score from all previous attempts
            best_attempt = max(existing_attempts.data, key=lambda x: x.get('score', 0))
            previous_best_score = best_attempt.get('score', 0)
            previous_best_attempt_id = best_attempt.get('id', None)
        
        # Calculate total points based on items and points per item
        points_per_item = quiz_data.get('points_per_item', 5)  # Default to 5 if not specified
        total_points = quiz_data['total_items'] * points_per_item
        
        # Validate score doesn't exceed maximum
        if score > total_points:
            return jsonify({'error': 'Invalid score'}), 400
        
        # Calculate points earned based on retake logic
        bonus_points_earned = 0
        is_best_score = True
        
        if attempt_number > 1:
            # This is a retake
            if score > previous_best_score:
                # Only get bonus points for improvement
                bonus_points_earned = score - previous_best_score
            else:
                # No penalty, no bonus if score stayed same or decreased
                bonus_points_earned = 0
            
            # Mark if this is the best score
            is_best_score = score > previous_best_score
        else:
            # First attempt - get full points
            bonus_points_earned = score
            is_best_score = True
        
        # Save results
        result_data = {
            'teacher_quiz_id': teacher_quiz_id,
            'student_id': student_id,
            'student_name': student_name,
            'score': score,
            'total_items': quiz_data['total_items'],
            'total_points': total_points,
            'topic': quiz_data['topic'],
            'quarter': quiz_data['quarter'],
            'grade_level': student_data['year_level'],
            'section': student_data['section'],
            'teacher_id': quiz_data['teacher_id'],
            'teacher_name': quiz_data['teacher_name'],
            'answers': answers,
            'time_spent': time_spent,
            'attempt_number': attempt_number,
            'previous_best_score': previous_best_score if attempt_number > 1 else None,
            'is_best_score': is_best_score,
            'points_earned': bonus_points_earned,
            'bonus_points_earned': bonus_points_earned
        }
        
        result = supabase.table('student_quiz_results').insert(result_data).execute()
        
        if not result.data:
            return jsonify({'error': 'Failed to save quiz results'}), 500
        
        # If this is a better score, update the previous best to mark it as no longer best
        if is_best_score and attempt_number > 1 and previous_best_attempt_id:
            try:
                supabase.table('student_quiz_results')\
                    .update({'is_best_score': False})\
                    .eq('id', previous_best_attempt_id)\
                    .execute()
            except Exception as e:
                logger.warning(f"Failed to update previous best score flag: {e}")
        
        # Insert points earned into points table
        points_earned = bonus_points_earned  # Only earn bonus points
        points_data = {
            'student_id': student_id,
            'teacher_id': quiz_data['teacher_id'],
            'points': points_earned,
            'point_category': 'Quiz',
            'note': f"Quiz: {quiz_data['topic']} - Score: {score}/{total_points}",
            'status': 'approved'
        }
        
        points_result = supabase.table('points').insert(points_data).execute()
        
        if not points_result.data:
            logger.warning(f"Failed to insert points for student {student_id} on quiz {teacher_quiz_id}")
        
        # Update student's total_points in user_info table
        try:
            # Get current total_points
            current_user = supabase.table('user_info')\
                .select('total_points')\
                .eq('id', student_id)\
                .execute()
            
            current_total = current_user.data[0]['total_points'] if current_user.data and current_user.data[0]['total_points'] else 0
            new_total = current_total + points_earned
            
            # Update total_points in user_info
            update_result = supabase.table('user_info')\
                .update({'total_points': new_total})\
                .eq('id', student_id)\
                .execute()
            
            if update_result.data:
                logger.info(f"Updated student {student_id} total_points to {new_total}")
            else:
                logger.warning(f"Failed to update total_points for student {student_id}")
        except Exception as e:
            logger.error(f"Error updating total_points: {e}")
        
        # Log the submission
        logger.info(f"Student {student_id} submitted quiz {teacher_quiz_id} (attempt {attempt_number}) with score {score}/{total_points} and earned {points_earned} bonus points")
        
        message = f'Quiz submitted! Score: {score}/{total_points}'
        if attempt_number > 1:
            if is_best_score:
                message += f' (+{bonus_points_earned} bonus points - New best score!)'
            else:
                message += f' (No bonus points - keep trying!)'
        
        return jsonify({
            'success': True,
            'result_id': result.data[0]['id'],
            'message': message,
            'percentage': (score / total_points) * 100 if total_points > 0 else 0,
            'attempt_number': attempt_number,
            'previous_best_score': previous_best_score if attempt_number > 1 else None,
            'bonus_points_earned': bonus_points_earned,
            'is_best_score': is_best_score
        })
        
    except Exception as e:
        logger.error(f"Error submitting quiz: {e}")
        return jsonify({'error': str(e)}), 500


# Student Get Their Performance Route
@app.route('/api/student/my-performance', methods=['GET'])
def get_my_performance():
    """Get authenticated student's performance"""
    try:
        student_id = session.get('user_id')
        
        if not student_id:
            return jsonify({'error': 'Not authenticated'}), 401
        
        # Verify student exists
        student = supabase.table('user_info')\
            .select('first_name, last_name, year_level, section')\
            .eq('id', student_id)\
            .execute()
        
        if not student.data:
            return jsonify({'error': 'Student not found'}), 404
        
        # Get all quiz results for this student
        results = supabase.table('student_quiz_results')\
            .select('*')\
            .eq('student_id', student_id)\
            .order('submitted_date', desc=True)\
            .execute()
        
        # Calculate statistics
        total_quizzes = len(results.data)
        if total_quizzes > 0:
            # Handle None values safely
            total_score = sum(r.get('score') or 0 for r in results.data)
            total_items = sum(r.get('total_items') or 0 for r in results.data)
            overall_percentage = (total_score / total_items * 100) if total_items > 0 else 0
            
            # Topic breakdown
            topics = {}
            for r in results.data:
                topic = r.get('topic') or 'General'
                if topic not in topics:
                    topics[topic] = {
                        'topic': topic,
                        'quizzes': 0,
                        'total_score': 0,
                        'total_items': 0,
                        'average': 0
                    }
                topics[topic]['quizzes'] += 1
                topics[topic]['total_score'] += r.get('score') or 0
                topics[topic]['total_items'] += r.get('total_items') or 0
            
            for t in topics.values():
                t['average'] = (t['total_score'] / t['total_items'] * 100) if t['total_items'] > 0 else 0
                
            # Quarterly breakdown
            quarterly = {}
            for r in results.data:
                quarter = r.get('quarter') or 'Unknown'
                if quarter not in quarterly:
                    quarterly[quarter] = {
                        'quarter': quarter,
                        'quizzes': 0,
                        'total_score': 0,
                        'total_items': 0,
                        'average': 0
                    }
                quarterly[quarter]['quizzes'] += 1
                quarterly[quarter]['total_score'] += r.get('score') or 0
                quarterly[quarter]['total_items'] += r.get('total_items') or 0
            
            for q in quarterly.values():
                q['average'] = (q['total_score'] / q['total_items'] * 100) if q['total_items'] > 0 else 0
        else:
            overall_percentage = 0
            topics = {}
            quarterly = {}
        
        return jsonify({
            'success': True,
            'statistics': {
                'total_quizzes_taken': total_quizzes,
                'overall_average': round(overall_percentage, 2),
                'total_score': total_score if total_quizzes > 0 else 0,
                'total_items': total_items if total_quizzes > 0 else 0
            },
            'topic_performance': list(topics.values()),
            'quarterly_performance': list(quarterly.values()),
            'recent_results': results.data[:10]  # Last 10 submissions
        })
        
    except Exception as e:
        logger.error(f"Error fetching student performance: {e}")
        return jsonify({'error': str(e)}), 500

# Get Student Answers for a Specific Quiz
@app.route('/api/quiz/<int:quiz_id>/student/<int:student_id>/answers', methods=['GET'])
def get_student_quiz_answers(quiz_id, student_id):
    """Get a specific student's answers for a quiz"""
    try:
        user_id = session.get('user_id')
        
        if not user_id:
            return jsonify({'error': 'Not authenticated'}), 401
        
        # Get user role
        user = supabase.table('user_info')\
            .select('role')\
            .eq('id', user_id)\
            .execute()
        
        if not user.data:
            return jsonify({'error': 'User not found'}), 404
            
        role = user.data[0]['role']
        
        # Get the quiz details
        quiz = supabase.table('teacher_quizzes')\
            .select('*')\
            .eq('id', quiz_id)\
            .execute()
        
        if not quiz.data:
            return jsonify({'error': 'Quiz not found'}), 404
        
        quiz_data = quiz.data[0]
        
        # Check permissions
        if role == 'Student':
            # Students can only see their own answers
            if student_id != user_id:
                return jsonify({'error': 'Access denied'}), 403
        elif role == 'Teacher':
            # Teachers can only see answers for their quizzes
            if quiz_data['teacher_id'] != user_id:
                return jsonify({'error': 'Access denied'}), 403
        else:
            return jsonify({'error': 'Invalid role'}), 403
        
        # Get student's results
        result = supabase.table('student_quiz_results')\
            .select('*')\
            .eq('teacher_quiz_id', quiz_id)\
            .eq('student_id', student_id)\
            .execute()
        
        if not result.data:
            return jsonify({'error': 'Student results not found'}), 404
        
        # Get student info
        student = supabase.table('user_info')\
            .select('first_name, last_name')\
            .eq('id', student_id)\
            .execute()
        
        student_name = f"{student.data[0]['first_name']} {student.data[0]['last_name']}" if student.data else "Unknown"
        
        # Get student profile picture
        student_pic = '/static/image/default-avatar.png'
        pic_result = supabase.table('profile_pictures')\
            .select('file_path')\
            .eq('user_id', student_id)\
            .execute()
        
        if pic_result.data and pic_result.data[0].get('file_path'):
            student_pic = pic_result.data[0]['file_path']
        
        return jsonify({
            'success': True,
            'quiz': quiz_data,
            'result': result.data[0],
            'student_name': student_name,
            'student_pic': student_pic
        })
        
    except Exception as e:
        logger.error(f"Error fetching student answers: {e}")
        return jsonify({'error': str(e)}), 500


# Get Student's Previous Quiz Attempts
@app.route('/api/student/<int:student_id>/quiz/<int:quiz_id>/attempts', methods=['GET'])
def get_student_quiz_attempts(student_id, quiz_id):
    """Get all attempts for a student on a specific quiz (for retake comparison)"""
    try:
        user_id = session.get('user_id')
        
        if not user_id:
            return jsonify({'error': 'Not authenticated'}), 401
        
        # Get user role
        user = supabase.table('user_info')\
            .select('role')\
            .eq('id', user_id)\
            .execute()
        
        if not user.data:
            return jsonify({'error': 'User not found'}), 404
            
        role = user.data[0]['role']
        
        # Get the quiz details to verify ownership (for teachers)
        quiz = supabase.table('teacher_quizzes')\
            .select('*')\
            .eq('id', quiz_id)\
            .execute()
        
        if not quiz.data:
            return jsonify({'error': 'Quiz not found'}), 404
        
        quiz_data = quiz.data[0]
        
        # Check permissions
        if role == 'Student':
            # Students can only see their own attempts
            if student_id != user_id:
                return jsonify({'error': 'Access denied'}), 403
        elif role == 'Teacher':
            # Teachers can only see attempts for their quizzes
            if quiz_data['teacher_id'] != user_id:
                return jsonify({'error': 'Access denied'}), 403
        else:
            return jsonify({'error': 'Invalid role'}), 403
        
        # Get all student's results for this quiz, ordered by date descending (most recent first)
        results = supabase.table('student_quiz_results')\
            .select('*')\
            .eq('teacher_quiz_id', quiz_id)\
            .eq('student_id', student_id)\
            .order('created_at', desc=True)\
            .execute()
        
        if not results.data:
            return jsonify({
                'success': True,
                'attempts': []
            })
        
        # Return all attempts
        return jsonify({
            'success': True,
            'attempts': results.data
        })
        
    except Exception as e:
        logger.error(f"Error fetching student quiz attempts: {e}")
        return jsonify({'error': str(e)}), 500


# Get Quiz Statistics for Student Dashboard
@app.route('/api/student/quiz-statistics', methods=['GET'])
def get_student_quiz_statistics():
    """Get detailed statistics for student dashboard"""
    try:
        student_id = session.get('user_id')
        
        if not student_id:
            return jsonify({'error': 'Not authenticated'}), 401
        
        # Get student info
        student = supabase.table('user_info')\
            .select('first_name, last_name, year_level, section')\
            .eq('id', student_id)\
            .execute()
        
        if not student.data:
            return jsonify({'error': 'Student not found'}), 404
            
        student_data = student.data[0]
        
        # Get all quiz results
        results = supabase.table('student_quiz_results')\
            .select('*')\
            .eq('student_id', student_id)\
            .execute()
        
        # Get all available quizzes for this student's class
        all_quizzes = supabase.table('teacher_quizzes')\
            .select('*')\
            .eq('grade_level', student_data['year_level'])\
            .eq('section', student_data['section'])\
            .eq('status', 'active')\
            .execute()
        
        # Calculate detailed statistics
        completed_ids = [r['teacher_quiz_id'] for r in results.data]
        
        # Performance over time (last 5 quizzes)
        recent_results = sorted(results.data, key=lambda x: x['submitted_date'], reverse=True)[:5]
        performance_trend = []
        for r in recent_results:
            performance_trend.append({
                'date': r['submitted_date'],
                'score': (r['score'] / r['total_items']) * 100,
                'quiz_title': get_quiz_title(r['teacher_quiz_id'], all_quizzes.data)
            })
        
        # Subject/topic mastery
        topic_mastery = {}
        for r in results.data:
            topic = r['topic']
            if topic not in topic_mastery:
                topic_mastery[topic] = {
                    'topic': topic,
                    'quizzes_taken': 0,
                    'total_percentage': 0,
                    'average': 0
                }
            topic_mastery[topic]['quizzes_taken'] += 1
            topic_mastery[topic]['total_percentage'] += (r['score'] / r['total_items']) * 100
        
        for topic in topic_mastery.values():
            topic['average'] = topic['total_percentage'] / topic['quizzes_taken']
        
        # Quarterly performance
        quarter_performance = {}
        for r in results.data:
            quarter = r['quarter']
            if quarter not in quarter_performance:
                quarter_performance[quarter] = {
                    'quarter': quarter,
                    'quizzes_taken': 0,
                    'total_percentage': 0,
                    'average': 0
                }
            quarter_performance[quarter]['quizzes_taken'] += 1
            quarter_performance[quarter]['total_percentage'] += (r['score'] / r['total_items']) * 100
        
        for quarter in quarter_performance.values():
            quarter['average'] = quarter['total_percentage'] / quarter['quizzes_taken']
        
        return jsonify({
            'success': True,
            'statistics': {
                'total_quizzes_available': len(all_quizzes.data),
                'total_completed': len(results.data),
                'completion_rate': (len(results.data) / len(all_quizzes.data) * 100) if all_quizzes.data else 0,
                'overall_average': sum((r['score'] / r['total_items']) * 100 for r in results.data) / len(results.data) if results.data else 0
            },
            'performance_trend': performance_trend,
            'topic_mastery': list(topic_mastery.values()),
            'quarter_performance': list(quarter_performance.values()),
            'recent_activities': recent_results
        })
        
    except Exception as e:
        logger.error(f"Error fetching student statistics: {e}")
        return jsonify({'error': str(e)}), 500


# Get Quiz Questions for Taking
@app.route('/api/quiz/<int:quiz_id>/take', methods=['GET'])
def take_quiz(quiz_id):
    """Get quiz questions for taking (with shuffled options for fairness)"""
    try:
        student_id = session.get('user_id')
        
        if not student_id:
            return jsonify({'error': 'Not authenticated'}), 401
        
        # Check if student has already taken this quiz (to get retake info)
        existing_attempts = supabase.table('student_quiz_results')\
            .select('*')\
            .eq('teacher_quiz_id', quiz_id)\
            .eq('student_id', student_id)\
            .order('attempt_number', desc=True)\
            .execute()
        
        # Get previous best score if this is a retake
        previous_best_score = None
        attempt_number = 1
        if existing_attempts.data:
            # Find the best score from previous attempts
            best_attempt = next((a for a in existing_attempts.data if a.get('is_best_score', False)), existing_attempts.data[0])
            previous_best_score = best_attempt.get('score', 0)
            attempt_number = max([a.get('attempt_number', 1) for a in existing_attempts.data]) + 1
        
        # Get quiz details
        quiz = supabase.table('teacher_quizzes')\
            .select('*')\
            .eq('id', quiz_id)\
            .execute()
        
        if not quiz.data:
            return jsonify({'error': 'Quiz not found'}), 404
            
        quiz_data = quiz.data[0]
        
        # Verify student is in the correct class
        student = supabase.table('user_info')\
            .select('year_level, section')\
            .eq('id', student_id)\
            .execute()
        
        if not student.data:
            return jsonify({'error': 'Student not found'}), 404
            
        student_data = student.data[0]
        
        if quiz_data['grade_level'] != student_data['year_level'] or quiz_data['section'] != student_data['section']:
            return jsonify({'error': 'This quiz is not assigned to your class'}), 403
        
        # Check if quiz is still active
        if quiz_data['status'] != 'active':
            return jsonify({'error': 'This quiz is no longer active'}), 400
        
        # Prepare quiz questions (shuffle choices for fairness, but keep correct answer reference)
        questions = quiz_data['quiz_data']
        for q in questions:
            # Store correct answer
            correct = q['correct_answer']
            # Create choices with letters
            choices = q['choices']
            # Store both the choices and which one is correct
            q['choices'] = choices
            q['correct_answer'] = correct  # Keep as index reference
        
        # Remove correct answers from response (will be used only for grading)
        quiz_response = {
            'id': quiz_data['id'],
            'quiz_title': quiz_data['quiz_title'],
            'instructions': quiz_data['instructions'],
            'time_limit_minutes': quiz_data['time_limit_minutes'],
            'total_items': quiz_data['total_items'],
            'total_points': quiz_data['total_points'],
            'points_per_item': quiz_data['points_per_item'],
            'questions': [{
                'id': idx,
                'question': q['question'],
                'choices': q['choices'],
                'topic': q.get('topic', '')
            } for idx, q in enumerate(questions)],
            'attempt_number': attempt_number,
            'previous_best_score': previous_best_score,
            'is_retake': previous_best_score is not None
        }
        
        return jsonify({
            'success': True,
            'quiz': quiz_response
        })
        
    except Exception as e:
        logger.error(f"Error preparing quiz: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/quiz/<int:quiz_id>/review', methods=['GET'])
def review_quiz_answers(quiz_id):
    """Get quiz questions with student's answers and correct answers for review"""
    try:
        student_id = session.get('user_id')
        
        if not student_id:
            return jsonify({'error': 'Not authenticated'}), 401
        
        # Get the quiz details with correct answers
        quiz = supabase.table('teacher_quizzes')\
            .select('*')\
            .eq('id', quiz_id)\
            .execute()
        
        if not quiz.data:
            return jsonify({'error': 'Quiz not found'}), 404
            
        quiz_data = quiz.data[0]
        
        # Get teacher info for this quiz
        teacher = supabase.table('user_info')\
            .select('first_name, last_name, subject')\
            .eq('id', quiz_data['teacher_id'])\
            .execute()
        
        teacher_name = "Unknown Teacher"
        teacher_subject = "General"
        teacher_avatar = '/static/image/default-avatar.png'
        
        if teacher.data:
            teacher_data = teacher.data[0]
            teacher_name = f"{teacher_data.get('first_name', '')} {teacher_data.get('last_name', '')}".strip() or "Unknown Teacher"
            teacher_subject = teacher_data.get('subject', 'General')
            # You might want to add avatar URL to user_info table or use a default
        
        # Get student's results with their answers
        result = supabase.table('student_quiz_results')\
            .select('*')\
            .eq('teacher_quiz_id', quiz_id)\
            .eq('student_id', student_id)\
            .execute()
        
        if not result.data:
            return jsonify({'error': 'Results not found'}), 404
            
        result_data = result.data[0]
        
        # Get student info
        student = supabase.table('user_info')\
            .select('first_name, last_name')\
            .eq('id', student_id)\
            .execute()
        
        student_name = f"{student.data[0]['first_name']} {student.data[0]['last_name']}" if student.data else "Unknown"
        
        # Prepare questions with student answers and correct answers
        questions = quiz_data['quiz_data']
        student_answers = result_data.get('answers', {})
        
        # Format questions for review
        review_data = []
        for idx, q in enumerate(questions):
            question_key = f"q{idx + 1}"
            student_answer = student_answers.get(question_key)
            
            # Handle case where student_answer might be string or number
            if student_answer is not None and student_answer != '':
                try:
                    student_answer = int(student_answer)
                except (ValueError, TypeError):
                    student_answer = None
            
            review_data.append({
                'question_number': idx + 1,
                'question': q['question'],
                'choices': q['choices'],
                'student_answer': student_answer,
                'correct_answer': q['correct_answer'],
                'is_correct': student_answer == q['correct_answer'] if student_answer is not None else False,
                'was_answered': student_answer is not None and student_answer != ''
            })
        
        # Calculate points per item
        points_per_item = quiz_data.get('points_per_item', 5)  # Default to 5
        total_items = result_data.get('total_items', 0) or 0
        total_points = result_data.get('total_points') or quiz_data.get('total_points') or (total_items * points_per_item)
        
        # Ensure total_points is at least 1 to avoid division by zero
        if not total_points or total_points <= 0:
            total_points = total_items * points_per_item
            if total_points <= 0:
                total_points = 1
        
        # Ensure score is a number (handle None)
        score = result_data.get('score', 0)
        if score is None:
            score = 0
        
        # Calculate percentage safely
        percentage = (score / total_points * 100) if total_points > 0 else 0
        
        return jsonify({
            'success': True,
            'quiz_title': quiz_data['quiz_title'],
            'student_name': student_name,
            'teacher_id': quiz_data['teacher_id'],
            'teacher_name': teacher_name,
            'teacher_subject': teacher_subject,
            'teacher_avatar': teacher_avatar,
            'total_score': score,
            'total_items': total_items,
            'total_points': total_points,
            'percentage': percentage,
            'questions': review_data,
            'submitted_date': result_data['submitted_date'],
            'points_per_item': points_per_item
        })
        
    except Exception as e:
        logger.error(f"Error reviewing quiz answers: {e}")
        return jsonify({'error': str(e)}), 500
 
@app.route('/api/quiz/<int:quiz_id>/allow-retake/<int:student_id>', methods=['POST'])
def allow_student_retake(quiz_id, student_id):
    """Teacher allows a student to retake a quiz - sends notification to student"""
    try:
        teacher_id = session.get('user_id')
        
        if not teacher_id:
            return jsonify({'error': 'Not authenticated'}), 401
        
        # Verify teacher owns this quiz
        quiz = supabase.table('teacher_quizzes')\
            .select('*')\
            .eq('id', quiz_id)\
            .eq('teacher_id', teacher_id)\
            .execute()
        
        if not quiz.data:
            return jsonify({'error': 'Quiz not found or access denied'}), 404
        
        quiz_data = quiz.data[0]
        
        # Get student info
        student = supabase.table('user_info')\
            .select('*')\
            .eq('id', student_id)\
            .execute()
        
        if not student.data:
            return jsonify({'error': 'Student not found'}), 404
        
        student_data = student.data[0]
        
        # Get teacher info
        teacher = supabase.table('user_info')\
            .select('*')\
            .eq('id', teacher_id)\
            .execute()
        
        teacher_info = teacher.data[0] if teacher.data else {}
        teacher_name = f"{teacher_info.get('first_name', '')} {teacher_info.get('last_name', '')}".strip()
        
        # Create notification for student - clarify they have 1 retake opportunity
        notification_message = f"Your teacher {teacher_name} has allowed you to retake the quiz: {quiz_data['quiz_title']}. You have 1 retake attempt available."
        
        notification_data = {
            'user_id': student_id,
            'notif_type': 'Task',
            'title': 'Quiz Retake Allowed',
            'message': notification_message
        }
        
        # Insert notification
        notif_result = supabase.table('notifications')\
            .insert(notification_data)\
            .execute()
        
        if not notif_result.data:
            logger.warning(f"Failed to create notification for student {student_id}")
        
        # Log activity
        record_admin_activity_log(
            user_id=teacher_id,
            action='Allow Quiz Retake',
            activity=f'Allowed student {student_data["first_name"]} {student_data["last_name"]} to retake quiz: {quiz_data["quiz_title"]}',
            description=f'Quiz ID: {quiz_id}, Student ID: {student_id}',
            user_role='Teacher'
        )
        
        logger.info(f"Teacher {teacher_id} allowed student {student_id} to retake quiz {quiz_id}")
        
        return jsonify({
            'success': True,
            'message': f'Notification sent to {student_data["first_name"]} {student_data["last_name"]}',
            'student_name': f"{student_data['first_name']} {student_data['last_name']}"
        }), 200
        
    except Exception as e:
        logger.error(f"Error allowing student retake: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/quiz/<int:quiz_id>/retake-history', methods=['GET'])
def get_retake_history(quiz_id):
    """Get all attempts and retake history for a specific quiz by the student"""
    try:
        student_id = session.get('user_id')
        user_role = session.get('user_role', 'Student')
        
        if not student_id:
            return jsonify({'error': 'Not authenticated'}), 401
        
        # Get all attempts for this quiz
        if user_role == 'Student':
            # Students can only see their own attempts
            attempts = supabase.table('student_quiz_results')\
                .select('*')\
                .eq('teacher_quiz_id', quiz_id)\
                .eq('student_id', student_id)\
                .order('attempt_number')\
                .execute()
        else:
            # Teachers would get this endpoint to view student history
            student_id_param = request.args.get('student_id')
            if not student_id_param:
                return jsonify({'error': 'student_id required for teachers'}), 400
            
            attempts = supabase.table('student_quiz_results')\
                .select('*')\
                .eq('teacher_quiz_id', quiz_id)\
                .eq('student_id', int(student_id_param))\
                .order('attempt_number')\
                .execute()
        
        if not attempts.data:
            return jsonify({
                'success': True,
                'attempts': [],
                'has_attempts': False,
                'best_score': None,
                'best_score_attempt': None,
                'can_retake': False,
                'total_attempts': 0,
                'retake_allowed_by_teacher': False
            }), 200
        
        # Find best score
        best_attempt = max(attempts.data, key=lambda x: x['score'])
        total_attempts = len(attempts.data)
        
        # Get quiz title for checking notification message
        quiz_title_result = supabase.table('teacher_quizzes')\
            .select('quiz_title')\
            .eq('id', quiz_id)\
            .execute()
        
        quiz_title = quiz_title_result.data[0]['quiz_title'] if quiz_title_result.data else ''
        
        # Check if teacher has sent "Allow Retake" notification for this quiz
        retake_allowed_by_teacher = False
        if user_role == 'Student':
            # Check for retake allowed notification from teacher
            notifications = supabase.table('notifications')\
                .select('*')\
                .eq('user_id', student_id)\
                .eq('title', 'Quiz Retake Allowed')\
                .execute()
            
            # Check if there's a notification for this specific quiz
            if notifications.data and quiz_title:
                for notif in notifications.data:
                    # Check if the message mentions this quiz title
                    if quiz_title in notif.get('message', ''):
                        retake_allowed_by_teacher = True
                        break
        
        # Only allow retake if:
        # 1. Teacher has explicitly allowed it (sent notification)
        # 2. Student has exactly 1 attempt (first attempt only)
        can_retake = retake_allowed_by_teacher and total_attempts == 1
        
        return jsonify({
            'success': True,
            'attempts': attempts.data,
            'has_attempts': len(attempts.data) > 0,
            'best_score': best_attempt['score'],
            'best_score_attempt': best_attempt['attempt_number'],
            'total_attempts': total_attempts,
            'can_retake': can_retake,
            'retake_allowed_by_teacher': retake_allowed_by_teacher
        }), 200
        
    except Exception as e:
        logger.error(f"Error fetching retake history: {e}")
        return jsonify({'error': str(e)}), 500


# Helper function to get quiz title
def get_quiz_title(quiz_id, quizzes_list):
    """Helper function to get quiz title from list"""
    for quiz in quizzes_list:
        if quiz['id'] == quiz_id:
            return quiz['quiz_title']
    return "Unknown Quiz"


# Error handler for quiz routes
@app.errorhandler(404)
def not_found_error(error):
    return jsonify({'error': 'Resource not found'}), 404


@app.errorhandler(500)
def internal_error(error):
    logger.error(f"Internal server error: {error}")
    return jsonify({'error': 'Internal server error'}), 500

#############################STUDENT_REWARDS.HTML#############################
@app.route('/api/student-teachers-rewards', methods=['GET'])
def api_student_teachers_rewards():
    if 'user_id' not in session or session.get('role') != 'Student':
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401

    student_id = session['user_id']
    
    # Get student's grade and section
    user_result = supabase.table('user_info').select('year_level', 'section').eq('id', student_id).execute()
    if not user_result.data:
        return jsonify({'success': False, 'message': 'Student not found'}), 404
    
    grade = user_result.data[0]['year_level']
    section = user_result.data[0]['section']

    # Get teachers assigned to this grade/section
    # NOTE: Only fetch ACTIVE teachers for the dropdown (legacy view pattern)
    assignments = supabase.table('teacher_class_assignments') \
        .select('teacher_id, subject') \
        .eq('grade_level', grade) \
        .eq('section', section) \
        .eq('status', 'active') \
        .execute()
    
    teacher_ids = [a['teacher_id'] for a in assignments.data] if assignments.data else []

    # Get teacher info
    teachers = []
    teacher_map = {}
    
    if teacher_ids:
        teacher_result = supabase.table('user_info').select('id, first_name, last_name, gender, subject').in_('id', teacher_ids).execute()
        
        # Fetch avatars
        pics_map = {}
        pics_result = supabase.table('profile_pictures').select('user_id, file_path').in_('user_id', teacher_ids).execute()
        
        if pics_result.data:
            for p in pics_result.data:
                pics_map[p['user_id']] = p['file_path']
        
        for t in teacher_result.data:
            prefix = 'Mr.' if (t.get('gender', '').lower() == 'male') else 'Ms.'
            teacher_name = f"{prefix} {t['last_name']}"
            teachers.append({
                'id': t['id'],
                'name': teacher_name,
                'subject': t.get('subject', ''),
                'avatar': pics_map.get(t['id'], '/static/image/default-avatar.png')
            })
            teacher_map[t['id']] = {
                'name': teacher_name,
                'subject': t.get('subject', ''),
                'avatar': pics_map.get(t['id'], '/static/image/default-avatar.png')
            }

    # Helper: normalize stored grade_level/section into list of strings
    def to_list(val):
        if val is None:
            return []
        if isinstance(val, list):
            return [str(v).strip() for v in val if v is not None]
        if isinstance(val, int):
            return [str(val)]
        if isinstance(val, str):
            s = val.strip()
            # try parse JSON array like '["Faith"]' or '[7]'
            if s.startswith('['):
                try:
                    parsed = json.loads(s)
                    if isinstance(parsed, list):
                        return [str(v).strip() for v in parsed if v is not None]
                except Exception:
                    pass
            # comma-separated fallback
            if ',' in s:
                return [p.strip().strip('"').strip("'") for p in s.split(',') if p.strip()]
            if s == '':
                return []
            return [s]
        return [str(val)]

    # Get rewards created by these teachers
    rewards = []
    if teacher_ids:
        rewards_result = supabase.table('rewards').select('*').in_('created_by', teacher_ids).execute()
        
        # Get student's redemption history to check which rewards are already redeemed
        redemption_result = supabase.table('reward_redemptions') \
            .select('reward_id') \
            .eq('student_id', student_id) \
            .execute()
        
        redeemed_reward_ids = [r['reward_id'] for r in redemption_result.data] if redemption_result.data else []
        
        if rewards_result.data:
            for r in rewards_result.data:
                # Only include rewards that are available and have quantity > 0
                if r.get('status', 'Available') != 'Available':
                    continue
                if int(r.get('available_quantity', 0)) <= 0:
                    continue
                
                # Filter rewards to only those that apply to this student's grade & section
                reward_grade_list = [g.lower() for g in to_list(r.get('grade_level'))]
                reward_section_list = [s_.lower() for s_ in to_list(r.get('section'))]

                # Treat empty or 'all' as global (match everything)
                grade_match = (not reward_grade_list) or ('all' in reward_grade_list) or (str(grade).lower() in reward_grade_list)
                section_match = (not reward_section_list) or ('all' in reward_section_list) or (str(section).lower() in reward_section_list)

                if grade_match and section_match:
                    teacher_info = teacher_map.get(r['created_by'], {})
                    
                    # Get quarterly redemption info
                    redemption_info = check_can_redeem_reward(student_id, r['reward_id'])
                    
                    rewards.append({
                        **r,
                        'teacher_id': r['created_by'],
                        'teacher_name': teacher_info.get('name', ''),
                        'teacher_avatar': teacher_info.get('avatar', '/static/image/default-avatar.png'),
                        'teacher_subject': teacher_info.get('subject', ''),
                        'grade_level': r.get('grade_level', ''),
                        'section': r.get('section', ''),
                        'is_redeemed': r['reward_id'] in redeemed_reward_ids,  # Add redemption status
                        'redemption_count': redemption_info['count'],
                        'remaining_redeems': redemption_info['remaining'],
                        'can_redeem': redemption_info['can_redeem'],
                        'is_maxed_out': not redemption_info['can_redeem'],  # Add maxed_out flag for UI
                        'quarter_name': redemption_info['quarter_name']
                    })

    return jsonify({'success': True, 'teachers': teachers, 'rewards': rewards})

@app.route('/api/check-reward-limit/<int:reward_id>', methods=['GET'])
def api_check_reward_limit(reward_id):
    if 'user_id' not in session or session.get('role') != 'Student':
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401

    student_id = session['user_id']
    
    try:
        # Check if reward exists
        reward_result = supabase.table('rewards').select('*').eq('reward_id', reward_id).execute()
        if not reward_result.data:
            return jsonify({'success': False, 'message': 'Reward not found'}), 404
        
        # Check quarterly redemption limit
        redemption_info = check_can_redeem_reward(student_id, reward_id)
        
        return jsonify({
            'success': True,
            'reward_id': reward_id,
            'can_redeem': redemption_info['can_redeem'],
            'maxed_out': not redemption_info['can_redeem'],
            'redemption_count': redemption_info['count'],
            'remaining_redeems': redemption_info['remaining'],
            'quarter_name': redemption_info['quarter_name']
        })
    
    except Exception as e:
        app.logger.error(f"Error checking reward limit: {str(e)}")
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/api/redeem-reward', methods=['POST'])
def api_redeem_reward():
    if 'user_id' not in session or session.get('role') != 'Student':
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401

    data = request.get_json()
    reward_id = data.get('reward_id')
    student_id = session['user_id']

    # Fetch reward info
    reward_result = supabase.table('rewards').select('*').eq('reward_id', reward_id).execute()
    if not reward_result.data:
        return jsonify({'success': False, 'message': 'Reward not found'}), 404
    reward = reward_result.data[0]

    # Fetch student points
    user_result = supabase.table('user_info').select('total_points, first_name, last_name, year_level, section').eq('id', student_id).execute()
    if not user_result.data:
        return jsonify({'success': False, 'message': 'Student not found'}), 404
    user_info = user_result.data[0]
    current_points = int(user_info.get('total_points', 0))
    student_name = f"{user_info.get('first_name', '')} {user_info.get('last_name', '')}"
    grade_level = user_info.get('year_level', '')
    section = user_info.get('section', '')

    # Check if enough points
    if current_points < int(reward['point_cost']):
        return jsonify({'success': False, 'message': 'Insufficient points'}), 400

    # Check if reward is available
    if int(reward.get('available_quantity', 0)) <= 0:
        return jsonify({'success': False, 'message': 'Reward out of stock'}), 400

    # Check quarterly redemption limit (3 per quarter)
    redemption_info = check_can_redeem_reward(student_id, reward_id)
    if not redemption_info['can_redeem']:
        return jsonify({
            'success': False, 
            'message': f'You have already redeemed this reward 3 times this quarter ({redemption_info["quarter_name"]}). Please try again next quarter.'
        }), 400

    # Deduct points and update reward quantity (transactional logic)
    new_points = current_points - int(reward['point_cost'])
    new_quantity = int(reward['available_quantity']) - 1

    # --- Helper for safe insert ---
    def safe_execute(query):
        try:
            query.execute()
        except Exception as e:
            app.logger.error(f"Safe execute error: {str(e)}")

    try:
        # Begin transaction-like operations
        # Update student points first
        user_update_result = supabase.table('user_info').update({
            'total_points': new_points
        }).eq('id', student_id).execute()
        
        if not user_update_result.data:
            raise Exception("Failed to update user points")

        # Update reward quantity and status - automatically set to Unavailable if quantity is 0
        reward_update_result = supabase.table('rewards').update({
            'available_quantity': new_quantity,
            'status': 'Unavailable' if new_quantity == 0 else reward.get('status', 'Available')
        }).eq('reward_id', reward_id).execute()
        
        if not reward_update_result.data:
            # Rollback user points update
            supabase.table('user_info').update({
                'total_points': current_points
            }).eq('id', student_id).execute()
            raise Exception("Failed to update reward quantity")

        # Log redemption in reward_redemptions table (use Philippines timezone for consistency with quarters)
        now = datetime.now(philippines_tz).isoformat()
        redemption_result = supabase.table('reward_redemptions').insert({
            'student_id': student_id,
            'reward_id': reward_id,
            'processed_at': now,
            'points_deducted': reward['point_cost'],
            'status': 'Unused',
            'grade_level': grade_level,
            'section': section,
            'remarks': data.get('remarks', ''),
            'teacher_id': reward.get('created_by')  # <-- Add this line
        }).execute()

        if not redemption_result.data:
            # Rollback both updates
            supabase.table('user_info').update({
                'total_points': current_points
            }).eq('id', student_id).execute()
            supabase.table('rewards').update({
                'available_quantity': int(reward['available_quantity'])
            }).eq('reward_id', reward_id).execute()
            raise Exception("Failed to log redemption")

        # --- ACTIVITY LOG & NOTIFICATION (your requested format) ---
        redemption_id = redemption_result.data[0].get('redemption_id')
        reward_name = reward['reward_name']
        points_deducted = reward['point_cost']
        teacher_id = reward.get('created_by')

        safe_execute(supabase.table('admin_activity_log').insert({
            'user_id': int(student_id),
            'user_role': 'Student',
            'action': 'Redeem Reward',
            'activity': 'Reward Redemption',
            'description': f"{student_name} redeemed the reward '{reward_name}'.",
            'details': f"Reward: {reward_name}, Points Deducted: {points_deducted}, Grade: {grade_level}, Section: {section}",
        }))

        # --- Insert into notifications table for teacher/admin ---
        if teacher_id:
            safe_execute(supabase.table('notifications').insert({
                'user_id': teacher_id,  # teacher/admin who created the reward
                'sender_id': int(student_id),
                'title': 'Reward Redemption',
                'message': f"{student_name} redeemed the reward '{reward_name}'. Please process the claim.",
                'redemption_id': redemption_id,
                'notif_type': 'Reward',
                'status': 'Unread',
            }))

        # Notify teacher if quantity becomes 0
        if new_quantity == 0 and teacher_id:
            safe_execute(supabase.table('notifications').insert({
                'user_id': teacher_id,
                'sender_id': student_id,
                'title': "Reward Out of Stock",
                'message': f"Your reward '{reward_name}' is now out of stock and has been automatically unlisted.",
                'notif_type': 'Reward',
                'status': 'Unread',
                'reward_id': reward_id
            }))

        return jsonify({
            'success': True, 
            'message': 'Reward redeemed successfully!', 
            'new_points': new_points
        })

    except Exception as e:
        print(f"Error in reward redemption: {str(e)}")
        return jsonify({
            'success': False, 
            'message': 'Failed to process reward redemption. Please try again.'
        }), 500

@app.route('/api/teacher-quarterly-reward-stats', methods=['GET'])
def api_teacher_quarterly_reward_stats():
    """Get quarterly redemption statistics for teacher's rewards"""
    if 'user_id' not in session or session.get('role') != 'Teacher':
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401
    
    teacher_id = session['user_id']
    
    try:
        # Get current quarter
        current_quarter = get_current_quarter()
        if not current_quarter:
            return jsonify({'success': True, 'stats': [], 'current_quarter': None})
        
        quarter_id = current_quarter['quarter_id']
        quarter_name = current_quarter.get('quarter_name', 'N/A')
        start_date = current_quarter['start_date']
        end_date = current_quarter['end_date']
        
        # Get all rewards created by this teacher
        rewards_result = supabase.table('rewards').select('*').eq('created_by', teacher_id).execute()
        if not rewards_result.data:
            return jsonify({'success': True, 'stats': [], 'current_quarter': current_quarter})
        
        stats = []
        for reward in rewards_result.data:
            reward_id = reward['reward_id']
            
            # Get all redemptions for this reward in the current quarter
            redemptions = supabase.table('reward_redemptions') \
                .select('student_id, reward_id, processed_at', count='exact') \
                .eq('reward_id', reward_id) \
                .gte('processed_at', start_date) \
                .lte('processed_at', end_date) \
                .execute()
            
            # Count redemptions per student
            student_redemption_counts = {}
            if redemptions.data:
                for redemption in redemptions.data:
                    student_id = redemption['student_id']
                    student_redemption_counts[student_id] = student_redemption_counts.get(student_id, 0) + 1
            
            # Count students by their redemption count
            students_hit_limit = sum(1 for count in student_redemption_counts.values() if count >= 3)
            total_redemptions = sum(student_redemption_counts.values()) if student_redemption_counts else 0
            
            stats.append({
                'reward_id': reward_id,
                'reward_name': reward.get('reward_name', ''),
                'category': reward.get('category', ''),
                'point_cost': reward.get('point_cost', 0),
                'available_quantity': reward.get('available_quantity', 0),
                'status': reward.get('status', 'Available'),
                'total_redemptions_this_quarter': total_redemptions,
                'students_redeemed': len(student_redemption_counts),
                'students_hit_limit': students_hit_limit,
                'redemption_breakdown': student_redemption_counts
            })
        
        return jsonify({
            'success': True,
            'current_quarter': current_quarter,
            'quarter_name': quarter_name,
            'stats': stats
        })
    
    except Exception as e:
        app.logger.error(f"Error getting quarterly reward stats: {str(e)}")
        return jsonify({
            'success': False,
            'message': 'Failed to load quarterly stats'
        }), 500

@app.route('/api/student-redemption-history')
def api_student_redemption_history():
    if 'user_id' not in session or session.get('role') != 'Student':
        return jsonify({'success': False, 'history': []})
    student_id = session['user_id']

    # Fetch redemptions and include nested reward creator info
    # ✅ Include 'remarks' in the SELECT query
    result = supabase.table('reward_redemptions') \
        .select('reward_id, points_deducted, processed_at, status, teacher_id, notes, used_at, remarks, rewards(reward_name, description, category, point_cost, created_by)') \
        .eq('student_id', student_id) \
        .order('processed_at', desc=True) \
        .execute()
    redemptions = result.data if result.data else []

    # Collect teacher IDs from either redemption.teacher_id OR rewards.created_by
    teacher_ids_set = set()
    for r in redemptions:
        if r.get('teacher_id'):
            teacher_ids_set.add(r['teacher_id'])
        else:
            reward = r.get('rewards') or {}
            if reward.get('created_by'):
                teacher_ids_set.add(reward.get('created_by'))
    teacher_ids = [tid for tid in teacher_ids_set if tid is not None]

    # Fetch teacher info + avatars in bulk
    teacher_map = {}
    if teacher_ids:
        teacher_result = supabase.table('user_info').select('id, first_name, last_name, gender, subject').in_('id', teacher_ids).execute()
        pics_map = {}
        pics_result = supabase.table('profile_pictures').select('user_id, file_path').in_('user_id', teacher_ids).execute()
        if pics_result.data:
            for p in pics_result.data:
                pics_map[p['user_id']] = p.get('file_path')
        for t in teacher_result.data or []:
            full_name = f"{t.get('first_name','').strip()} {t.get('last_name','').strip()}".strip()
            teacher_map[t['id']] = {
                'name': full_name if full_name else 'Unknown',
                'subject': t.get('subject', '') or '',
                'avatar': pics_map.get(t['id'], '/static/image/default-avatar.png')
            }

    # Build history payload, preferring redemption.teacher_id then reward.created_by
    history = []
    for r in redemptions:
        reward = r.get('rewards', {}) or {}
        # prefer explicit teacher_id stored on redemption; fallback to reward.created_by
        teacher_ref_id = r.get('teacher_id') or reward.get('created_by')
        teacher_info = teacher_map.get(teacher_ref_id, {
            'name': 'Unknown',
            'subject': '',
            'avatar': '/static/image/default-avatar.png'
        })

        history.append({
            'reward_name': reward.get('reward_name', ''),
            'description': reward.get('description', ''),
            'category': reward.get('category', ''),
            'point_cost': reward.get('point_cost', ''),
            'teacher': teacher_info['name'],
            'teacher_id': teacher_ref_id,
            'teacher_avatar': teacher_info['avatar'],
            'teacher_subject': teacher_info['subject'],
            'status': r.get('status', 'Unused'),
            'date': format_date(r.get('processed_at', '')),
            'notes': r.get('notes', ''),
            'remarks': r.get('remarks', ''),  # ✅ This will now have data
            'used_at': r.get('used_at', ''),
        })

    return jsonify({'success': True, 'history': history})

#############################STUDENT_LEADERBOARDS.HTML#############################
@app.route('/api/student-leaderboard', methods=['GET'])
def api_student_leaderboard():
    # Check authentication
    if 'user_id' not in session:
        print("No user_id in session")
        return jsonify({'success': False, 'error': 'Not authenticated'}), 401
    
    if session.get('role') != 'Student':
        print(f"Wrong role: {session.get('role')}")
        return jsonify({'success': False, 'error': 'Access denied - Students only'}), 403

    student_id = session['user_id']
    period = request.args.get('period', 'current')
    school_year = request.args.get('school_year', None)
    
    print(f"Fetching leaderboard for student_id: {student_id}, period: {period}, school_year: {school_year}")

    try:
        # Get all unique school years from quarters to determine current school year
        quarters_resp = supabase.table('quarters')\
            .select('school_year')\
            .order('school_year', desc=True)\
            .limit(1)\
            .execute()
        
        current_school_year = quarters_resp.data[0]['school_year'] if quarters_resp.data else None
        
        # If no school_year specified, use current
        if not school_year:
            school_year = current_school_year
        
        print(f"Current school year: {current_school_year}, Requested: {school_year}")
        
        # If requesting a past school year, use archived data
        if school_year != current_school_year:
            print(f"Loading archived data for school year {school_year}")
            
            # Get archived student data for this school year
            archived_resp = supabase.table('student_archives')\
                .select('*')\
                .eq('student_id', student_id)\
                .eq('school_year', school_year)\
                .execute()
            
            if not archived_resp.data:
                print(f"No archived data found for student {student_id} in {school_year}")
                return jsonify({
                    'success': True, 
                    'students': [], 
                    'stats': {'total': 0, 'avg': 0, 'top': 0},
                    'school_year': school_year,
                    'is_archived': True
                })
            
            # Get the archived student's grade and section from the archive
            archived_student = archived_resp.data[0]
            grade = archived_student.get('year_level')
            section = archived_student.get('section')
            
            print(f"Archived student grade: {grade}, section: {section}")
            
            # Get all students who were archived in the same classroom during this school year
            archived_students_resp = supabase.table('student_archives')\
                .select('student_id, year_level, section, total_points, archived_at')\
                .eq('year_level', grade)\
                .eq('section', section)\
                .eq('school_year', school_year)\
                .execute()
            
            students_with_names = []
            if archived_students_resp.data:
                # Get names from user_info for each archived student
                for archived_student in archived_students_resp.data:
                    student_id = archived_student.get('student_id')
                    user_info_resp = supabase.table('user_info')\
                        .select('id, first_name, last_name')\
                        .eq('id', student_id)\
                        .execute()
                    
                    if user_info_resp.data:
                        user_info = user_info_resp.data[0]
                        # Merge archived data with current user info
                        merged = {
                            **archived_student,
                            'id': student_id,
                            'first_name': user_info.get('first_name', ''),
                            'last_name': user_info.get('last_name', ''),
                        }
                        students_with_names.append(merged)
            
            students = students_with_names
            print(f"Found {len(students)} archived classmates for {school_year}")
        
        else:
            # Current school year - use active students
            # Get student's classroom info
            user_result = supabase.table('user_info')\
                .select('year_level, section')\
                .eq('id', student_id)\
                .execute()
            
            if not user_result.data:
                print(f"Student not found: {student_id}")
                return jsonify({'success': False, 'error': 'Student not found'}), 404
            
            student_info = user_result.data[0]
            grade = student_info.get('year_level')
            section = student_info.get('section')
            
            print(f"Student grade: {grade}, section: {section}")
            
            if not grade or not section:
                print("Missing grade or section")
                return jsonify({'success': False, 'error': 'Student grade or section not set'}), 400

            # Get all students in the same classroom
            students_resp = supabase.table('user_info')\
                .select('id, first_name, last_name, section, year_level, total_points')\
                .eq('section', section)\
                .eq('year_level', grade)\
                .eq('role', 'Student')\
                .eq('status', 'Active')\
                .execute()
            
            students = students_resp.data if students_resp.data else []
            print(f"Found {len(students)} classmates")
        
        if not students:
            return jsonify({
                'success': True, 
                'students': [], 
                'stats': {'total': 0, 'avg': 0, 'top': 0},
                'school_year': school_year,
                'is_archived': school_year != current_school_year
            })

        # Get profile pictures for all students
        student_ids = [s['id'] if 'id' in s else s['student_id'] for s in students]
        pics_map = {}
        
        if student_ids:
            pics_result = supabase.table('profile_pictures')\
                .select('user_id, file_path')\
                .in_('user_id', student_ids)\
                .execute()
            
            if pics_result.data:
                for pic in pics_result.data:
                    pics_map[pic['user_id']] = pic['file_path']

        # Prepare leaderboard data based on period
        from datetime import datetime, timedelta, timezone
        
        leaderboard = []
        
        if period == 'current':
            # Use total_points
            for student in students:
                user_id = student['id'] if 'id' in student else student['student_id']
                first_name = student.get('first_name', '').strip()
                last_name = student.get('last_name', '').strip()
                name = f"{first_name} {last_name}".strip()
                
                if not name:
                    name = f"Student {user_id}"
                
                profile_picture = pics_map.get(user_id, '/static/image/default-avatar.png')
                
                leaderboard.append({
                    'id': user_id,
                    'name': name,
                    'section': student.get('section', ''),
                    'year_level': student.get('year_level', ''),
                    'avatar': profile_picture,
                    'points': int(student.get('total_points', 0) or 0)
                })
        else:
            # Handle quarter periods and other time periods
            now = datetime.now(timezone.utc)
            start_date = None
            end_date = None
            period_not_started = False
            
            if period == 'week':
                # Get points from last 7 days (including today)
                start_date = now - timedelta(days=7)
                period_not_started = False
                print(f"Week period: {start_date} to {now}")
            elif period == 'month':
                start_date = now.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
                period_not_started = False
                print(f"Month period: {start_date} to {now}")
            elif period == 'all':
                period_not_started = False
                print("All time period: no date filtering")
            elif period.startswith('quarter'):
                # Handle quarter periods
                # Extract the quarter ID (everything after 'quarter' prefix)
                quarter_id = period.replace('quarter', '')
                print(f"Looking for quarter ID: {quarter_id}")
                
                # Get the specific quarter by ID
                quarter_result = supabase.table('quarters')\
                    .select('*')\
                    .eq('quarter_id', quarter_id)\
                    .execute()
                
                print(f"Quarter lookup result: {quarter_result.data}")
                
                target_quarter = None
                if quarter_result.data:
                    target_quarter = quarter_result.data[0]
                    print(f"Found target quarter: {target_quarter['quarter_name']}")
                
                if target_quarter:
                    start_date_str = target_quarter.get('start_date')
                    end_date_str = target_quarter.get('end_date')
                    
                    print(f"Found target quarter: {target_quarter['quarter_name']}")
                    print(f"Start date: {start_date_str}, End date: {end_date_str}")
                    
                    if start_date_str and end_date_str:
                        try:
                            # Parse the dates
                            start_date = datetime.fromisoformat(start_date_str.replace('Z', '+00:00'))
                            end_date = datetime.fromisoformat(end_date_str.replace('Z', '+00:00'))
                            print(f"Quarter {target_quarter.get('quarter_name')} date range: {start_date} to {end_date}")
                            
                            # Check if quarter has started - if not, don't return any points
                            if start_date > now:
                                print(f"⏳ Quarter {target_quarter.get('quarter_name')} hasn't started yet (starts on {start_date})")
                                # Quarter hasn't started - no points should be shown
                                start_date = None
                                end_date = None
                                # Set a flag to skip points lookup
                                period_not_started = True
                            else:
                                period_not_started = False
                        except Exception as e:
                            print(f"Error parsing quarter dates: {e}")
                else:
                    print(f"No quarter found for quarter ID {quarter_id}")
                    period_not_started = True
            
            # For 'all', start_date remains None (get all points)
            
            for student in students:
                user_id = student['id'] if 'id' in student else student['student_id']
                first_name = student.get('first_name', '').strip()
                last_name = student.get('last_name', '').strip()
                name = f"{first_name} {last_name}".strip()
                
                if not name:
                    name = f"Student {user_id}"
                
                # Always query points table (not points_archives)
                # BUT skip if period hasn't started yet
                if period_not_started:
                    total_points = 0
                    print(f"Student {user_id} - Period hasn't started, returning 0 points")
                else:
                    points_query = supabase.table('points')\
                        .select('points, received_at')\
                        .eq('student_id', user_id)\
                        .eq('status', 'approved')
                    
                    if start_date:
                        points_query = points_query.gte('received_at', start_date.isoformat())
                    
                    if end_date:
                        points_query = points_query.lte('received_at', end_date.isoformat())
                    
                    points_resp = points_query.execute()
                    
                    # For archived data, milestone_claims may not exist, so only query for current year
                    milestone_resp_data = []
                    if school_year == current_school_year:
                        # Query milestone claims table only for current school year
                        milestone_query = supabase.table('milestone_claims')\
                            .select('points_awarded, claimed_at')\
                            .eq('student_id', user_id)
                        
                        if start_date:
                            milestone_query = milestone_query.gte('claimed_at', start_date.isoformat())
                        
                        if end_date:
                            milestone_query = milestone_query.lte('claimed_at', end_date.isoformat())
                        
                        milestone_resp = milestone_query.execute()
                        milestone_resp_data = milestone_resp.data if milestone_resp.data else []
                    
                    # Calculate total points from both sources
                    total_points = 0
                    
                    # Add regular points
                    if points_resp.data:
                        regular_points = sum([p['points'] for p in points_resp.data])
                        total_points += regular_points
                        print(f"Student {user_id} - {len(points_resp.data)} points records: {regular_points}")
                    
                    # Add milestone points
                    if milestone_resp_data:
                        milestone_points = sum([m['points_awarded'] for m in milestone_resp_data])
                        total_points += milestone_points
                        print(f"Student {user_id} - {len(milestone_resp_data)} milestone records: {milestone_points}")
                
                print(f"Student {user_id} - Total points: {total_points}")
                
                profile_picture = pics_map.get(user_id, '/static/image/default-avatar.png')
                
                leaderboard.append({
                    'id': user_id,
                    'name': name,
                    'section': student.get('section', ''),
                    'year_level': student.get('year_level', ''),
                    'avatar': profile_picture,
                    'points': total_points
                })

        # Sort by points descending
        leaderboard.sort(key=lambda x: x['points'], reverse=True)
        
        # Calculate stats
        total_students = len(leaderboard)
        total_points = sum(s['points'] for s in leaderboard)
        avg_points = total_points // total_students if total_students > 0 else 0
        top_points = leaderboard[0]['points'] if leaderboard else 0
        
        print(f"Returning {total_students} students, avg: {avg_points}, top: {top_points}")
        
        return jsonify({
            'success': True, 
            'students': leaderboard,
            'stats': {
                'total': total_students,
                'avg': avg_points,
                'top': top_points
            },
            'school_year': school_year,
            'current_school_year': current_school_year,
            'is_archived': school_year != current_school_year
        })

    except Exception as e:
        print(f"Error in leaderboard API: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)}), 500
    
@app.route('/api/weekly-points/<int:student_id>')
def api_weekly_points(student_id):
    from datetime import datetime, timedelta, timezone
    import pytz
    
    # Get current time in Philippines timezone
    philippines_tz = pytz.timezone('Asia/Manila')
    now_ph = datetime.now(philippines_tz)
    
    # Calculate start of current week (Monday in PH time)
    start_of_week = now_ph - timedelta(days=now_ph.weekday())
    start_of_week = start_of_week.replace(hour=0, minute=0, second=0, microsecond=0)
    
    # Calculate end of current week (Sunday in PH time)
    end_of_week = start_of_week + timedelta(days=7)
    
    # Prepare 7 days (Mon-Sun)
    daily_points = [0] * 7
    
    # Convert to UTC strings for Supabase query
    start_str = start_of_week.astimezone(pytz.utc).isoformat()
    end_str = end_of_week.astimezone(pytz.utc).isoformat()
    
    print(f"Weekly points calculation for student {student_id}")
    print(f"PH Time: {now_ph.strftime('%Y-%m-%d %H:%M:%S %Z')}")
    print(f"Week range (PH): {start_of_week.strftime('%Y-%m-%d')} to {end_of_week.strftime('%Y-%m-%d')}")
    print(f"Query range (UTC): {start_str} to {end_str}")
    
    # Fetch all regular points for this student THIS WEEK ONLY
    points_result = supabase.table('points').select('points, received_at').eq('student_id', student_id).eq('status', 'approved').gte('received_at', start_str).lt('received_at', end_str).execute()
    if points_result.data:
        for row in points_result.data:
            try:
                received = datetime.fromisoformat(str(row['received_at']).replace('Z', '+00:00'))
                if received.tzinfo is None:
                    received = received.replace(tzinfo=timezone.utc)
                
                # Convert to PH time to get the correct day of week
                received_ph = received.astimezone(philippines_tz)
                day_index = (received_ph.weekday()) % 7  # Monday=0, Sunday=6
                daily_points[day_index] += row['points']
                
                print(f"Points record: {received_ph.strftime('%A %Y-%m-%d')} - {row['points']} pts (day_index: {day_index})")
            except Exception as e:
                print(f"Error processing point: {e}")
                continue
    
    # Fetch all milestone claims for this student THIS WEEK ONLY
    milestone_result = supabase.table('milestone_claims').select('points_awarded, claimed_at').eq('student_id', student_id).gte('claimed_at', start_str).lt('claimed_at', end_str).execute()
    if milestone_result.data:
        for row in milestone_result.data:
            try:
                claimed = datetime.fromisoformat(str(row['claimed_at']).replace('Z', '+00:00'))
                if claimed.tzinfo is None:
                    claimed = claimed.replace(tzinfo=timezone.utc)
                
                # Convert to PH time to get the correct day of week
                claimed_ph = claimed.astimezone(philippines_tz)
                day_index = (claimed_ph.weekday()) % 7  # Monday=0, Sunday=6
                daily_points[day_index] += row['points_awarded']
                
                print(f"Milestone record: {claimed_ph.strftime('%A %Y-%m-%d')} - {row['points_awarded']} pts (day_index: {day_index})")
            except Exception as e:
                print(f"Error processing milestone: {e}")
                continue
    
    print(f"Final daily_points: {daily_points}")
    
    return jsonify({'success': True, 'weekly_points': daily_points})

@app.route('/api/activities-completed/<int:student_id>', methods=['GET'])
def api_activities_completed(student_id):
    # Check authentication
    if 'user_id' not in session:
        return jsonify({'success': False, 'error': 'Not authenticated'}), 401
    
    if session.get('role') != 'Student':
        return jsonify({'success': False, 'error': 'Access denied - Students only'}), 403

    period = request.args.get('period', 'current')
    
    print(f"Fetching completed activities for student_id: {student_id}, period: {period}")

    try:
        from datetime import datetime, timedelta, timezone
        
        # Base query for completed tasks
        query = supabase.table('task_assignments')\
            .select('*', count='exact')\
            .eq('student_id', student_id)\
            .eq('status', 'Completed')
        
        # Apply time period filter
        if period == 'current':
            # Get current quarter and filter by its date range
            quarter_response = supabase.table('quarters')\
                .select('*')\
                .eq('quarter_name', '2nd Quarter')\
                .eq('school_year', '2025-2026')\
                .execute()
            
            if quarter_response.data and len(quarter_response.data) > 0:
                current_quarter = quarter_response.data[0]
                quarter_start = current_quarter['start_date']
                quarter_end = current_quarter['end_date']
                
                print(f"Current quarter: {current_quarter['quarter_name']} from {quarter_start} to {quarter_end}")
                
                # Filter tasks completed within the current quarter
                query = query.gte('completed_at', quarter_start)\
                            .lte('completed_at', quarter_end)
            else:
                # If no current quarter found, return 0
                print("No current quarter found")
                return jsonify({
                    'success': True, 
                    'activities_completed': 0
                })
                
        elif period != 'all':
            now = datetime.now(timezone.utc)
            
            if period == 'week':
                # Get tasks completed in the last 7 days
                start_date = now - timedelta(days=7)
                query = query.gte('completed_at', start_date.isoformat())
            elif period == 'month':
                # Get tasks completed this month
                start_of_month = now.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
                query = query.gte('completed_at', start_of_month.isoformat())
            elif period.startswith('quarter'):
                # Handle specific quarter selection (quarterUUID format)
                quarter_id = period.replace('quarter', '')
                
                quarter_response = supabase.table('quarters')\
                    .select('*')\
                    .eq('quarter_id', quarter_id)\
                    .execute()
                
                if quarter_response.data and len(quarter_response.data) > 0:
                    quarter = quarter_response.data[0]
                    quarter_start = quarter['start_date']
                    quarter_end = quarter['end_date']
                    
                    query = query.gte('completed_at', quarter_start)\
                                .lte('completed_at', quarter_end)
        
        # Execute query
        result = query.execute()
        
        completed_count = len(result.data) if result.data else 0
        
        print(f"Found {completed_count} completed activities for student {student_id} in period {period}")
        
        return jsonify({
            'success': True, 
            'activities_completed': completed_count
        })

    except Exception as e:
        print(f"Error fetching completed activities: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)}), 500


#############################LEADERBOARD SCHOOL YEAR ENDPOINTS#############################
@app.route('/api/student-leaderboard-school-years', methods=['GET'])
def api_student_leaderboard_school_years():
    """Get all available school years for the leaderboard (current and archived)"""
    if 'user_id' not in session or session.get('role') != 'Student':
        return jsonify({'success': False, 'error': 'Not authenticated'}), 403
    
    try:
        # Get all unique school years from quarters table
        quarters_resp = supabase.table('quarters')\
            .select('school_year')\
            .execute()
        
        if not quarters_resp.data:
            return jsonify({'success': True, 'school_years': []})
        
        # Extract unique school years and sort in descending order
        school_years = sorted(list(set([q['school_year'] for q in quarters_resp.data])), reverse=True)
        
        # Get current school year
        current_year_resp = supabase.table('quarters')\
            .select('school_year')\
            .order('school_year', desc=True)\
            .limit(1)\
            .execute()
        
        current_school_year = current_year_resp.data[0]['school_year'] if current_year_resp.data else None
        
        return jsonify({
            'success': True,
            'school_years': school_years,
            'current_school_year': current_school_year
        })
    
    except Exception as e:
        logger.error(f"Error getting school years: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/student-leaderboard-quarters', methods=['GET'])
def api_student_leaderboard_quarters():
    """Get quarters for a specific school year"""
    if 'user_id' not in session or session.get('role') != 'Student':
        return jsonify({'success': False, 'error': 'Not authenticated'}), 403
    
    school_year = request.args.get('school_year', None)
    
    # DEBUG: COMPREHENSIVE LOGGING (ASCII-only to avoid encoding issues on Windows)
    logger.info("\n" + "="*80)
    logger.info("[API] /api/student-leaderboard-quarters ENDPOINT CALLED")
    logger.info("="*80)
    logger.info(f"[API] Received school_year: '{school_year}'")
    logger.info(f"[API]   Type: {type(school_year)}")
    logger.info(f"[API]   Length: {len(school_year) if school_year else 0}")
    logger.info(f"[API]   Repr: {repr(school_year)}")
    
    try:
        if not school_year:
            logger.error("[API] Error: school_year parameter is missing")
            return jsonify({'success': False, 'error': 'school_year parameter required'}), 400
        
        # FIRST: Get ALL quarters to see what's in the database
        logger.info(f"[API] Fetching ALL quarters from database to check what exists...")
        all_quarters_resp = supabase.table('quarters').select('quarter_id, quarter_name, school_year').execute()
        
        if all_quarters_resp.data:
            logger.info(f"[API] Total quarters in database: {len(all_quarters_resp.data)}")
            logger.info(f"[API] All quarters by school_year:")
            for q in all_quarters_resp.data:
                logger.info(f"[API]   - ID: {q.get('quarter_id')}, Name: {q.get('quarter_name')}, Year: '{q.get('school_year')}'")
        else:
            logger.warning(f"[API] WARNING: NO QUARTERS FOUND IN DATABASE AT ALL!")
        
        # SECOND: Get quarters for this school year
        logger.info(f"[API] Now querying for school_year='{school_year}'...")
        quarters_resp = supabase.table('quarters')\
            .select('*')\
            .eq('school_year', school_year)\
            .order('start_date')\
            .execute()
        
        result_count = len(quarters_resp.data) if quarters_resp.data else 0
        logger.info(f"[API] Query result: {result_count} quarters found")
        
        if quarters_resp.data:
            for q in quarters_resp.data:
                logger.info(f"[API]   - Quarter {q.get('quarter_id')}: {q.get('quarter_name')}")
        
        if not quarters_resp.data:
            logger.info(f"[API] Returning empty quarters array for school_year='{school_year}'")
            return jsonify({'success': True, 'quarters': []})
        
        logger.info(f"[API] Processing {result_count} quarters...")
        
        # Add current status for each quarter
        quarters_with_status = []
        now = datetime.now(timezone.utc)
        
        for q in quarters_resp.data:
            try:
                # Parse start and end dates from database
                start_str = str(q['start_date']).replace('Z', '+00:00')
                end_str = str(q['end_date']).replace('Z', '+00:00')
                
                # Parse with timezone info
                start = datetime.fromisoformat(start_str)
                end = datetime.fromisoformat(end_str)
                
                # If dates are naive (no timezone), make them UTC-aware
                if start.tzinfo is None:
                    start = start.replace(tzinfo=timezone.utc)
                if end.tzinfo is None:
                    end = end.replace(tzinfo=timezone.utc)
                
                # Now safe to compare with timezone-aware 'now'
                if now < start:
                    status = 'upcoming'
                    days_until = (start - now).days
                elif now > end:
                    status = 'finished'
                    days_until = None
                else:
                    status = 'active'
                    days_until = None
                
                quarters_with_status.append({
                    'quarter_id': q['quarter_id'],
                    'quarter_name': q['quarter_name'],
                    'start_date': q['start_date'],
                    'end_date': q['end_date'],
                    'school_year': q['school_year'],
                    'status': status,
                    'days_until': days_until
                })
            except Exception as e:
                logger.error(f"[API] Error processing quarter: {e}")
                continue
        
        logger.info(f"[API] Returning {len(quarters_with_status)} quarters with status")
        
        return jsonify({
            'success': True,
            'quarters': quarters_with_status
        })
    
    except Exception as e:
        logger.error(f"[API] EXCEPTION in /api/student-leaderboard-quarters: {str(e)}")
        logger.error(f"[API] Exception type: {type(e).__name__}")
        return jsonify({'success': False, 'error': str(e)}), 500


#############################STUDENT_PROFILE.HTML#############################
@app.route('/api/student-profile-summary', methods=['GET'])
def api_student_profile_summary():
    if 'user_id' not in session or session.get('role') != 'Student':
        return jsonify({'success': False, 'message': 'Unauthorized'}), 403

    student_id = session['user_id']

    try:
        # Get student's classroom and current totals
        user_resp = supabase.table('user_info') \
            .select('year_level, section, total_points, streak') \
            .eq('id', student_id) \
            .execute()
        if not user_resp.data:
            return jsonify({'success': False, 'message': 'User not found'}), 404

        user = user_resp.data[0]
        grade = user.get('year_level')
        section = user.get('section')
        total_points = int(user.get('total_points') or 0)
        streak = int(user.get('streak') or 0)

        # If classroom info missing, still return points/streak with no rank
        if not grade or not section:
            return jsonify({
                'success': True,
                'streak': streak,
                'points': total_points,
                'rank': None,
                'total_students': 0,
                'top_points': total_points
            })

        # Fetch classmates (active students in same grade & section)
        classmates_resp = supabase.table('user_info') \
            .select('id, total_points') \
            .eq('role', 'Student') \
            .eq('year_level', grade) \
            .eq('section', section) \
            .eq('status', 'Active') \
            .execute()
        classmates = classmates_resp.data if classmates_resp.data else []

        # Build leaderboard (points based on total_points for current period)
        leaderboard = []
        for s in classmates:
            leaderboard.append({
                'id': s['id'],
                'points': int(s.get('total_points') or 0)
            })

        # Sort descending and compute rank
        leaderboard.sort(key=lambda x: x['points'], reverse=True)
        rank = next((idx + 1 for idx, s in enumerate(leaderboard) if s['id'] == student_id), None)
        total_students = len(leaderboard)
        top_points = leaderboard[0]['points'] if leaderboard else total_points

        return jsonify({
            'success': True,
            'streak': streak,
            'points': total_points,
            'rank': rank,
            'total_students': total_students,
            'top_points': top_points
        })

    except Exception as e:
        app.logger.error(f"Error fetching student profile summary: {str(e)}")
        return jsonify({'success': False, 'message': 'Server error'}), 500

@app.route('/api/student-all-activities', methods=['GET'])
def api_student_all_activities():
    if 'user_id' not in session or session.get('role') != 'Student':
        return jsonify({'success': False, 'activities': []}), 401
    
    user_id = session['user_id']
    page = request.args.get('page', 1, type=int)
    activity_type = request.args.get('type', 'all')
    date_filter = request.args.get('date', 'all')
    
    limit = 20  # Activities per page
    offset = (page - 1) * limit
    
    def get_teacher_name_with_prefix(teacher_data):
        """Get teacher name with appropriate gender prefix"""
        prefix = 'Mr.'
        if teacher_data.get('gender', '').lower() == 'female':
            prefix = 'Mrs.'
        elif teacher_data.get('gender', '').lower() == 'other':
            prefix = 'Mx.'
        
        return f"{prefix} {teacher_data['first_name']} {teacher_data['last_name']}".strip()

    try:
        all_activities = []
        
        # --- Completed/denied tasks ---
        tasks_resp = safe_execute(
            supabase.table('task_assignments').select(
                'task, points, status, completed_at, due_date, awarded_by'
            ).eq('student_id', user_id).in_('status', ['Completed', 'Denied']).order('completed_at', desc=True)
        )
        if tasks_resp.data:
            # Get teacher names for tasks
            task_teacher_ids = list(set([t['awarded_by'] for t in tasks_resp.data if t.get('awarded_by')]))
            teachers_map = {}
            
            if task_teacher_ids:
                teachers_resp = safe_execute(
                    supabase.table('user_info').select('id, first_name, last_name, gender').in_('id', task_teacher_ids)
                )
                if teachers_resp.data:
                    for teacher in teachers_resp.data:
                        teacher_name = get_teacher_name_with_prefix(teacher)
                        teachers_map[teacher['id']] = teacher_name
            
            for t in tasks_resp.data:
                teacher_name = "Teacher"
                if t.get('awarded_by') and t['awarded_by'] in teachers_map:
                    teacher_name = teachers_map[t['awarded_by']]
                
                raw_time = t.get('completed_at', t.get('due_date', ''))
                all_activities.append({
                    'type': 'task',
                    'title': t['task'],
                    'points': t['points'],
                    'status': f'Awarded by {teacher_name}',
                    'time': format_date(raw_time),
                    'raw_time': raw_time  # Keep for filtering/sorting
                })

        # --- Reward redemptions ---
        rewards_resp = safe_execute(
            supabase.table('reward_redemptions').select(
                'reward_id, points_deducted, processed_at, rewards(reward_name)'
            ).eq('student_id', user_id).order('processed_at', desc=True)
        )
        if rewards_resp.data:
            for r in rewards_resp.data:
                reward_name = ''
                if r.get('rewards') and r['rewards'].get('reward_name'):
                    reward_name = r['rewards']['reward_name']
                
                raw_time = r.get('processed_at', '')
                all_activities.append({
                    'type': 'reward',
                    'title': f"Redeemed: {reward_name}",
                    'points': -abs(r['points_deducted']),
                    'status': 'Redeemed',
                    'time': format_date(raw_time),
                    'raw_time': raw_time  # Keep for filtering/sorting
                })

        # --- Star awards ---
        stars_resp = safe_execute(
            supabase.table('points').select(
                'points, received_at, note, point_category, teacher_id'
            ).eq('student_id', user_id).eq('status', 'approved').order('received_at', desc=True)
        )
        if stars_resp.data:
            # Get teacher names for star awards
            star_teacher_ids = list(set([s['teacher_id'] for s in stars_resp.data if s.get('teacher_id')]))
            star_teachers_map = {}
            
            if star_teacher_ids:
                star_teachers_resp = safe_execute(
                    supabase.table('user_info').select('id, first_name, last_name, gender').in_('id', star_teacher_ids)
                )
                if star_teachers_resp.data:
                    for teacher in star_teachers_resp.data:
                        teacher_name = get_teacher_name_with_prefix(teacher)
                        star_teachers_map[teacher['id']] = teacher_name
            
            for s in stars_resp.data:
                # Prefer note, then point_category, then fallback
                if s.get('note'):
                    title = s['note']
                elif s.get('point_category'):
                    title = f"Awarded for {s['point_category'].capitalize()}"
                else:
                    title = "Teacher Award"
                
                # Get teacher name who gave the award
                teacher_name = "Teacher"
                if s.get('teacher_id') and s['teacher_id'] in star_teachers_map:
                    teacher_name = star_teachers_map[s['teacher_id']]
                
                raw_time = s.get('received_at', '')
                all_activities.append({
                    'type': 'star',
                    'title': title,
                    'points': s['points'],
                    'status': f'Awarded by {teacher_name}',
                    'time': format_date(raw_time),
                    'raw_time': raw_time  # Keep for filtering/sorting
                })

        # --- Milestone claims ---
        milestone_claims_resp = safe_execute(
            supabase.table('milestone_claims')
            .select('milestone_type, milestone, points_awarded, claimed_at')
            .eq('student_id', user_id)
            .order('claimed_at', desc=True)
        )
        if milestone_claims_resp.data:
            for m in milestone_claims_resp.data:
                raw_time = m.get('claimed_at', '')
                all_activities.append({
                    'type': 'milestone',
                    'title': f"Claimed {m['milestone_type'].capitalize()} Milestone: {m['milestone']}",
                    'points': m.get('points_awarded', 0),
                    'status': 'Claimed',
                    'time': format_date(raw_time),
                    'raw_time': raw_time  # Keep for filtering/sorting
                })

        # --- Streak milestones ---
        streak_milestone_notifs = safe_execute(
            supabase.table('notifications')
            .select('title, message, created_at')
            .eq('user_id', user_id)
            .eq('title', 'Streak Milestone Unlocked!')
            .order('created_at', desc=True)
        )

        if streak_milestone_notifs.data:
            for notif in streak_milestone_notifs.data:
                import re
                milestone_match = re.search(r'the (\d+) streak milestone', notif.get('message', ''))
                milestone_num = milestone_match.group(1) if milestone_match else '?'
                
                raw_time = notif.get('created_at', '')
                all_activities.append({
                    'type': 'streak',
                    'title': f"Streak Milestone: {milestone_num} days!",
                    'points': 0,
                    'status': 'Milestone',
                    'time': format_date(raw_time),
                    'raw_time': raw_time  # Keep for filtering/sorting
                })

        # Sort all activities by raw_time (newest first)
        all_activities.sort(key=lambda x: x.get('raw_time', ''), reverse=True)

        # Apply filters
        filtered_activities = []
        for activity in all_activities:
            # Type filter
            if activity_type != 'all' and activity['type'] != activity_type:
                continue
            
            # Date filter
            if date_filter != 'all':
                activity_time = activity.get('raw_time')
                if activity_time:
                    from datetime import datetime, timedelta
                    activity_date = datetime.fromisoformat(activity_time.replace('Z', '+00:00'))
                    now = datetime.now(activity_date.tzinfo)
                    
                    if date_filter == 'today':
                        start_date = now.replace(hour=0, minute=0, second=0, microsecond=0)
                        if activity_date < start_date:
                            continue
                    elif date_filter == 'week':
                        start_date = now - timedelta(days=now.weekday())
                        start_date = start_date.replace(hour=0, minute=0, second=0, microsecond=0)
                        if activity_date < start_date:
                            continue
                    elif date_filter == 'month':
                        start_date = now.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
                        if activity_date < start_date:
                            continue
                
            filtered_activities.append(activity)
        
        # Paginate results
        start_idx = offset
        end_idx = offset + limit
        paginated_activities = filtered_activities[start_idx:end_idx]
        
        # Remove raw_time from final response (keep only formatted time)
        for activity in paginated_activities:
            activity.pop('raw_time', None)
        
        has_more = len(filtered_activities) > end_idx
        
        return jsonify({
            'success': True, 
            'activities': paginated_activities,
            'has_more': has_more,
            'total': len(filtered_activities)
        })
        
    except Exception as e:
        return jsonify({'success': False, 'activities': [], 'message': str(e)}), 500

@app.route('/api/parents/<int:student_id>', methods=['GET'])
def get_parents_by_student(student_id):
    """Get all parents for a specific student"""
    try:
        # Get student info from user_info table (not users)
        student_response = supabase.table('user_info')\
            .select('id, first_name, last_name')\
            .eq('id', student_id)\
            .execute()
        
        if not student_response.data:
            # Return empty parents list if student not found, rather than 404
            return jsonify({
                'success': True,
                'student': {
                    'student_id': student_id,
                    'first_name': '',
                    'last_name': ''
                },
                'parents': []
            })
        
        student = student_response.data[0]
        
        # Get parents data
        parents_response = supabase.table('parents')\
            .select('*')\
            .eq('student_id', student_id)\
            .order('relationship')\
            .execute()
        
        parents = parents_response.data
        
        return jsonify({
            'success': True,
            'student': {
                'student_id': student['id'],
                'first_name': student['first_name'],
                'last_name': student['last_name']
            },
            'parents': parents
        })
        
    except Exception as e:
        print(f"Error fetching parents: {str(e)}")
        return jsonify({'success': False, 'message': 'Failed to fetch parent information'}), 500

# Add new parent
@app.route('/api/parents', methods=['POST'])
def add_parent():
    try:
        data = request.get_json()
        
        # Validate required fields
        required_fields = ['student_id', 'relationship', 'first_name', 'last_name', 'gender', 'email', 'mobile_no']
        for field in required_fields:
            if not data.get(field):
                return jsonify({'success': False, 'message': f'Missing required field: {field}'}), 400
        
        # Insert parent
        response = supabase.table('parents').insert({
            'student_id': data['student_id'],
            'relationship': data['relationship'],
            'first_name': data['first_name'],
            'middle_name': data.get('middle_name', ''),
            'last_name': data['last_name'],
            'gender': data['gender'],
            'email': data['email'],
            'mobile_no': data['mobile_no'],
            'occupation': data.get('occupation', ''),
            'created_at': 'now()'
        }).execute()
        
        if response.data:
            return jsonify({'success': True, 'message': 'Parent added successfully', 'parent': response.data[0]})
        else:
            return jsonify({'success': False, 'message': 'Failed to add parent'}), 500
            
    except Exception as e:
        print(f"Error adding parent: {str(e)}")
        return jsonify({'success': False, 'message': 'Failed to add parent'}), 500

# Update parent
@app.route('/api/parents/<int:parent_id>', methods=['PUT'])
def update_parent(parent_id):
    try:
        data = request.get_json()
        
        # Validate required fields
        required_fields = ['relationship', 'first_name', 'last_name', 'gender', 'email', 'mobile_no']
        for field in required_fields:
            if not data.get(field):
                return jsonify({'success': False, 'message': f'Missing required field: {field}'}), 400
        
        # Update parent
        response = supabase.table('parents').update({
            'relationship': data['relationship'],
            'first_name': data['first_name'],
            'middle_name': data.get('middle_name', ''),
            'last_name': data['last_name'],
            'gender': data['gender'],
            'email': data['email'],
            'mobile_no': data['mobile_no'],
            'occupation': data.get('occupation', '')
        }).eq('parent_id', parent_id).execute()
        
        if response.data:
            return jsonify({'success': True, 'message': 'Parent updated successfully'})
        else:
            return jsonify({'success': False, 'message': 'Parent not found'}), 404
            
    except Exception as e:
        print(f"Error updating parent: {str(e)}")
        return jsonify({'success': False, 'message': 'Failed to update parent'}), 500

# Delete parent
@app.route('/api/parents/<int:parent_id>', methods=['DELETE'])
def delete_parent(parent_id):
    try:
        response = supabase.table('parents').delete().eq('parent_id', parent_id).execute()
        
        if response.data:
            return jsonify({'success': True, 'message': 'Parent deleted successfully'})
        else:
            return jsonify({'success': False, 'message': 'Parent not found'}), 404
            
    except Exception as e:
        print(f"Error deleting parent: {str(e)}")
        return jsonify({'success': False, 'message': 'Failed to delete parent'}), 500

# ==================== ACTIVITY FILE & LINK ATTACHMENT FUNCTIONS ====================

def upload_file_to_supabase(file, task_id, teacher_id):
    """Upload a file to Supabase Storage and return file metadata as JSON string"""
    try:
        if not file or file.filename == '':
            return None
        
        # Secure filename
        original_filename = file.filename
        filename = secure_filename(file.filename)
        timestamp = datetime.now().timestamp()
        
        # Create unique filename: task-files/teacher_{teacher_id}_{timestamp}_{filename}
        file_path = f"task-files/teacher_{teacher_id}_{timestamp}_{filename}"
        
        # Read file content
        file_content = file.read()
        file_size = len(file_content)
        
        # Determine file type
        file_ext = original_filename.rsplit('.', 1)[-1].lower() if '.' in original_filename else 'unknown'
        
        # Map extension to file type
        file_type_map = {
            'pdf': 'pdf',
            'doc': 'word', 'docx': 'word',
            'xls': 'excel', 'xlsx': 'excel',
            'ppt': 'powerpoint', 'pptx': 'powerpoint',
            'jpg': 'image', 'jpeg': 'image', 'png': 'image', 'gif': 'image',
            'mp4': 'video', 'mov': 'video', 'avi': 'video',
            'mp3': 'audio', 'wav': 'audio'
        }
        file_type = file_type_map.get(file_ext, 'file')
        
        # Upload to Supabase Storage bucket 'task-files'
        try:
            response = supabase.storage.from_('task-files').upload(
                file_path,
                file_content,
                file_options={"content-type": file.content_type}
            )
        except Exception as upload_err:
            print(f"Upload error: {str(upload_err)}")
            return None
        
        # Get public URL
        public_url = supabase.storage.from_('task-files').get_public_url(file_path)
        
        # Return as JSON string matching the expected format
        attachment_json = {
            "file_url": public_url,
            "original_filename": original_filename,
            "file_type": file_type,
            "file_size": file_size
        }
        
        return json.dumps(attachment_json)
    except Exception as e:
        print(f"Error uploading file to Supabase: {str(e)}")
        return None


@app.route('/create_activity', methods=['POST'])
def create_activity():
    """Create a new activity with file and link attachments"""
    if 'user_id' not in session or session.get('role') != 'Teacher':
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401
    
    try:
        # Get form data
        teacher_id = session['user_id']
        grade_level = request.form.get('grade_level')
        section = request.form.get('section')
        task_title = request.form.get('task')
        points = int(request.form.get('points', 0))
        due_date = request.form.get('due_date')
        priority = request.form.get('priority', 'medium')
        description = request.form.get('description', '')
        template = request.form.get('template', 'default')
        links_json = request.form.get('links', '[]')
        requested_school_year = request.form.get('school_year')
        requested_quarter = request.form.get('quarter')
        
        # Validate required fields
        if not all([grade_level, section, task_title]):
            return jsonify({'success': False, 'message': 'Missing required fields: grade level, section, or activity title'}), 400
        
        # Parse links with better error handling
        links = []
        if links_json:
            try:
                links = json.loads(links_json)
                if not isinstance(links, list):
                    links = []
            except json.JSONDecodeError as e:
                logger.warning(f"Failed to parse links JSON: {str(e)}")
                links = []
        
        # Normalize links to ensure consistent format
        links = normalize_links(links)
        
        # Generate activity group ID
        activity_group_id = str(uuid.uuid4())

        # Resolve school year and quarter (prefer server-side authoritative values)
        resolved_school_year = get_current_school_year() or requested_school_year
        current_quarter = get_current_quarter()
        resolved_quarter = None
        if current_quarter and isinstance(current_quarter, dict):
            resolved_quarter = current_quarter.get('quarter_name')
        if not resolved_quarter:
            resolved_quarter = requested_quarter

        # Block activity creation when there is no active quarter yet
        if not current_quarter:
            upcoming_quarter = None
            try:
                upcoming_query = supabase.table('quarters').select('*').order('start_date', desc=False)
                if resolved_school_year:
                    upcoming_query = upcoming_query.eq('school_year', resolved_school_year)
                upcoming_result = upcoming_query.execute()

                if upcoming_result.data:
                    now_dt = datetime.now(philippines_tz).date()
                    for quarter in upcoming_result.data:
                        start_str = str(quarter.get('start_date', '')).split('T')[0]
                        if not start_str:
                            continue
                        try:
                            start_dt = datetime.strptime(start_str, '%Y-%m-%d').date()
                        except ValueError:
                            continue
                        if start_dt > now_dt:
                            days_away = (start_dt - now_dt).days
                            upcoming_quarter = {
                                'quarter_name': quarter.get('quarter_name', 'Upcoming Quarter'),
                                'days_away': days_away
                            }
                            break
            except Exception as quarter_lookup_error:
                app.logger.warning(f"Unable to resolve upcoming quarter: {str(quarter_lookup_error)}")

            if upcoming_quarter:
                return jsonify({
                    'success': False,
                    'message': f"{upcoming_quarter['quarter_name']} has not started yet ({upcoming_quarter['days_away']} day{'s' if upcoming_quarter['days_away'] != 1 else ''} away). You cannot create an activity yet."
                }), 400

            return jsonify({
                'success': False,
                'message': 'No active quarter right now. You cannot create an activity until a quarter starts.'
            }), 400
        
        # Get all students in the classroom
        try:
            students_result = supabase.table('user_info') \
                .select('id') \
                .eq('role', 'Student') \
                .eq('year_level', grade_level) \
                .eq('section', section) \
                .execute()
            
            students = students_result.data if students_result.data else []
        except Exception as db_err:
            logger.error(f"Database error fetching students: {str(db_err)}")
            return jsonify({'success': False, 'message': f'Database error: {str(db_err)}'}), 500
        
        if not students:
            return jsonify({'success': False, 'message': 'No students found in this classroom'}), 400
        
        # Process file uploads for Supabase
        file_attachments = []  # Text array of JSON strings
        if 'files' in request.files:
            files = request.files.getlist('files')
            for idx, file in enumerate(files):
                if file and file.filename != '':
                    try:
                        # Create a temporary task_id for file path organization
                        temp_task_id = f"temp_{activity_group_id}_{idx}"
                        file_json_str = upload_file_to_supabase(file, temp_task_id, teacher_id)
                        if file_json_str:
                            file_attachments.append(file_json_str)  # Append JSON string directly
                    except Exception as file_err:
                        logger.error(f"Error uploading file {idx}: {str(file_err)}")
                        # Continue with other files, don't fail the entire activity creation
        
        # Convert links to JSON strings for storage
        links_data = []
        if links:
            try:
                links_data = [json.dumps(link) if isinstance(link, dict) else link for link in links]
            except Exception as err:
                logger.warning(f"Error processing links: {str(err)}")
                links_data = []
        
        # Create task assignments for all students
        assignments = []
        for student in students:
            assignments.append({
                'teacher_id': teacher_id,
                'student_id': student['id'],
                'grade_level': grade_level,
                'section': section,
                'task': task_title,
                'points': points,
                'description': description,
                'due_date': due_date if due_date else None,
                'priority': priority,
                'template': template,
                'status': 'Assigned',
                'attachments': file_attachments if file_attachments else None,  # Text array of JSON strings
                'links': links_data if links_data else None,  # Text array of JSON strings
                'activity_group_id': activity_group_id,
                'school_year': resolved_school_year,
                'quarter': resolved_quarter
            })
        
        # Insert assignments
        if assignments:
            logger.info(f"Creating activity {activity_group_id} for teacher {teacher_id}: {task_title} ({len(assignments)} students)")
            
            try:
                result = supabase.table('task_assignments').insert(assignments).execute()
                
                if hasattr(result, 'error') and result.error:
                    return jsonify({'success': False, 'message': str(result.error)}), 400
            except Exception as db_err:
                logger.error(f"Database insert error: {str(db_err)}")
                return jsonify({'success': False, 'message': f'Database error: {str(db_err)}'}), 500
            
            # Log activity
            try:
                record_admin_activity_log(
                    user_id=teacher_id,
                    action='CREATE_ACTIVITY',
                    activity='Activity Created',
                    description=f"Created activity '{task_title}' for Grade {grade_level} - Section {section} with {len(assignments)} students. Files: {len(file_attachments)}, Links: {len(links_data)}",
                    user_role='Teacher'
                )
            except Exception as log_err:
                logger.error(f"Failed to log activity: {str(log_err)}")
                # Don't fail the activity creation if logging fails
            
            return jsonify({
                'success': True,
                'message': 'Activity created successfully',
                'activity_group_id': activity_group_id,
                'files_uploaded': len(file_attachments),
                'links_attached': len(links_data),
                'students_assigned': len(assignments)
            }), 200
        else:
            return jsonify({'success': False, 'message': 'Failed to create activity - no students to assign'}), 400
            
    except Exception as e:
        logger.error(f"Error creating activity: {str(e)}", exc_info=True)
        return jsonify({'success': False, 'message': f'Failed to create activity: {str(e)}'}), 500


@app.route('/update_activity', methods=['POST'])
def update_activity():
    """Update an existing activity with file and link attachments"""
    if 'user_id' not in session or session.get('role') != 'Teacher':
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401
    
    try:
        # Get form data
        teacher_id = session['user_id']
        activity_group_id = request.form.get('activity_group_id')
        grade_level = request.form.get('grade_level')
        section = request.form.get('section')
        task_title = request.form.get('task')
        points = int(request.form.get('points', 0))
        due_date = request.form.get('due_date')
        priority = request.form.get('priority', 'medium')
        description = request.form.get('description', '')
        template = request.form.get('template', 'default')
        links_json = request.form.get('links', '[]')
        existing_files_json = request.form.get('existing_files', '[]')
        
        # Validate required fields
        if not activity_group_id:
            return jsonify({'success': False, 'message': 'Activity group ID is required'}), 400
        
        if not all([grade_level, section, task_title]):
            return jsonify({'success': False, 'message': 'Missing required fields: grade level, section, or activity title'}), 400
        
        # Parse JSON data with better error handling
        links = []
        if links_json:
            try:
                links = json.loads(links_json)
                if not isinstance(links, list):
                    links = []
            except json.JSONDecodeError as e:
                logger.warning(f"Failed to parse links JSON: {str(e)}")
                links = []
        
        # Normalize links to ensure consistent format
        links = normalize_links(links)
        
        existing_files = []
        if existing_files_json:
            try:
                existing_files = json.loads(existing_files_json)
                if not isinstance(existing_files, list):
                    existing_files = []
            except json.JSONDecodeError as e:
                logger.warning(f"Failed to parse existing files JSON: {str(e)}")
                existing_files = []
        
        # Process new file uploads - returns JSON strings
        new_file_attachments = []
        if 'files' in request.files:
            files = request.files.getlist('files')
            for idx, file in enumerate(files):
                if file and file.filename != '':
                    try:
                        temp_task_id = f"update_{activity_group_id}_{idx}"
                        file_json_str = upload_file_to_supabase(file, temp_task_id, teacher_id)
                        if file_json_str:
                            new_file_attachments.append(file_json_str)
                    except Exception as file_err:
                        logger.error(f"Error uploading file {idx}: {str(file_err)}")
                        # Continue with other files, don't fail the entire update
        
        # Convert existing files objects to JSON strings for database storage
        existing_files_json_strings = []
        for file_obj in existing_files:
            try:
                if isinstance(file_obj, dict):
                    # Build a standardized attachment object
                    attachment = {
                        "file_url": file_obj.get('file_url') or file_obj.get('url'),
                        "original_filename": file_obj.get('original_filename') or file_obj.get('name'),
                        "file_type": file_obj.get('file_type') or file_obj.get('type'),
                        "file_size": file_obj.get('file_size') or file_obj.get('size', 0)
                    }
                    existing_files_json_strings.append(json.dumps(attachment))
                elif isinstance(file_obj, str):
                    # Already a JSON string
                    existing_files_json_strings.append(file_obj)
            except Exception as err:
                logger.warning(f"Error processing existing file: {str(err)}")
                # Skip this file and continue
        
        # Combine existing and new files (text array of JSON strings)
        all_file_attachments = existing_files_json_strings + new_file_attachments
        
        # Convert links to JSON strings for storage
        links_data = []
        if links:
            try:
                links_data = [json.dumps(link) if isinstance(link, dict) else link for link in links]
            except Exception as err:
                logger.warning(f"Error processing links: {str(err)}")
                links_data = []
        
        # Update all task assignments with this activity_group_id
        update_data = {
            'task': task_title,
            'points': points,
            'description': description,
            'due_date': due_date if due_date else None,
            'priority': priority,
            'template': template,
            'attachments': all_file_attachments if all_file_attachments else None,  # Text array of JSON strings
            'links': links_data if links_data else None  # Text array of JSON strings
        }
        
        logger.info(f"Updating activity {activity_group_id} for teacher {teacher_id}: {task_title}")
        
        try:
            result = supabase.table('task_assignments') \
                .update(update_data) \
                .eq('activity_group_id', activity_group_id) \
                .execute()
        except Exception as db_err:
            logger.error(f"Database update error: {str(db_err)}")
            return jsonify({'success': False, 'message': f'Database error: {str(db_err)}'}), 500
        
        if result.data:
            # Log activity
            try:
                record_admin_activity_log(
                    user_id=teacher_id,
                    action='UPDATE_ACTIVITY',
                    activity='Activity Updated',
                    description=f"Updated activity '{task_title}'. Files: {len(all_file_attachments)}, Links: {len(links_data)}",
                    user_role='Teacher'
                )
            except Exception as log_err:
                logger.error(f"Failed to log activity: {str(log_err)}")
                # Don't fail the update if logging fails
            
            return jsonify({
                'success': True,
                'message': 'Activity updated successfully',
                'total_files': len(all_file_attachments),
                'new_files': len(new_file_attachments),
                'total_links': len(links_data)
            }), 200
        else:
            return jsonify({'success': False, 'message': 'Activity not found or no changes made'}), 404
            
    except Exception as e:
        logger.error(f"Error updating activity: {str(e)}", exc_info=True)
        return jsonify({'success': False, 'message': f'Failed to update activity: {str(e)}'}), 500


@app.route('/api/activity-files/<activity_group_id>', methods=['GET'])
def get_activity_files(activity_group_id):
    """Get all files and links for an activity"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401
    
    try:
        # Get the activity details
        result = supabase.table('task_assignments') \
            .select('attachments, links') \
            .eq('activity_group_id', activity_group_id) \
            .limit(1) \
            .execute()
        
        if result.data:
            activity = result.data[0]
            return jsonify({
                'success': True,
                'files': deserialize_attachment_array(activity.get('attachments', [])),
                'links': deserialize_link_array(activity.get('links', []))
            }), 200
        else:
            return jsonify({'success': False, 'message': 'Activity not found'}), 404
            
    except Exception as e:
        print(f"Error retrieving activity files: {str(e)}")
        return jsonify({'success': False, 'message': 'Failed to retrieve files'}), 500

if __name__ == '__main__':
    app.run(debug=True)